#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成面相流式接口文档并追加到Word文档
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_interface_doc(doc, interface_num, interface_path, interface_alias, method, description, remarks):
    title = doc.add_paragraph()
    title.add_run(f"{interface_num}. {interface_path} ({interface_alias})").bold = True
    doc.add_paragraph("接口详情")
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Light Grid Accent 1'
    details = [
        ("接口路径", interface_path),
        ("接口别名", interface_alias),
        ("请求方法", method),
        ("接口描述", description),
        ("备注", remarks)
    ]
    for i, (label, value) in enumerate(details):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = value
    doc.add_paragraph()

def add_request_params_table(doc, params):
    doc.add_paragraph("请求参数")
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = "字段名"
    header_cells[1].text = "类型"
    header_cells[2].text = "必填"
    header_cells[3].text = "描述"
    header_cells[4].text = "示例"
    for param in params:
        row_cells = table.add_row().cells
        row_cells[0].text = param['name']
        row_cells[1].text = param['type']
        row_cells[2].text = param['required']
        row_cells[3].text = param['description']
        row_cells[4].text = param.get('example', '')
    doc.add_paragraph()

def add_response_table(doc, response_fields):
    doc.add_paragraph("响应格式")
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = "字段名"
    header_cells[1].text = "类型"
    header_cells[2].text = "描述"
    for field in response_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field['name']
        row_cells[1].text = field['type']
        row_cells[2].text = field['description']
    doc.add_paragraph()

# 打开现有文档
doc = Document('/Users/zhoudt/Desktop/66666.docx')
doc.add_page_break()

# 面相流式接口
add_interface_doc(doc, 4, "/api/v2/face/analyze/stream", "AI智能面相分析V2-流式", "POST",
    "上传人脸照片进行AI智能面相分析，先返回基础分析数据，然后流式返回大模型整合分析结果。使用Server-Sent Events (SSE)格式实时推送分析内容",
    "流式接口，响应格式为text/event-stream。先返回完整的基础分析数据（type: 'data'），然后流式返回大模型整合分析（type: 'progress'），最后返回完成信号（type: 'complete'）")

# 请求参数
face_stream_request_params = [
    {'name': 'image', 'type': 'File (multipart/form-data)', 'required': '是', 'description': '面相照片文件（必填），支持JPG、PNG格式', 'example': 'face.jpg'},
    {'name': 'analysis_types', 'type': 'str', 'required': '否', 'description': '分析类型（可选），逗号分隔，默认 "gongwei,liuqin,shishen"。可选值：gongwei(宫位)、liuqin(六亲)、shishen(十神)、features(特征)', 'example': 'gongwei,liuqin,shishen'},
    {'name': 'birth_year', 'type': 'int', 'required': '否', 'description': '出生年份（可选），用于结合八字分析', 'example': '1990'},
    {'name': 'birth_month', 'type': 'int', 'required': '否', 'description': '出生月份（可选），用于结合八字分析', 'example': '1'},
    {'name': 'birth_day', 'type': 'int', 'required': '否', 'description': '出生日期（可选），用于结合八字分析', 'example': '15'},
    {'name': 'birth_hour', 'type': 'int', 'required': '否', 'description': '出生时辰（可选），用于结合八字分析，范围0-23', 'example': '12'},
    {'name': 'gender', 'type': 'str', 'required': '否', 'description': '性别（可选），用于结合八字分析。可选值：male(男)、female(女)', 'example': 'male'},
    {'name': 'bot_id', 'type': 'str', 'required': '否', 'description': 'Bot ID（可选），用于指定使用的大模型Bot，默认从数据库配置读取', 'example': '7597406985550282787'}
]
add_request_params_table(doc, face_stream_request_params)

# 响应格式说明
doc.add_paragraph("响应说明")
doc.add_paragraph("本接口使用Server-Sent Events (SSE)格式，响应类型为text/event-stream。响应数据以'data: '开头的行，每行包含一个JSON对象。")
doc.add_paragraph()

# 响应字段
face_stream_response_fields = [
    {'name': 'type: data', 'type': 'object', 'description': '基础分析数据（第一个返回的数据块），包含完整的面相分析结果'},
    {'name': 'type: data.content.success', 'type': 'bool', 'description': '是否成功'},
    {'name': 'type: data.content.data', 'type': 'dict', 'description': '基础分析结果数据，结构与普通接口相同'},
    {'name': 'type: data.content.data.face_detected', 'type': 'bool', 'description': '是否检测到人脸'},
    {'name': 'type: data.content.data.landmarks', 'type': 'dict', 'description': '关键点信息，包含MediaPipe关键点数量和映射特征点数量'},
    {'name': 'type: data.content.data.santing_analysis', 'type': 'dict', 'description': '三停分析结果，包含上停、中停、下停的比例和评价'},
    {'name': 'type: data.content.data.wuyan_analysis', 'type': 'dict', 'description': '五眼分析结果'},
    {'name': 'type: data.content.data.gongwei_analysis', 'type': 'list', 'description': '十三宫位分析列表'},
    {'name': 'type: data.content.data.liuqin_analysis', 'type': 'list', 'description': '六亲分析列表'},
    {'name': 'type: data.content.data.shishen_analysis', 'type': 'list', 'description': '十神分析列表'},
    {'name': 'type: data.content.data.overall_summary', 'type': 'str', 'description': '综合面相分析总结'},
    {'name': 'type: data.content.data.birth_info', 'type': 'dict', 'description': '生辰信息（如果提供了生辰参数）'},
    {'name': 'type: progress', 'type': 'object', 'description': '流式LLM输出数据块（多个），实时推送大模型生成的分析内容'},
    {'name': 'type: progress.content', 'type': 'str', 'description': 'LLM生成的分析文本片段（增量内容）'},
    {'name': 'type: progress.statusText', 'type': 'str', 'description': '状态文本（可选），用于显示当前处理状态'},
    {'name': 'type: complete', 'type': 'object', 'description': '完成信号（最后一个数据块），表示流式输出已完成'},
    {'name': 'type: complete.content', 'type': 'str', 'description': '完整的LLM分析内容（可选），包含所有流式输出的完整文本'},
    {'name': 'type: error', 'type': 'object', 'description': '错误信息（如果发生错误）'},
    {'name': 'type: error.content', 'type': 'str', 'description': '错误描述信息'}
]
add_response_table(doc, face_stream_response_fields)

# 添加使用示例
doc.add_paragraph("使用示例")
doc.add_paragraph("1. 使用curl调用（multipart/form-data）：")
example_code = doc.add_paragraph()
example_code.add_run('curl -X POST "http://localhost:8001/api/v2/face/analyze/stream" \\\n  -F "image=@face.jpg" \\\n  -F "analysis_types=gongwei,liuqin,shishen" \\\n  -F "birth_year=1990" \\\n  -F "birth_month=1" \\\n  -F "birth_day=15" \\\n  -F "gender=male"')
example_code.style = 'No Spacing'
example_code_font = example_code.runs[0].font
example_code_font.name = 'Courier New'
example_code_font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph("2. 前端JavaScript调用示例：")
js_example = doc.add_paragraph()
js_example.add_run('const formData = new FormData();\nformData.append("image", file);\nformData.append("analysis_types", "gongwei,liuqin,shishen");\n\nconst response = await fetch("/api/v2/face/analyze/stream", {\n  method: "POST",\n  body: formData\n});\n\nconst reader = response.body.getReader();\nconst decoder = new TextDecoder();\nlet buffer = "";\n\nwhile (true) {\n  const { done, value } = await reader.read();\n  if (done) break;\n  \n  buffer += decoder.decode(value, { stream: true });\n  const lines = buffer.split("\\n");\n  buffer = lines.pop() || "";\n  \n  for (const line of lines) {\n    if (line.startsWith("data: ")) {\n      const data = JSON.parse(line.substring(6));\n      if (data.type === "data") {\n        // 处理基础分析数据\n        console.log("基础数据:", data.content);\n      } else if (data.type === "progress") {\n        // 处理流式LLM输出\n        console.log("LLM内容:", data.content);\n      } else if (data.type === "complete") {\n        // 流式输出完成\n        console.log("完成");\n      }\n    }\n  }\n}')
js_example.style = 'No Spacing'
js_example_font = js_example.runs[0].font
js_example_font.name = 'Courier New'
js_example_font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph("注意事项：")
notes = [
    "1. 本接口使用SSE格式，需要客户端支持流式读取",
    "2. 响应数据以'data: '开头，每行一个JSON对象",
    "3. 第一个数据块type为'data'，包含完整的基础分析结果",
    "4. 后续数据块type为'progress'，包含流式LLM输出片段",
    "5. 最后一个数据块type为'complete'，表示流式输出完成",
    "6. 如果发生错误，会返回type为'error'的数据块",
    "7. 基础分析数据与普通接口(/api/v2/face/analyze)返回的数据结构相同",
    "8. LLM整合分析基于基础分析数据生成，提供更深入的综合解读"
]
for note in notes:
    doc.add_paragraph(note, style='List Bullet')

doc.save('/Users/zhoudt/Desktop/66666.docx')
print("✅ 完成！面相流式接口文档已追加到文档")
