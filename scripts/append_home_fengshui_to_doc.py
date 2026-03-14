#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将居家风水接口文档追加到 66666.docx，格式与现有接口一致"""

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

DOC_PATH = '/Users/zhoudt/Desktop/公司/other/66666.docx'

def append_home_fengshui_section(doc: Document):
    # 分隔线
    doc.add_paragraph('============================================================')
    doc.add_paragraph()

    # 接口详情
    doc.add_paragraph('接口详情3. /api/v2/home-fengshui/analyze/stream（居家风水流式分析）')
    doc.add_paragraph('接口路径：/api/v2/home-fengshui/analyze/stream')
    doc.add_paragraph('接口别名：居家风水流式分析')
    doc.add_paragraph('请求方法：POST')
    doc.add_paragraph(
        '接口描述：上传 1-4 张房间照片，支持多房间（客厅/卧室/书房等）自动识别或手动指定，'
        'SSE 流式返回视觉识别、规则匹配、命卦分析及 LLM 风水报告。支持八字命卦个性化。'
    )
    doc.add_paragraph('请求参数：')
    doc.add_paragraph()

    # 请求参数表（5列：字段名、类型、必填、描述、示例）
    req_table = doc.add_table(rows=9, cols=5)
    req_table.style = 'Table Grid'
    req_headers = ['字段名', '类型', '必填', '描述', '示例']
    req_rows = [
        ['photos', 'File[]', '是', '房间照片，1-4 张', 'multipart 文件'],
        ['room_types', 'list[str]', '否', '每张照片对应的房间类型，与 photos 一一对应；留空或 auto 则自动识别', 'living_room,bedroom,study'],
        ['room_type', 'str', '否', '[兼容] 单一房间类型，多张照片时建议用 room_types', 'bedroom'],
        ['door_direction', 'str', '否', '大门朝向：北/东北/东/东南/南/西南/西/西北', '南'],
        ['solar_date', 'str', '否', '出生日期（用于命卦）', '1990-05-15'],
        ['solar_time', 'str', '否', '出生时间', '08:30'],
        ['gender', 'str', '否', '性别：male/female', 'male'],
        ['bot_id', 'str', '否', '百炼报告智能体 App ID，默认从配置读取', ''],
    ]
    for c, h in enumerate(req_headers):
        req_table.rows[0].cells[c].text = h
    for r, row in enumerate(req_rows, 1):
        for c, val in enumerate(row):
            req_table.rows[r].cells[c].text = val

    doc.add_paragraph()
    doc.add_paragraph('响应格式：SSE 流式（text/event-stream），每行 data: {JSON}')
    doc.add_paragraph()

    # 响应/SSE 事件表（3列：字段名、类型、描述）
    resp_table = doc.add_table(rows=21, cols=3)
    resp_table.style = 'Table Grid'
    resp_headers = ['字段名', '类型', '描述']
    resp_rows = [
        ['type', 'str', '事件类型，见下表'],
        ['request_id', 'str', 'type=request_id 时，请求唯一ID'],
        ['content', 'str', 'type=progress_msg/error/room_full_report 时，文本内容'],
        ['room_index', 'int', 'type=room_result/room_annotated_image/room_progress 时，房间序号(0-based)'],
        ['room_type', 'str', '房间类型：bedroom/living_room/study/kitchen/dining_room'],
        ['room_label', 'str', '房间中文名'],
        ['auto_detected', 'bool', '是否自动识别房间类型'],
        ['content', 'object', 'type=room_result 时，结构化分析结果'],
        ['content.furnitures', 'array', '家具列表，含 name/label/position_zone/state/element'],
        ['content.critical_issues', 'array', '风水问题'],
        ['content.suggestions', 'array', '改善建议'],
        ['content.tips', 'array', '小贴士'],
        ['content.overall_score', 'int', '该房间综合评分 0-100'],
        ['content.mingua_info', 'object', '命卦信息（含 mingua_name/mingua_type）'],
        ['content.summary', 'str', '简要总结'],
        ['image_base64', 'str', 'type=room_annotated_image 时，标注图 base64'],
        ['overall_score', 'int', 'type=room_annotated_image/home_score 时，评分'],
        ['room_count', 'int', 'type=home_score 时，房间数量'],
        ['room_scores', 'array', 'type=home_score 时，各房间评分 [{room_index,room_type,room_label,score}]'],
    ]
    for c, h in enumerate(resp_headers):
        resp_table.rows[0].cells[c].text = h
    for r, row in enumerate(resp_rows, 1):
        for c, val in enumerate(row):
            resp_table.rows[r].cells[c].text = val

    doc.add_paragraph()
    doc.add_paragraph('SSE 事件类型：request_id | progress_msg | room_result | room_annotated_image | mingua_result | room_progress | room_complete | room_full_report | home_score | error')
    doc.add_paragraph()
    doc.add_paragraph('示例：')
    doc.add_paragraph(
        'curl -X POST "http://localhost:8001/api/v2/home-fengshui/analyze/stream" '
        '-F "photos=@客厅.png" -F "photos=@卧室.png" -F "photos=@书房.png" '
        '-F "room_types=living_room" -F "room_types=bedroom" -F "room_types=study" '
        '-F "door_direction=南" -F "solar_date=1990-05-15" -F "gender=male" -N'
    )
    doc.add_paragraph()


def main():
    doc = Document(DOC_PATH)
    append_home_fengshui_section(doc)
    doc.save(DOC_PATH)
    print(f'✅ 已追加居家风水接口文档到 {DOC_PATH}')


if __name__ == '__main__':
    main()
