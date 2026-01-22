#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证支付接口文档内容
"""

from docx import Document

def verify_doc(doc_path):
    """验证文档内容"""
    doc = Document(doc_path)
    
    print("=" * 60)
    print("支付接口文档验证报告")
    print("=" * 60)
    print(f"\n文档路径: {doc_path}")
    print(f"段落总数: {len(doc.paragraphs)}")
    print(f"表格总数: {len(doc.tables)}")
    print("\n" + "-" * 60)
    
    # 查找支付接口段落
    payment_paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if 'payment' in text.lower() or '统一' in text or '支付' in text:
            payment_paragraphs.append((i+1, text))
    
    if payment_paragraphs:
        print("\n✅ 找到支付接口相关段落:")
        for para_num, text in payment_paragraphs:
            print(f"  段落 {para_num}: {text}")
    else:
        print("\n❌ 未找到支付接口相关段落")
    
    # 查找支付接口表格
    print("\n" + "-" * 60)
    print("\n支付接口相关表格:")
    
    payment_tables = []
    for idx, table in enumerate(doc.tables):
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            first_cell = table.rows[0].cells[0].text
            
            # 检查是否是接口详情表格
            if '接口路径' in first_cell:
                if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
                    path = table.rows[1].cells[1].text
                    if 'payment' in path.lower():
                        payment_tables.append((idx+1, '接口详情', path))
            
            # 检查是否是参数表格
            elif first_cell == '字段名':
                if len(table.rows) > 1 and len(table.rows[1].cells) > 0:
                    first_param = table.rows[1].cells[0].text
                    if first_param == 'provider':
                        payment_tables.append((idx+1, '请求参数', f'{len(table.rows)-1}个参数'))
                    elif first_param == 'success':
                        payment_tables.append((idx+1, '响应格式', f'{len(table.rows)-1}个字段'))
    
    if payment_tables:
        print("\n✅ 找到支付接口相关表格:")
        for table_num, table_type, info in payment_tables:
            print(f"  表格 {table_num}: {table_type} - {info}")
    else:
        print("\n❌ 未找到支付接口相关表格")
    
    # 详细显示接口1的内容
    print("\n" + "=" * 60)
    print("\n接口1: /payment/unified/create 详细信息")
    print("=" * 60)
    
    # 查找接口1的接口详情表格（应该是表格13）
    if len(doc.tables) >= 13:
        table = doc.tables[12]  # 索引12
        if len(table.rows) > 0 and '接口路径' in table.rows[0].cells[0].text:
            if len(table.rows) > 1 and 'payment/unified/create' in table.rows[1].cells[1].text:
                print("\n接口详情:")
                for i in range(1, len(table.rows)):
                    if len(table.rows[i].cells) >= 2:
                        label = table.rows[i].cells[0].text
                        value = table.rows[i].cells[1].text
                        print(f"  {label}: {value}")
    
    # 查找接口1的请求参数表格（应该是表格14）
    if len(doc.tables) >= 14:
        table = doc.tables[13]  # 索引13
        if len(table.rows) > 0 and table.rows[0].cells[0].text == '字段名':
            if len(table.rows) > 1 and table.rows[1].cells[0].text == 'provider':
                print(f"\n请求参数 ({len(table.rows)-1}个):")
                for i in range(1, min(len(table.rows), 6)):  # 显示前5个
                    if len(table.rows[i].cells) >= 1:
                        param = table.rows[i].cells[0].text
                        print(f"  {i}. {param}")
                if len(table.rows) > 6:
                    print(f"  ... 还有 {len(table.rows)-6} 个参数")
    
    # 查找接口1的响应格式表格（应该是表格15）
    if len(doc.tables) >= 15:
        table = doc.tables[14]  # 索引14
        if len(table.rows) > 0 and table.rows[0].cells[0].text == '字段名':
            if len(table.rows) > 1 and table.rows[1].cells[0].text == 'success':
                print(f"\n响应格式 ({len(table.rows)-1}个字段):")
                for i in range(1, min(len(table.rows), 6)):  # 显示前5个
                    if len(table.rows[i].cells) >= 1:
                        field = table.rows[i].cells[0].text
                        print(f"  {i}. {field}")
                if len(table.rows) > 6:
                    print(f"  ... 还有 {len(table.rows)-6} 个字段")
    
    print("\n" + "=" * 60)
    print("\n✅ 验证完成！")
    print("\n如果Word文档中没有显示，请尝试：")
    print("1. 关闭并重新打开Word文档")
    print("2. 在Word中按 Ctrl+F 搜索 'payment' 或 '统一'")
    print("3. 滚动到文档末尾查看")
    print("=" * 60)

if __name__ == '__main__':
    doc_path = '/Users/zhoudt/Desktop/66666.docx'
    verify_doc(doc_path)
