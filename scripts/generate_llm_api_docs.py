#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成大模型接口参数文档（完整版）
包含统一接口和完整响应参数
将所有给大模型提供参数的接口整理成Word文档
"""

import sys
import os
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)


# 统一接口定义
UNIFIED_INTERFACE = {
    'name': '统一数据获取接口',
    'path': '/api/v1/bazi/data',
    'method': 'POST',
    'description': '统一数据获取接口，通过配置参数按需获取数据，支持并行计算、多级缓存和性能优化',
    'request_fields': [
        {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
        {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
        {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
        {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
        {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
        {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
        {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
        {'name': 'modules', 'type': 'dict', 'required': '可选', 'description': '数据模块配置，例如: {"dayun": {"mode": "current_with_neighbors"}, "rules": {"types": ["shishen"]}}', 'example': '{}'},
        {'name': 'use_cache', 'type': 'boolean', 'required': '可选', 'description': '是否使用缓存，默认True', 'example': 'True'},
        {'name': 'parallel', 'type': 'boolean', 'required': '可选', 'description': '是否并行获取数据，默认True', 'example': 'True'},
        {'name': 'timeout', 'type': 'integer', 'required': '可选', 'description': '超时时间（秒），默认30', 'example': '30'},
        {'name': 'verify_consistency', 'type': 'boolean', 'required': '可选', 'description': '是否验证数据一致性（与前端页面数据对比），默认True', 'example': 'True'},
    ],
    'response_fields': [
        # 响应结构顶层字段
        {'path': 'success', 'type': 'boolean', 'description': '是否成功', 'source': 'BaziDataResponse模型'},
        {'path': 'data', 'type': 'dict', 'description': '数据内容（包含所有请求模块的数据）', 'source': 'BaziDataOrchestrator.fetch_data()'},
        {'path': 'message', 'type': 'string', 'description': '消息', 'source': 'BaziDataResponse模型'},
        {'path': 'error', 'type': 'string', 'description': '错误信息', 'source': 'BaziDataResponse模型'},
        {'path': 'validation_errors', 'type': 'list', 'description': '数据一致性验证错误列表', 'source': 'BaziDataValidator.validate_consistency()'},
        
        # data字段下的模块（展开所有嵌套字段）
        # bazi模块
        {'path': 'data.bazi', 'type': 'dict', 'description': '基础八字数据', 'source': 'BaziService.calculate_bazi_full()'},
        {'path': 'data.bazi.basic_info', 'type': 'dict', 'description': '基本信息', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.basic_info.solar_date', 'type': 'string', 'description': '阳历日期', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.basic_info.solar_time', 'type': 'string', 'description': '出生时间', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.basic_info.lunar_date', 'type': 'dict', 'description': '农历日期', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.basic_info.gender', 'type': 'string', 'description': '性别', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars', 'type': 'dict', 'description': '四柱八字', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.year', 'type': 'dict', 'description': '年柱', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.year.stem', 'type': 'string', 'description': '年干', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.year.branch', 'type': 'string', 'description': '年支', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.month', 'type': 'dict', 'description': '月柱', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.month.stem', 'type': 'string', 'description': '月干', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.month.branch', 'type': 'string', 'description': '月支', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.day', 'type': 'dict', 'description': '日柱', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.day.stem', 'type': 'string', 'description': '日干', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.day.branch', 'type': 'string', 'description': '日支', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.hour', 'type': 'dict', 'description': '时柱', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.hour.stem', 'type': 'string', 'description': '时干', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.bazi_pillars.hour.branch', 'type': 'string', 'description': '时支', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.details', 'type': 'dict', 'description': '详细信息（包含各柱的详细信息）', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.ten_gods_stats', 'type': 'dict', 'description': '十神统计', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.elements', 'type': 'dict', 'description': '五行元素', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.element_counts', 'type': 'dict', 'description': '五行计数', 'source': 'BaziService._format_bazi_result()'},
        {'path': 'data.bazi.relationships', 'type': 'dict', 'description': '关系信息（包含地支关系等）', 'source': 'BaziService._format_bazi_result()'},
        
        # wangshuai模块
        {'path': 'data.wangshuai', 'type': 'dict', 'description': '旺衰数据', 'source': 'WangShuaiService.calculate_wangshuai()'},
        {'path': 'data.wangshuai.success', 'type': 'boolean', 'description': '是否成功', 'source': 'WangShuaiService.calculate_wangshuai()'},
        {'path': 'data.wangshuai.data', 'type': 'dict', 'description': '旺衰分析结果', 'source': 'WangShuaiAnalyzer.analyze()'},
        {'path': 'data.wangshuai.data.wangshuai', 'type': 'string', 'description': '旺衰结果（如：身旺、身弱等）', 'source': 'WangShuaiAnalyzer.analyze()'},
        {'path': 'data.wangshuai.data.total_score', 'type': 'float', 'description': '总分', 'source': 'WangShuaiAnalyzer.analyze()'},
        {'path': 'data.wangshuai.data.tiaohou', 'type': 'dict', 'description': '调候信息', 'source': 'WangShuaiAnalyzer.calculate_tiaohou()'},
        
        # xishen_jishen模块
        {'path': 'data.xishen_jishen', 'type': 'dict', 'description': '喜神忌神数据', 'source': 'get_xishen_jishen()'},
        
        # wuxing模块
        {'path': 'data.wuxing', 'type': 'dict', 'description': '五行数据', 'source': '从bazi数据中提取'},
        {'path': 'data.wuxing.elements', 'type': 'dict', 'description': '五行元素', 'source': 'data.bazi.elements'},
        {'path': 'data.wuxing.element_counts', 'type': 'dict', 'description': '五行计数', 'source': 'data.bazi.element_counts'},
        
        # dayun模块
        {'path': 'data.dayun', 'type': 'list', 'description': '大运列表', 'source': 'BaziDetailService.calculate_detail_full() -> dayun_sequence'},
        {'path': 'data.dayun[].step', 'type': 'integer', 'description': '大运步数', 'source': 'BaziCalculator计算'},
        {'path': 'data.dayun[].ganzhi', 'type': 'string', 'description': '大运干支', 'source': 'BaziCalculator计算'},
        {'path': 'data.dayun[].year_start', 'type': 'integer', 'description': '起始年份', 'source': 'BaziCalculator计算'},
        {'path': 'data.dayun[].year_end', 'type': 'integer', 'description': '结束年份', 'source': 'BaziCalculator计算'},
        {'path': 'data.dayun[].age_range', 'type': 'dict', 'description': '年龄范围', 'source': 'BaziCalculator计算'},
        {'path': 'data.dayun[].liunian_sequence', 'type': 'list', 'description': '该大运下的流年序列', 'source': '从全局流年序列中提取'},
        
        # liunian模块
        {'path': 'data.liunian', 'type': 'list', 'description': '流年列表', 'source': 'BaziDetailService.calculate_detail_full() -> liunian_sequence'},
        {'path': 'data.liunian[].year', 'type': 'integer', 'description': '流年年份', 'source': 'BaziCalculator计算'},
        {'path': 'data.liunian[].ganzhi', 'type': 'string', 'description': '流年干支', 'source': 'BaziCalculator计算'},
        
        # liuyue模块
        {'path': 'data.liuyue', 'type': 'list', 'description': '流月列表', 'source': 'BaziDetailService.calculate_detail_full() -> liuyue_sequence'},
        
        # special_liunians模块
        {'path': 'data.special_liunians', 'type': 'dict', 'description': '特殊流年数据', 'source': 'SpecialLiunianService.get_special_liunians_batch()'},
        {'path': 'data.special_liunians.list', 'type': 'list', 'description': '原始列表格式', 'source': 'SpecialLiunianService.get_special_liunians_batch()'},
        {'path': 'data.special_liunians.by_dayun', 'type': 'dict', 'description': '按大运分组格式', 'source': '按dayun_step分组'},
        {'path': 'data.special_liunians.formatted', 'type': 'string', 'description': '格式化后的提示词', 'source': 'SpecialLiunianService.format_special_liunians_for_prompt()'},
        
        # rules模块
        {'path': 'data.rules', 'type': 'list', 'description': '匹配的规则列表', 'source': 'RuleService.match_rules()'},
        {'path': 'data.rules[].id', 'type': 'integer', 'description': '规则ID', 'source': '数据库'},
        {'path': 'data.rules[].type', 'type': 'string', 'description': '规则类型', 'source': '数据库'},
        {'path': 'data.rules[].content', 'type': 'string', 'description': '规则内容', 'source': '数据库'},
        {'path': 'data.rules[].match_score', 'type': 'float', 'description': '匹配分数', 'source': 'RuleService.match_rules()'},
        
        # health模块
        {'path': 'data.health', 'type': 'dict', 'description': '健康分析数据', 'source': 'HealthAnalysisService.analyze()'},
        
        # personality模块
        {'path': 'data.personality', 'type': 'dict', 'description': '性格分析数据', 'source': 'RizhuGenderAnalyzer.analyze_rizhu_gender()'},
        
        # rizhu模块
        {'path': 'data.rizhu', 'type': 'dict', 'description': '日柱六十甲子数据', 'source': 'RizhuLiujiaziService.get_rizhu_analysis()'},
        
        # wuxing_proportion模块
        {'path': 'data.wuxing_proportion', 'type': 'dict', 'description': '五行占比数据', 'source': 'WuxingProportionService.calculate_proportion()'},
        
        # liunian_enhanced模块
        {'path': 'data.liunian_enhanced', 'type': 'dict', 'description': '流年增强分析数据', 'source': 'LiunianEnhancedService.analyze_liunian_enhanced()'},
        
        # detail模块
        {'path': 'data.detail', 'type': 'dict', 'description': '详细八字数据（包含所有大运流年信息）', 'source': 'BaziDetailService.calculate_detail_full()'},
        {'path': 'data.detail.dayun_sequence', 'type': 'list', 'description': '完整大运序列', 'source': 'BaziCalculator计算'},
        {'path': 'data.detail.liunian_sequence', 'type': 'list', 'description': '完整流年序列', 'source': 'BaziCalculator计算'},
        {'path': 'data.detail.liuyue_sequence', 'type': 'list', 'description': '流月序列', 'source': 'BaziCalculator计算'},
        {'path': 'data.detail.matched_rules', 'type': 'list', 'description': '匹配的规则', 'source': 'RuleService.match_rules()'},
        {'path': 'data.detail.wangshuai', 'type': 'dict', 'description': '旺衰数据', 'source': 'WangShuaiService.calculate_wangshuai()'},
        {'path': 'data.detail.wuxing_proportion', 'type': 'dict', 'description': '五行占比数据', 'source': 'WuxingProportionService.calculate_proportion()'},
        {'path': 'data.detail.rizhu_liujiazi', 'type': 'dict', 'description': '日柱六十甲子数据', 'source': 'RizhuLiujiaziService.get_rizhu_analysis()'},
        
        # fortune_display模块
        {'path': 'data.fortune_display', 'type': 'dict', 'description': '大运流年流月统一展示数据', 'source': 'BaziDisplayService.get_fortune_display()'},
        
        # 其他模块
        {'path': 'data.deities', 'type': 'dict', 'description': '神煞数据', 'source': '从bazi.details中提取'},
        {'path': 'data.branch_relations', 'type': 'dict', 'description': '地支关系数据', 'source': '从bazi.relationships中提取'},
        {'path': 'data.career_star', 'type': 'list', 'description': '事业星数据', 'source': '从bazi.ten_gods中提取'},
        {'path': 'data.wealth_star', 'type': 'list', 'description': '财富星数据', 'source': '从bazi.ten_gods中提取'},
        {'path': 'data.children_star', 'type': 'list', 'description': '子女星数据', 'source': '从bazi.ten_gods中提取'},
        {'path': 'data.shengong_minggong', 'type': 'dict', 'description': '身宫命宫数据', 'source': 'BaziInterfaceService.generate_interface_full()'},
        {'path': 'data.daily_fortune', 'type': 'dict', 'description': '每日运势数据', 'source': 'DailyFortuneService.calculate_daily_fortune()'},
        {'path': 'data.monthly_fortune', 'type': 'dict', 'description': '每月运势数据', 'source': 'MonthlyFortuneService.calculate_monthly_fortune()'},
        {'path': 'data.daily_fortune_calendar', 'type': 'dict', 'description': '每日运势日历数据', 'source': 'DailyFortuneCalendarService.get_daily_fortune_calendar()'},
        {'path': 'data.yigua', 'type': 'dict', 'description': '易卦数据', 'source': 'YiGuaService.divinate()'},
        {'path': 'data.bazi_interface', 'type': 'dict', 'description': '八字界面数据', 'source': 'BaziInterfaceService.generate_interface_full()'},
        {'path': 'data.bazi_ai', 'type': 'dict', 'description': '八字AI分析数据', 'source': 'BaziAIService.analyze_bazi_with_ai()'},
    ],
}

# 接口定义列表（手动定义所有参数，确保完整）
INTERFACES = [
    {
        'name': 'LLM生成完整报告',
        'path': '/api/v1/bazi/llm-generate',
        'method': 'POST',
        'description': '使用 LLM 直接生成完整的命理报告（类似 FateTell 的实时生成模式）',
        'stream': False,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'user_question', 'type': 'string', 'required': '可选', 'description': '用户问题或分析需求（可选）', 'example': '我想了解我的事业运势'},
            {'name': 'access_token', 'type': 'string', 'required': '可选', 'description': 'Coze Access Token（可选）', 'example': ''},
            {'name': 'bot_id', 'type': 'string', 'required': '可选', 'description': 'Coze Bot ID（可选）', 'example': ''},
            {'name': 'api_base', 'type': 'string', 'required': '可选', 'description': 'Coze API 基础URL（可选）', 'example': ''},
        ],
        'response_fields': [
            {'path': 'success', 'type': 'boolean', 'description': '是否成功', 'source': 'LLMGenerateResponse模型'},
            {'path': 'report', 'type': 'string', 'description': 'LLM生成的完整报告', 'source': 'Coze Bot API响应 -> analysis字段'},
            {'path': 'bazi_data', 'type': 'dict', 'description': '八字数据', 'source': '统一接口 -> data.bazi'},
            {'path': 'prompt_length', 'type': 'integer', 'description': '提示词长度', 'source': 'len(prompt)'},
            {'path': 'report_length', 'type': 'integer', 'description': '报告长度', 'source': 'len(report)'},
            {'path': 'error', 'type': 'string', 'description': '错误信息', 'source': 'LLMGenerateResponse模型'},
            {'path': 'error_detail', 'type': 'string', 'description': '原始错误信息（供调试）', 'source': 'LLMGenerateResponse模型'},
            {'path': 'suggestion', 'type': 'string', 'description': '建议的替代方案', 'source': 'LLMGenerateResponse模型'},
        ],
        'frontend': ['index.html'],
        'bot_id_config': 'COZE_BOT_ID (环境变量)',
    },
    {
        'name': 'Coze AI分析八字',
        'path': '/api/v1/bazi/ai-analyze',
        'method': 'POST',
        'description': '调用八字接口获取数据，然后使用Coze AI进行分析',
        'stream': False,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
            {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
            {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
            {'name': 'user_question', 'type': 'string', 'required': '可选', 'description': '用户的问题或分析需求', 'example': '请分析我的财运和事业'},
            {'name': 'access_token', 'type': 'string', 'required': '可选', 'description': 'Coze Access Token，如果不提供则使用环境变量 COZE_ACCESS_TOKEN', 'example': 'pat_...'},
            {'name': 'bot_id', 'type': 'string', 'required': '可选', 'description': 'Coze Bot ID，如果不提供则使用环境变量 COZE_BOT_ID', 'example': '1234567890'},
            {'name': 'api_base', 'type': 'string', 'required': '可选', 'description': 'Coze API 基础URL，默认 https://api.coze.cn/v1', 'example': 'https://api.coze.cn/v1'},
            {'name': 'include_rizhu_analysis', 'type': 'boolean', 'required': '可选', 'description': '是否包含日柱性别分析结果', 'example': 'True'},
        ],
        'response_fields': [
            {'path': 'success', 'type': 'boolean', 'description': '是否成功', 'source': 'BaziAIResponse模型'},
            {'path': 'bazi_data', 'type': 'dict', 'description': '八字数据', 'source': '统一接口 -> data.bazi'},
            {'path': 'ai_analysis', 'type': 'dict', 'description': 'AI分析结果', 'source': 'Coze Bot API响应'},
            {'path': 'ai_analysis.success', 'type': 'boolean', 'description': 'AI分析是否成功', 'source': 'BaziAIAnalyzer._call_coze_api()'},
            {'path': 'ai_analysis.analysis', 'type': 'string', 'description': 'AI分析内容', 'source': 'Coze Bot API响应 -> analysis字段'},
            {'path': 'rizhu_analysis', 'type': 'string', 'description': '日柱性别分析结果', 'source': 'RizhuGenderAnalyzer.get_formatted_output()'},
            {'path': 'polished_rules', 'type': 'string', 'description': '大模型润色后的规则内容', 'source': 'BaziAIAnalyzer.polish_character_analysis()'},
            {'path': 'polished_rules_info', 'type': 'dict', 'description': '润色前后的对比信息', 'source': 'BaziAIAnalyzer.polish_character_analysis()'},
            {'path': 'polished_rules_info.original', 'type': 'string', 'description': '原始内容', 'source': 'BaziAIAnalyzer.polish_character_analysis()'},
            {'path': 'polished_rules_info.polished', 'type': 'string', 'description': '润色后内容', 'source': 'BaziAIAnalyzer.polish_character_analysis()'},
            {'path': 'polished_rules_info.changes', 'type': 'list', 'description': '修改列表', 'source': 'BaziAIAnalyzer.polish_character_analysis()'},
            {'path': 'polished_rules_info.changes_count', 'type': 'integer', 'description': '修改数量', 'source': 'BaziAIAnalyzer.polish_character_analysis()'},
            {'path': 'error', 'type': 'string', 'description': '错误信息', 'source': 'BaziAIResponse模型'},
        ],
        'frontend': ['index.html'],
        'bot_id_config': 'COZE_BOT_ID (环境变量)',
    },
    {
        'name': '创建新对话',
        'path': '/api/v1/bazi/chat/create',
        'method': 'POST',
        'description': '创建新的对话会话，保存用户的八字信息',
        'stream': False,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'user_name', 'type': 'string', 'required': '可选', 'description': '用户名称（可选）', 'example': ''},
        ],
        'response_fields': [
            {'path': 'success', 'type': 'boolean', 'description': '是否成功', 'source': 'ChatService.create_conversation()'},
            {'path': 'conversation_id', 'type': 'string', 'description': '对话ID', 'source': '生成的UUID'},
            {'path': 'error', 'type': 'string', 'description': '错误信息', 'source': 'ChatService.create_conversation()'},
        ],
        'frontend': ['qa-conversation.html', 'ai-qa.html'],
        'bot_id_config': 'COZE_BOT_ID (环境变量)',
    },
    {
        'name': '发送消息',
        'path': '/api/v1/bazi/chat/send',
        'method': 'POST',
        'description': '发送消息并获取 AI 回复',
        'stream': False,
        'fields': [
            {'name': 'conversation_id', 'type': 'string', 'required': '必填', 'description': '对话 ID', 'example': '550e8400-e29b-41d4-a716-446655440000'},
            {'name': 'user_message', 'type': 'string', 'required': '必填', 'description': '用户消息', 'example': '我想了解我的事业运势'},
            {'name': 'access_token', 'type': 'string', 'required': '可选', 'description': 'Coze Access Token（可选）', 'example': ''},
            {'name': 'bot_id', 'type': 'string', 'required': '可选', 'description': 'Coze Bot ID（可选）', 'example': ''},
            {'name': 'api_base', 'type': 'string', 'required': '可选', 'description': 'Coze API 基础URL（可选）', 'example': ''},
        ],
        'response_fields': [
            {'path': 'success', 'type': 'boolean', 'description': '是否成功', 'source': 'ChatService.chat()'},
            {'path': 'reply', 'type': 'string', 'description': 'AI回复内容', 'source': 'Coze Bot API响应 -> analysis字段'},
            {'path': 'conversation_id', 'type': 'string', 'description': '对话ID', 'source': 'ChatService.chat()'},
            {'path': 'message_count', 'type': 'integer', 'description': '消息数量', 'source': 'len(conversation["messages"])'},
            {'path': 'error', 'type': 'string', 'description': '错误信息', 'source': 'ChatService.chat()'},
        ],
        'frontend': ['qa-conversation.html', 'ai-qa.html'],
        'bot_id_config': 'COZE_BOT_ID (环境变量)',
    },
    {
        'name': '智能运势分析',
        'path': '/api/v1/smart-fortune/smart-analyze',
        'method': 'GET',
        'description': '自动识别用户问题意图，返回针对性的分析结果',
        'stream': False,
        'fields': [
            {'name': 'question', 'type': 'string', 'required': '必填', 'description': '用户问题', 'example': '我想了解我的事业运势'},
            {'name': 'year', 'type': 'integer', 'required': '必填', 'description': '出生年份', 'example': '1990'},
            {'name': 'month', 'type': 'integer', 'required': '必填', 'description': '出生月份', 'example': '5'},
            {'name': 'day', 'type': 'integer', 'required': '必填', 'description': '出生日期', 'example': '15'},
            {'name': 'hour', 'type': 'integer', 'required': '可选', 'description': '出生时辰（0-23）', 'example': '14'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别（male/female）', 'example': 'male'},
            {'name': 'user_id', 'type': 'string', 'required': '可选', 'description': '用户ID', 'example': ''},
            {'name': 'include_fortune_context', 'type': 'boolean', 'required': '可选', 'description': '是否包含流年大运分析（实验性功能，默认关闭）', 'example': 'false'},
        ],
        'response_fields': [
            {'path': 'success', 'type': 'boolean', 'description': '是否成功', 'source': 'SmartAnalyzeResponse模型'},
            {'path': 'intent', 'type': 'string', 'description': '识别的意图', 'source': 'IntentServiceClient.classify()'},
            {'path': 'analysis', 'type': 'string', 'description': '分析结果', 'source': 'FortuneLLMClient.analyze_fortune()'},
            {'path': 'bazi_data', 'type': 'dict', 'description': '八字数据', 'source': '统一接口 -> data.bazi'},
            {'path': 'rules', 'type': 'list', 'description': '匹配的规则', 'source': '统一接口 -> data.rules'},
            {'path': 'error', 'type': 'string', 'description': '错误信息', 'source': 'SmartAnalyzeResponse模型'},
        ],
        'frontend': ['smart-fortune.html'],
        'bot_id_config': 'FORTUNE_ANALYSIS_BOT_ID 或 COZE_BOT_ID (环境变量)',
    },
    {
        'name': '智能运势分析（流式）',
        'path': '/api/v1/smart-fortune/smart-analyze-stream',
        'method': 'GET',
        'description': '智能运势分析（流式输出版），支持两种场景：场景1（点击选择项）和场景2（点击预设问题/输入问题）',
        'stream': True,
        'fields': [
            {'name': 'question', 'type': 'string', 'required': '可选', 'description': '用户问题', 'example': '我想了解我的事业运势'},
            {'name': 'year', 'type': 'integer', 'required': '可选', 'description': '出生年份（场景1必填）', 'example': '1990'},
            {'name': 'month', 'type': 'integer', 'required': '可选', 'description': '出生月份（场景1必填）', 'example': '5'},
            {'name': 'day', 'type': 'integer', 'required': '可选', 'description': '出生日期（场景1必填）', 'example': '15'},
            {'name': 'hour', 'type': 'integer', 'required': '可选', 'description': '出生时辰（0-23）', 'example': '14'},
            {'name': 'gender', 'type': 'string', 'required': '可选', 'description': '性别（场景1必填，male/female）', 'example': 'male'},
            {'name': 'user_id', 'type': 'string', 'required': '必填', 'description': '用户ID', 'example': 'user123'},
            {'name': 'category', 'type': 'string', 'required': '必填', 'description': '分类名称（如：事业财富、婚姻、健康、子女、流年运势、年运报告）', 'example': '事业财富'},
        ],
        'response_fields': [
            {'path': 'SSE事件类型', 'type': 'string', 'description': '流式响应使用SSE格式', 'source': 'smart_fortune_stream()'},
            {'path': 'event: status', 'type': 'string', 'description': '状态更新事件', 'source': 'smart_fortune_stream()'},
            {'path': 'event: status.stage', 'type': 'string', 'description': '当前阶段（如：intent, bazi, rules, llm）', 'source': '性能监控'},
            {'path': 'event: status.message', 'type': 'string', 'description': '状态消息', 'source': '性能监控'},
            {'path': 'event: brief_response_start', 'type': 'string', 'description': '简要回答开始', 'source': 'FortuneLLMClient.generate_brief_response()'},
            {'path': 'event: brief_response_chunk', 'type': 'string', 'description': '简要回答内容块', 'source': 'FortuneLLMClient.generate_brief_response()'},
            {'path': 'event: brief_response_chunk.content', 'type': 'string', 'description': '简要回答内容', 'source': 'Coze Bot API响应'},
            {'path': 'event: brief_response_end', 'type': 'string', 'description': '简要回答结束', 'source': 'FortuneLLMClient.generate_brief_response()'},
            {'path': 'event: brief_response_end.content', 'type': 'string', 'description': '完整简要回答', 'source': 'Coze Bot API响应'},
            {'path': 'event: brief_response_error', 'type': 'string', 'description': '简要回答错误', 'source': 'FortuneLLMClient.generate_brief_response()'},
            {'path': 'event: brief_response_error.message', 'type': 'string', 'description': '错误信息', 'source': '错误详情'},
            {'path': 'event: preset_questions', 'type': 'string', 'description': '预设问题', 'source': 'FortuneLLMClient.generate_preset_questions()'},
            {'path': 'event: preset_questions.questions', 'type': 'list', 'description': '预设问题列表', 'source': 'Coze Bot API响应'},
            {'path': 'event: basic_analysis', 'type': 'string', 'description': '基础分析结果', 'source': '规则匹配结果'},
            {'path': 'event: basic_analysis.analysis', 'type': 'string', 'description': '基础分析内容', 'source': '规则匹配结果格式化'},
            {'path': 'event: llm_start', 'type': 'string', 'description': 'LLM分析开始', 'source': 'FortuneLLMClient.analyze_fortune()'},
            {'path': 'event: llm_chunk', 'type': 'string', 'description': 'LLM分析内容块', 'source': 'FortuneLLMClient.analyze_fortune()'},
            {'path': 'event: llm_chunk.content', 'type': 'string', 'description': 'LLM分析内容', 'source': 'Coze Bot API响应'},
            {'path': 'event: llm_end', 'type': 'string', 'description': 'LLM分析结束', 'source': 'FortuneLLMClient.analyze_fortune()'},
            {'path': 'event: llm_error', 'type': 'string', 'description': 'LLM分析错误', 'source': 'FortuneLLMClient.analyze_fortune()'},
            {'path': 'event: llm_error.message', 'type': 'string', 'description': '错误信息', 'source': '错误详情'},
            {'path': 'event: related_questions', 'type': 'string', 'description': '相关问题', 'source': 'FortuneLLMClient.generate_related_questions()'},
            {'path': 'event: related_questions.questions', 'type': 'list', 'description': '相关问题列表', 'source': 'Coze Bot API响应'},
            {'path': 'event: performance', 'type': 'string', 'description': '性能统计', 'source': '性能监控'},
            {'path': 'event: performance.summary', 'type': 'dict', 'description': '性能统计摘要', 'source': '性能监控'},
            {'path': 'event: error', 'type': 'string', 'description': '错误事件', 'source': '异常处理'},
            {'path': 'event: error.message', 'type': 'string', 'description': '错误信息', 'source': '异常详情'},
            {'path': 'event: error.performance', 'type': 'dict', 'description': '性能统计（错误时）', 'source': '性能监控'},
            {'path': 'event: end', 'type': 'string', 'description': '流结束', 'source': 'smart_fortune_stream()'},
        ],
        'frontend': ['smart-fortune-stream.html'],
        'bot_id_config': 'FORTUNE_ANALYSIS_BOT_ID 或 COZE_BOT_ID (环境变量)',
    },
    {
        'name': '婚姻分析（流式）',
        'path': '/api/v1/bazi/marriage-analysis/stream',
        'method': 'POST',
        'description': '流式生成感情婚姻分析，包含5个部分：命盘总论、配偶特征、感情走势、神煞点睛和建议方向',
        'stream': True,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
            {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
            {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
            {'name': 'bot_id', 'type': 'string', 'required': '可选', 'description': 'Coze Bot ID（可选，优先级：参数 > MARRIAGE_ANALYSIS_BOT_ID 环境变量）', 'example': ''},
        ],
        'response_fields': [
            {'path': 'SSE事件类型', 'type': 'string', 'description': '流式响应使用SSE格式', 'source': 'CozeStreamService.stream_custom_analysis()'},
            {'path': 'event: progress', 'type': 'string', 'description': '进度更新事件，包含分析内容', 'source': 'CozeStreamService.stream_custom_analysis() -> type: progress'},
            {'path': 'event: progress.content', 'type': 'string', 'description': '分析内容文本', 'source': 'Coze Bot API响应'},
            {'path': 'event: complete', 'type': 'string', 'description': '完成事件，包含完整分析', 'source': 'CozeStreamService.stream_custom_analysis() -> type: complete'},
            {'path': 'event: complete.content', 'type': 'string', 'description': '完整分析内容', 'source': 'Coze Bot API响应'},
            {'path': 'event: error', 'type': 'string', 'description': '错误事件', 'source': 'CozeStreamService.stream_custom_analysis() -> type: error'},
            {'path': 'event: error.content', 'type': 'string', 'description': '错误信息', 'source': '错误详情'},
        ],
        'frontend': ['marriage-analysis.html'],
        'bot_id_config': 'MARRIAGE_ANALYSIS_BOT_ID 或 COZE_BOT_ID (环境变量)',
    },
    {
        'name': '健康分析（流式）',
        'path': '/api/v1/health/stream',
        'method': 'POST',
        'description': '流式生成健康分析，包含4个部分：命盘体质总论、五行病理推演、大运流年健康警示、体质调理建议',
        'stream': True,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
            {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
            {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
            {'name': 'bot_id', 'type': 'string', 'required': '可选', 'description': 'Coze Bot ID（可选，默认使用环境变量配置）', 'example': ''},
        ],
        'response_fields': [
            {'path': 'SSE事件类型', 'type': 'string', 'description': '流式响应使用SSE格式', 'source': 'CozeStreamService.stream_custom_analysis()'},
            {'path': 'event: progress', 'type': 'string', 'description': '进度更新事件，包含分析内容', 'source': 'CozeStreamService.stream_custom_analysis() -> type: progress'},
            {'path': 'event: progress.content', 'type': 'string', 'description': '分析内容文本', 'source': 'Coze Bot API响应'},
            {'path': 'event: complete', 'type': 'string', 'description': '完成事件，包含完整分析', 'source': 'CozeStreamService.stream_custom_analysis() -> type: complete'},
            {'path': 'event: complete.content', 'type': 'string', 'description': '完整分析内容', 'source': 'Coze Bot API响应'},
            {'path': 'event: error', 'type': 'string', 'description': '错误事件', 'source': 'CozeStreamService.stream_custom_analysis() -> type: error'},
            {'path': 'event: error.content', 'type': 'string', 'description': '错误信息', 'source': '错误详情'},
        ],
        'frontend': ['health-analysis.html'],
        'bot_id_config': 'HEALTH_ANALYSIS_BOT_ID 或 COZE_BOT_ID (环境变量)',
    },
    {
        'name': '子女学习分析（流式）',
        'path': '/api/v1/children-study/stream',
        'method': 'POST',
        'description': '流式生成子女学习分析',
        'stream': True,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
            {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
            {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
            {'name': 'bot_id', 'type': 'string', 'required': '可选', 'description': 'Coze Bot ID（可选，默认使用环境变量配置）', 'example': ''},
        ],
        'response_fields': [
            {'path': 'SSE事件类型', 'type': 'string', 'description': '流式响应使用SSE格式', 'source': 'CozeStreamService.stream_custom_analysis()'},
            {'path': 'event: progress', 'type': 'string', 'description': '进度更新事件，包含分析内容', 'source': 'CozeStreamService.stream_custom_analysis() -> type: progress'},
            {'path': 'event: progress.content', 'type': 'string', 'description': '分析内容文本', 'source': 'Coze Bot API响应'},
            {'path': 'event: complete', 'type': 'string', 'description': '完成事件，包含完整分析', 'source': 'CozeStreamService.stream_custom_analysis() -> type: complete'},
            {'path': 'event: complete.content', 'type': 'string', 'description': '完整分析内容', 'source': 'Coze Bot API响应'},
            {'path': 'event: error', 'type': 'string', 'description': '错误事件', 'source': 'CozeStreamService.stream_custom_analysis() -> type: error'},
            {'path': 'event: error.content', 'type': 'string', 'description': '错误信息', 'source': '错误详情'},
        ],
        'frontend': ['children-study-analysis.html'],
        'bot_id_config': 'COZE_BOT_ID (环境变量)',
    },
    {
        'name': '总评分析（流式）',
        'path': '/api/v1/general-review/stream',
        'method': 'POST',
        'description': '流式生成总评分析',
        'stream': True,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
            {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
            {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
            {'name': 'bot_id', 'type': 'string', 'required': '可选', 'description': 'Coze Bot ID（可选，默认使用环境变量配置）', 'example': ''},
        ],
        'response_fields': [
            {'path': 'SSE事件类型', 'type': 'string', 'description': '流式响应使用SSE格式', 'source': 'CozeStreamService.stream_custom_analysis()'},
            {'path': 'event: progress', 'type': 'string', 'description': '进度更新事件，包含分析内容', 'source': 'CozeStreamService.stream_custom_analysis() -> type: progress'},
            {'path': 'event: progress.content', 'type': 'string', 'description': '分析内容文本', 'source': 'Coze Bot API响应'},
            {'path': 'event: complete', 'type': 'string', 'description': '完成事件，包含完整分析', 'source': 'CozeStreamService.stream_custom_analysis() -> type: complete'},
            {'path': 'event: complete.content', 'type': 'string', 'description': '完整分析内容', 'source': 'Coze Bot API响应'},
            {'path': 'event: error', 'type': 'string', 'description': '错误事件', 'source': 'CozeStreamService.stream_custom_analysis() -> type: error'},
            {'path': 'event: error.content', 'type': 'string', 'description': '错误信息', 'source': '错误详情'},
        ],
        'frontend': ['general-review-analysis.html'],
        'bot_id_config': 'GENERAL_REVIEW_BOT_ID 或 COZE_BOT_ID (环境变量)',
    },
    {
        'name': '五行占比分析',
        'path': '/api/v1/bazi/wuxing-proportion',
        'method': 'POST',
        'description': '查询五行占比分析（包含LLM分析功能）',
        'stream': False,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
            {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
            {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
        ],
        'response_fields': [
            {'path': 'success', 'type': 'boolean', 'description': '是否成功', 'source': 'WuxingProportionResponse模型'},
            {'path': 'data', 'type': 'dict', 'description': '五行占比数据', 'source': 'WuxingProportionService.calculate_proportion()'},
            {'path': 'data.proportions', 'type': 'dict', 'description': '五行占比', 'source': 'WuxingProportionService.calculate_proportion()'},
            {'path': 'data.llm_analysis', 'type': 'string', 'description': 'LLM分析结果', 'source': 'Coze Bot API响应'},
            {'path': 'error', 'type': 'string', 'description': '错误信息', 'source': 'WuxingProportionResponse模型'},
        ],
        'frontend': ['wuxing-proportion.html'],
        'bot_id_config': 'WUXING_PROPORTION_BOT_ID (硬编码: 7585498208202473523)',
    },
    {
        'name': '喜神忌神分析（流式）',
        'path': '/api/v1/bazi/xishen-jishen/stream',
        'method': 'POST',
        'description': '流式生成喜神忌神大模型分析',
        'stream': True,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
            {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
            {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
            {'name': 'bot_id', 'type': 'string', 'required': '可选', 'description': 'Coze Bot ID（可选）', 'example': ''},
        ],
        'response_fields': [
            {'path': 'SSE事件类型', 'type': 'string', 'description': '流式响应使用SSE格式', 'source': 'CozeStreamService.stream_custom_analysis()'},
            {'path': 'event: progress', 'type': 'string', 'description': '进度更新事件，包含分析内容', 'source': 'CozeStreamService.stream_custom_analysis() -> type: progress'},
            {'path': 'event: progress.content', 'type': 'string', 'description': '分析内容文本', 'source': 'Coze Bot API响应'},
            {'path': 'event: complete', 'type': 'string', 'description': '完成事件，包含完整分析', 'source': 'CozeStreamService.stream_custom_analysis() -> type: complete'},
            {'path': 'event: complete.content', 'type': 'string', 'description': '完整分析内容', 'source': 'Coze Bot API响应'},
            {'path': 'event: error', 'type': 'string', 'description': '错误事件', 'source': 'CozeStreamService.stream_custom_analysis() -> type: error'},
            {'path': 'event: error.content', 'type': 'string', 'description': '错误信息', 'source': '错误详情'},
        ],
        'frontend': ['xishen-jishen.html'],
        'bot_id_config': 'COZE_BOT_ID (环境变量)',
    },
    {
        'name': '开始QA对话',
        'path': '/api/v1/qa/start',
        'method': 'POST',
        'description': '开始新对话',
        'stream': False,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
            {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
            {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
            {'name': 'user_id', 'type': 'string', 'required': '可选', 'description': '用户ID', 'example': 'user123'},
        ],
        'response_fields': [
            {'path': 'success', 'type': 'boolean', 'description': '是否成功', 'source': 'QAConversationService.start_conversation()'},
            {'path': 'session_id', 'type': 'string', 'description': '会话ID', 'source': '生成的UUID'},
            {'path': 'initial_question', 'type': 'string', 'description': '初始问题', 'source': 'QAConversationService.start_conversation()'},
            {'path': 'categories', 'type': 'list', 'description': '分类列表', 'source': 'QAConversationService.start_conversation()'},
            {'path': 'error', 'type': 'string', 'description': '错误信息', 'source': 'QAConversationService.start_conversation()'},
        ],
        'frontend': ['qa-conversation.html', 'ai-qa.html'],
        'bot_id_config': 'QA_BOT_ID 或 COZE_BOT_ID (环境变量)',
    },
    {
        'name': 'QA提问（流式）',
        'path': '/api/v1/qa/ask',
        'method': 'POST',
        'description': '提问并获取流式回答',
        'stream': True,
        'fields': [
            {'name': 'session_id', 'type': 'string', 'required': '必填', 'description': '会话 ID', 'example': '550e8400-e29b-41d4-a716-446655440000'},
            {'name': 'question', 'type': 'string', 'required': '必填', 'description': '用户问题', 'example': '我想了解一下我的事业运势'},
        ],
        'response_fields': [
            {'path': 'SSE事件类型', 'type': 'string', 'description': '流式响应使用SSE格式', 'source': 'QAConversationService.ask_question_stream()'},
            {'path': 'event: progress', 'type': 'string', 'description': '进度更新事件', 'source': 'CozeStreamService.stream_custom_analysis()'},
            {'path': 'event: complete', 'type': 'string', 'description': '完成事件', 'source': 'CozeStreamService.stream_custom_analysis()'},
            {'path': 'event: error', 'type': 'string', 'description': '错误事件', 'source': 'CozeStreamService.stream_custom_analysis()'},
        ],
        'frontend': ['qa-conversation.html', 'ai-qa.html'],
        'bot_id_config': 'QA_BOT_ID 或 COZE_BOT_ID (环境变量)',
    },
    {
        'name': '事业财富分析（流式）',
        'path': '/api/v1/career-wealth/stream',
        'method': 'POST',
        'description': '流式生成事业财富分析',
        'stream': True,
        'fields': [
            {'name': 'solar_date', 'type': 'string', 'required': '必填', 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'string', 'required': '必填', 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'string', 'required': '必填', 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'string', 'required': '可选', 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar'},
            {'name': 'location', 'type': 'string', 'required': '可选', 'description': '出生地点（用于时区转换，优先级1）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': '可选', 'description': '纬度（用于时区转换，优先级2）', 'example': '39.90'},
            {'name': 'longitude', 'type': 'float', 'required': '可选', 'description': '经度（用于时区转换和真太阳时计算，优先级2）', 'example': '116.40'},
            {'name': 'bot_id', 'type': 'string', 'required': '可选', 'description': 'Coze Bot ID（可选，默认使用环境变量配置）', 'example': ''},
        ],
        'response_fields': [
            {'path': 'SSE事件类型', 'type': 'string', 'description': '流式响应使用SSE格式', 'source': 'CozeStreamService.stream_custom_analysis()'},
            {'path': 'event: progress', 'type': 'string', 'description': '进度更新事件，包含分析内容', 'source': 'CozeStreamService.stream_custom_analysis() -> type: progress'},
            {'path': 'event: progress.content', 'type': 'string', 'description': '分析内容文本', 'source': 'Coze Bot API响应'},
            {'path': 'event: complete', 'type': 'string', 'description': '完成事件，包含完整分析', 'source': 'CozeStreamService.stream_custom_analysis() -> type: complete'},
            {'path': 'event: complete.content', 'type': 'string', 'description': '完整分析内容', 'source': 'Coze Bot API响应'},
            {'path': 'event: error', 'type': 'string', 'description': '错误事件', 'source': 'CozeStreamService.stream_custom_analysis() -> type: error'},
            {'path': 'event: error.content', 'type': 'string', 'description': '错误信息', 'source': '错误详情'},
        ],
        'frontend': ['career-wealth-analysis.html'],
        'bot_id_config': 'CAREER_WEALTH_BOT_ID 或 COZE_BOT_ID (环境变量)',
    },
]


def create_word_document():
    """创建Word文档"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # 添加标题
    title = doc.add_heading('HiFate-bazi 大模型接口参数文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    
    # 添加文档信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f'文档生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}').font.size = Pt(12)
    
    doc.add_paragraph('')
    doc.add_paragraph('本文档详细列出了项目中所有给大模型提供参数的接口，包括每个字段的详细说明、类型、是否必填、示例值以及对应的前端界面和服务。')
    doc.add_paragraph('')
    
    # 添加目录占位
    doc.add_heading('目录', 1)
    doc.add_paragraph('（注：请在Word中使用"引用"->"目录"->"自动目录1"来生成目录）')
    doc.add_page_break()
    
    # 添加概述
    doc.add_heading('1. 概述', 1)
    doc.add_paragraph('本文档包含以下类型的接口：')
    
    categories = [
        ('LLM生成模式接口', ['/api/v1/bazi/llm-generate', '/api/v1/bazi/ai-analyze']),
        ('对话接口', ['/api/v1/bazi/chat/create', '/api/v1/bazi/chat/send']),
        ('智能运势分析接口', ['/api/v1/smart-fortune/smart-analyze', '/api/v1/smart-fortune/smart-analyze-stream']),
        ('专项分析接口（流式）', [
            '/api/v1/bazi/marriage-analysis/stream',
            '/api/v1/health/stream',
            '/api/v1/children-study/stream',
            '/api/v1/general-review/stream',
            '/api/v1/career-wealth/stream',
            '/api/v1/bazi/wuxing-proportion',
            '/api/v1/bazi/xishen-jishen/stream',
        ]),
        ('QA对话接口', ['/api/v1/qa/start', '/api/v1/qa/ask']),
    ]
    
    for cat_name, paths in categories:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(cat_name)
        run.bold = True
        run.font.size = Pt(12)
        for path in paths:
            doc.add_paragraph('  • ' + path, style='List Bullet 2')
    
    doc.add_paragraph('')
    
    # 添加统一接口章节
    doc.add_heading('2. 统一数据获取接口', 1)
    section_num = 3
    
    # 统一接口基本信息
    doc.add_heading('2.1 接口基本信息', 2)
    
    unified_info_table = doc.add_table(rows=4, cols=2)
    unified_info_table.style = 'Light Grid Accent 1'
    unified_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    unified_info_table.cell(0, 0).text = '接口路径'
    unified_info_table.cell(0, 1).text = UNIFIED_INTERFACE['path']
    unified_info_table.cell(1, 0).text = '请求方法'
    unified_info_table.cell(1, 1).text = UNIFIED_INTERFACE['method']
    unified_info_table.cell(2, 0).text = '是否流式'
    unified_info_table.cell(2, 1).text = '否'
    unified_info_table.cell(3, 0).text = '接口描述'
    unified_info_table.cell(3, 1).text = UNIFIED_INTERFACE['description']
    
    # 设置表格样式
    for row_idx, row in enumerate(unified_info_table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
                    if row_idx == 0:  # 表头加粗
                        run.font.bold = True
    
    doc.add_paragraph('')
    
    # 统一接口请求参数
    doc.add_heading('2.2 请求参数', 2)
    
    unified_request_fields = UNIFIED_INTERFACE.get('request_fields', [])
    if unified_request_fields:
        unified_param_table = doc.add_table(rows=1, cols=5)
        unified_param_table.style = 'Light Grid Accent 1'
        unified_param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_cells = unified_param_table.rows[0].cells
        headers = ['字段名', '类型', '必填/可选', '描述', '示例']
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
        
        # 添加数据行
        for field in unified_request_fields:
            row_cells = unified_param_table.add_row().cells
            row_cells[0].text = field['name']
            row_cells[1].text = field['type']
            row_cells[2].text = field['required']
            row_cells[3].text = field.get('description', '')
            row_cells[4].text = str(field.get('example', ''))
            
            # 设置字体
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)
    
    doc.add_paragraph('')
    
    # 统一接口响应参数
    doc.add_heading('2.3 响应参数', 2)
    doc.add_paragraph('统一接口返回的数据结构包含所有请求模块的数据，以下是完整的响应字段列表（展开所有嵌套字段）：')
    doc.add_paragraph('')
    
    unified_response_fields = UNIFIED_INTERFACE.get('response_fields', [])
    if unified_response_fields:
        unified_response_table = doc.add_table(rows=1, cols=4)
        unified_response_table.style = 'Light Grid Accent 1'
        unified_response_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_cells = unified_response_table.rows[0].cells
        headers = ['字段路径', '类型', '描述', '来源']
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
        
        # 添加数据行
        for field in unified_response_fields:
            row_cells = unified_response_table.add_row().cells
            row_cells[0].text = field['path']
            row_cells[1].text = field['type']
            row_cells[2].text = field.get('description', '')
            row_cells[3].text = field.get('source', '')
            
            # 设置字体
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)
    
    doc.add_paragraph('')
    doc.add_page_break()
    
    # 为每个接口添加详细章节
    for interface in INTERFACES:
        doc.add_heading(f'{section_num}. {interface["name"]}', 1)
        section_num += 1
        
        # 接口基本信息
        doc.add_heading(f'{section_num - 2}.1 接口基本信息', 2)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_table.cell(0, 0).text = '接口路径'
        info_table.cell(0, 1).text = interface['path']
        info_table.cell(1, 0).text = '请求方法'
        info_table.cell(1, 1).text = interface['method']
        info_table.cell(2, 0).text = '是否流式'
        info_table.cell(2, 1).text = '是' if interface.get('stream') else '否'
        info_table.cell(3, 0).text = '接口描述'
        info_table.cell(3, 1).text = interface['description']
        
        # 设置表格样式
        for row_idx, row in enumerate(info_table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)
                        if row_idx == 0:  # 表头加粗
                            run.font.bold = True
        
        doc.add_paragraph('')
        
        # 请求参数
        doc.add_heading(f'{section_num - 2}.2 请求参数', 2)
        
        fields = interface.get('fields', [])
        if fields:
            # 创建参数表格
            param_table = doc.add_table(rows=1, cols=5)
            param_table.style = 'Light Grid Accent 1'
            param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            header_cells = param_table.rows[0].cells
            headers = ['字段名', '类型', '必填/可选', '描述', '示例']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)
            
            # 添加数据行
            for field in fields:
                row_cells = param_table.add_row().cells
                row_cells[0].text = field['name']
                row_cells[1].text = field['type']
                row_cells[2].text = field['required']
                row_cells[3].text = field.get('description', '')
                row_cells[4].text = str(field.get('example', ''))
                
                # 设置字体
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.size = Pt(10.5)
        else:
            doc.add_paragraph('（该接口无请求参数）')
        
        doc.add_paragraph('')
        
        # 前端界面
        doc.add_heading(f'{section_num - 2}.3 前端界面', 2)
        frontend_pages = interface.get('frontend', [])
        if frontend_pages:
            for page in frontend_pages:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(page)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            doc.add_paragraph('（未找到对应的前端页面）')
        
        doc.add_paragraph('')
        
        # 响应参数
        doc.add_heading(f'{section_num - 2}.4 响应参数', 2)
        response_fields = interface.get('response_fields', [])
        if response_fields:
            doc.add_paragraph('以下是该接口的完整响应字段列表（展开所有嵌套字段）：')
            doc.add_paragraph('')
            
            response_table = doc.add_table(rows=1, cols=4)
            response_table.style = 'Light Grid Accent 1'
            response_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            header_cells = response_table.rows[0].cells
            headers = ['字段路径', '类型', '描述', '来源（统一接口字段）']
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)
            
            # 添加数据行
            for field in response_fields:
                row_cells = response_table.add_row().cells
                row_cells[0].text = field['path']
                row_cells[1].text = field['type']
                row_cells[2].text = field.get('description', '')
                row_cells[3].text = field.get('source', '')
                
                # 设置字体
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.size = Pt(10.5)
        else:
            doc.add_paragraph('（该接口无响应参数定义）')
        
        doc.add_paragraph('')
        
        # SSE事件类型（仅流式接口）
        if interface.get('stream'):
            doc.add_heading(f'{section_num - 2}.5 SSE事件类型', 2)
            doc.add_paragraph('该接口使用Server-Sent Events (SSE)流式输出，以下是可能的事件类型：')
            doc.add_paragraph('')
            
            # 根据接口路径确定SSE事件类型
            sse_events = []
            if 'smart-analyze-stream' in interface['path']:
                sse_events = [
                    {'event': 'status', 'description': '状态更新', 'data_fields': ['stage', 'message']},
                    {'event': 'brief_response_start', 'description': '简要回答开始', 'data_fields': []},
                    {'event': 'brief_response_chunk', 'description': '简要回答内容块', 'data_fields': ['content']},
                    {'event': 'brief_response_end', 'description': '简要回答结束', 'data_fields': ['content']},
                    {'event': 'brief_response_error', 'description': '简要回答错误', 'data_fields': ['message']},
                    {'event': 'preset_questions', 'description': '预设问题', 'data_fields': ['questions']},
                    {'event': 'basic_analysis', 'description': '基础分析结果', 'data_fields': ['analysis']},
                    {'event': 'llm_start', 'description': 'LLM分析开始', 'data_fields': []},
                    {'event': 'llm_chunk', 'description': 'LLM分析内容块', 'data_fields': ['content']},
                    {'event': 'llm_end', 'description': 'LLM分析结束', 'data_fields': []},
                    {'event': 'llm_error', 'description': 'LLM分析错误', 'data_fields': ['message']},
                    {'event': 'related_questions', 'description': '相关问题', 'data_fields': ['questions']},
                    {'event': 'performance', 'description': '性能统计', 'data_fields': ['summary']},
                    {'event': 'error', 'description': '错误', 'data_fields': ['message', 'performance']},
                    {'event': 'end', 'description': '流结束', 'data_fields': []},
                ]
            elif '/stream' in interface['path']:
                sse_events = [
                    {'event': 'progress', 'description': '进度更新', 'data_fields': ['content']},
                    {'event': 'complete', 'description': '完成', 'data_fields': ['content']},
                    {'event': 'error', 'description': '错误', 'data_fields': ['content']},
                ]
            
            if sse_events:
                sse_table = doc.add_table(rows=1, cols=3)
                sse_table.style = 'Light Grid Accent 1'
                sse_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # 表头
                header_cells = sse_table.rows[0].cells
                headers = ['事件类型', '描述', '数据字段']
                for i, header_text in enumerate(headers):
                    header_cells[i].text = header_text
                    for paragraph in header_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.size = Pt(10.5)
                
                # 添加数据行
                for event in sse_events:
                    row_cells = sse_table.add_row().cells
                    row_cells[0].text = event['event']
                    row_cells[1].text = event['description']
                    data_fields_str = ', '.join(event.get('data_fields', [])) if event.get('data_fields') else '无'
                    row_cells[2].text = data_fields_str
                    
                    # 设置字体
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                run.font.size = Pt(10.5)
            else:
                doc.add_paragraph('（该接口的SSE事件类型待补充）')
        
        doc.add_paragraph('')
        
        # 前端界面
        doc.add_heading(f'{section_num - 2}.{5 if interface.get("stream") else 4} 前端界面', 2)
        frontend_pages = interface.get('frontend', [])
        if frontend_pages:
            for page in frontend_pages:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(page)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            doc.add_paragraph('（未找到对应的前端页面）')
        
        doc.add_paragraph('')
        
        # Bot ID配置
        doc.add_heading(f'{section_num - 2}.{6 if interface.get("stream") else 5} Coze Bot ID配置', 2)
        bot_id_config = interface.get('bot_id_config', 'COZE_BOT_ID (环境变量，默认)')
        doc.add_paragraph(bot_id_config)
        
        doc.add_paragraph('')
        doc.add_page_break()
    
    # 添加字段索引
    doc.add_heading('附录：字段索引', 1)
    doc.add_paragraph('按字母顺序列出所有字段：')
    
    all_fields = []
    for interface in INTERFACES:
        fields = interface.get('fields', [])
        for field in fields:
            all_fields.append({
                'name': field['name'],
                'interface': interface['name'],
                'type': field['type'],
                'description': field.get('description', '')
            })
    
    # 按字段名排序
    all_fields.sort(key=lambda x: x['name'])
    
    if all_fields:
        index_table = doc.add_table(rows=1, cols=4)
        index_table.style = 'Light Grid Accent 1'
        index_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        header_cells = index_table.rows[0].cells
        headers = ['字段名', '类型', '所属接口', '描述']
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
        
        for field in all_fields:
            row_cells = index_table.add_row().cells
            row_cells[0].text = field['name']
            row_cells[1].text = field['type']
            row_cells[2].text = field['interface']
            desc = field['description']
            row_cells[3].text = desc[:100] + '...' if len(desc) > 100 else desc
            
            # 设置字体
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)
    
    return doc


def main():
    """主函数"""
    print("开始生成大模型接口参数文档...")
    
    try:
        doc = create_word_document()
        
        # 保存到桌面
        desktop_path = Path.home() / 'Desktop'
        output_file = desktop_path / 'HiFate-bazi大模型接口参数文档.docx'
        
        doc.save(str(output_file))
        print(f"✓ 文档已成功生成：{output_file}")
        print(f"✓ 共包含 1 个统一接口")
        print(f"✓ 共包含 {len(INTERFACES)} 个大模型接口")
        
        total_request_fields = sum(len(iface.get('fields', [])) for iface in INTERFACES)
        total_response_fields = sum(len(iface.get('response_fields', [])) for iface in INTERFACES)
        unified_response_fields = len(UNIFIED_INTERFACE.get('response_fields', []))
        unified_request_fields = len(UNIFIED_INTERFACE.get('request_fields', []))
        
        print(f"✓ 统一接口：{unified_request_fields} 个请求字段，{unified_response_fields} 个响应字段")
        print(f"✓ 大模型接口：{total_request_fields} 个请求字段，{total_response_fields} 个响应字段")
        print(f"✓ 总计：{unified_request_fields + total_request_fields} 个请求字段，{unified_response_fields + total_response_fields} 个响应字段")
        
    except Exception as e:
        import traceback
        print(f"✗ 生成文档时出错：{e}")
        print(traceback.format_exc())
        sys.exit(1)


if __name__ == '__main__':
    main()
