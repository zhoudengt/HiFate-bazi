#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成支付接口前端文档并追加到 Word 文档
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_payment_api_docs(doc_path):
    """添加支付接口文档到 Word 文档"""
    doc = Document(doc_path)
    
    # 添加接口1: 统一创建支付
    add_interface_1(doc)
    
    # 添加接口2: 统一验证支付
    add_interface_2(doc)
    
    # 保存文档
    doc.save(doc_path)
    print(f"✅ 文档已更新: {doc_path}")

def add_interface_1(doc):
    """添加接口1: 统一创建支付"""
    # 添加标题
    para = doc.add_paragraph()
    para.add_run("接口详情").bold = True
    para.style = 'Heading 2'
    
    # 添加接口标题
    para = doc.add_paragraph()
    para.add_run("1、/payment/unified/create (统一创建支付)").bold = True
    
    # 添加接口详情表格（2列）
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # 填充接口详情
    details = [
        ("接口路径", "/payment/unified/create"),
        ("接口别名", "统一创建支付"),
        ("请求方法", "POST"),
        ("接口描述", "根据支付渠道创建支付订单，支持 Stripe、PayPal、支付宝、微信、Line Pay、蓝新金流"),
        ("备注", "不同支付渠道返回的支付URL字段不同：Stripe返回checkout_url，PayPal返回approval_url，支付宝/微信/Line Pay/蓝新金流返回payment_url。根据provider参数自动路由到对应支付渠道。Stripe支付需要提供customer_email，微信JSAPI支付需要提供openid。")
    ]
    
    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    # 添加请求参数标题
    para = doc.add_paragraph()
    para.add_run("请求参数").bold = True
    
    # 添加请求参数表格（5列）
    table = doc.add_table(rows=10, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ["字段名", "类型", "必填", "描述", "示例"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # 请求参数数据
    params = [
        ("provider", "str", "是", "支付渠道，可选值：stripe/paypal/alipay/wechat/linepay/newebpay", "stripe"),
        ("amount", "str", "是", "金额，格式：19.90（TWD/JPY/THB必须为整数）", "19.90"),
        ("currency", "str", "否", "货币代码，默认USD。支持：USD/EUR/HKD/CNY/TWD/JPY/THB/PHP", "USD"),
        ("product_name", "str", "是", "产品名称", "月订阅会员"),
        ("customer_email", "str", "否", "客户邮箱。Stripe支付必需，NewebPay可选", "user@example.com"),
        ("openid", "str", "否", "微信用户openid，微信JSAPI支付必需", "oUpF8uMuAJO_M2pxb1Q9zNjWeS6o"),
        ("payment_type", "str", "否", "微信支付类型：native(扫码支付)或jsapi(公众号/小程序支付)，默认native", "native"),
        ("metadata", "dict", "否", "元数据，键值对字典，用于传递额外信息", '{"source": "web", "order_no": "12345"}'),
        ("order_id", "str", "否", "商户订单号，可选。如不提供则自动生成（格式：ORDER_时间戳或ALIPAY_时间戳等）", "ORDER_1234567890"),
    ]
    
    for i, (field, type_, required, desc, example) in enumerate(params, start=1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = type_
        table.rows[i].cells[2].text = required
        table.rows[i].cells[3].text = desc
        table.rows[i].cells[4].text = example
    
    # 添加响应格式标题
    para = doc.add_paragraph()
    para.add_run("响应格式").bold = True
    
    # 添加响应格式表格（3列）
    table = doc.add_table(rows=14, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ["字段名", "类型", "描述"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # 响应参数数据
    responses = [
        ("success", "bool", "是否成功"),
        ("provider", "str", "支付渠道（stripe/paypal/alipay/wechat/linepay/newebpay）"),
        ("payment_id", "str", "支付ID（Stripe返回session_id，PayPal返回payment_id）"),
        ("transaction_id", "str", "交易ID（Line Pay使用）"),
        ("order_id", "str", "订单号（支付宝/微信/蓝新金流使用）"),
        ("payment_url", "str", "支付URL（支付宝/微信/Line Pay/蓝新金流返回）"),
        ("checkout_url", "str", "支付URL（Stripe返回，在浏览器中打开完成支付）"),
        ("approval_url", "str", "支付URL（PayPal返回，在浏览器中打开完成支付）"),
        ("code_url", "str", "二维码URL（微信Native支付返回，用于扫码支付）"),
        ("jsapi_params", "dict", "JSAPI参数（微信JSAPI支付返回，用于调起微信支付）"),
        ("form_data", "dict", "表单数据（蓝新金流返回，用于提交支付表单）"),
        ("status", "str", "订单状态（created/CREATED/success等）"),
        ("message", "str", "消息提示"),
    ]
    
    for i, (field, type_, desc) in enumerate(responses, start=1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = type_
        table.rows[i].cells[2].text = desc

def add_interface_2(doc):
    """添加接口2: 统一验证支付"""
    # 添加空行
    doc.add_paragraph()
    
    # 添加标题
    para = doc.add_paragraph()
    para.add_run("接口详情").bold = True
    para.style = 'Heading 2'
    
    # 添加接口标题
    para = doc.add_paragraph()
    para.add_run("2、/payment/unified/verify (统一验证支付)").bold = True
    
    # 添加接口详情表格（2列）
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # 填充接口详情
    details = [
        ("接口路径", "/payment/unified/verify"),
        ("接口别名", "统一验证支付"),
        ("请求方法", "POST"),
        ("接口描述", "根据支付渠道验证支付状态，不同渠道使用不同的ID进行验证"),
        ("备注", "根据provider参数，必须提供对应的验证ID：Stripe使用session_id，PayPal使用payment_id，Line Pay使用transaction_id，支付宝/微信/蓝新金流使用order_id。返回的status字段表示支付状态：Stripe返回success/pending/failed，PayPal返回COMPLETED/CREATED，Line Pay返回PAYMENT/null等。")
    ]
    
    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    # 添加请求参数标题
    para = doc.add_paragraph()
    para.add_run("请求参数").bold = True
    
    # 添加请求参数表格（5列）
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ["字段名", "类型", "必填", "描述", "示例"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # 请求参数数据
    params = [
        ("provider", "str", "是", "支付渠道，可选值：stripe/paypal/alipay/wechat/linepay/newebpay", "stripe"),
        ("session_id", "str", "否", "Session ID（Stripe验证使用，与payment_id二选一）", "cs_test_a1B2c3D4..."),
        ("payment_id", "str", "否", "支付ID（PayPal验证使用）", "5O190127TN364715T"),
        ("transaction_id", "str", "否", "交易ID（Line Pay验证使用）", "2026012102333767510"),
        ("order_id", "str", "否", "订单号（支付宝/微信/蓝新金流验证使用）", "ORDER_1234567890"),
    ]
    
    # 确保表格有足够的行（包括表头）
    if len(table.rows) < len(params) + 1:
        for _ in range(len(params) + 1 - len(table.rows)):
            table.add_row()
    
    for i, (field, type_, required, desc, example) in enumerate(params, start=1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = type_
        table.rows[i].cells[2].text = required
        table.rows[i].cells[3].text = desc
        table.rows[i].cells[4].text = example
    
    # 添加响应格式标题
    para = doc.add_paragraph()
    para.add_run("响应格式").bold = True
    
    # 添加响应格式表格（3列）
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ["字段名", "类型", "描述"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # 响应参数数据
    responses = [
        ("success", "bool", "是否成功"),
        ("provider", "str", "支付渠道（stripe/paypal/alipay/wechat/linepay/newebpay）"),
        ("status", "str", "支付状态。Stripe：success/pending/failed；PayPal：COMPLETED/CREATED；Line Pay：PAYMENT/null；支付宝/微信：success/failed等"),
        ("payment_id", "str", "支付ID（Stripe/PayPal返回）"),
        ("order_id", "str", "订单号（支付宝/微信/蓝新金流返回）"),
        ("amount", "str", "支付金额"),
        ("currency", "str", "货币代码"),
        ("customer_email", "str", "客户邮箱（Stripe/PayPal返回）"),
        ("paid_time", "str", "支付时间（时间戳或时间字符串）"),
        ("message", "str", "消息提示"),
    ]
    
    # 确保表格有足够的行（包括表头）
    if len(table.rows) < len(responses) + 1:
        for _ in range(len(responses) + 1 - len(table.rows)):
            table.add_row()
    
    for i, (field, type_, desc) in enumerate(responses, start=1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = type_
        table.rows[i].cells[2].text = desc

if __name__ == '__main__':
    doc_path = '/Users/zhoudt/Desktop/66666.docx'
    add_payment_api_docs(doc_path)
