#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成接口文档（.docx格式）
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_api_document():
    """创建接口文档"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('八字系统接口文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    doc.add_paragraph('本文档包含八字系统的主要接口说明，包括请求参数、响应格式和使用示例。')
    doc.add_paragraph('')
    
    # ========== 1. 基本信息接口 ==========
    doc.add_heading('1. 基本信息接口', 1)
    
    # 1.1 /bazi/interface
    doc.add_heading('1.1 基本信息 - /bazi/interface', 2)
    doc.add_paragraph('接口说明：生成八字界面信息（包含命宫、身宫、胎元、胎息、命卦等）', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求方式：POST', style='Normal')
    doc.add_paragraph('接口地址：/api/v1/bazi/interface', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求参数：', style='Normal')
    params_table = doc.add_table(rows=1, cols=4)
    params_table.style = 'Light Grid Accent 1'
    hdr_cells = params_table.rows[0].cells
    hdr_cells[0].text = '参数名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '必填'
    hdr_cells[3].text = '说明'
    
    params = [
        ('solar_date', 'string', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）'),
        ('solar_time', 'string', '是', '出生时间，格式：HH:MM'),
        ('gender', 'string', '是', '性别：male(男) 或 female(女)'),
        ('name', 'string', '否', '姓名'),
        ('calendar_type', 'string', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar'),
        ('location', 'string', '否', '出生地点（用于时区转换）'),
        ('latitude', 'float', '否', '纬度（用于时区转换）'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算）'),
    ]
    
    for param_name, param_type, required, desc in params:
        row_cells = params_table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_type
        row_cells[2].text = required
        row_cells[3].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('请求示例：', style='Normal')
    doc.add_paragraph('```json', style='Normal')
    doc.add_paragraph('{', style='Normal')
    doc.add_paragraph('  "solar_date": "1990-05-15",', style='Normal')
    doc.add_paragraph('  "solar_time": "14:30",', style='Normal')
    doc.add_paragraph('  "gender": "male",', style='Normal')
    doc.add_paragraph('  "name": "张三",', style='Normal')
    doc.add_paragraph('  "calendar_type": "solar",', style='Normal')
    doc.add_paragraph('  "location": "北京"', style='Normal')
    doc.add_paragraph('}', style='Normal')
    doc.add_paragraph('```', style='Normal')
    doc.add_paragraph('')
    doc.add_paragraph('响应说明：返回完整的八字界面信息，包括命宫、身宫、胎元、胎息、命卦等。', style='Normal')
    doc.add_paragraph('备注：支持农历输入和时区转换，支持夏令时自动处理。', style='Normal')
    doc.add_paragraph('')
    
    # ========== 2. 基本排盘接口 ==========
    doc.add_heading('2. 基本排盘接口', 1)
    
    # 2.1 /bazi/pan/display
    doc.add_heading('2.1 基本排盘 - /bazi/pan/display', 2)
    doc.add_paragraph('接口说明：获取排盘数据（前端优化格式）', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求方式：POST', style='Normal')
    doc.add_paragraph('接口地址：/api/v1/bazi/pan/display', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求参数：', style='Normal')
    params_table = doc.add_table(rows=1, cols=4)
    params_table.style = 'Light Grid Accent 1'
    hdr_cells = params_table.rows[0].cells
    hdr_cells[0].text = '参数名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '必填'
    hdr_cells[3].text = '说明'
    
    params = [
        ('solar_date', 'string', '是', '阳历日期，格式：YYYY-MM-DD'),
        ('solar_time', 'string', '是', '出生时间，格式：HH:MM'),
        ('gender', 'string', '是', '性别：male(男) 或 female(女)'),
    ]
    
    for param_name, param_type, required, desc in params:
        row_cells = params_table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_type
        row_cells[2].text = required
        row_cells[3].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('请求示例：', style='Normal')
    doc.add_paragraph('```json', style='Normal')
    doc.add_paragraph('{', style='Normal')
    doc.add_paragraph('  "solar_date": "1990-05-15",', style='Normal')
    doc.add_paragraph('  "solar_time": "14:30",', style='Normal')
    doc.add_paragraph('  "gender": "male"', style='Normal')
    doc.add_paragraph('}', style='Normal')
    doc.add_paragraph('```', style='Normal')
    doc.add_paragraph('')
    doc.add_paragraph('响应格式：', style='Normal')
    response_table = doc.add_table(rows=1, cols=3)
    response_table.style = 'Light Grid Accent 1'
    hdr_cells = response_table.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('pan.basic', 'dict', '基本信息'),
        ('pan.pillars', 'list', '四柱数组（便于前端循环渲染）'),
        ('pan.wuxing', 'dict', '五行统计（包含百分比）'),
        ('pan.rizhu_analysis', 'dict', '日柱分析'),
        ('pan.marriage_rules', 'list', '婚姻规则'),
        ('conversion_info', 'dict', '转换信息（如果进行了农历转换或时区转换）'),
    ]
    
    for field_name, field_type, desc in response_fields:
        row_cells = response_table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = field_type
        row_cells[2].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('响应说明：返回前端友好的排盘数据，包括基本信息、四柱数组（便于前端循环渲染）、五行统计（包含百分比）。', style='Normal')
    doc.add_paragraph('备注：返回数组格式的四柱数据，便于前端 v-for 渲染。支持农历输入和时区转换，支持夏令时自动处理。', style='Normal')
    doc.add_paragraph('')
    
    # ========== 3. 专业排盘接口 ==========
    doc.add_heading('3. 专业排盘接口', 1)
    
    # 3.1 /bazi/fortune/display
    doc.add_heading('3.1 专业排盘-大运流年流月 - /bazi/fortune/display', 2)
    doc.add_paragraph('接口说明：获取大运流年流月数据（统一接口，一次返回所有数据）', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求方式：POST', style='Normal')
    doc.add_paragraph('接口地址：/api/v1/bazi/fortune/display', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求参数：', style='Normal')
    params_table = doc.add_table(rows=1, cols=4)
    params_table.style = 'Light Grid Accent 1'
    hdr_cells = params_table.rows[0].cells
    hdr_cells[0].text = '参数名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '必填'
    hdr_cells[3].text = '说明'
    
    params = [
        ('solar_date', 'string', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）'),
        ('solar_time', 'string', '是', '出生时间，格式：HH:MM'),
        ('gender', 'string', '是', '性别：male(男) 或 female(女)'),
        ('calendar_type', 'string', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar'),
        ('location', 'string', '否', '出生地点（用于时区转换，优先级1）'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）'),
        ('current_time', 'string', '否', '当前时间，格式：YYYY-MM-DD HH:MM'),
        ('dayun_index', 'int', '否', '大运索引（已废弃，优先使用dayun_year_start和dayun_year_end）'),
        ('dayun_year_start', 'int', '否', '大运起始年份，指定要显示的大运的起始年份'),
        ('dayun_year_end', 'int', '否', '大运结束年份，指定要显示的大运的结束年份'),
        ('target_year', 'int', '否', '目标年份，用于计算该年份的流月'),
    ]
    
    for param_name, param_type, required, desc in params:
        row_cells = params_table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_type
        row_cells[2].text = required
        row_cells[3].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('请求示例：', style='Normal')
    doc.add_paragraph('```json', style='Normal')
    doc.add_paragraph('{', style='Normal')
    doc.add_paragraph('  "solar_date": "1990-05-15",', style='Normal')
    doc.add_paragraph('  "solar_time": "14:30",', style='Normal')
    doc.add_paragraph('  "gender": "male",', style='Normal')
    doc.add_paragraph('  "current_time": "2025-01-15 12:00",', style='Normal')
    doc.add_paragraph('  "target_year": 2025', style='Normal')
    doc.add_paragraph('}', style='Normal')
    doc.add_paragraph('```', style='Normal')
    doc.add_paragraph('')
    doc.add_paragraph('响应格式：', style='Normal')
    response_table = doc.add_table(rows=1, cols=3)
    response_table.style = 'Light Grid Accent 1'
    hdr_cells = response_table.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('dayun.current', 'dict', '当前大运'),
        ('dayun.list', 'list', '大运列表'),
        ('dayun.qiyun', 'dict', '起运交运信息'),
        ('liunian.current', 'dict', '当前流年'),
        ('liunian.list', 'list', '流年列表'),
        ('liuyue.current', 'dict', '当前流月'),
        ('liuyue.list', 'list', '流月列表'),
        ('conversion_info', 'dict', '转换信息（如果进行了农历转换或时区转换）'),
    ]
    
    for field_name, field_type, desc in response_fields:
        row_cells = response_table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = field_type
        row_cells[2].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('响应说明：返回大运数据（当前大运、大运列表、起运交运信息）、流年数据（当前流年、流年列表）、流月数据（当前流月、流月列表）。', style='Normal')
    doc.add_paragraph('备注：性能优化，只计算一次，避免重复调用。支持农历输入和时区转换，支持夏令时自动处理。', style='Normal')
    doc.add_paragraph('')
    
    # 3.2 /bazi/shengong-minggong
    doc.add_heading('3.2 专业排盘-身宫命宫胎元 - /bazi/shengong-minggong', 2)
    doc.add_paragraph('接口说明：获取身宫和命宫的详细信息（主星、藏干、星运、自坐、空亡、纳音、神煞等）', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求方式：POST', style='Normal')
    doc.add_paragraph('接口地址：/api/v1/bazi/shengong-minggong', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求参数：', style='Normal')
    params_table = doc.add_table(rows=1, cols=4)
    params_table.style = 'Light Grid Accent 1'
    hdr_cells = params_table.rows[0].cells
    hdr_cells[0].text = '参数名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '必填'
    hdr_cells[3].text = '说明'
    
    params = [
        ('solar_date', 'string', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）'),
        ('solar_time', 'string', '是', '出生时间，格式：HH:MM'),
        ('gender', 'string', '是', '性别：male(男) 或 female(女)'),
        ('calendar_type', 'string', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar'),
        ('location', 'string', '否', '出生地点（用于时区转换）'),
        ('latitude', 'float', '否', '纬度（用于时区转换）'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算）'),
    ]
    
    for param_name, param_type, required, desc in params:
        row_cells = params_table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_type
        row_cells[2].text = required
        row_cells[3].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('请求示例：', style='Normal')
    doc.add_paragraph('```json', style='Normal')
    doc.add_paragraph('{', style='Normal')
    doc.add_paragraph('  "solar_date": "1990-05-15",', style='Normal')
    doc.add_paragraph('  "solar_time": "14:30",', style='Normal')
    doc.add_paragraph('  "gender": "male",', style='Normal')
    doc.add_paragraph('  "calendar_type": "solar",', style='Normal')
    doc.add_paragraph('  "location": "北京"', style='Normal')
    doc.add_paragraph('}', style='Normal')
    doc.add_paragraph('```', style='Normal')
    doc.add_paragraph('')
    doc.add_paragraph('响应格式：', style='Normal')
    response_table = doc.add_table(rows=1, cols=3)
    response_table.style = 'Light Grid Accent 1'
    hdr_cells = response_table.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('data.shengong', 'dict', '身宫信息（包含干支、主星、藏干、星运、自坐、空亡、纳音、神煞等）'),
        ('data.minggong', 'dict', '命宫信息（包含干支、主星、藏干、星运、自坐、空亡、纳音、神煞等）'),
        ('data.taiyuan', 'dict', '胎元信息（包含干支、主星、藏干、星运、自坐、空亡、纳音、神煞等）'),
        ('data.pillars', 'dict', '四柱详细信息（年、月、日、时）'),
        ('data.dayun.current', 'dict', '当前大运（包含 index, ganzhi, stem, branch, age_range, year_start, year_end, wuxing, nayin, main_star, hidden_stems, hidden_stars, star_fortune, self_sitting, kongwang, deities, is_current）'),
        ('data.dayun.list', 'list', '大运列表（所有大运，每个大运包含 index, ganzhi, stem, branch, age_range, year_start, year_end, wuxing, nayin, main_star, hidden_stems, hidden_stars, star_fortune, self_sitting, kongwang, deities, is_current）'),
        ('data.dayun.qiyun.date', 'string', '起运日期'),
        ('data.dayun.qiyun.age_display', 'string', '起运年龄显示（如：10岁）'),
        ('data.dayun.qiyun.description', 'string', '起运描述'),
        ('data.dayun.jiaoyun.date', 'string', '交运日期'),
        ('data.dayun.jiaoyun.age_display', 'string', '交运年龄显示（如：20岁）'),
        ('data.dayun.jiaoyun.description', 'string', '交运描述'),
        ('data.liunian.current', 'dict', '当前流年（包含 year, age, age_display, ganzhi, stem, branch, nayin, main_star, hidden_stems, hidden_stars, star_fortune, self_sitting, kongwang, deities, is_current）'),
        ('data.liunian.list', 'list', '流年列表（当前大运范围内的流年，每个流年包含 year, age, age_display, ganzhi, stem, branch, nayin, main_star, hidden_stems, hidden_stars, star_fortune, self_sitting, kongwang, deities, is_current）'),
        ('data.liuyue.current', 'dict', '当前流月（包含 month, solar_term, term_date, ganzhi, stem, branch, nayin, is_current）'),
        ('data.liuyue.list', 'list', '流月列表（当前年份的12个月，每个流月包含 month, solar_term, term_date, ganzhi, stem, branch, nayin, is_current）'),
        ('data.liuyue.target_year', 'int', '目标年份（流月对应的年份）'),
        ('data.conversion_info', 'dict', '转换信息（如果进行了农历转换或时区转换）'),
    ]
    
    for field_name, field_type, desc in response_fields:
        row_cells = response_table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = field_type
        row_cells[2].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('响应说明：返回身宫、命宫、胎元和四柱的详细信息，包括主星、藏干、星运、自坐、空亡、纳音、神煞等。同时返回大运流年流月数据（当前大运、大运列表、起运交运信息、流年列表、流月列表），数据格式与 /bazi/fortune/display 接口一致。', style='Normal')
    doc.add_paragraph('备注：支持农历输入和时区转换，支持夏令时自动处理。大运流年流月数据使用系统当前时间自动计算。', style='Normal')
    doc.add_paragraph('')
    
    # ========== 4. 八字命理接口 ==========
    doc.add_heading('4. 八字命理接口', 1)
    
    # 4.1 /daily-fortune-calendar/query
    doc.add_heading('4.1 八字命理-每日运势 - /daily-fortune-calendar/query', 2)
    doc.add_paragraph('接口说明：查询每日运势日历信息', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求方式：POST', style='Normal')
    doc.add_paragraph('接口地址：/api/v1/daily-fortune-calendar/query', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求参数：', style='Normal')
    params_table = doc.add_table(rows=1, cols=4)
    params_table.style = 'Light Grid Accent 1'
    hdr_cells = params_table.rows[0].cells
    hdr_cells[0].text = '参数名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '必填'
    hdr_cells[3].text = '说明'
    
    params = [
        ('date', 'string', '否', '查询日期（可选，默认为今天），格式：YYYY-MM-DD'),
        ('solar_date', 'string', '否', '用户生辰阳历日期（可选，用于十神提示），格式：YYYY-MM-DD 或农历日期（当calendar_type=lunar时）'),
        ('solar_time', 'string', '否', '用户生辰时间（可选），格式：HH:MM'),
        ('gender', 'string', '否', '用户性别（可选），male/female'),
        ('calendar_type', 'string', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar'),
        ('location', 'string', '否', '出生地点（可选，用于时区转换）'),
        ('latitude', 'float', '否', '纬度（可选，用于时区转换）'),
        ('longitude', 'float', '否', '经度（可选，用于时区转换和真太阳时计算）'),
    ]
    
    for param_name, param_type, required, desc in params:
        row_cells = params_table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_type
        row_cells[2].text = required
        row_cells[3].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('请求示例：', style='Normal')
    doc.add_paragraph('```json', style='Normal')
    doc.add_paragraph('{', style='Normal')
    doc.add_paragraph('  "date": "2025-01-15",', style='Normal')
    doc.add_paragraph('  "solar_date": "1990-05-15",', style='Normal')
    doc.add_paragraph('  "solar_time": "14:30",', style='Normal')
    doc.add_paragraph('  "gender": "male"', style='Normal')
    doc.add_paragraph('}', style='Normal')
    doc.add_paragraph('```', style='Normal')
    doc.add_paragraph('')
    doc.add_paragraph('响应格式：', style='Normal')
    response_table = doc.add_table(rows=1, cols=3)
    response_table.style = 'Light Grid Accent 1'
    hdr_cells = response_table.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('solar_date', 'string', '当前阳历日期'),
        ('lunar_date', 'string', '当前阴历日期'),
        ('weekday', 'string', '星期几（中文）'),
        ('yi', 'list', '宜'),
        ('ji', 'list', '忌'),
        ('luck_level', 'string', '吉凶等级'),
        ('deities', 'dict', '神煞方位（喜神、财神、福神、胎神）'),
        ('jianchu', 'dict', '建除信息（包含名称、分数、小结）'),
        ('taishen', 'string', '胎神方位'),
        ('jiazi_fortune', 'string', '整体运势（六十甲子）'),
        ('shishen_hint', 'string', '十神提示（需要用户生辰）'),
        ('master_info', 'dict', '命主信息（日主、今日十神）'),
        ('wuxing_wear', 'string', '五行穿搭'),
        ('conversion_info', 'dict', '转换信息（如果进行了农历转换或时区转换）'),
    ]
    
    for field_name, field_type, desc in response_fields:
        row_cells = response_table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = field_type
        row_cells[2].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('响应说明：返回完整的每日运势信息，包括基础万年历信息、胎神信息、整体运势、十神提示、生肖简运、命主信息、五行穿搭、贵人方位、瘟神方位等。', style='Normal')
    doc.add_paragraph('备注：如果未提供用户生辰信息，十神提示将为空。所有运势数据从数据库读取。支持农历输入和时区转换，支持夏令时自动处理。', style='Normal')
    doc.add_paragraph('')
    
    # 4.2 /bazi/wuxing-proportion
    doc.add_heading('4.2 八字命理-五行占比 - /bazi/wuxing-proportion', 2)
    doc.add_paragraph('接口说明：查询五行占比分析', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求方式：POST', style='Normal')
    doc.add_paragraph('接口地址：/api/v1/bazi/wuxing-proportion', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求参数：', style='Normal')
    params_table = doc.add_table(rows=1, cols=4)
    params_table.style = 'Light Grid Accent 1'
    hdr_cells = params_table.rows[0].cells
    hdr_cells[0].text = '参数名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '必填'
    hdr_cells[3].text = '说明'
    
    params = [
        ('solar_date', 'string', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）'),
        ('solar_time', 'string', '是', '出生时间，格式：HH:MM'),
        ('gender', 'string', '是', '性别：male(男) 或 female(女)'),
        ('calendar_type', 'string', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar'),
        ('location', 'string', '否', '出生地点（用于时区转换，优先级1）'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）'),
    ]
    
    for param_name, param_type, required, desc in params:
        row_cells = params_table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_type
        row_cells[2].text = required
        row_cells[3].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('请求示例：', style='Normal')
    doc.add_paragraph('```json', style='Normal')
    doc.add_paragraph('{', style='Normal')
    doc.add_paragraph('  "solar_date": "1990-05-15",', style='Normal')
    doc.add_paragraph('  "solar_time": "14:30",', style='Normal')
    doc.add_paragraph('  "gender": "male",', style='Normal')
    doc.add_paragraph('  "calendar_type": "solar",', style='Normal')
    doc.add_paragraph('  "location": "北京"', style='Normal')
    doc.add_paragraph('}', style='Normal')
    doc.add_paragraph('```', style='Normal')
    doc.add_paragraph('')
    doc.add_paragraph('响应格式：', style='Normal')
    response_table = doc.add_table(rows=1, cols=3)
    response_table.style = 'Light Grid Accent 1'
    hdr_cells = response_table.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('data.wuxing_proportion', 'dict', '五行占比统计（天干+地支，8个位置）'),
        ('data.bazi_data', 'dict', '八字数据'),
        ('data.matched_rules', 'dict', '匹配的规则'),
        ('data.conversion_info', 'dict', '转换信息（如果进行了农历转换或时区转换）'),
    ]
    
    for field_name, field_type, desc in response_fields:
        row_cells = response_table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = field_type
        row_cells[2].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('响应说明：返回五行占比分析数据，包括五行占比统计（天干+地支，8个位置）、四柱十神信息（主星和副星）、旺衰分析结果、相生相克关系。', style='Normal')
    doc.add_paragraph('备注：基于生辰八字统计五行占比（金木水火土）。支持农历输入和时区转换，支持夏令时自动处理。', style='Normal')
    doc.add_paragraph('')
    
    # 4.3 /bazi/rizhu-liujiazi
    doc.add_heading('4.3 八字命理-日元-六十甲子 - /bazi/rizhu-liujiazi', 2)
    doc.add_paragraph('接口说明：根据用户生辰查询日柱对应的六十甲子解析', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求方式：POST', style='Normal')
    doc.add_paragraph('接口地址：/api/v1/bazi/rizhu-liujiazi', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求参数：', style='Normal')
    params_table = doc.add_table(rows=1, cols=4)
    params_table.style = 'Light Grid Accent 1'
    hdr_cells = params_table.rows[0].cells
    hdr_cells[0].text = '参数名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '必填'
    hdr_cells[3].text = '说明'
    
    params = [
        ('solar_date', 'string', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）'),
        ('solar_time', 'string', '是', '出生时间，格式：HH:MM'),
        ('gender', 'string', '是', '性别：male(男) 或 female(女)'),
        ('calendar_type', 'string', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar'),
        ('location', 'string', '否', '出生地点（用于时区转换，优先级1）'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）'),
    ]
    
    for param_name, param_type, required, desc in params:
        row_cells = params_table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_type
        row_cells[2].text = required
        row_cells[3].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('请求示例：', style='Normal')
    doc.add_paragraph('```json', style='Normal')
    doc.add_paragraph('{', style='Normal')
    doc.add_paragraph('  "solar_date": "1990-05-15",', style='Normal')
    doc.add_paragraph('  "solar_time": "14:30",', style='Normal')
    doc.add_paragraph('  "gender": "male",', style='Normal')
    doc.add_paragraph('  "calendar_type": "solar",', style='Normal')
    doc.add_paragraph('  "location": "北京"', style='Normal')
    doc.add_paragraph('}', style='Normal')
    doc.add_paragraph('```', style='Normal')
    doc.add_paragraph('')
    doc.add_paragraph('响应格式：', style='Normal')
    response_table = doc.add_table(rows=1, cols=3)
    response_table.style = 'Light Grid Accent 1'
    hdr_cells = response_table.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('data.id', 'int', '日柱ID'),
        ('data.rizhu', 'string', '日柱（如：甲子）'),
        ('data.analysis', 'dict', '解析内容（基础信息、深度解读、断语展示等）'),
        ('data.conversion_info', 'dict', '转换信息（如果进行了农历转换或时区转换）'),
    ]
    
    for field_name, field_type, desc in response_fields:
        row_cells = response_table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = field_type
        row_cells[2].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('响应说明：返回日柱解析结果，包含基础信息、深度解读、断语展示等。', style='Normal')
    doc.add_paragraph('备注：流程：1. 调用八字排盘服务获取日柱；2. 根据日柱查询数据库获取解析内容；3. 返回ID、日柱、解析内容。支持农历输入和时区转换，支持夏令时自动处理。', style='Normal')
    doc.add_paragraph('')
    
    # 4.4 /bazi/xishen-jishen
    doc.add_heading('4.4 八字命理-喜神忌神 - /bazi/xishen-jishen', 2)
    doc.add_paragraph('接口说明：获取喜神五行、忌神五行和十神命格', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求方式：POST', style='Normal')
    doc.add_paragraph('接口地址：/api/v1/bazi/xishen-jishen', style='Normal')
    doc.add_paragraph('')
    
    doc.add_paragraph('请求参数：', style='Normal')
    params_table = doc.add_table(rows=1, cols=4)
    params_table.style = 'Light Grid Accent 1'
    hdr_cells = params_table.rows[0].cells
    hdr_cells[0].text = '参数名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '必填'
    hdr_cells[3].text = '说明'
    
    params = [
        ('solar_date', 'string', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）'),
        ('solar_time', 'string', '是', '出生时间，格式：HH:MM'),
        ('gender', 'string', '是', '性别：male(男) 或 female(女)'),
        ('calendar_type', 'string', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar'),
        ('location', 'string', '否', '出生地点（用于时区转换，优先级1）'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）'),
    ]
    
    for param_name, param_type, required, desc in params:
        row_cells = params_table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_type
        row_cells[2].text = required
        row_cells[3].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('请求示例：', style='Normal')
    doc.add_paragraph('```json', style='Normal')
    doc.add_paragraph('{', style='Normal')
    doc.add_paragraph('  "solar_date": "1990-05-15",', style='Normal')
    doc.add_paragraph('  "solar_time": "14:30",', style='Normal')
    doc.add_paragraph('  "gender": "male",', style='Normal')
    doc.add_paragraph('  "calendar_type": "solar",', style='Normal')
    doc.add_paragraph('  "location": "北京"', style='Normal')
    doc.add_paragraph('}', style='Normal')
    doc.add_paragraph('```', style='Normal')
    doc.add_paragraph('')
    doc.add_paragraph('响应格式：', style='Normal')
    response_table = doc.add_table(rows=1, cols=3)
    response_table.style = 'Light Grid Accent 1'
    hdr_cells = response_table.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('data.xi_shen_elements', 'list', '喜神五行列表（包含名称和ID）'),
        ('data.ji_shen_elements', 'list', '忌神五行列表（包含名称和ID）'),
        ('data.shishen_mingge', 'list', '十神命格列表（包含名称和ID）'),
        ('data.wangshuai', 'string', '旺衰状态'),
        ('data.total_score', 'int', '总分'),
        ('data.conversion_info', 'dict', '转换信息（如果进行了农历转换或时区转换）'),
    ]
    
    for field_name, field_type, desc in response_fields:
        row_cells = response_table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = field_type
        row_cells[2].text = desc
    
    doc.add_paragraph('')
    doc.add_paragraph('响应说明：返回喜神五行列表（包含名称和ID）、忌神五行列表（包含名称和ID）、十神命格列表（包含名称和ID）、旺衰状态、总分。', style='Normal')
    doc.add_paragraph('备注：根据用户的生辰：1. 从旺衰分析中获取喜神五行和忌神五行；2. 从公式分析中获取十神命格；3. 查询配置表获取对应的ID。支持农历输入和时区转换，支持夏令时自动处理。', style='Normal')
    doc.add_paragraph('')
    
    # 保存文档
    output_path = '/Users/zhoudt/Desktop/接口文档.docx'
    doc.save(output_path)
    print(f'✅ 接口文档已生成：{output_path}')

if __name__ == '__main__':
    create_api_document()

