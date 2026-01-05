#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
追加 /bazi/wuxing-proportion/stream 和 /bazi/xishen-jishen/stream 接口文档到 Word 文档
字段命名与普通接口完全一致
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), kwargs[edge].get('sz', '4'))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """设置表格边框"""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'}
            )


def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading


def add_table_with_data(doc, headers, data, col_widths=None):
    """添加带数据的表格"""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10)
    
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data) if cell_data else ''
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)
    
    set_table_borders(table)
    return table


def add_wuxing_proportion_stream_doc(doc):
    """添加五行占比流式接口文档"""
    
    add_heading(doc, '五行占比流式分析 - /bazi/wuxing-proportion/stream', level=1)
    
    # 接口详情
    doc.add_paragraph()
    add_heading(doc, '接口详情', level=2)
    
    interface_details = [
        ('接口路径', '/bazi/wuxing-proportion/stream'),
        ('接口别名', '五行占比流式分析'),
        ('请求方法', 'POST'),
        ('接口描述', '流式查询五行占比分析（SSE流式响应）'),
        ('备注', '与 /bazi/wuxing-proportion 接口相同的输入，但以SSE流式方式返回数据：1. 首先返回完整的五行占比数据（type: "data"）2. 然后流式返回大模型分析（type: "progress"）3. 最后返回完成标记（type: "complete"）'),
    ]
    add_table_with_data(doc, ['属性', '值'], interface_details, col_widths=[3, 12])
    
    # 请求参数
    doc.add_paragraph()
    add_heading(doc, '请求参数', level=2)
    
    request_params = [
        ('solar_date', 'str', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', '1990-05-15'),
        ('solar_time', 'str', '是', '出生时间，格式：HH:MM', '14:30'),
        ('gender', 'str', '是', '性别：male(男) 或 female(女)', 'male'),
        ('calendar_type', 'str', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'solar'),
        ('location', 'str', '否', '出生地点（用于时区转换，优先级1）', '北京'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）', '39.90'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）', '116.40'),
    ]
    add_table_with_data(doc, ['字段名', '类型', '必填', '描述', '示例'], request_params, col_widths=[2.5, 1.5, 1, 7, 2])
    
    # 响应格式
    doc.add_paragraph()
    add_heading(doc, '响应格式', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('SSE流式响应，每行格式：data: {"type": "data|progress|complete|error", "content": ...}')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    
    # SSE事件类型
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('SSE事件类型：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    sse_types = [
        ('data', '完整的五行占比数据（JSON对象，与普通接口响应一致）'),
        ('progress', '大模型分析进度（字符串片段）'),
        ('complete', '大模型分析完成（完整的分析内容）'),
        ('error', '错误信息'),
    ]
    add_table_with_data(doc, ['类型', '描述'], sse_types, col_widths=[2, 13])
    
    # data类型的content字段结构（与普通接口一致）
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('data类型的content字段结构（与普通接口响应一致）：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('data', 'dict', '五行占比数据'),
        ('  data.bazi_pillars', 'dict', '四柱信息'),
        ('    bazi_pillars.year', 'dict', '年柱（stem: 天干, branch: 地支）'),
        ('    bazi_pillars.month', 'dict', '月柱（stem: 天干, branch: 地支）'),
        ('    bazi_pillars.day', 'dict', '日柱（stem: 天干, branch: 地支）'),
        ('    bazi_pillars.hour', 'dict', '时柱（stem: 天干, branch: 地支）'),
        ('  data.proportions', 'dict', '五行占比'),
        ('    proportions.金', 'dict', '金的占比（count: 数量, percentage: 百分比, details: 详情）'),
        ('    proportions.木', 'dict', '木的占比（count: 数量, percentage: 百分比, details: 详情）'),
        ('    proportions.水', 'dict', '水的占比（count: 数量, percentage: 百分比, details: 详情）'),
        ('    proportions.火', 'dict', '火的占比（count: 数量, percentage: 百分比, details: 详情）'),
        ('    proportions.土', 'dict', '土的占比（count: 数量, percentage: 百分比, details: 详情）'),
        ('  data.ten_gods', 'dict', '四柱十神信息'),
        ('    ten_gods.year', 'dict', '年柱十神（main_star: 主星, hidden_stars: 副星列表）'),
        ('    ten_gods.month', 'dict', '月柱十神（main_star: 主星, hidden_stars: 副星列表）'),
        ('    ten_gods.day', 'dict', '日柱十神'),
        ('    ten_gods.hour', 'dict', '时柱十神（main_star: 主星, hidden_stars: 副星列表）'),
        ('  data.wangshuai', 'dict', '旺衰分析结果'),
        ('  data.element_relations', 'dict', '相生相克关系'),
        ('    element_relations.produces', 'list', '生（我生）'),
        ('    element_relations.controls', 'list', '克（我克）'),
        ('    element_relations.produced_by', 'list', '被生（生我）'),
        ('    element_relations.controlled_by', 'list', '被克（克我）'),
        ('  data.bazi_data', 'dict', '完整八字数据'),
        ('  data.conversion_info', 'dict', '转换信息（可选，农历或时区转换时返回）'),
        ('error', 'str', '错误信息（可选，失败时返回）'),
    ]
    add_table_with_data(doc, ['字段名', '类型', '描述'], response_fields, col_widths=[4, 1.5, 9.5])
    
    # progress/complete/error 类型
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('progress/complete/error类型的content字段：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    other_types = [
        ('progress.content', 'str', '大模型分析的文本片段'),
        ('complete.content', 'str', '完整的大模型分析内容'),
        ('error.content', 'str', '错误信息'),
    ]
    add_table_with_data(doc, ['字段名', '类型', '描述'], other_types, col_widths=[4, 1.5, 9.5])


def add_xishen_jishen_stream_doc(doc):
    """添加喜神忌神流式接口文档"""
    
    add_heading(doc, '喜神忌神流式分析 - /bazi/xishen-jishen/stream', level=1)
    
    # 接口详情
    doc.add_paragraph()
    add_heading(doc, '接口详情', level=2)
    
    interface_details = [
        ('接口路径', '/bazi/xishen-jishen/stream'),
        ('接口别名', '喜神忌神流式分析'),
        ('请求方法', 'POST'),
        ('接口描述', '流式查询喜神忌神分析（SSE流式响应）'),
        ('备注', '与 /bazi/xishen-jishen 接口相同的输入，但以SSE流式方式返回数据：1. 首先返回完整的喜神忌神数据（type: "data"）2. 然后流式返回大模型分析（type: "progress"）3. 最后返回完成标记（type: "complete"）'),
    ]
    add_table_with_data(doc, ['属性', '值'], interface_details, col_widths=[3, 12])
    
    # 请求参数
    doc.add_paragraph()
    add_heading(doc, '请求参数', level=2)
    
    request_params = [
        ('solar_date', 'str', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', '1990-05-15'),
        ('solar_time', 'str', '是', '出生时间，格式：HH:MM', '14:30'),
        ('gender', 'str', '是', '性别：male(男) 或 female(女)', 'male'),
        ('calendar_type', 'str', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'solar'),
        ('location', 'str', '否', '出生地点（用于时区转换，优先级1）', '北京'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）', '39.90'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）', '116.40'),
    ]
    add_table_with_data(doc, ['字段名', '类型', '必填', '描述', '示例'], request_params, col_widths=[2.5, 1.5, 1, 7, 2])
    
    # 响应格式
    doc.add_paragraph()
    add_heading(doc, '响应格式', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('SSE流式响应，每行格式：data: {"type": "data|progress|complete|error", "content": ...}')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    
    # SSE事件类型
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('SSE事件类型：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    sse_types = [
        ('data', '完整的喜神忌神数据（JSON对象，与普通接口响应一致）'),
        ('progress', '大模型分析进度（字符串片段）'),
        ('complete', '大模型分析完成（完整的分析内容）'),
        ('error', '错误信息'),
    ]
    add_table_with_data(doc, ['类型', '描述'], sse_types, col_widths=[2, 13])
    
    # data类型的content字段结构（与普通接口一致）
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('data类型的content字段结构（与普通接口响应一致）：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('data', 'dict', '喜神忌神数据'),
        ('  data.solar_date', 'str', '阳历日期'),
        ('  data.solar_time', 'str', '出生时间'),
        ('  data.gender', 'str', '性别'),
        ('  data.xi_shen_elements', 'list', '喜神五行列表'),
        ('    xi_shen_elements[].name', 'str', '五行名称（如：金、木、水、火、土）'),
        ('    xi_shen_elements[].id', 'int', '五行ID'),
        ('  data.ji_shen_elements', 'list', '忌神五行列表'),
        ('    ji_shen_elements[].name', 'str', '五行名称'),
        ('    ji_shen_elements[].id', 'int', '五行ID'),
        ('  data.shishen_mingge', 'list', '十神命格列表'),
        ('    shishen_mingge[].name', 'str', '命格名称（如：正官格、偏财格）'),
        ('    shishen_mingge[].id', 'int', '命格ID'),
        ('  data.wangshuai', 'str', '旺衰状态（如：旺、弱、中和）'),
        ('  data.total_score', 'int', '总分'),
        ('error', 'str', '错误信息（可选，失败时返回）'),
    ]
    add_table_with_data(doc, ['字段名', '类型', '描述'], response_fields, col_widths=[4, 1.5, 9.5])
    
    # progress/complete/error 类型
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('progress/complete/error类型的content字段：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    other_types = [
        ('progress.content', 'str', '大模型分析的文本片段'),
        ('complete.content', 'str', '完整的大模型分析内容'),
        ('error.content', 'str', '错误信息'),
    ]
    add_table_with_data(doc, ['字段名', '类型', '描述'], other_types, col_widths=[4, 1.5, 9.5])


def main():
    doc_path = '/Users/zhoudt/Desktop/66666.docx'
    doc = Document(doc_path)
    
    # 删除之前追加的错误文档（通过添加新页覆盖）
    # 注意：由于无法删除已有内容，我们直接追加正确的版本
    # 用户可以手动删除之前错误的页面
    
    doc.add_page_break()
    
    # 添加五行占比流式接口文档（字段与普通接口一致）
    add_wuxing_proportion_stream_doc(doc)
    
    doc.add_page_break()
    
    # 添加喜神忌神流式接口文档（字段与普通接口一致）
    add_xishen_jishen_stream_doc(doc)
    
    doc.save(doc_path)
    print(f'✅ 两个流式接口文档（字段与普通接口一致）已成功追加到 {doc_path}')
    print('   - /bazi/wuxing-proportion/stream (五行占比流式分析)')
    print('   - /bazi/xishen-jishen/stream (喜神忌神流式分析)')
    print('')
    print('⚠️  注意：请手动删除之前追加的错误版本（如有）')


if __name__ == '__main__':
    main()

