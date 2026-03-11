#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 5 个接口的 docx 文档，请求参数与响应格式为表格"""

import os
import sys

py313 = os.path.join(os.path.dirname(__file__), '..', '.venv', 'bin', 'python3.13')
if os.path.exists(py313) and sys.version_info < (3, 12):
    os.execv(py313, [py313] + sys.argv)

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = "/Users/zhoudt/Desktop/公司/other/66666.docx"


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            if ci < len(t.rows[ri + 1].cells):
                t.rows[ri + 1].cells[ci].text = str(cell)
    doc.add_paragraph()


def main():
    doc = Document()
    doc.add_paragraph('六爻与办公桌风水接口文档').bold = True
    doc.add_paragraph()

    # ========== 1. 六爻起卦-铜钱法 ==========
    add_para(doc, '接口详情1. /api/v1/liuyao/divinate（六爻起卦-铜钱法）', bold=True)
    add_para(doc, '接口路径：/api/v1/liuyao/divinate')
    add_para(doc, '接口别名：六爻起卦-铜钱法')
    add_para(doc, '请求方法：POST')
    add_para(doc, '接口描述：铜钱法起卦，根据 6 次摇卦结果（每项 2/3/6/9）计算卦象，仅返回卦象数据。')
    add_para(doc, '请求参数：')
    add_table(doc,
        ['字段名', '类型', '必填', '描述', '示例'],
        [
            ['question', 'str', '是', '占卜问题', '这次投资能成功吗？'],
            ['method', 'str', '是', '固定为 coin', 'coin'],
            ['coin_results', 'list[int]', '是', '6 个数字，每项 2/3/6/9', '[2,3,6,3,2,9]'],
        ])
    add_para(doc, '响应格式：')
    add_table(doc,
        ['字段名', '类型', '描述'],
        [
            ['success', 'bool', '是否成功'],
            ['data', 'object', '卦象数据'],
            ['data.question', 'str', '占卜问题'],
            ['data.method', 'str', '起卦方式'],
            ['data.ben_gua', 'object', '本卦，含 name、lines'],
            ['data.bian_gua', 'object', '变卦'],
            ['data.gua_ci', 'str', '卦辞'],
            ['data.shi_ying', 'object', '世应 shi_yao、ying_yao'],
        ])
    add_para(doc, '示例：')
    add_para(doc, 'curl -X POST http://localhost:8001/api/v1/liuyao/divinate -H "Content-Type: application/json" -d \'{"question":"这次投资能成功吗？","method":"coin","coin_results":[2,3,6,3,2,9]}\'')
    doc.add_paragraph('=' * 60)

    # ========== 2. 六爻起卦-数字法 ==========
    add_para(doc, '接口详情2. /api/v1/liuyao/divinate（六爻起卦-数字法）', bold=True)
    add_para(doc, '接口路径：/api/v1/liuyao/divinate')
    add_para(doc, '接口别名：六爻起卦-数字法')
    add_para(doc, '请求方法：POST')
    add_para(doc, '接口描述：数字法起卦，根据 3 个数字计算卦象，仅返回卦象数据。')
    add_para(doc, '请求参数：')
    add_table(doc,
        ['字段名', '类型', '必填', '描述', '示例'],
        [
            ['question', 'str', '是', '占卜问题', '问事业'],
            ['method', 'str', '是', '固定为 number', 'number'],
            ['number', 'list[int]', '是', '3 个数字', '[12,34,56]'],
        ])
    add_para(doc, '响应格式：同上（与铜钱法一致）')
    add_para(doc, '示例：')
    add_para(doc, 'curl -X POST http://localhost:8001/api/v1/liuyao/divinate -H "Content-Type: application/json" -d \'{"question":"问事业","method":"number","number":[12,34,56]}\'')
    doc.add_paragraph('=' * 60)

    # ========== 3. 六爻起卦-时间法 ==========
    add_para(doc, '接口详情3. /api/v1/liuyao/divinate（六爻起卦-时间法）', bold=True)
    add_para(doc, '接口路径：/api/v1/liuyao/divinate')
    add_para(doc, '接口别名：六爻起卦-时间法')
    add_para(doc, '请求方法：POST')
    add_para(doc, '接口描述：时间法起卦，根据占卜时间计算卦象，仅返回卦象数据。')
    add_para(doc, '请求参数：')
    add_table(doc,
        ['字段名', '类型', '必填', '描述', '示例'],
        [
            ['question', 'str', '是', '占卜问题', '问健康'],
            ['method', 'str', '是', '固定为 time', 'time'],
            ['divination_time', 'str', '是', '占卜时间 YYYY-MM-DD HH:mm', '2025-03-06 14:30'],
        ])
    add_para(doc, '响应格式：同上（与铜钱法一致）')
    add_para(doc, '示例：')
    add_para(doc, 'curl -X POST http://localhost:8001/api/v1/liuyao/divinate -H "Content-Type: application/json" -d \'{"question":"问健康","method":"time","divination_time":"2025-03-06 14:30"}\'')
    doc.add_paragraph('=' * 60)

    # ========== 4. 六爻占卜-流式 ==========
    add_para(doc, '接口详情4. /api/v1/liuyao/stream（六爻占卜-流式解读）', bold=True)
    add_para(doc, '接口路径：/api/v1/liuyao/stream')
    add_para(doc, '接口别名：六爻占卜流式解读')
    add_para(doc, '请求方法：POST')
    add_para(doc, '接口描述：起卦后先返回完整卦象，再流式返回大模型解读。SSE 格式 text/event-stream。')
    add_para(doc, '备注：先 type: data，再 type: progress，最后 type: complete。')
    add_para(doc, '请求参数：')
    add_table(doc,
        ['字段名', '类型', '必填', '描述', '示例'],
        [
            ['question', 'str', '是', '占卜问题', '这次投资能成功吗？'],
            ['method', 'str', '是', 'coin/number/time', 'coin'],
            ['coin_results', 'list[int]', 'method=coin 时必填', '6 个 2/3/6/9', '[2,3,6,3,2,9]'],
            ['number', 'list[int]', 'method=number 时必填', '3 个数字', '[12,34,56]'],
            ['divination_time', 'str', 'method=time 时必填', 'YYYY-MM-DD HH:mm', '2025-03-06 14:30'],
            ['bot_id', 'str', '否', 'Bot ID（可选）', ''],
        ])
    add_para(doc, '响应格式：')
    add_table(doc,
        ['字段名', '类型', '描述'],
        [
            ['type: data', 'object', '首条，完整卦象 content.success、content.data'],
            ['type: progress', 'object', '流式 AI 解读片段，content 为增量文本'],
            ['type: complete', 'object', '完成信号，content 为完整解读'],
            ['type: error', 'object', '错误时，content 为错误描述'],
        ])
    add_para(doc, '示例：')
    add_para(doc, 'curl -X POST http://localhost:8001/api/v1/liuyao/stream -H "Content-Type: application/json" -d \'{"question":"这次投资能成功吗？","method":"coin","coin_results":[2,3,6,3,2,9]}\'')
    doc.add_paragraph('=' * 60)

    # ========== 5. 办公桌风水分析-流式 ==========
    add_para(doc, '接口详情5. /api/v2/desk-fengshui/analyze/stream（办公桌风水分析-流式）', bold=True)
    add_para(doc, '接口路径：/api/v2/desk-fengshui/analyze/stream')
    add_para(doc, '接口别名：办公桌风水分析-流式')
    add_para(doc, '请求方法：POST')
    add_para(doc, '接口描述：上传办公桌照片，先返回基础风水分析，再流式返回大模型整合分析。multipart/form-data，SSE 格式。')
    add_para(doc, '备注：先 type: request_id，再 type: data，再 type: progress，最后 type: complete。')
    add_para(doc, '请求参数：')
    add_table(doc,
        ['字段名', '类型', '必填', '描述', '示例'],
        [
            ['image', 'File(multipart/form-data)', '是', '办公桌照片，JPG/PNG', 'desk.jpg'],
            ['bot_id', 'str', '否', 'Bot ID（可选）', ''],
        ])
    add_para(doc, '响应格式：')
    add_table(doc,
        ['字段名', '类型', '描述'],
        [
            ['type: request_id', 'object', '首条，含 request_id'],
            ['type: data', 'object', '基础分析，content 含以下字段'],
            ['type: data.content.items', 'list', '识别到的物品列表'],
            ['type: data.content.rules', 'list', '规则匹配结果'],
            ['type: data.content.scene_info', 'object', '场景信息'],
            ['type: data.content.annotated_image', 'str', '标注图（原图+风水标记），Base64 编码 PNG'],
            ['type: data.content.layout_image', 'str', '新生成布局图（AI 生成），Base64 编码 PNG'],
            ['type: progress', 'object', '流式 AI 整合分析，content 为增量文本'],
            ['type: complete', 'object', '完成信号，content 为完整分析'],
            ['type: error', 'object', '错误时，content 为错误描述'],
        ])

    doc.save(DOC_PATH)
    print(f"已生成 5 个接口文档（表格格式）到 {DOC_PATH}")


if __name__ == "__main__":
    main()
