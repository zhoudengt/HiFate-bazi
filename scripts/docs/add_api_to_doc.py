#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将API接口信息追加到Word文档中
"""

import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_api_table(doc, api_info):
    """添加API接口表格到文档"""
    
    # 添加接口路径标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(f"接口路径 (Interface Path): {api_info['path']}")
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    
    # 添加请求方法
    method = doc.add_paragraph()
    method.alignment = WD_ALIGN_PARAGRAPH.LEFT
    method_run = method.add_run(f"请求方法 (Request Method): {api_info['method']}")
    method_run.font.size = Pt(12)
    method_run.font.bold = True
    
    # 添加接口描述
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    desc_run = desc.add_run(f"接口描述 (Interface Description): {api_info['description']}")
    desc_run.font.size = Pt(12)
    desc_run.font.bold = True
    
    # 添加空行
    doc.add_paragraph()
    
    # 创建请求参数表格
    req_title = doc.add_paragraph()
    req_title_run = req_title.add_run("请求参数 (Request Parameters):")
    req_title_run.font.size = Pt(11)
    req_title_run.font.bold = True
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    header_cells = table.rows[0].cells
    headers = ['字段名 (Field Name)', '类型 (Type)', '必填 (Required)', '描述 (Description)', '示例 (Example)']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # 添加数据行
    for param in api_info['params']:
        row = table.add_row()
        row.cells[0].text = param['name']
        row.cells[1].text = param['type']
        row.cells[2].text = param['required']
        row.cells[3].text = param['description']
        row.cells[4].text = param['example']
        
        # 设置字体大小
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    # 添加空行
    doc.add_paragraph()
    
    # 创建响应参数表格
    resp_title = doc.add_paragraph()
    resp_title_run = resp_title.add_run("响应参数 (Response Parameters):")
    resp_title_run.font.size = Pt(11)
    resp_title_run.font.bold = True
    
    resp_table = doc.add_table(rows=1, cols=4)
    resp_table.style = 'Light Grid Accent 1'
    
    # 响应表头
    resp_header_cells = resp_table.rows[0].cells
    resp_headers = ['字段名 (Field Name)', '类型 (Type)', '描述 (Description)', '示例 (Example)']
    for i, header in enumerate(resp_headers):
        resp_header_cells[i].text = header
        for paragraph in resp_header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # 添加响应数据行
    if 'response' in api_info:
        for resp in api_info['response']:
            row = resp_table.add_row()
            row.cells[0].text = resp['name']
            row.cells[1].text = resp['type']
            row.cells[2].text = resp['description']
            row.cells[3].text = resp.get('example', '')
            
            # 设置字体大小
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    
    # 添加空行
    doc.add_paragraph()
    doc.add_paragraph()


def main():
    # API接口信息
    api_1 = {
        'path': '/bazi/wuxing-proportion',
        'method': 'POST',
        'description': '查询五行占比 (基于八字)',
        'params': [
            {
                'name': 'solar_date',
                'type': 'str',
                'required': '是',
                'description': '阳历出生日期, 格式: YYYY-MM-DD (Solar birth date, format: YYYY-MM-DD)',
                'example': '1990-01-15'
            },
            {
                'name': 'solar_time',
                'type': 'str',
                'required': '是',
                'description': '出生时间, 格式: HH:MM (Birth time, format: HH:MM)',
                'example': '12:00'
            },
            {
                'name': 'gender',
                'type': 'str',
                'required': '是',
                'description': '性别: male/female (Gender: male/female)',
                'example': 'male'
            }
        ],
        'response': [
            {
                'name': 'success',
                'type': 'bool',
                'description': '请求是否成功 (Whether the request is successful)',
                'example': 'true'
            },
            {
                'name': 'data',
                'type': 'object',
                'description': '五行占比数据 (Wuxing proportion data)',
                'example': '{...}'
            },
            {
                'name': 'data.proportions',
                'type': 'object',
                'description': '五行占比统计, 包含金、木、水、火、土 (Each element contains count, percentage, details)',
                'example': '{"金": {"count": 2, "percentage": 25.0, "details": ["庚", "申"]}, ...}'
            },
            {
                'name': 'data.ten_gods',
                'type': 'object',
                'description': '四柱十神信息, 包含主星和副星 (Ten gods info with main_star and hidden_stars for each pillar)',
                'example': '{"year": {"main_star": "正官", "hidden_stars": ["正财"]}, ...}'
            },
            {
                'name': 'data.wangshuai',
                'type': 'object',
                'description': '旺衰分析结果 (Wangshuai analysis result)',
                'example': '{...}'
            },
            {
                'name': 'data.element_relations',
                'type': 'object',
                'description': '五行相生相克关系 (Element relations: produces, controls, etc.)',
                'example': '{...}'
            },
            {
                'name': 'error',
                'type': 'string',
                'description': '错误信息 (Error message, present when success is false)',
                'example': '计算失败'
            }
        ]
    }
    
    api_2 = {
        'path': '/bazi/rizhu-liujiazi',
        'method': 'POST',
        'description': '查询日元-六十甲子解析 (基于八字)',
        'params': [
            {
                'name': 'solar_date',
                'type': 'str',
                'required': '是',
                'description': '阳历出生日期, 格式: YYYY-MM-DD (Solar birth date, format: YYYY-MM-DD)',
                'example': '1990-05-15'
            },
            {
                'name': 'solar_time',
                'type': 'str',
                'required': '是',
                'description': '出生时间, 格式: HH:MM (Birth time, format: HH:MM)',
                'example': '14:30'
            },
            {
                'name': 'gender',
                'type': 'str',
                'required': '是',
                'description': '性别: male/female (Gender: male/female)',
                'example': 'male'
            }
        ],
        'response': [
            {
                'name': 'success',
                'type': 'bool',
                'description': '请求是否成功 (Whether the request is successful)',
                'example': 'true'
            },
            {
                'name': 'data',
                'type': 'object',
                'description': '日柱解析数据 (Rizhu analysis data)',
                'example': '{...}'
            },
            {
                'name': 'data.id',
                'type': 'int',
                'description': '记录ID (Record ID)',
                'example': '1'
            },
            {
                'name': 'data.rizhu',
                'type': 'string',
                'description': '日柱 (如: 乙丑) (Day pillar, e.g., 乙丑)',
                'example': '乙丑'
            },
            {
                'name': 'data.analysis',
                'type': 'string',
                'description': '解析内容, 包含【基础信息】、【深度解读】、【断语展示】等 (Analysis content with basic info, deep interpretation, etc.)',
                'example': '【基础信息】...【深度解读】...【断语展示】...'
            },
            {
                'name': 'error',
                'type': 'string',
                'description': '错误信息 (Error message, present when success is false)',
                'example': '未找到日柱解析内容'
            }
        ]
    }
    
    api_3 = {
        'path': '/daily-fortune-calendar/query',
        'method': 'POST',
        'description': '查询每日运势日历 (万年历)',
        'params': [
            {
                'name': 'date',
                'type': 'str',
                'required': '否',
                'description': '查询日期, 格式: YYYY-MM-DD, 默认为今天 (Query date, format: YYYY-MM-DD, defaults to today)',
                'example': '2025-01-15'
            },
            {
                'name': 'solar_date',
                'type': 'str',
                'required': '否',
                'description': '用户生辰阳历日期, 格式: YYYY-MM-DD, 用于十神提示 (User solar birth date, format: YYYY-MM-DD, for shishen hint)',
                'example': '1990-01-15'
            },
            {
                'name': 'solar_time',
                'type': 'str',
                'required': '否',
                'description': '用户生辰时间, 格式: HH:MM (User birth time, format: HH:MM)',
                'example': '12:00'
            },
            {
                'name': 'gender',
                'type': 'str',
                'required': '否',
                'description': '用户性别: male/female (User gender: male/female)',
                'example': 'male'
            }
        ],
        'response': [
            {
                'name': 'success',
                'type': 'bool',
                'description': '请求是否成功 (Whether the request is successful)',
                'example': 'true'
            },
            {
                'name': 'solar_date',
                'type': 'string',
                'description': '当前阳历日期 (Current solar date)',
                'example': '2025年1月15日'
            },
            {
                'name': 'lunar_date',
                'type': 'string',
                'description': '当前阴历日期 (Current lunar date)',
                'example': '甲子年十二月十六'
            },
            {
                'name': 'weekday',
                'type': 'string',
                'description': '星期几（中文） (Weekday in Chinese)',
                'example': '星期三'
            },
            {
                'name': 'yi',
                'type': 'array<string>',
                'description': '宜 (Things that are suitable to do)',
                'example': '["祭祀", "祈福", "求嗣"]'
            },
            {
                'name': 'ji',
                'type': 'array<string>',
                'description': '忌 (Things that should be avoided)',
                'example': '["开仓", "出财"]'
            },
            {
                'name': 'luck_level',
                'type': 'string',
                'description': '吉凶等级 (Luck level)',
                'example': '大吉'
            },
            {
                'name': 'jianchu',
                'type': 'object',
                'description': '建除信息, 包含name(名称), energy(能量), summary(小结) (Jianchu info: name, energy, summary)',
                'example': '{"name": "收", "energy": 90, "summary": "..."}'
            },
            {
                'name': 'jiazi_fortune',
                'type': 'string',
                'description': '整体运势（六十甲子）(Overall fortune based on jiazi)',
                'example': '...'
            },
            {
                'name': 'shishen_hint',
                'type': 'string',
                'description': '十神提示（需要用户生辰）(Shishen hint, requires user birth info)',
                'example': '...'
            },
            {
                'name': 'master_info',
                'type': 'object',
                'description': '命主信息, 包含rizhu(日主), today_shishen(今日十神) (Master info: rizhu, today_shishen)',
                'example': '{"rizhu": "甲木", "today_shishen": "比肩"}'
            },
            {
                'name': 'wuxing_wear',
                'type': 'string',
                'description': '五行穿搭（逗号分隔）(Wuxing wear colors, comma-separated)',
                'example': '绿色,蓝色'
            },
            {
                'name': 'error',
                'type': 'string',
                'description': '错误信息 (Error message, present when success is false)',
                'example': '查询失败'
            }
        ]
    }
    
    # 打开文档
    doc_path = '/Users/zhoudt/Desktop/前端接口文档 1.docx'
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"错误：无法打开文档: {e}")
        sys.exit(1)
    
    # 添加三个接口
    add_api_table(doc, api_1)
    add_api_table(doc, api_2)
    add_api_table(doc, api_3)
    
    # 保存文档
    doc.save(doc_path)
    print(f"✅ 成功追加3个接口到文档: {doc_path}")


if __name__ == '__main__':
    main()

