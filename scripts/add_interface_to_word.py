#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 smart-analyze-stream 接口文档追加到 Word 文档
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, color=None, font_size=11):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_code_block(doc, code, language=""):
    """添加代码块"""
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 添加灰色背景
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def main():
    doc_path = '/Users/zhoudt/Desktop/66666.docx'
    
    try:
        # 打开现有文档
        doc = Document(doc_path)
    except Exception as e:
        print(f"打开文档失败: {e}")
        return
    
    # 添加分页符
    doc.add_page_break()
    
    # ========== 开始添加接口文档 ==========
    
    # 主标题
    add_heading(doc, '智能运势流式分析接口文档', level=1)
    add_paragraph(doc, '接口路径：/api/v1/smart-fortune/smart-analyze-stream', bold=True, font_size=12)
    add_paragraph(doc, '接口类型：GET（流式响应）/ gRPC-Web POST', font_size=11)
    add_paragraph(doc, '最后更新：2026-01-07', font_size=10, color=(128, 128, 128))
    
    # 一、接口背景和来龙去脉
    add_heading(doc, '一、接口背景和来龙去脉', level=2)
    add_paragraph(doc, '1. 接口目的：', bold=True)
    add_paragraph(doc, '   这是一个智能运势分析接口，采用流式响应技术，可以实时返回分析结果，提升用户体验。接口支持两种业务场景：', font_size=11)
    add_paragraph(doc, '   • 场景1：用户首次点击选择项（如"事业财富"），需要输入生辰信息，系统返回简短答复和预设问题列表', font_size=11)
    add_paragraph(doc, '   • 场景2：用户点击预设问题或输入自定义问题，系统从会话缓存中获取生辰信息，返回详细流式回答和相关问题', font_size=11)
    
    add_paragraph(doc, '2. 技术实现：', bold=True)
    add_paragraph(doc, '   接口采用 SSE（Server-Sent Events）流式技术，可以实时推送数据到前端，无需等待完整响应。同时支持 gRPC-Web 网关调用，方便前端统一处理。', font_size=11)
    
    add_paragraph(doc, '3. 业务流程：', bold=True)
    add_paragraph(doc, '   场景1流程：用户输入生辰信息 → 计算八字 → 生成简短答复（流式）→ 生成预设问题列表 → 完成', font_size=11)
    add_paragraph(doc, '   场景2流程：用户提问 → 从缓存获取生辰信息 → 基础分析 → LLM深度解读（流式）→ 生成相关问题 → 完成', font_size=11)
    
    # 二、接口参数说明
    add_heading(doc, '二、接口参数说明', level=2)
    
    add_paragraph(doc, '（一）场景1参数（点击选择项）', bold=True, font_size=12)
    add_paragraph(doc, '使用场景：用户首次使用，需要输入完整的生辰信息', font_size=11)
    
    params_table1 = doc.add_table(rows=1, cols=3)
    params_table1.style = 'Light Grid Accent 1'
    header_cells = params_table1.rows[0].cells
    header_cells[0].text = '参数名'
    header_cells[1].text = '类型'
    header_cells[2].text = '说明'
    
    scenario1_params = [
        ('category', '必填', '分类，如"事业财富"、"健康"、"婚姻"等'),
        ('year', '必填', '出生年份，如1990'),
        ('month', '必填', '出生月份，1-12'),
        ('day', '必填', '出生日期，1-31'),
        ('hour', '可选', '出生时辰，0-23，默认12'),
        ('gender', '必填', '性别，male（男）或 female（女）'),
        ('user_id', '必填', '用户ID，用于会话缓存'),
    ]
    
    for param_name, param_type, param_desc in scenario1_params:
        row = params_table1.add_row()
        row.cells[0].text = param_name
        row.cells[1].text = param_type
        row.cells[2].text = param_desc
    
    add_paragraph(doc, '（二）场景2参数（点击预设问题）', bold=True, font_size=12)
    add_paragraph(doc, '使用场景：用户已经完成场景1，系统已缓存生辰信息，现在提问', font_size=11)
    
    params_table2 = doc.add_table(rows=1, cols=3)
    params_table2.style = 'Light Grid Accent 1'
    header_cells = params_table2.rows[0].cells
    header_cells[0].text = '参数名'
    header_cells[1].text = '类型'
    header_cells[2].text = '说明'
    
    scenario2_params = [
        ('category', '必填', '分类，如"事业财富"'),
        ('question', '必填', '问题内容，如"我今年的事业运势如何？"'),
        ('user_id', '必填', '用户ID，系统从缓存中获取该用户的生辰信息'),
    ]
    
    for param_name, param_type, param_desc in scenario2_params:
        row = params_table2.add_row()
        row.cells[0].text = param_name
        row.cells[1].text = param_type
        row.cells[2].text = param_desc
    
    # 三、调用方式
    add_heading(doc, '三、调用方式', level=2)
    
    add_paragraph(doc, '（一）方式1：直接 GET 请求（推荐，前端实时显示）', bold=True, font_size=12)
    add_paragraph(doc, '优点：实时流式输出，用户体验好，代码简单', font_size=11)
    add_paragraph(doc, '缺点：需要浏览器支持 EventSource', font_size=11)
    
    add_paragraph(doc, '场景1调用示例：', bold=True, font_size=11)
    add_code_block(doc, '''// 前端 JavaScript 代码
const url = 'http://8.210.52.217:8001/api/v1/smart-fortune/smart-analyze-stream?' +
    'category=' + encodeURIComponent('事业财富') +
    '&year=1990&month=5&day=15&hour=14&gender=male&user_id=test_user_001';

const eventSource = new EventSource(url);

// 监听流式数据
eventSource.addEventListener('brief_response_chunk', function(e) {
    const data = JSON.parse(e.data);
    console.log('收到内容:', data.content);
});''')
    
    add_paragraph(doc, '场景2调用示例：', bold=True, font_size=11)
    add_code_block(doc, '''// 前端 JavaScript 代码
const url = 'http://8.210.52.217:8001/api/v1/smart-fortune/smart-analyze-stream?' +
    'category=' + encodeURIComponent('事业财富') +
    '&question=' + encodeURIComponent('我今年的事业运势如何？') +
    '&user_id=test_user_001';

const eventSource = new EventSource(url);

// 监听流式数据
eventSource.addEventListener('llm_chunk', function(e) {
    const data = JSON.parse(e.data);
    console.log('收到内容:', data.content);
});''')
    
    add_paragraph(doc, '（二）方式2：gRPC-Web 网关（POST 请求）', bold=True, font_size=12)
    add_paragraph(doc, '优点：统一调用方式，返回完整JSON数据', font_size=11)
    add_paragraph(doc, '缺点：需要等待所有数据收集完成后才返回，响应时间较长', font_size=11)
    add_paragraph(doc, '注意：gRPC-Web 需要 protobuf 编码，建议使用提供的 Python 脚本', font_size=11, color=(255, 0, 0))
    
    add_paragraph(doc, '调用方法：', bold=True, font_size=11)
    add_code_block(doc, '''# 使用提供的测试脚本
python3 scripts/test_grpc_web_smart_analyze_stream.py http://8.210.52.217:8001''')
    
    # 四、返回值说明
    add_heading(doc, '四、返回值说明', level=2)
    
    add_paragraph(doc, '（一）GET 请求返回（SSE 流式）', bold=True, font_size=12)
    add_paragraph(doc, 'GET 请求返回的是 SSE（Server-Sent Events）流式响应，每个事件都有类型和数据。', font_size=11)
    
    add_paragraph(doc, '场景1事件序列：', bold=True, font_size=11)
    event_list1 = [
        '1. status 事件：状态更新，如"正在计算八字..."',
        '2. brief_response_start 事件：简短答复开始',
        '3. brief_response_chunk 事件：简短答复内容块（多个，流式输出）',
        '4. brief_response_end 事件：简短答复结束，包含完整内容',
        '5. preset_questions 事件：预设问题列表（10-15个问题）',
        '6. performance 事件：性能数据（总耗时、各阶段耗时）',
        '7. end 事件：流结束',
    ]
    for event in event_list1:
        add_paragraph(doc, event, font_size=11)
    
    add_paragraph(doc, '场景2事件序列：', bold=True, font_size=11)
    event_list2 = [
        '1. basic_analysis 事件：基础分析结果（八字信息、规则匹配等）',
        '2. status 事件：状态更新，如"正在生成深度解读..."',
        '3. llm_start 事件：LLM开始生成',
        '4. llm_chunk 事件：LLM内容块（多个，流式输出）',
        '5. llm_end 事件：LLM结束',
        '6. related_questions 事件：相关问题列表（2-3个问题）',
        '7. performance 事件：性能数据',
        '8. end 事件：流结束',
    ]
    for event in event_list2:
        add_paragraph(doc, event, font_size=11)
    
    add_paragraph(doc, '（二）gRPC-Web 返回（JSON 格式）', bold=True, font_size=12)
    add_paragraph(doc, 'gRPC-Web 网关会将所有流式事件收集后，统一返回一个 JSON 对象。', font_size=11)
    
    add_paragraph(doc, '场景1返回示例：', bold=True, font_size=11)
    add_code_block(doc, '''{
  "success": true,
  "data": {
    "brief_response": "庚金日主，柱中金旺得火炼，事业有官杀制劫...",
    "preset_questions": [
      "我的八字适合从事什么类型的行业或职业？",
      "今年事业发展是否有贵人相助？",
      ...
    ],
    "performance": {
      "total_time_ms": 50720,
      "stages": [...]
    },
    "last_status": {
      "stage": "preset_questions",
      "message": "正在生成预设问题..."
    }
  },
  "stream_content": "庚金日主，柱中金旺得火炼，事业有官杀制劫..."
}''')
    
    add_paragraph(doc, '场景2返回示例：', bold=True, font_size=11)
    add_code_block(doc, '''{
  "success": true,
  "data": {
    "basic_analysis": {
      "intent": {...},
      "bazi_info": {...},
      "matched_rules_count": 10
    },
    "related_questions": [
      "问题1",
      "问题2"
    ],
    "performance": {
      "total_time_ms": 95627,
      "stages": [...]
    },
    "last_status": {
      "stage": "related_questions",
      "message": "正在生成相关问题..."
    },
    "llm_completed": true
  },
  "stream_content": "所有llm_chunk内容合并后的完整文本"
}''')
    
    add_paragraph(doc, '字段说明：', bold=True, font_size=11)
    fields = [
        ('success', 'boolean', '是否成功'),
        ('data', 'object', '结构化数据'),
        ('data.brief_response', 'string', '场景1：简短答复完整内容'),
        ('data.preset_questions', 'array', '场景1：预设问题列表'),
        ('data.basic_analysis', 'object', '场景2：基础分析结果'),
        ('data.related_questions', 'array', '场景2：相关问题列表'),
        ('data.performance', 'object', '性能数据（总耗时、各阶段耗时）'),
        ('data.last_status', 'object', '最后一个状态信息'),
        ('stream_content', 'string', '所有流式内容块的合并文本'),
    ]
    
    fields_table = doc.add_table(rows=1, cols=3)
    fields_table.style = 'Light Grid Accent 1'
    header_cells = fields_table.rows[0].cells
    header_cells[0].text = '字段名'
    header_cells[1].text = '类型'
    header_cells[2].text = '说明'
    
    for field_name, field_type, field_desc in fields:
        row = fields_table.add_row()
        row.cells[0].text = field_name
        row.cells[1].text = field_type
        row.cells[2].text = field_desc
    
    # 五、完整代码示例
    add_heading(doc, '五、完整代码示例', level=2)
    
    add_paragraph(doc, '（一）前端完整示例（场景1）', bold=True, font_size=12)
    add_code_block(doc, '''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>智能运势分析</title>
</head>
<body>
    <button onclick="testScenario1()">测试场景1</button>
    <div id="result"></div>

    <script>
        function testScenario1() {
            const resultDiv = document.getElementById('result');
            resultDiv.innerHTML = '开始请求...';

            // 构建 URL
            const url = 'http://8.210.52.217:8001/api/v1/smart-fortune/smart-analyze-stream?' +
                'category=' + encodeURIComponent('事业财富') +
                '&year=1990&month=5&day=15&hour=14&gender=male&user_id=test_user_001';

            // 创建 EventSource 连接
            const eventSource = new EventSource(url);

            // 监听状态更新
            eventSource.addEventListener('status', function(e) {
                const data = JSON.parse(e.data);
                resultDiv.innerHTML += '<p>状态: ' + data.message + '</p>';
            });

            // 监听简短答复内容（流式）
            eventSource.addEventListener('brief_response_chunk', function(e) {
                const data = JSON.parse(e.data);
                resultDiv.innerHTML += data.content;
            });

            // 监听预设问题列表
            eventSource.addEventListener('preset_questions', function(e) {
                const data = JSON.parse(e.data);
                resultDiv.innerHTML += '<h3>预设问题:</h3><ul>';
                data.questions.forEach(function(q) {
                    resultDiv.innerHTML += '<li>' + q + '</li>';
                });
                resultDiv.innerHTML += '</ul>';
            });

            // 监听结束
            eventSource.addEventListener('end', function(e) {
                resultDiv.innerHTML += '<p>完成</p>';
                eventSource.close();
            });

            // 错误处理
            eventSource.onerror = function(error) {
                resultDiv.innerHTML += '<p style="color:red">连接错误</p>';
                eventSource.close();
            };
        }
    </script>
</body>
</html>''')
    
    # 六、常见问题
    add_heading(doc, '六、常见问题', level=2)
    
    qa_list = [
        ('Q: 返回 "Invalid HTTP request received"', 'A: 可能是 URL 编码问题，中文参数必须使用 encodeURIComponent() 编码'),
        ('Q: gRPC-Web 返回 "Not Found"', 'A: 检查路径是否正确，必须是 /api/grpc-web/frontend.gateway.FrontendGateway/Call'),
        ('Q: 没有返回数据', 'A: 检查服务器是否运行，参数是否完整，使用 -v 参数查看详细错误信息'),
        ('Q: 如何查看流式输出', 'A: 使用 curl -N 参数禁用缓冲，或使用浏览器 EventSource API'),
        ('Q: 字段为 null', 'A: 可能是流式数据还未收集完成（需要等待），或没有对应的事件被发送'),
    ]
    
    for question, answer in qa_list:
        add_paragraph(doc, question, bold=True, font_size=11)
        add_paragraph(doc, answer, font_size=11)
        add_paragraph(doc, '', font_size=11)  # 空行
    
    # 七、注意事项
    add_heading(doc, '七、注意事项', level=2)
    notes = [
        '1. 中文参数必须使用 URL 编码（encodeURIComponent）',
        '2. EventSource 是浏览器原生支持，无需额外库',
        '3. 记得在完成后调用 eventSource.close() 关闭连接',
        '4. 务必监听 error 和 onerror 事件进行错误处理',
        '5. 如果跨域，需要服务器配置 CORS',
        '6. gRPC-Web 方式需要 protobuf 编码，建议使用提供的脚本',
        '7. 生产环境地址：http://8.210.52.217:8001',
        '8. 本地环境地址：http://localhost:8001',
    ]
    
    for note in notes:
        add_paragraph(doc, note, font_size=11)
    
    # 保存文档
    try:
        doc.save(doc_path)
        print(f"✅ 接口文档已成功追加到: {doc_path}")
    except Exception as e:
        print(f"❌ 保存文档失败: {e}")

if __name__ == '__main__':
    main()





