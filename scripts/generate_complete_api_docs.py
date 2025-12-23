#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成完整接口文档（.docx格式）
按照用户提供的格式，包含所有请求参数和响应格式（不省略）
"""

import sys
import os
import ast
import inspect
from pathlib import Path
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 添加项目根目录到路径
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, project_root)

# 不直接导入模型，而是通过读取文件内容提取信息


def extract_field_info_from_file(file_path: str, class_name: str) -> List[Dict[str, Any]]:
    """从Python文件中提取Pydantic模型字段信息"""
    fields = []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            tree = ast.parse(content)
            
            for node in ast.walk(tree):
                if isinstance(node, ast.ClassDef) and node.name == class_name:
                    # 找到类定义
                    for item in node.body:
                        if isinstance(item, ast.AnnAssign) and isinstance(item.target, ast.Name):
                            # 字段定义：field_name: type = Field(...)
                            field_name = item.target.id
                            field_type = 'Any'
                            if item.annotation:
                                # 尝试解析类型
                                if isinstance(item.annotation, ast.Name):
                                    field_type = item.annotation.id
                                elif isinstance(item.annotation, ast.Subscript):
                                    # Optional[str] 或 Union[str, int] 等
                                    if isinstance(item.annotation.value, ast.Name):
                                        base_type = item.annotation.value.id
                                        if isinstance(item.annotation.slice, ast.Name):
                                            inner_type = item.annotation.slice.id
                                            field_type = f"{base_type}[{inner_type}]"
                                        elif isinstance(item.annotation.slice, ast.Constant):
                                            inner_type = str(item.annotation.slice.value)
                                            field_type = f"{base_type}[{inner_type}]"
                                        else:
                                            field_type = base_type
                                else:
                                    field_type = str(item.annotation)
                            
                            # 提取Field信息
                            desc = ''
                            example = ''
                            required = True
                            default_value = None
                            
                            if item.value:
                                if isinstance(item.value, ast.Call) and isinstance(item.value.func, ast.Name) and item.value.func.id == 'Field':
                                    # Field(...)
                                    for keyword in item.value.keywords:
                                        if keyword.arg == 'description':
                                            if isinstance(keyword.value, ast.Constant):
                                                desc = keyword.value.value
                                        elif keyword.arg == 'example':
                                            if isinstance(keyword.value, ast.Constant):
                                                example = str(keyword.value.value)
                                        elif keyword.arg == 'default':
                                            required = False
                                            if isinstance(keyword.value, ast.Constant):
                                                default_value = keyword.value.value
                                elif isinstance(item.value, ast.Constant):
                                    # 有默认值
                                    required = False
                                    default_value = item.value.value
                            
                            # 简化类型显示
                            if 'Optional' in field_type:
                                field_type = field_type.replace('Optional[', '').replace(']', '')
                                required = False
                            if 'Union' in field_type:
                                field_type = field_type.split(',')[0].replace('Union[', '').strip()
                            
                            fields.append({
                                'name': field_name,
                                'type': field_type,
                                'required': '是' if required else '否',
                                'description': desc,
                                'example': str(example) if example else (str(default_value) if default_value else '')
                            })
                    break
    except Exception as e:
        print(f"警告：无法解析 {file_path} 中的 {class_name}: {e}")
    
    return fields


def get_response_structure(api_path: str) -> List[Dict[str, Any]]:
    """根据API路径获取响应结构（从服务层代码中提取）"""
    response_fields = []
    
    # 基础响应字段
    response_fields.append({
        'name': 'success',
        'type': 'bool',
        'description': '是否成功'
    })
    
    # 根据不同的API路径返回不同的响应结构
    if api_path == '/bazi/interface':
        response_fields.extend([
            {'name': 'data', 'type': 'dict', 'description': '界面信息数据'},
            {'name': '  data.basic_info', 'type': 'dict', 'description': '基本信息'},
            {'name': '    basic_info.name', 'type': 'str', 'description': '姓名'},
            {'name': '    basic_info.gender', 'type': 'str', 'description': '性别'},
            {'name': '    basic_info.solar_date', 'type': 'str', 'description': '阳历日期'},
            {'name': '    basic_info.solar_time', 'type': 'str', 'description': '出生时间'},
            {'name': '    basic_info.lunar_date', 'type': 'str', 'description': '农历日期'},
            {'name': '    basic_info.location', 'type': 'str', 'description': '出生地点'},
            {'name': '  data.bazi_pillars', 'type': 'dict', 'description': '四柱信息'},
            {'name': '    bazi_pillars.year', 'type': 'str', 'description': '年柱'},
            {'name': '    bazi_pillars.month', 'type': 'str', 'description': '月柱'},
            {'name': '    bazi_pillars.day', 'type': 'str', 'description': '日柱'},
            {'name': '    bazi_pillars.hour', 'type': 'str', 'description': '时柱'},
            {'name': '  data.astrology', 'type': 'dict', 'description': '占星信息'},
            {'name': '    astrology.zodiac', 'type': 'str', 'description': '生肖'},
            {'name': '    astrology.constellation', 'type': 'str', 'description': '星座'},
            {'name': '    astrology.mansion', 'type': 'str', 'description': '二十八宿'},
            {'name': '    astrology.bagua', 'type': 'str', 'description': '命卦'},
            {'name': '  data.palaces', 'type': 'dict', 'description': '宫位信息'},
            {'name': '    palaces.life_palace', 'type': 'dict', 'description': '命宫'},
            {'name': '      life_palace.ganzhi', 'type': 'str', 'description': '干支'},
            {'name': '      life_palace.nayin', 'type': 'str', 'description': '纳音'},
            {'name': '    palaces.body_palace', 'type': 'dict', 'description': '身宫'},
            {'name': '      body_palace.ganzhi', 'type': 'str', 'description': '干支'},
            {'name': '      body_palace.nayin', 'type': 'str', 'description': '纳音'},
            {'name': '    palaces.fetal_origin', 'type': 'dict', 'description': '胎元'},
            {'name': '      fetal_origin.ganzhi', 'type': 'str', 'description': '干支'},
            {'name': '      fetal_origin.nayin', 'type': 'str', 'description': '纳音'},
            {'name': '    palaces.fetal_breath', 'type': 'dict', 'description': '胎息'},
            {'name': '      fetal_breath.ganzhi', 'type': 'str', 'description': '干支'},
            {'name': '      fetal_breath.nayin', 'type': 'str', 'description': '纳音'},
            {'name': '  data.solar_terms', 'type': 'dict', 'description': '节气信息'},
            {'name': '  data.commander', 'type': 'dict', 'description': '人元司令'},
            {'name': '  data.void_emptiness', 'type': 'str', 'description': '空亡'},
            {'name': 'message', 'type': 'str', 'description': '消息（可选）'},
        ])
    elif api_path == '/bazi/shengong-minggong':
        response_fields.extend([
            {'name': 'data', 'type': 'dict', 'description': '身宫命宫数据'},
            {'name': '  data.shengong', 'type': 'dict', 'description': '身宫信息'},
            {'name': '    shengong.stem', 'type': 'dict', 'description': '天干'},
            {'name': '      stem.char', 'type': 'str', 'description': '天干字符'},
            {'name': '    shengong.branch', 'type': 'dict', 'description': '地支'},
            {'name': '      branch.char', 'type': 'str', 'description': '地支字符'},
            {'name': '    shengong.main_star', 'type': 'str', 'description': '主星'},
            {'name': '    shengong.hidden_stems', 'type': 'list', 'description': '藏干'},
            {'name': '    shengong.star_fortune', 'type': 'str', 'description': '星运'},
            {'name': '    shengong.self_sitting', 'type': 'str', 'description': '自坐'},
            {'name': '    shengong.kongwang', 'type': 'str', 'description': '空亡'},
            {'name': '    shengong.nayin', 'type': 'str', 'description': '纳音'},
            {'name': '    shengong.deities', 'type': 'list', 'description': '神煞'},
            {'name': '  data.minggong', 'type': 'dict', 'description': '命宫信息（字段同身宫）'},
            {'name': '  data.taiyuan', 'type': 'dict', 'description': '胎元信息（字段同身宫）'},
            {'name': '  data.pillars', 'type': 'dict', 'description': '四柱详细信息'},
            {'name': '    pillars.year', 'type': 'dict', 'description': '年柱（包含stem, branch, main_star, hidden_stems等）'},
            {'name': '    pillars.month', 'type': 'dict', 'description': '月柱（包含stem, branch, main_star, hidden_stems等）'},
            {'name': '    pillars.day', 'type': 'dict', 'description': '日柱（包含stem, branch, main_star, hidden_stems等）'},
            {'name': '    pillars.hour', 'type': 'dict', 'description': '时柱（包含stem, branch, main_star, hidden_stems等）'},
            {'name': '  data.dayun', 'type': 'dict', 'description': '大运数据'},
            {'name': '    dayun.current', 'type': 'dict', 'description': '当前大运'},
            {'name': '      current.index', 'type': 'int', 'description': '大运索引'},
            {'name': '      current.ganzhi', 'type': 'str', 'description': '大运干支'},
            {'name': '      current.stem', 'type': 'dict', 'description': '天干（char, wuxing）'},
            {'name': '      current.branch', 'type': 'dict', 'description': '地支（char, wuxing）'},
            {'name': '      current.age_range', 'type': 'dict', 'description': '年龄范围（start, end）'},
            {'name': '      current.year_range', 'type': 'dict', 'description': '年份范围（start, end）'},
            {'name': '      current.nayin', 'type': 'str', 'description': '纳音'},
            {'name': '      current.main_star', 'type': 'str', 'description': '主星'},
            {'name': '      current.hidden_stems', 'type': 'list', 'description': '藏干'},
            {'name': '      current.hidden_stars', 'type': 'list', 'description': '副星'},
            {'name': '      current.star_fortune', 'type': 'str', 'description': '星运'},
            {'name': '      current.self_sitting', 'type': 'str', 'description': '自坐'},
            {'name': '      current.kongwang', 'type': 'str', 'description': '空亡'},
            {'name': '      current.deities', 'type': 'list', 'description': '神煞'},
            {'name': '      current.is_current', 'type': 'bool', 'description': '是否当前大运'},
            {'name': '    dayun.list', 'type': 'list', 'description': '大运列表（每个大运包含上述所有字段）'},
            {'name': '    dayun.qiyun', 'type': 'dict', 'description': '起运信息'},
            {'name': '      qiyun.date', 'type': 'str', 'description': '起运日期'},
            {'name': '      qiyun.age_display', 'type': 'str', 'description': '起运年龄显示（如：10岁）'},
            {'name': '      qiyun.description', 'type': 'str', 'description': '起运描述'},
            {'name': '    dayun.jiaoyun', 'type': 'dict', 'description': '交运信息'},
            {'name': '      jiaoyun.date', 'type': 'str', 'description': '交运日期'},
            {'name': '      jiaoyun.age_display', 'type': 'str', 'description': '交运年龄显示（如：20岁）'},
            {'name': '      jiaoyun.description', 'type': 'str', 'description': '交运描述'},
            {'name': '  data.liunian', 'type': 'dict', 'description': '流年数据'},
            {'name': '    liunian.current', 'type': 'dict', 'description': '当前流年'},
            {'name': '      current.year', 'type': 'int', 'description': '年份'},
            {'name': '      current.age', 'type': 'int', 'description': '年龄'},
            {'name': '      current.age_display', 'type': 'str', 'description': '年龄显示'},
            {'name': '      current.ganzhi', 'type': 'str', 'description': '流年干支'},
            {'name': '      current.stem', 'type': 'dict', 'description': '天干（char, wuxing）'},
            {'name': '      current.branch', 'type': 'dict', 'description': '地支（char, wuxing）'},
            {'name': '      current.nayin', 'type': 'str', 'description': '纳音'},
            {'name': '      current.main_star', 'type': 'str', 'description': '主星'},
            {'name': '      current.hidden_stems', 'type': 'list', 'description': '藏干'},
            {'name': '      current.hidden_stars', 'type': 'list', 'description': '副星'},
            {'name': '      current.star_fortune', 'type': 'str', 'description': '星运'},
            {'name': '      current.self_sitting', 'type': 'str', 'description': '自坐'},
            {'name': '      current.kongwang', 'type': 'str', 'description': '空亡'},
            {'name': '      current.deities', 'type': 'list', 'description': '神煞'},
            {'name': '      current.relations', 'type': 'list', 'description': '关系列表'},
            {'name': '      current.is_current', 'type': 'bool', 'description': '是否当前流年'},
            {'name': '    liunian.list', 'type': 'list', 'description': '流年列表（每个流年包含上述所有字段）'},
            {'name': '  data.liuyue', 'type': 'dict', 'description': '流月数据'},
            {'name': '    liuyue.current', 'type': 'dict', 'description': '当前流月'},
            {'name': '      current.month', 'type': 'int', 'description': '月份（1-12）'},
            {'name': '      current.solar_term', 'type': 'str', 'description': '节气'},
            {'name': '      current.term_date', 'type': 'str', 'description': '节气日期'},
            {'name': '      current.ganzhi', 'type': 'str', 'description': '流月干支'},
            {'name': '      current.stem', 'type': 'dict', 'description': '天干（char, wuxing）'},
            {'name': '      current.branch', 'type': 'dict', 'description': '地支（char, wuxing）'},
            {'name': '      current.nayin', 'type': 'str', 'description': '纳音'},
            {'name': '      current.is_current', 'type': 'bool', 'description': '是否当前流月'},
            {'name': '    liuyue.list', 'type': 'list', 'description': '流月列表（每个流月包含上述所有字段）'},
            {'name': '    liuyue.target_year', 'type': 'int', 'description': '目标年份（流月对应的年份）'},
            {'name': '  data.conversion_info', 'type': 'dict', 'description': '转换信息（如果进行了农历转换或时区转换）'},
        ])
    elif api_path == '/bazi/fortune/display':
        response_fields.extend([
            {'name': 'dayun', 'type': 'dict', 'description': '大运数据'},
            {'name': '  dayun.current', 'type': 'dict', 'description': '当前大运（字段同shengong-minggong）'},
            {'name': '  dayun.list', 'type': 'list', 'description': '大运列表'},
            {'name': '  dayun.qiyun', 'type': 'dict', 'description': '起运信息（date, age_display, description）'},
            {'name': '  dayun.jiaoyun', 'type': 'dict', 'description': '交运信息（date, age_display, description）'},
            {'name': 'liunian', 'type': 'dict', 'description': '流年数据'},
            {'name': '  liunian.current', 'type': 'dict', 'description': '当前流年（字段同shengong-minggong）'},
            {'name': '  liunian.list', 'type': 'list', 'description': '流年列表'},
            {'name': 'liuyue', 'type': 'dict', 'description': '流月数据'},
            {'name': '  liuyue.current', 'type': 'dict', 'description': '当前流月（字段同shengong-minggong）'},
            {'name': '  liuyue.list', 'type': 'list', 'description': '流月列表'},
            {'name': '  liuyue.target_year', 'type': 'int', 'description': '目标年份'},
            {'name': 'pillars', 'type': 'dict', 'description': '四柱信息'},
            {'name': 'conversion_info', 'type': 'dict', 'description': '转换信息'},
        ])
    else:
        # 默认响应结构
        response_fields.extend([
            {'name': 'data', 'type': 'dict', 'description': '响应数据'},
            {'name': 'message', 'type': 'str', 'description': '消息（可选）'},
        ])
    
    return response_fields


# 定义所有需要文档化的接口
API_DEFINITIONS = [
    {
        'path': '/bazi/interface',
        'alias': '基本信息',
        'method': 'POST',
        'description': '生成八字界面信息（包含命宫、身宫、胎元、胎息、命卦等）',
        'note': '返回完整的八字界面信息，包含命宫、身宫、胎元、胎息、命卦、生肖、星座等详细信息',
        'request_file': 'server/api/v1/bazi.py',
        'request_class': 'BaziInterfaceRequest',
        'response_model': None
    },
    {
        'path': '/bazi/pan/display',
        'alias': '基本排盘',
        'method': 'POST',
        'description': '获取排盘数据（前端优化格式）',
        'note': '返回前端友好的排盘数据，包括基本信息、四柱数组（便于前端循环渲染）、五行统计（包含百分比）',
        'request_file': 'server/api/v1/bazi_display.py',
        'request_class': 'BaziDisplayRequest',
        'response_model': None
    },
    {
        'path': '/bazi/fortune/display',
        'alias': '专业排盘-大运流年流月',
        'method': 'POST',
        'description': '获取大运流年流月数据（统一接口，一次返回所有数据）',
        'note': '性能优化，只计算一次，避免重复调用。支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/bazi_display.py',
        'request_class': 'FortuneDisplayRequest',
        'response_model': None
    },
    {
        'path': '/bazi/shengong-minggong',
        'alias': '专业排盘-身宫命宫胎元',
        'method': 'POST',
        'description': '获取身宫和命宫的详细信息（主星、藏干、星运、自坐、空亡、纳音、神煞等）',
        'note': '支持农历输入和时区转换，支持夏令时自动处理。大运流年流月数据使用系统当前时间自动计算。',
        'request_file': 'server/api/v1/bazi.py',
        'request_class': 'ShengongMinggongRequest',
        'response_model': None
    },
    {
        'path': '/daily-fortune-calendar/query',
        'alias': '八字命理-每日运势',
        'method': 'POST',
        'description': '查询每日运势日历信息',
        'note': '如果未提供用户生辰信息，十神提示将为空。所有运势数据从数据库读取。',
        'request_file': 'server/api/v1/daily_fortune_calendar.py',
        'request_class': 'DailyFortuneCalendarRequest',
        'response_model': None
    },
    {
        'path': '/bazi/wuxing-proportion',
        'alias': '八字命理-五行占比',
        'method': 'POST',
        'description': '查询五行占比分析',
        'note': '基于生辰八字统计五行占比（金木水火土）。支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/wuxing_proportion.py',
        'request_class': 'WuxingProportionRequest',
        'response_model': None
    },
    {
        'path': '/bazi/rizhu-liujiazi',
        'alias': '八字命理-日元-六十甲子',
        'method': 'POST',
        'description': '根据用户生辰查询日柱对应的六十甲子解析',
        'note': '流程：1. 调用八字排盘服务获取日柱；2. 根据日柱查询数据库获取解析内容；3. 返回ID、日柱、解析内容。支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/rizhu_liujiazi.py',
        'request_class': 'RizhuLiujiaziRequest',
        'response_model': None
    },
    {
        'path': '/bazi/xishen-jishen',
        'alias': '八字命理-喜神忌神',
        'method': 'POST',
        'description': '获取喜神五行、忌神五行和十神命格',
        'note': '根据用户的生辰：1. 从旺衰分析中获取喜神五行和忌神五行；2. 从公式分析中获取十神命格；3. 查询配置表获取对应的ID。支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/xishen_jishen.py',
        'request_class': 'XishenJishenRequest',
        'response_model': None
    },
    {
        'path': '/bazi/formula-analysis',
        'alias': '八字命理-公式分析',
        'method': 'POST',
        'description': '算法公式规则分析',
        'note': '根据八字信息匹配相应的规则，返回匹配的规则详情。支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/formula_analysis.py',
        'request_class': 'FormulaAnalysisRequest',
        'response_model': None
    },
    {
        'path': '/bazi/wangshuai',
        'alias': '八字命理-旺衰分析',
        'method': 'POST',
        'description': '计算命局旺衰',
        'note': '支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/wangshuai.py',
        'request_class': 'WangShuaiRequest',
        'response_model': None
    },
    {
        'path': '/bazi/monthly-fortune',
        'alias': '八字命理-月运势',
        'method': 'POST',
        'description': '获取月运势分析',
        'note': '支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/monthly_fortune.py',
        'request_class': 'MonthlyFortuneRequest',
        'response_model': None
    },
    {
        'path': '/bazi/daily-fortune',
        'alias': '八字命理-日运势',
        'method': 'POST',
        'description': '今日运势分析（类似 FateTell 的日运日签）',
        'note': '结合用户的八字信息和当前日期（或指定日期），分析该日的运势。支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/daily_fortune.py',
        'request_class': 'DailyFortuneRequest',
        'response_model': None
    },
    {
        'path': '/bazi/liunian-enhanced',
        'alias': '流年大运增强分析',
        'method': 'POST',
        'description': '流年大运增强分析',
        'note': '功能包括：流年吉凶量化评分（0-100分）、大运转折点识别、流年与命局互动分析、关键时间节点预测。支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/liunian_enhanced.py',
        'request_class': 'LiunianEnhancedRequest',
        'response_model': None
    },
    {
        'path': '/bazi/ai-analyze',
        'alias': 'AI分析',
        'method': 'POST',
        'description': 'Coze AI分析八字',
        'note': '支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/bazi_ai.py',
        'request_class': 'BaziAIRequest',
        'response_model': None
    },
    {
        'path': '/bazi/llm-generate',
        'alias': 'LLM生成报告',
        'method': 'POST',
        'description': 'LLM 生成完整报告（类似 FateTell）',
        'note': '支持农历输入和时区转换，支持夏令时自动处理。',
        'request_file': 'server/api/v1/llm_generate.py',
        'request_class': 'LLMGenerateRequest',
        'response_model': None
    },
]


def create_complete_api_document():
    """创建完整的接口文档"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('八字系统接口文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    doc.add_paragraph('本文档包含八字系统的主要接口说明，包括请求参数、响应格式和使用示例。')
    doc.add_paragraph('所有接口支持农历输入和时区转换，支持夏令时自动处理。')
    doc.add_paragraph('')
    
    # 为每个接口生成文档
    for api_def in API_DEFINITIONS:
        # 接口详情
        doc.add_heading(f'{api_def["alias"]} - {api_def["path"]}', 1)
        doc.add_paragraph('')
        
        # 接口详情表格
        doc.add_paragraph('接口详情', style='Normal')
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Light Grid Accent 1'
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = '属性'
        hdr_cells[1].text = '值'
        
        info_rows = [
            ('接口路径', api_def['path']),
            ('接口别名', api_def['alias']),
            ('请求方法', api_def['method']),
            ('接口描述', api_def['description']),
            ('备注', api_def['note']),
        ]
        
        for attr, value in info_rows:
            row_cells = info_table.add_row().cells
            row_cells[0].text = attr
            row_cells[1].text = value
        
        doc.add_paragraph('')
        
        # 请求参数
        doc.add_paragraph('请求参数', style='Normal')
        params_table = doc.add_table(rows=1, cols=5)
        params_table.style = 'Light Grid Accent 1'
        hdr_cells = params_table.rows[0].cells
        hdr_cells[0].text = '字段名'
        hdr_cells[1].text = '类型'
        hdr_cells[2].text = '必填'
        hdr_cells[3].text = '描述'
        hdr_cells[4].text = '示例'
        
        # 提取请求参数
        request_file_path = os.path.join(project_root, api_def['request_file'])
        request_fields = extract_field_info_from_file(request_file_path, api_def['request_class'])
        
        # 如果继承自BaziBaseRequest，需要添加基础字段
        if api_def['request_class'] != 'BaziInterfaceRequest':
            # 检查是否继承自BaziBaseRequest
            base_file_path = os.path.join(project_root, 'server/api/v1/models/bazi_base_models.py')
            base_fields = extract_field_info_from_file(base_file_path, 'BaziBaseRequest')
            # 合并字段（去重）
            existing_names = {f['name'] for f in request_fields}
            for base_field in base_fields:
                if base_field['name'] not in existing_names:
                    request_fields.append(base_field)
        
        for field in request_fields:
            row_cells = params_table.add_row().cells
            row_cells[0].text = field['name']
            row_cells[1].text = field['type']
            row_cells[2].text = field['required']
            row_cells[3].text = field['description']
            row_cells[4].text = field['example']
        
        doc.add_paragraph('')
        
        # 响应格式
        doc.add_paragraph('响应格式', style='Normal')
        response_table = doc.add_table(rows=1, cols=3)
        response_table.style = 'Light Grid Accent 1'
        hdr_cells = response_table.rows[0].cells
        hdr_cells[0].text = '字段名'
        hdr_cells[1].text = '类型'
        hdr_cells[2].text = '描述'
        
        # 获取响应结构
        response_fields = get_response_structure(api_def['path'])
        
        for field in response_fields:
            row_cells = response_table.add_row().cells
            row_cells[0].text = field['name']
            row_cells[1].text = field['type']
            row_cells[2].text = field['description']
        
        doc.add_paragraph('')
        doc.add_page_break()
    
    # 保存文档
    output_path = '/Users/zhoudt/Desktop/接口文档.docx'
    doc.save(output_path)
    print(f'✅ 接口文档已生成：{output_path}')


if __name__ == '__main__':
    create_complete_api_document()

