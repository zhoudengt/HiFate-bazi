#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 /bazi/shengong-minggong 接口文档（docx格式）
包含完整的描述信息、请求参数、响应格式，不能有省略
"""

import sys
import os
from datetime import datetime

# 添加项目根目录到路径
project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, project_root)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


def create_document() -> Document:
    """创建文档"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('专业排盘-身宫命宫胎元', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 接口路径
    doc.add_heading('接口路径', 1)
    path_para = doc.add_paragraph()
    path_para.add_run('/bazi/shengong-minggong').bold = True
    
    # 请求方法
    doc.add_heading('请求方法', 1)
    method_para = doc.add_paragraph()
    method_para.add_run('POST').bold = True
    
    # 接口描述
    doc.add_heading('接口描述', 1)
    desc_para = doc.add_paragraph()
    desc_text = """获取身宫和命宫的详细信息（主星、藏干、星运、自坐、空亡、纳音、神煞等）

本接口支持大运流年联动功能：
- 默认返回当前大运及其范围内的流年和流月数据
- 支持通过 dayun_year_start 和 dayun_year_end 参数指定要显示的大运范围
- 支持通过 target_year 参数指定要计算的流月年份
- 用户在前端点击某个大运时，会传递该大运的年份范围，显示该大运下的流年与关系
- 用户在前端点击某个流年时，会传递 target_year 参数，显示该年份的流月

注意：
- 细盘数据（身宫、命宫、胎元、四柱）是固定的，不会随大运流年变化
- 大运展示和算法与 /bazi/fortune/display 接口一致
- 两个接口的四柱都不会变化（因为八字由生辰决定）"""
    desc_para.add_run(desc_text)
    
    # 请求参数
    doc.add_heading('请求参数', 1)
    req_table = doc.add_table(rows=1, cols=5)
    req_table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ['字段名', '类型', '必填', '描述', '示例']
    for i, header in enumerate(headers):
        req_table.cell(0, i).text = header
        req_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    # 请求参数字段
    request_fields = [
        {
            'name': 'solar_date',
            'type': 'str',
            'required': True,
            'description': '阳历日期，格式：YYYY-MM-DD。当 calendar_type=lunar 时，此字段为农历日期',
            'example': '1990-05-15'
        },
        {
            'name': 'solar_time',
            'type': 'str',
            'required': True,
            'description': '出生时间，格式：HH:MM',
            'example': '14:30'
        },
        {
            'name': 'gender',
            'type': 'str',
            'required': True,
            'description': '性别：male(男) 或 female(女)',
            'example': 'male'
        },
        {
            'name': 'calendar_type',
            'type': 'str',
            'required': False,
            'description': '历法类型：solar(阳历) 或 lunar(农历)，默认 solar',
            'example': 'solar'
        },
        {
            'name': 'location',
            'type': 'str',
            'required': False,
            'description': '出生地点（可选），用于时区转换',
            'example': '北京'
        },
        {
            'name': 'latitude',
            'type': 'float',
            'required': False,
            'description': '纬度（可选），用于时区转换',
            'example': 39.90
        },
        {
            'name': 'longitude',
            'type': 'float',
            'required': False,
            'description': '经度（可选），用于时区转换和真太阳时计算',
            'example': 116.40
        },
        {
            'name': 'current_time',
            'type': 'str',
            'required': False,
            'description': '当前时间（可选），格式：YYYY-MM-DD HH:MM。用于确定当前大运，如果不提供则使用系统当前时间',
            'example': '2024-01-01 12:00'
        },
        {
            'name': 'dayun_year_start',
            'type': 'int',
            'required': False,
            'description': '大运起始年份（可选），指定要显示的大运的起始年份。用于大运联动，当用户点击某个大运时传递此参数',
            'example': 2020
        },
        {
            'name': 'dayun_year_end',
            'type': 'int',
            'required': False,
            'description': '大运结束年份（可选），指定要显示的大运的结束年份。用于大运联动，当用户点击某个大运时传递此参数',
            'example': 2030
        },
        {
            'name': 'target_year',
            'type': 'int',
            'required': False,
            'description': '目标年份（可选），用于计算该年份的流月。用于流年联动，当用户点击某个流年时传递此参数',
            'example': 2024
        },
    ]
    
    # 数据行
    for field in request_fields:
        row = req_table.add_row()
        row.cells[0].text = field['name']
        row.cells[1].text = field['type']
        row.cells[2].text = '是' if field['required'] else '否'
        row.cells[3].text = field['description']
        row.cells[4].text = str(field['example']) if field.get('example') else ''
    
    doc.add_paragraph()  # 空行
    
    # 响应格式
    doc.add_heading('响应格式', 1)
    resp_table = doc.add_table(rows=1, cols=3)
    resp_table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ['字段名', '类型', '描述']
    for i, header in enumerate(headers):
        resp_table.cell(0, i).text = header
        resp_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    # 递归添加字段（包括嵌套字段）
    def add_field_to_table(field, prefix='', level=0):
        """递归添加字段到表格"""
        indent = '  ' * level
        field_name = f"{indent}{prefix}{field['name']}" if prefix else field['name']
        
        row = resp_table.add_row()
        row.cells[0].text = field_name
        row.cells[1].text = field['type']
        row.cells[2].text = field.get('description', '')
        
        # 如果有嵌套字段，递归添加
        if 'nested' in field and field['nested']:
            # 如果是列表类型，说明嵌套字段是列表元素的结构
            if field['type'].startswith('list'):
                prefix_str = f"{field['name']}[]."
            else:
                prefix_str = f"{field['name']}."
            
            for nested_field in field['nested']:
                add_field_to_table(nested_field, prefix=prefix_str, level=level+1)
    
    # 响应字段定义
    response_fields = [
        {
            'name': 'success',
            'type': 'bool',
            'description': '是否成功'
        },
        {
            'name': 'data',
            'type': 'dict',
            'description': '身宫命宫详细信息',
            'nested': [
                {
                    'name': 'shengong',
                    'type': 'dict',
                    'description': '身宫信息（固定，不随大运流年变化）',
                    'nested': [
                        {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                            {'name': 'char', 'type': 'str', 'description': '天干字符'}
                        ]},
                        {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                            {'name': 'char', 'type': 'str', 'description': '地支字符'}
                        ]},
                        {'name': 'main_star', 'type': 'str', 'description': '主星'},
                        {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表'},
                        {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                        {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                        {'name': 'kongwang', 'type': 'bool', 'description': '空亡'},
                        {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                        {'name': 'deities', 'type': 'list', 'description': '神煞列表'}
                    ]
                },
                {
                    'name': 'minggong',
                    'type': 'dict',
                    'description': '命宫信息（固定，不随大运流年变化）',
                    'nested': [
                        {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                            {'name': 'char', 'type': 'str', 'description': '天干字符'}
                        ]},
                        {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                            {'name': 'char', 'type': 'str', 'description': '地支字符'}
                        ]},
                        {'name': 'main_star', 'type': 'str', 'description': '主星'},
                        {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表'},
                        {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                        {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                        {'name': 'kongwang', 'type': 'bool', 'description': '空亡'},
                        {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                        {'name': 'deities', 'type': 'list', 'description': '神煞列表'}
                    ]
                },
                {
                    'name': 'taiyuan',
                    'type': 'dict',
                    'description': '胎元信息（固定，不随大运流年变化）',
                    'nested': [
                        {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                            {'name': 'char', 'type': 'str', 'description': '天干字符'}
                        ]},
                        {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                            {'name': 'char', 'type': 'str', 'description': '地支字符'}
                        ]},
                        {'name': 'main_star', 'type': 'str', 'description': '主星'},
                        {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表'},
                        {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                        {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                        {'name': 'kongwang', 'type': 'bool', 'description': '空亡'},
                        {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                        {'name': 'deities', 'type': 'list', 'description': '神煞列表'}
                    ]
                },
                {
                    'name': 'pillars',
                    'type': 'dict',
                    'description': '四柱详细信息（固定，不随大运流年变化）',
                    'nested': [
                        {
                            'name': 'year',
                            'type': 'dict',
                            'description': '年柱详细信息',
                            'nested': [
                                {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '天干字符'}
                                ]},
                                {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '地支字符'}
                                ]},
                                {'name': 'main_star', 'type': 'str', 'description': '主星'},
                                {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表'},
                                {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                                {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                                {'name': 'kongwang', 'type': 'str', 'description': '空亡'},
                                {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                                {'name': 'deities', 'type': 'list', 'description': '神煞列表'}
                            ]
                        },
                        {
                            'name': 'month',
                            'type': 'dict',
                            'description': '月柱详细信息',
                            'nested': [
                                {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '天干字符'}
                                ]},
                                {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '地支字符'}
                                ]},
                                {'name': 'main_star', 'type': 'str', 'description': '主星'},
                                {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表'},
                                {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                                {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                                {'name': 'kongwang', 'type': 'str', 'description': '空亡'},
                                {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                                {'name': 'deities', 'type': 'list', 'description': '神煞列表'}
                            ]
                        },
                        {
                            'name': 'day',
                            'type': 'dict',
                            'description': '日柱详细信息',
                            'nested': [
                                {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '天干字符'}
                                ]},
                                {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '地支字符'}
                                ]},
                                {'name': 'main_star', 'type': 'str', 'description': '主星'},
                                {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表'},
                                {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                                {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                                {'name': 'kongwang', 'type': 'str', 'description': '空亡'},
                                {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                                {'name': 'deities', 'type': 'list', 'description': '神煞列表'}
                            ]
                        },
                        {
                            'name': 'hour',
                            'type': 'dict',
                            'description': '时柱详细信息',
                            'nested': [
                                {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '天干字符'}
                                ]},
                                {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '地支字符'}
                                ]},
                                {'name': 'main_star', 'type': 'str', 'description': '主星'},
                                {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表'},
                                {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                                {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                                {'name': 'kongwang', 'type': 'str', 'description': '空亡'},
                                {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                                {'name': 'deities', 'type': 'list', 'description': '神煞列表'}
                            ]
                        }
                    ]
                },
                {
                    'name': 'dayun',
                    'type': 'dict',
                    'description': '大运数据（支持大运联动）',
                    'nested': [
                        {
                            'name': 'current',
                            'type': 'dict',
                            'description': '当前大运',
                            'nested': [
                                {'name': 'index', 'type': 'int', 'description': '大运序号'},
                                {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                                {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '天干字符'},
                                    {'name': 'wuxing', 'type': 'str', 'description': '五行属性'}
                                ]},
                                {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '地支字符'},
                                    {'name': 'wuxing', 'type': 'str', 'description': '五行属性'}
                                ]},
                                {'name': 'age_range', 'type': 'dict', 'description': '年龄范围', 'nested': [
                                    {'name': 'start', 'type': 'int', 'description': '起始年龄'},
                                    {'name': 'end', 'type': 'int', 'description': '结束年龄'}
                                ]},
                                {'name': 'year_range', 'type': 'dict', 'description': '年份范围', 'nested': [
                                    {'name': 'start', 'type': 'int', 'description': '起始年份'},
                                    {'name': 'end', 'type': 'int', 'description': '结束年份'}
                                ]},
                                {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                                {'name': 'main_star', 'type': 'str', 'description': '主星'},
                                {'name': 'is_current', 'type': 'bool', 'description': '是否为当前大运'}
                            ]
                        },
                        {
                            'name': 'list',
                            'type': 'list',
                            'description': '大运列表（数组格式，每个元素结构同current）。包含所有大运，用于前端展示',
                            'nested': [
                                {'name': 'index', 'type': 'int', 'description': '大运序号'},
                                {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                                {'name': 'stem', 'type': 'dict', 'description': '天干信息'},
                                {'name': 'branch', 'type': 'dict', 'description': '地支信息'},
                                {'name': 'age_range', 'type': 'dict', 'description': '年龄范围'},
                                {'name': 'year_range', 'type': 'dict', 'description': '年份范围'},
                                {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                                {'name': 'main_star', 'type': 'str', 'description': '主星'},
                                {'name': 'is_current', 'type': 'bool', 'description': '是否为当前大运'}
                            ]
                        },
                        {
                            'name': 'qiyun',
                            'type': 'dict',
                            'description': '起运信息',
                            'nested': [
                                {'name': 'date', 'type': 'str', 'description': '起运日期'},
                                {'name': 'age_display', 'type': 'str', 'description': '起运年龄显示'},
                                {'name': 'description', 'type': 'str', 'description': '描述'}
                            ]
                        },
                        {
                            'name': 'jiaoyun',
                            'type': 'dict',
                            'description': '交运信息',
                            'nested': [
                                {'name': 'date', 'type': 'str', 'description': '交运日期'},
                                {'name': 'age_display', 'type': 'str', 'description': '交运年龄显示'},
                                {'name': 'description', 'type': 'str', 'description': '描述'}
                            ]
                        }
                    ]
                },
                {
                    'name': 'liunian',
                    'type': 'dict',
                    'description': '流年数据（支持流年联动）',
                    'nested': [
                        {
                            'name': 'current',
                            'type': 'dict',
                            'description': '当前流年',
                            'nested': [
                                {'name': 'year', 'type': 'int', 'description': '年份'},
                                {'name': 'age', 'type': 'int', 'description': '年龄'},
                                {'name': 'age_display', 'type': 'str', 'description': '年龄显示'},
                                {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                                {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '天干字符'},
                                    {'name': 'wuxing', 'type': 'str', 'description': '五行属性'}
                                ]},
                                {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                                    {'name': 'char', 'type': 'str', 'description': '地支字符'},
                                    {'name': 'wuxing', 'type': 'str', 'description': '五行属性'}
                                ]},
                                {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                                {'name': 'main_star', 'type': 'str', 'description': '主星'},
                                {'name': 'is_current', 'type': 'bool', 'description': '是否为当前流年'}
                            ]
                        },
                        {
                            'name': 'list',
                            'type': 'list',
                            'description': '流年列表（数组格式，每个元素结构同current）。根据 dayun_year_start 和 dayun_year_end 参数返回对应大运范围内的流年',
                            'nested': [
                                {'name': 'year', 'type': 'int', 'description': '年份'},
                                {'name': 'age', 'type': 'int', 'description': '年龄'},
                                {'name': 'age_display', 'type': 'str', 'description': '年龄显示'},
                                {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                                {'name': 'stem', 'type': 'dict', 'description': '天干信息'},
                                {'name': 'branch', 'type': 'dict', 'description': '地支信息'},
                                {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                                {'name': 'main_star', 'type': 'str', 'description': '主星'},
                                {'name': 'is_current', 'type': 'bool', 'description': '是否为当前流年'}
                            ]
                        }
                    ]
                },
                {
                    'name': 'liuyue',
                    'type': 'dict',
                    'description': '流月数据（支持流月联动）',
                    'nested': [
                        {
                            'name': 'current',
                            'type': 'dict',
                            'description': '当前流月',
                            'nested': [
                                {'name': 'month', 'type': 'int', 'description': '月份（1-12）'},
                                {'name': 'solar_term', 'type': 'str', 'description': '节气'},
                                {'name': 'term_date', 'type': 'str', 'description': '节气日期'},
                                {'name': 'ganzhi', 'type': 'str', 'description': '干支'}
                            ]
                        },
                        {
                            'name': 'list',
                            'type': 'list',
                            'description': '流月列表（数组格式，每个元素结构同current）。根据 target_year 参数返回对应年份的流月',
                            'nested': [
                                {'name': 'month', 'type': 'int', 'description': '月份（1-12）'},
                                {'name': 'solar_term', 'type': 'str', 'description': '节气'},
                                {'name': 'term_date', 'type': 'str', 'description': '节气日期'},
                                {'name': 'ganzhi', 'type': 'str', 'description': '干支'}
                            ]
                        },
                        {
                            'name': 'target_year',
                            'type': 'int',
                            'description': '目标年份（用于标识当前流月列表对应的年份）'
                        }
                    ]
                },
                {
                    'name': 'conversion_info',
                    'type': 'dict',
                    'required': False,
                    'description': '转换信息（可选，当输入农历或进行时区转换时返回）',
                    'nested': [
                        {'name': 'converted', 'type': 'bool', 'description': '是否进行了转换'},
                        {'name': 'timezone_info', 'type': 'dict', 'description': '时区转换信息'}
                    ]
                }
            ]
        },
        {
            'name': 'error',
            'type': 'str',
            'description': '错误信息（失败时）'
        }
    ]
    
    # 数据行
    for field in response_fields:
        add_field_to_table(field)
    
    doc.add_paragraph()  # 空行
    
    # 注意事项
    doc.add_heading('注意事项', 1)
    notes = [
        '通过 gRPC-Web 网关调用，路径：/grpc-web/frontend.gateway.FrontendGateway/Call',
        '请求体格式：{"endpoint": "/bazi/shengong-minggong", "payload_json": "{...}"}',
        '所有接口均返回 JSON 格式数据',
        '细盘数据（身宫、命宫、胎元、四柱）是固定的，不会随大运流年变化',
        '大运展示和算法与 /bazi/fortune/display 接口一致',
        '两个接口的四柱都不会变化（因为八字由生辰决定）',
        '当不传递 dayun_year_start 和 dayun_year_end 时，默认返回当前大运',
        '当传递 dayun_year_start 和 dayun_year_end 时，返回指定大运范围内的流年',
        '当传递 target_year 时，返回指定年份的流月',
        '前端可以通过点击大运和流年实现联动效果'
    ]
    
    for note in notes:
        para = doc.add_paragraph(note, style='List Bullet')
    
    return doc


def main():
    """主函数"""
    print("正在生成 /bazi/shengong-minggong 接口文档...")
    
    # 创建文档
    doc = create_document()
    
    # 保存到桌面
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
    output_path = os.path.join(desktop_path, 'bazi_shengong_minggong_api.docx')
    doc.save(output_path)
    
    print(f"文档已生成：{output_path}")
    print("文档包含完整的描述信息、请求参数、响应格式，无省略")


if __name__ == '__main__':
    main()

