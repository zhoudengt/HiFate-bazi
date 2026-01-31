#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用接口文档生成工具 - 自动解析 FastAPI 接口并追加到飞书文档

功能：
1. 自动解析 FastAPI 路由和 Pydantic 模型
2. 生成标准格式的接口文档（表格形式）
3. 直接追加到飞书文档（不覆盖现有内容）- 需要配置 App ID 和 App Secret
4. 支持任何 FastAPI 接口

使用方法（直接写入飞书）：
    # 方式1：使用环境变量
    export FEISHU_APP_ID="your_app_id"
    export FEISHU_APP_SECRET="your_app_secret"
    python3 scripts/tools/generate_api_doc_to_feishu.py \
        --file server/api/v1/unified_payment.py \
        --route /payment/unified/create \
        --provider stripe \
        --feishu-url https://kgo2k5dye3.feishu.cn/docx/IHTKdY4rvop4BHx0N8Zc69N5nRb

    # 方式2：使用命令行参数
    python3 scripts/tools/generate_api_doc_to_feishu.py \
        --file server/api/v1/unified_payment.py \
        --route /payment/unified/create \
        --provider stripe \
        --feishu-url https://kgo2k5dye3.feishu.cn/docx/IHTKdY4rvop4BHx0N8Zc69N5nRb \
        --feishu-app-id your_app_id \
        --feishu-app-secret your_app_secret

使用方法（导出文件，未配置 API）：
    python3 scripts/tools/generate_api_doc_to_feishu.py \
        --file server/api/v1/unified_payment.py \
        --route /payment/unified/create \
        --provider stripe \
        --feishu-url https://kgo2k5dye3.feishu.cn/docx/IHTKdY4rvop4BHx0N8Zc69N5nRb

或通过别名：
    @generate_api_doc --file server/api/v1/unified_payment.py --route /payment/unified/create --provider stripe

注意：
- 如果配置了飞书 App ID 和 App Secret，内容会直接写入飞书文档
- 如果未配置，会导出 Markdown 文件供手动复制
"""

import ast
import re
import sys
import os
import json
import argparse
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import inspect

# 尝试导入 requests（用于飞书 API）
try:
    import requests
except ImportError:
    requests = None
    print("⚠️  requests 库未安装，飞书 API 功能将不可用")
    print("💡 安装命令: pip install requests")

# 尝试导入 python-docx（用于 Word 文档）
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx 库未安装，Word 导出功能将不可用")
    print("💡 安装命令: pip install python-docx")

# 添加项目根目录到路径
project_root = Path(__file__).resolve().parent.parent.parent
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))


class FastAPIParser:
    """FastAPI 代码解析器"""
    
    def __init__(self, file_path: Path):
        self.file_path = file_path
        self.content = self._read_file()
        self.tree = ast.parse(self.content)
    
    def _read_file(self) -> str:
        """读取文件内容"""
        with open(self.file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def find_routes(self) -> List[Dict[str, Any]]:
        """查找所有路由定义"""
        routes = []
        
        for node in ast.walk(self.tree):
            if isinstance(node, ast.FunctionDef):
                # 查找路由装饰器
                for decorator in node.decorator_list:
                    if isinstance(decorator, ast.Call):
                        if isinstance(decorator.func, ast.Attribute):
                            if decorator.func.attr in ['get', 'post', 'put', 'delete', 'patch']:
                                route_info = self._extract_route_info(node, decorator)
                                if route_info:
                                    routes.append(route_info)
        
        return routes
    
    def _extract_route_info(self, func_node: ast.FunctionDef, decorator: ast.Call) -> Optional[Dict[str, Any]]:
        """提取路由信息"""
        try:
            # 提取 HTTP 方法
            method = decorator.func.attr.upper()
            
            # 提取路径
            if decorator.args and isinstance(decorator.args[0], ast.Constant):
                path = decorator.args[0].value
            elif decorator.args and isinstance(decorator.args[0], ast.Str):
                path = decorator.args[0].s
            else:
                return None
            
            # 提取 summary
            summary = None
            response_model = None
            for keyword in decorator.keywords:
                if keyword.arg == 'summary' and isinstance(keyword.value, ast.Constant):
                    summary = keyword.value.value
                elif keyword.arg == 'response_model':
                    if isinstance(keyword.value, ast.Name):
                        response_model = keyword.value.id
            
            # 提取函数文档字符串
            docstring = ast.get_docstring(func_node) or ""
            
            # 提取请求参数类型（跳过 Request 类型，通常是 FastAPI 的 Request 对象）
            request_model = None
            for arg in func_node.args.args:
                if arg.annotation:
                    if isinstance(arg.annotation, ast.Name):
                        # 跳过 FastAPI 的 Request 类型
                        if arg.annotation.id not in ['Request', 'HTTPException']:
                            request_model = arg.annotation.id
                            break
            
            return {
                'method': method,
                'path': path,
                'function_name': func_node.name,
                'summary': summary,
                'docstring': docstring,
                'request_model': request_model,
                'response_model': response_model
            }
        except Exception as e:
            print(f"解析路由信息失败: {e}")
            return None
    
    def find_model(self, model_name: str) -> Optional[Dict[str, Any]]:
        """查找 Pydantic 模型定义"""
        for node in ast.walk(self.tree):
            if isinstance(node, ast.ClassDef):
                if node.name == model_name:
                    model_info = self._extract_model_info(node)
                    if model_info:
                        return model_info
        
        # 如果 AST 解析失败，尝试使用正则表达式作为备选
        return self._find_model_by_regex(model_name)
    
    def _extract_model_info(self, class_node: ast.ClassDef) -> Dict[str, Any]:
        """提取模型信息"""
        fields = []
        docstring = ast.get_docstring(class_node) or ""
        
        for node in class_node.body:
            if isinstance(node, ast.AnnAssign):
                # 普通字段定义: field: type = value
                field_info = self._extract_field_from_annassign(node)
                if field_info:
                    fields.append(field_info)
            elif isinstance(node, ast.Assign):
                # Field 定义: field: type = Field(...)
                field_info = self._extract_field_from_assign(node)
                if field_info:
                    fields.append(field_info)
        
        return {
            'name': class_node.name,
            'docstring': docstring,
            'fields': fields
        }
    
    def _find_model_by_regex(self, model_name: str) -> Optional[Dict[str, Any]]:
        """使用正则表达式查找模型（AST 解析失败时的备选方案）"""
        fields = []
        
        # 查找类定义
        class_pattern = rf'class {model_name}\(BaseModel\):.*?(?=class |\Z)'
        match = re.search(class_pattern, self.content, re.DOTALL)
        if not match:
            return None
        
        class_content = match.group(0)
        
        # 提取文档字符串
        docstring_match = re.search(r'"""(.*?)"""', class_content, re.DOTALL)
        docstring = docstring_match.group(1).strip() if docstring_match else ""
        
        # 提取字段定义
        # 匹配: field_name: type = Field(...) 或 field_name: type = Field(...) 跨行的情况
        # 更宽松的匹配，允许换行
        field_pattern = r'(\w+):\s*([^=\n]+?)\s*=\s*Field\(([^)]+(?:\([^)]*\)[^)]*)*)\)'
        for match in re.finditer(field_pattern, class_content, re.MULTILINE):
            field_name = match.group(1)
            field_type = match.group(2).strip()
            field_args = match.group(3)
            
            # 解析 Field 参数
            is_required = '...' in field_args or 'Field(...)' in field_args
            description = None
            example = None
            default_value = None
            
            # 提取 description
            desc_match = re.search(r'description=["\']([^"\']+)["\']', field_args)
            if desc_match:
                description = desc_match.group(1)
            
            # 提取 example
            example_match = re.search(r'example=["\']([^"\']+)["\']', field_args)
            if example_match:
                example = example_match.group(1)
            
            # 提取 default
            default_match = re.search(r'default=([^,)]+)', field_args)
            if default_match:
                default_str = default_match.group(1).strip()
                if default_str.startswith('"') or default_str.startswith("'"):
                    default_value = default_str.strip('"\'')
                elif default_str == 'None':
                    default_value = None
                else:
                    default_value = default_str
                is_required = False
            
            # 判断 Optional
            if 'Optional' in field_type:
                is_required = False
                # 提取 Optional 内部类型
                optional_match = re.search(r'Optional\[(.+)\]', field_type)
                if optional_match:
                    field_type = optional_match.group(1).strip()
            
            fields.append({
                'name': field_name,
                'type': field_type,
                'required': is_required,
                'default': default_value,
                'description': description,
                'example': example
            })
        
        if fields:
            return {
                'name': model_name,
                'docstring': docstring,
                'fields': fields
            }
        
        return None
    
    def _extract_field_from_annassign(self, node: ast.AnnAssign) -> Optional[Dict[str, Any]]:
        """从 AnnAssign 节点提取字段信息"""
        if not isinstance(node.target, ast.Name):
            return None
        
        field_name = node.target.id
        field_type_str = self._get_type_string(node.annotation)
        
        # 检查是否是 Optional 类型
        is_optional = 'Optional' in field_type_str or field_type_str.startswith('Optional[')
        # 如果是 Optional，提取内部类型
        if is_optional:
            # 从 Optional[str] 中提取 str
            import re
            match = re.search(r'Optional\[(.+)\]', field_type_str)
            if match:
                field_type_str = match.group(1)
        
        # 检查是否有默认值
        is_required = not is_optional  # Optional 类型默认可选
        default_value = None
        description = None
        example = None
        
        if node.value:
            if isinstance(node.value, ast.Constant):
                default_value = node.value.value
                is_required = False
            elif isinstance(node.value, ast.Call):
                # 可能是 Field(...)
                result = self._extract_field_metadata(node.value)
                if result:
                    field_required, default_value, description, example = result
                    # Field 中的 required 优先级更高
                    if default_value is not None or not field_required:
                        is_required = False
                else:
                    # 如果解析失败，保持默认值
                    pass
            elif isinstance(node.value, ast.Name):
                # 可能是常量引用，如 None
                if node.value.id == 'None':
                    is_required = False
                    default_value = None
        
        return {
            'name': field_name,
            'type': field_type_str,
            'required': is_required,
            'default': default_value,
            'description': description,
            'example': example
        }
        
        return {
            'name': field_name,
            'type': field_type,
            'required': is_required,
            'default': default_value,
            'description': description,
            'example': example
        }
    
    def _extract_field_from_assign(self, node: ast.Assign) -> Optional[Dict[str, Any]]:
        """从 Assign 节点提取字段信息（带类型注解的情况）"""
        # 这种情况较少，先跳过
        return None
    
    def _extract_field_metadata(self, call_node: ast.Call) -> Optional[Tuple[bool, Optional[Any], Optional[str], Optional[str]]]:
        """从 Field() 调用中提取必填性、默认值、描述和示例"""
        # 检查是否是 Field 调用
        if not isinstance(call_node.func, ast.Name) or call_node.func.id != 'Field':
            return None
        
        is_required = True
        default_value = None
        description = None
        example = None
        
        # 检查第一个位置参数（可能是 ... 或默认值）
        if call_node.args:
            first_arg = call_node.args[0]
            arg_value = self._get_constant_value(first_arg)
            if arg_value == ...:
                is_required = True
            elif isinstance(first_arg, ast.Ellipsis):
                is_required = True
            elif arg_value is not None:
                is_required = False
                default_value = arg_value
        
        # 检查关键字参数
        for keyword in call_node.keywords:
            if keyword.arg == 'description':
                description = self._get_constant_value(keyword.value)
            elif keyword.arg == 'example':
                example = self._get_constant_value(keyword.value)
            elif keyword.arg == 'default':
                is_required = False
                default_value = self._get_constant_value(keyword.value)
        
        return (is_required, default_value, description, example)
    
    def _get_constant_value(self, node) -> Any:
        """获取常量值（兼容不同 Python 版本）"""
        if isinstance(node, ast.Constant):
            return node.value
        elif isinstance(node, ast.Str):  # Python 3.7
            return node.s
        elif isinstance(node, ast.Num):  # Python 3.7
            return node.n
        elif isinstance(node, ast.NameConstant):  # Python 3.7
            return node.value
        elif isinstance(node, ast.Name):
            if node.id == 'None':
                return None
        return None
        
        return is_required, default_value, description, example
    
    def _get_type_string(self, annotation) -> str:
        """获取类型字符串"""
        if isinstance(annotation, ast.Name):
            return annotation.id
        elif isinstance(annotation, ast.Subscript):
            # Optional[str], List[str], Dict[str, str] 等
            if isinstance(annotation.value, ast.Name):
                base = annotation.value.id
                # 处理 Optional[T] -> T
                if base == 'Optional':
                    if hasattr(annotation, 'slice'):
                        slice_node = annotation.slice
                    else:
                        # Python 3.9+
                        slice_node = annotation.slice if hasattr(annotation, 'slice') else None
                    
                    if slice_node:
                        if isinstance(slice_node, ast.Name):
                            return slice_node.id
                        elif isinstance(slice_node, ast.Subscript):
                            return self._get_type_string(slice_node)
                        elif isinstance(slice_node, ast.Tuple):
                            # Dict[str, str] 的情况
                            elts = []
                            for elt in slice_node.elts:
                                elts.append(self._get_type_string(elt))
                            return f"Dict[{', '.join(elts)}]"
                
                # 处理其他泛型
                if hasattr(annotation, 'slice'):
                    slice_node = annotation.slice
                else:
                    slice_node = None
                
                if slice_node:
                    if isinstance(slice_node, ast.Name):
                        return f"{base}[{slice_node.id}]"
                    elif isinstance(slice_node, ast.Constant):
                        return f"{base}[{slice_node.value}]"
                    elif isinstance(slice_node, ast.Tuple):
                        elts = []
                        for elt in slice_node.elts:
                            elts.append(self._get_type_string(elt))
                        return f"{base}[{', '.join(elts)}]"
                return base
        elif isinstance(annotation, ast.Constant):
            return str(annotation.value)
        return "Any"


class DocGenerator:
    """文档生成器"""
    
    def __init__(self, parser: FastAPIParser, route_info: Dict[str, Any], provider: Optional[str] = None):
        self.parser = parser
        self.route_info = route_info
        self.provider = provider
    
    def generate_doc(self) -> str:
        """生成完整文档"""
        lines = []
        
        # 1. 接口基本信息
        lines.extend(self._generate_basic_info())
        lines.append("")
        
        # 2. 接口关系说明（如果是支付相关接口）
        if self._is_payment_interface():
            lines.extend(self._generate_interface_relationship())
            lines.append("")
        
        # 3. 请求参数表格
        lines.extend(self._generate_request_params())
        lines.append("")
        
        # 4. 响应格式表格
        lines.extend(self._generate_response_format())
        lines.append("")
        
        # 5. 弱网/断网处理说明
        lines.extend(self._generate_network_handling())
        lines.append("")
        
        # 6. 30分钟待支付流程说明（如果是支付接口）
        if self._is_payment_interface():
            lines.extend(self._generate_payment_timeout_info())
            lines.append("")
        
        # 7. 使用说明
        lines.extend(self._generate_usage_examples())
        
        return "\n".join(lines)
    
    def _generate_basic_info(self) -> List[str]:
        """生成接口基本信息"""
        lines = []
        lines.append("## 接口基本信息")
        lines.append("")
        
        # 提取接口别名（从 summary 或 docstring）
        alias = self.route_info.get('summary') or self.route_info.get('path', '').split('/')[-1]
        description = self.route_info.get('docstring', '').split('\n')[0] if self.route_info.get('docstring') else ""
        
        # 如果是支付接口且有 provider，添加特定描述
        if self._is_payment_interface() and self.provider:
            if self.provider == 'stripe':
                description = "创建 Stripe 支付订单，支持全球信用卡支付，适合美洲、欧洲、香港等地区"
            elif self.provider == 'payermax':
                description = "创建 PayerMax 支付订单，支持全球 600+ 支付方式聚合，适合全球市场（除台湾 LINE Pay）"
        
        info_table = [
            ("接口路径", self.route_info.get('path', '')),
            ("接口别名", alias),
            ("请求方法", self.route_info.get('method', 'POST')),
            ("接口描述", description)
        ]
        
        lines.append("| 字段 | 值 |")
        lines.append("|------|-----|")
        for label, value in info_table:
            lines.append(f"| {label} | {value} |")
        
        return lines
    
    def _generate_request_params(self) -> List[str]:
        """生成请求参数表格"""
        lines = []
        lines.append("## 请求参数")
        lines.append("")
        
        request_model_name = self.route_info.get('request_model')
        if not request_model_name:
            lines.append("无请求参数")
            return lines
        
        model_info = self.parser.find_model(request_model_name)
        if not model_info:
            lines.append(f"无法找到请求模型: {request_model_name}")
            return lines
        
        if not model_info.get('fields'):
            lines.append(f"请求模型 {request_model_name} 没有字段定义")
            return lines
        
        lines.append("| 字段名 | 类型 | 必填 | 描述 | 示例 |")
        lines.append("|--------|------|------|------|------|")
        
        for field in model_info['fields']:
            name = field.get('name', '')
            field_type = field.get('type', 'Any')
            required = "是" if field.get('required', False) else "否"
            description = field.get('description', '') or ''
            example = field.get('example') or field.get('default')
            example_str = f"`{example}`" if example is not None else ""
            
            # 如果是支付接口且有 provider，添加特定说明
            if self._is_payment_interface() and self.provider:
                if name == 'provider':
                    example_str = f"`{self.provider}`"
                    field_type = "str"  # provider 实际是字符串
                elif name == 'customer_email' and self.provider == 'stripe':
                    required = "是"
                    if description and "（Stripe 必需）" not in description:
                        description += "（Stripe 必需）"
                elif name == 'customer_email' and self.provider == 'payermax':
                    description = description or "客户邮箱（可选，用于生成userId）"
            
            lines.append(f"| {name} | {field_type} | {required} | {description} | {example_str} |")
        
        # 添加备注
        if self._is_payment_interface() and self.provider:
            lines.append("")
            lines.append("**备注**:")
            if self.provider == 'stripe':
                lines.append("- Stripe 支付必须提供 `customer_email`")
                lines.append("- 系统自动设置订单过期时间为创建后 30 分钟")
                lines.append("- 支持货币自动转换（Adaptive Pricing）")
            elif self.provider == 'payermax':
                lines.append("- PayerMax 使用收银台模式，用户可选择多种支付方式")
                lines.append("- 如果指定 `payment_method`，则使用指定支付方式")
                lines.append("- 系统自动设置订单过期时间为创建后 30 分钟")
                lines.append("- `customer_email` 会自动转换为符合 PayerMax 要求的 userId（只含字母数字下划线）")
        
        return lines
    
    def _generate_response_format(self) -> List[str]:
        """生成响应格式表格"""
        lines = []
        lines.append("## 响应格式")
        lines.append("")
        
        response_model_name = self.route_info.get('response_model')
        if not response_model_name:
            lines.append("无响应模型定义")
            return lines
        
        model_info = self.parser.find_model(response_model_name)
        if not model_info or not model_info.get('fields'):
            lines.append("无法解析响应格式")
            return lines
        
        lines.append("| 字段名 | 类型 | 描述 |")
        lines.append("|--------|------|------|")
        
        for field in model_info['fields']:
            name = field.get('name', '')
            field_type = field.get('type', 'Any')
            description = field.get('description', '') or ''
            
            # 如果是支付接口，添加特定描述
            if self._is_payment_interface() and self.provider:
                if name == 'checkout_url' and self.provider == 'stripe':
                    description = "支付URL（Stripe返回，在浏览器中打开完成支付）"
                elif name == 'payment_url' and self.provider == 'payermax':
                    description = "支付URL（PayerMax返回，在浏览器中打开完成支付）"
                elif name == 'payment_id' and self.provider == 'stripe':
                    description = "支付ID（Stripe返回session_id）"
                elif name == 'payment_id' and self.provider == 'payermax':
                    description = "支付ID（PayerMax返回tradeToken）"
                elif name == 'expires_at':
                    description = "订单过期时间（ISO 8601格式，创建后30分钟）"
            
            lines.append(f"| {name} | {field_type} | {description} |")
        
        return lines
    
    def _generate_network_handling(self) -> List[str]:
        """生成弱网/断网处理说明"""
        lines = []
        lines.append("## 弱网/断网处理说明")
        lines.append("")
        lines.append("### 后端超时设置")
        lines.append("")
        
        if self._is_payment_interface():
            if self.provider == 'stripe':
                lines.append("- HTTP 请求超时：Stripe SDK 内部处理（默认约 30 秒）")
            elif self.provider == 'payermax':
                lines.append("- HTTP 请求超时：`timeout=30秒`")
            else:
                lines.append("- HTTP 请求超时：30秒")
        else:
            lines.append("- HTTP 请求超时：30秒")
        
        lines.append("- Nginx 代理超时：`proxy_read_timeout 30s`, `proxy_send_timeout 30s`")
        lines.append("- 连接超时：`proxy_connect_timeout 10s`")
        lines.append("")
        lines.append("### 建议前端处理")
        lines.append("")
        lines.append("- 实现重试机制：最大重试 3 次，间隔 1 秒递增（1s, 2s, 3s）")
        lines.append("- 检测网络错误：`fetch failed`, `network error`, `timeout`")
        lines.append("- 超时后提示用户检查网络连接")
        lines.append("")
        lines.append("### 错误处理")
        lines.append("")
        lines.append("- 网络超时：返回 HTTP 504 或 500，前端应重试")
        if self._is_payment_interface() and self.provider:
            if self.provider == 'stripe':
                lines.append("- Stripe API 错误：返回具体错误信息，前端显示给用户")
            elif self.provider == 'payermax':
                lines.append("- PayerMax API 错误：返回具体错误码和消息（如 `PARAMS_INVALID`, `SIGN_VERIFY_FAILED`），前端显示给用户")
                lines.append("- 签名验证失败：后端会记录详细日志，前端提示联系技术支持")
        else:
            lines.append("- API 错误：返回具体错误信息，前端显示给用户")
        
        return lines
    
    def _generate_payment_timeout_info(self) -> List[str]:
        """生成30分钟待支付流程说明"""
        lines = []
        lines.append("## 30分钟待支付流程说明")
        lines.append("")
        lines.append("### 订单过期机制")
        lines.append("")
        lines.append("- 系统在创建订单时自动设置 `expires_at = created_at + 30分钟`")
        if self.provider == 'stripe':
            lines.append("- Stripe Session 本身也有过期时间（30分钟），由 Stripe 管理")
        elif self.provider == 'payermax':
            lines.append("- PayerMax 订单本身也有过期时间（由 PayerMax 管理，通常 30 分钟）")
        lines.append("")
        lines.append("### 检查订单状态")
        lines.append("")
        lines.append("- 使用 `/api/v1/payment/unified/verify` 接口查询支付状态")
        if self.provider == 'payermax':
            lines.append("- 使用 `transaction_id` 或 `order_id` 进行查询")
        lines.append("- 检查 `expires_at` 字段判断订单是否过期")
        lines.append("- 如果当前时间 > `expires_at`，订单已过期，需要重新创建")
        lines.append("")
        lines.append("### 待支付订单处理")
        lines.append("")
        lines.append("- 前端应定期轮询订单状态（建议每 30 秒）")
        lines.append("- 如果订单过期，提示用户重新创建支付")
        if self.provider == 'payermax':
            lines.append("- 后端会记录订单状态到 `payment_transactions` 表，状态为 `pending`")
        else:
            lines.append("- 后端会记录订单状态到 `payment_transactions` 表")
        
        return lines
    
    def _generate_usage_examples(self) -> List[str]:
        """生成使用说明"""
        lines = []
        lines.append("## 使用说明")
        lines.append("")
        lines.append("### 后端 API 调用示例")
        lines.append("")
        
        method = self.route_info.get('method', 'POST')
        path = self.route_info.get('path', '')
        
        # 生成 curl 命令
        curl_lines = [f"curl -X {method} http://localhost:8001{path} \\"]
        curl_lines.append('  -H "Content-Type: application/json" \\')
        curl_lines.append("  -d '{")
        
        # 生成请求体示例
        request_model_name = self.route_info.get('request_model')
        if request_model_name:
            model_info = self.parser.find_model(request_model_name)
            if model_info and model_info.get('fields'):
                json_fields = []
                for field in model_info['fields']:
                    name = field.get('name', '')
                    example = field.get('example') or field.get('default')
                    is_required = field.get('required', False)
                    
                    # 只包含有示例值的字段，或者必填字段
                    if example is not None or (is_required and self._is_payment_interface() and self.provider):
                        if self._is_payment_interface() and self.provider and name == 'provider':
                            example = self.provider
                        elif example is None and is_required:
                            # 必填字段但没有示例，使用占位符
                            example = "required_value"
                        
                        if isinstance(example, str):
                            json_fields.append(f'    "{name}": "{example}"')
                        else:
                            json_fields.append(f'    "{name}": {json.dumps(example)}')
                
                if json_fields:
                    curl_lines.append(",\n".join(json_fields))
        
        curl_lines.append("  }'")
        lines.append("```bash")
        lines.extend(curl_lines)
        lines.append("```")
        lines.append("")
        
        # 生成响应示例
        lines.append("### 成功响应示例")
        lines.append("")
        lines.append("```json")
        lines.append("{")
        
        response_model_name = self.route_info.get('response_model')
        if response_model_name:
            model_info = self.parser.find_model(response_model_name)
            if model_info and model_info.get('fields'):
                json_fields = []
                for field in model_info['fields']:
                    name = field.get('name', '')
                    example = field.get('example') or field.get('default')
                    if example is not None:
                        if isinstance(example, str):
                            json_fields.append(f'  "{name}": "{example}"')
                        else:
                            json_fields.append(f'  "{name}": {json.dumps(example)}')
                    elif name == 'success':
                        json_fields.append(f'  "{name}": true')
                    elif name == 'provider' and self.provider:
                        json_fields.append(f'  "{name}": "{self.provider}"')
                
                if json_fields:
                    lines.append(",\n".join(json_fields))
        
        lines.append("}")
        lines.append("```")
        
        return lines
    
    def _generate_interface_relationship(self) -> List[str]:
        """生成接口关系说明"""
        lines = []
        path = self.route_info.get('path', '')
        
        if '/payment/unified/create' in path:
            lines.append("## 接口关系说明")
            lines.append("")
            lines.append("### 与验证接口的关系")
            lines.append("")
            lines.append("`/payment/unified/create` 接口用于创建支付订单，返回支付链接。")
            lines.append("")
            lines.append("**工作流程**：")
            lines.append("1. 调用 `/payment/unified/create` 创建支付订单，获取 `payment_url` 或 `checkout_url`")
            lines.append("2. 用户跳转到支付页面完成支付")
            lines.append("3. 调用 `/payment/unified/verify` 验证支付状态，查询订单是否已支付")
            lines.append("")
            lines.append("**关键字段关联**：")
            lines.append("- `create` 接口返回的 `payment_id`、`transaction_id` 或 `order_id` 用于 `verify` 接口查询")
            lines.append("- Stripe: 使用 `session_id` (即 `payment_id`) 进行验证")
            lines.append("- PayerMax: 使用 `transaction_id` 或 `order_id` 进行验证")
            lines.append("- 其他支付渠道: 根据返回的字段使用对应的 ID 进行验证")
            lines.append("")
            lines.append("**验证时机**：")
            lines.append("- 支付成功后（通过 webhook 或前端回调）")
            lines.append("- 定期轮询（建议每 30 秒）检查支付状态")
            lines.append("- 用户返回页面时主动查询")
        
        elif '/payment/unified/verify' in path:
            lines.append("## 接口关系说明")
            lines.append("")
            lines.append("### 与创建接口的关系")
            lines.append("")
            lines.append("`/payment/unified/verify` 接口用于验证支付状态，必须配合 `/payment/unified/create` 接口使用。")
            lines.append("")
            lines.append("**工作流程**：")
            lines.append("1. 先调用 `/payment/unified/create` 创建支付订单")
            lines.append("2. 获取返回的 `payment_id`、`transaction_id` 或 `order_id`")
            lines.append("3. 使用这些 ID 调用 `/payment/unified/verify` 查询支付状态")
            lines.append("")
            lines.append("**参数对应关系**：")
            lines.append("| 支付渠道 | create 返回字段 | verify 使用参数 |")
            lines.append("|---------|---------------|---------------|")
            lines.append("| Stripe | `payment_id` (session_id) | `session_id` |")
            lines.append("| PayerMax | `transaction_id` 或 `order_id` | `transaction_id` 或 `order_id` |")
            lines.append("| PayPal | `payment_id` | `payment_id` |")
            lines.append("| Alipay | `order_id` | `order_id` |")
            lines.append("| WeChat | `order_id` | `order_id` |")
            lines.append("| Line Pay | `transaction_id` | `transaction_id` |")
            lines.append("")
            lines.append("**验证结果**：")
            lines.append("- `success: true, status: 'paid'` - 支付成功")
            lines.append("- `success: true, status: 'pending'` - 待支付")
            lines.append("- `success: false` - 支付失败或订单过期")
            lines.append("")
            lines.append("**注意事项**：")
            lines.append("- 订单创建后 30 分钟内有效，超时后需要重新创建")
            lines.append("- 建议在支付完成后立即验证，避免状态延迟")
            lines.append("- 支持通过 webhook 异步通知，但建议同时使用 verify 接口确认")
        
        return lines
    
    def _is_payment_interface(self) -> bool:
        """判断是否是支付接口"""
        path = self.route_info.get('path', '')
        return 'payment' in path.lower()


class FeishuClient:
    """飞书文档客户端 - 支持直接追加内容到飞书文档"""
    
    def __init__(self, feishu_url: str, app_id: Optional[str] = None, app_secret: Optional[str] = None, provider: Optional[str] = None):
        self.feishu_url = feishu_url
        self.app_id = app_id or os.getenv("FEISHU_APP_ID")
        self.app_secret = app_secret or os.getenv("FEISHU_APP_SECRET")
        self.document_id = self._extract_document_id(feishu_url)
        self.token = None
        self.provider = provider  # 保存 provider 信息用于文件名
    
    def _extract_document_id(self, url: str) -> Optional[str]:
        """从飞书文档 URL 中提取 document_id"""
        # 飞书文档 URL 格式: https://xxx.feishu.cn/docx/IHTKdY4rvop4BHx0N8Zc69N5nRb
        # 或: https://xxx.feishu.cn/docx/IHTKdY4rvop4BHx0N8Zc69N5nRb?xxx
        import re
        match = re.search(r'/docx/([A-Za-z0-9]+)', url)
        if match:
            return match.group(1)
        return None
    
    def _get_tenant_access_token(self) -> Optional[str]:
        """获取 tenant_access_token"""
        if not self.app_id or not self.app_secret:
            return None
        
        try:
            import requests
            url = "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal/"
            data = {
                "app_id": self.app_id,
                "app_secret": self.app_secret
            }
            response = requests.post(url, json=data, timeout=10)
            result = response.json()
            
            if result.get("code") == 0:
                return result.get("tenant_access_token")
            else:
                print(f"❌ 获取飞书 token 失败: {result.get('msg')}")
                return None
        except Exception as e:
            print(f"❌ 获取飞书 token 异常: {e}")
            return None
    
    def _get_document_blocks(self) -> Optional[List[Dict[str, Any]]]:
        """获取文档的所有 blocks"""
        if not self.token:
            self.token = self._get_tenant_access_token()
            if not self.token:
                return None
        
        try:
            import requests
            url = f"https://open.feishu.cn/open-apis/docx/v1/documents/{self.document_id}/blocks"
            headers = {
                "Authorization": f"Bearer {self.token}",
                "Content-Type": "application/json"
            }
            response = requests.get(url, headers=headers, timeout=10)
            result = response.json()
            
            if result.get("code") == 0:
                return result.get("data", {}).get("items", [])
            else:
                print(f"❌ 获取文档 blocks 失败: {result.get('msg')}")
                return None
        except Exception as e:
            print(f"❌ 获取文档 blocks 异常: {e}")
            return None
    
    def _markdown_to_blocks(self, markdown_content: str) -> List[Dict[str, Any]]:
        """将 Markdown 内容转换为飞书 blocks 格式（简化版）"""
        blocks = []
        lines = markdown_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            if not stripped:
                # 空行，添加一个空文本块
                blocks.append({
                    "block_type": 1,
                    "text": {
                        "elements": [{
                            "text_run": {
                                "content": ""
                            }
                        }]
                    }
                })
                i += 1
                continue
            
            # 处理标题
            if stripped.startswith('## '):
                title_text = stripped[3:].strip()
                blocks.append({
                    "block_type": 2,
                    "heading1": {
                        "elements": [{
                            "text_run": {
                                "content": title_text
                            }
                        }]
                    }
                })
            elif stripped.startswith('### '):
                title_text = stripped[4:].strip()
                blocks.append({
                    "block_type": 3,
                    "heading2": {
                        "elements": [{
                            "text_run": {
                                "content": title_text
                            }
                        }]
                    }
                })
            elif stripped.startswith('#### '):
                title_text = stripped[5:].strip()
                blocks.append({
                    "block_type": 4,
                    "heading3": {
                        "elements": [{
                            "text_run": {
                                "content": title_text
                            }
                        }]
                    }
                })
            # 处理表格 - 简化处理，将表格作为文本块
            elif stripped.startswith('|'):
                # 收集整个表格
                table_lines = []
                j = i
                while j < len(lines) and lines[j].strip().startswith('|'):
                    table_lines.append(lines[j].rstrip())
                    j += 1
                
                # 将表格作为代码块格式插入（保持格式）
                table_text = '\n'.join(table_lines)
                blocks.append({
                    "block_type": 15,  # 代码块
                    "code": {
                        "language": 1,
                        "elements": [{
                            "text_run": {
                                "content": table_text
                            }
                        }]
                    }
                })
                i = j - 1
            # 处理代码块
            elif stripped.startswith('```'):
                code_lines = []
                j = i + 1
                while j < len(lines) and not lines[j].strip().startswith('```'):
                    code_lines.append(lines[j])
                    j += 1
                
                if code_lines:
                    code_text = '\n'.join(code_lines).rstrip()
                    blocks.append({
                        "block_type": 15,
                        "code": {
                            "language": 1,
                            "elements": [{
                                "text_run": {
                                    "content": code_text
                                }
                            }]
                        }
                    })
                    i = j
            # 处理列表项
            elif stripped.startswith('- ') or stripped.startswith('* '):
                content = stripped[2:].strip()
                # 移除粗体标记
                if content.startswith('**') and content.endswith('**'):
                    content = content[2:-2]
                
                blocks.append({
                    "block_type": 11,
                    "bullet": {
                        "elements": [{
                            "text_run": {
                                "content": content
                            }
                        }]
                    }
                })
            # 处理普通文本
            else:
                # 简单处理粗体
                content = stripped
                elements = []
                
                # 简单的粗体处理（**text**）
                import re
                parts = re.split(r'(\*\*[^*]+\*\*)', content)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # 粗体文本
                        text = part[2:-2]
                        elements.append({
                            "text_run": {
                                "content": text,
                                "text_element_style": {
                                    "bold": True
                                }
                            }
                        })
                    elif part:
                        # 普通文本
                        elements.append({
                            "text_run": {
                                "content": part
                            }
                        })
                
                if not elements:
                    elements = [{
                        "text_run": {
                            "content": content
                        }
                    }]
                
                blocks.append({
                    "block_type": 1,
                    "text": {
                        "elements": elements
                    }
                })
            
            i += 1
        
        return blocks
    
    def append_content(self, content: str) -> bool:
        """追加内容到飞书文档"""
        # 如果没有配置 App ID 和 App Secret，降级到导出文件
        if not self.app_id or not self.app_secret:
            print("⚠️  未配置飞书 App ID 和 App Secret，将导出 Markdown 文件")
            print("💡 提示：设置环境变量 FEISHU_APP_ID 和 FEISHU_APP_SECRET 可启用直接写入功能")
            return self._export_to_file(content)
        
        if not self.document_id:
            print("❌ 无法从 URL 中提取文档 ID")
            return self._export_to_file(content)
        
        # 获取 token
        self.token = self._get_tenant_access_token()
        if not self.token:
            print("❌ 无法获取飞书访问令牌，将导出 Markdown 文件")
            return self._export_to_file(content)
        
        # 获取文档 blocks
        blocks = self._get_document_blocks()
        if blocks is None:
            print("❌ 无法获取文档内容，将导出 Markdown 文件")
            return self._export_to_file(content)
        
        if not blocks:
            print("❌ 文档为空，将导出 Markdown 文件")
            return self._export_to_file(content)
        
        # 获取最后一个 block 的 ID
        last_block = blocks[-1]
        last_block_id = last_block.get("block_id")
        
        if not last_block_id:
            print("❌ 无法获取最后一个块的 ID，将导出 Markdown 文件")
            return self._export_to_file(content)
        
        # 将 Markdown 转换为 blocks
        new_blocks = self._markdown_to_blocks(content)
        
        if not new_blocks:
            print("❌ 无法将内容转换为 blocks，将导出 Markdown 文件")
            return self._export_to_file(content)
        
        # 追加内容
        try:
            import requests
            url = f"https://open.feishu.cn/open-apis/docx/v1/blocks/{last_block_id}/children"
            headers = {
                "Authorization": f"Bearer {self.token}",
                "Content-Type": "application/json"
            }
            
            # 分批插入 blocks（飞书 API 可能有限制）
            batch_size = 10
            for i in range(0, len(new_blocks), batch_size):
                batch = new_blocks[i:i+batch_size]
                data = {
                    "children": batch,
                    "index": -1  # 追加到末尾
                }
                
                response = requests.post(url, json=data, headers=headers, timeout=30)
                result = response.json()
                
                if result.get("code") != 0:
                    print(f"❌ 追加内容失败: {result.get('msg')}")
                    return self._export_to_file(content)
                
                # 更新 last_block_id 为刚插入的最后一个 block
                if result.get("data", {}).get("children"):
                    inserted_blocks = result.get("data", {}).get("children", [])
                    if inserted_blocks:
                        last_block_id = inserted_blocks[-1].get("block_id")
                        url = f"https://open.feishu.cn/open-apis/docx/v1/blocks/{last_block_id}/children"
            
            print(f"✅ 内容已成功追加到飞书文档: {self.feishu_url}")
            return True
            
        except Exception as e:
            print(f"❌ 追加内容异常: {e}")
            return self._export_to_file(content)
    
    def _export_to_file(self, content: str) -> bool:
        """导出到文件（降级方案）- 优先导出 Word 文档到桌面"""
        # 优先导出 Word 文档到桌面
        if DOCX_AVAILABLE:
            return self._export_to_word(content)
        
        # 如果没有 python-docx，导出 Markdown
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = project_root / "docs" / f"api_doc_append_{timestamp}.md"
        
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"✅ 文档已生成: {output_file}")
        print(f"📋 请手动复制内容到飞书文档: {self.feishu_url}")
        print(f"\n📄 文档内容预览（前500字符）:")
        print("-" * 80)
        print(content[:500])
        print("-" * 80)
        
        return True
    
    def _export_to_word(self, content: str) -> bool:
        """导出到 Word 文档（保存到桌面）"""
        try:
            # 获取桌面路径
            desktop_path = Path.home() / "Desktop"
            if not desktop_path.exists():
                # 尝试其他可能的桌面路径
                desktop_path = Path.home() / "桌面"
                if not desktop_path.exists():
                    desktop_path = project_root / "docs"
            
            # 生成更简单的文件名（包含 provider 信息）
            provider_name = ""
            if self.provider:
                provider_name = f"_{self.provider.upper()}"
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            word_file = desktop_path / f"API文档{provider_name}_{timestamp}.docx"
            
            # 创建 Word 文档
            doc = Document()
            
            # 设置中文字体
            doc.styles['Normal'].font.name = '微软雅黑'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            doc.styles['Normal'].font.size = Pt(10.5)
            
            # 解析 Markdown 并写入 Word
            self._markdown_to_word(doc, content)
            
            # 保存文档
            doc.save(str(word_file))
            
            print(f"✅ Word 文档已生成: {word_file}")
            print(f"📋 文档已保存到桌面，可以直接打开查看")
            
            return True
            
        except Exception as e:
            print(f"❌ 生成 Word 文档失败: {e}")
            # 降级到 Markdown
            return self._export_markdown_fallback(content)
    
    def _markdown_to_word(self, doc: Document, markdown_content: str):
        """将 Markdown 内容写入 Word 文档"""
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            if not stripped:
                doc.add_paragraph()
                i += 1
                continue
            
            # 处理标题
            if stripped.startswith('## '):
                title = doc.add_heading(stripped[3:].strip(), level=1)
                title.runs[0].font.size = Pt(16)
            elif stripped.startswith('### '):
                title = doc.add_heading(stripped[4:].strip(), level=2)
                title.runs[0].font.size = Pt(14)
            elif stripped.startswith('#### '):
                title = doc.add_heading(stripped[5:].strip(), level=3)
                title.runs[0].font.size = Pt(12)
            # 处理表格
            elif stripped.startswith('|'):
                # 收集表格行
                table_rows = []
                j = i
                while j < len(lines) and lines[j].strip().startswith('|'):
                    row_line = lines[j].strip()
                    if not row_line.startswith('|---'):  # 跳过分隔行
                        cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                        table_rows.append(cells)
                    j += 1
                
                if table_rows:
                    # 创建表格
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data
                                # 第一行设为粗体
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                    
                    i = j - 1
            # 处理代码块
            elif stripped.startswith('```'):
                code_lines = []
                j = i + 1
                while j < len(lines) and not lines[j].strip().startswith('```'):
                    code_lines.append(lines[j])
                    j += 1
                
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_para.style = 'No Spacing'
                    run = code_para.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                    # 添加灰色背景（通过段落格式）
                    code_para.paragraph_format.left_indent = Pt(20)
                    i = j
            # 处理列表项
            elif stripped.startswith('- ') or stripped.startswith('* '):
                content = stripped[2:].strip()
                # 移除粗体标记
                if content.startswith('**') and content.endswith('**'):
                    content = content[2:-2]
                    para = doc.add_paragraph(content, style='List Bullet')
                    para.runs[0].bold = True
                else:
                    doc.add_paragraph(content, style='List Bullet')
            # 处理普通文本
            else:
                para = doc.add_paragraph()
                # 处理粗体
                import re
                parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        para.add_run(part)
            
            i += 1
    
    def _export_markdown_fallback(self, content: str) -> bool:
        """降级方案：导出 Markdown 文件"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = project_root / "docs" / f"api_doc_append_{timestamp}.md"
        
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"✅ Markdown 文档已生成: {output_file}")
        return True


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='生成接口文档并追加到飞书文档')
    parser.add_argument('--file', required=True, help='接口文件路径，如 server/api/v1/unified_payment.py')
    parser.add_argument('--route', help='接口路径，如 /payment/unified/create（可选，默认解析所有）')
    parser.add_argument('--provider', help='特定提供者（可选，用于支付接口），如 stripe, payermax')
    parser.add_argument('--feishu-url', required=True, help='飞书文档URL')
    parser.add_argument('--feishu-app-id', help='飞书 App ID（可选，也可通过环境变量 FEISHU_APP_ID 设置）')
    parser.add_argument('--feishu-app-secret', help='飞书 App Secret（可选，也可通过环境变量 FEISHU_APP_SECRET 设置）')
    parser.add_argument('--output', default='markdown', choices=['markdown', 'json'], help='输出格式')
    
    args = parser.parse_args()
    
    # 解析文件路径
    file_path = project_root / args.file
    if not file_path.exists():
        print(f"❌ 文件不存在: {file_path}")
        return 1
    
    # 解析 FastAPI 文件
    print(f"📖 正在解析文件: {file_path}")
    fastapi_parser = FastAPIParser(file_path)
    routes = fastapi_parser.find_routes()
    
    if not routes:
        print("❌ 未找到任何路由定义")
        return 1
    
    # 如果指定了 route，筛选匹配的路由
    if args.route:
        routes = [r for r in routes if r.get('path') == args.route]
        if not routes:
            print(f"❌ 未找到路径为 {args.route} 的路由")
            return 1
    
    # 生成文档（如果指定了 route，只生成第一个匹配的）
    route_info = routes[0]
    print(f"📝 正在生成文档: {route_info.get('path')} ({route_info.get('method')})")
    
    doc_generator = DocGenerator(fastapi_parser, route_info, args.provider)
    doc_content = doc_generator.generate_doc()
    
    # 追加到飞书文档
    print(f"📤 正在追加到飞书文档...")
    feishu_client = FeishuClient(
        args.feishu_url,
        app_id=args.feishu_app_id,
        app_secret=args.feishu_app_secret,
        provider=args.provider
    )
    feishu_client.append_content(doc_content)
    
    print("\n✅ 完成！")
    return 0


if __name__ == '__main__':
    sys.exit(main())
