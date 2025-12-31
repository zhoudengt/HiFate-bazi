#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
清理文档中重复的接口（删除不带示例数据的旧版本）
"""

import sys
import os

# 添加项目根目录到路径
project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, project_root)

try:
    from docx import Document
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


def find_interface_sections(doc):
    """找到所有接口章节的位置"""
    interfaces = [
        '八字命理-感情婚姻',
        '八字命理-事业财富',
        '八字命理-子女学习',
        '八字命理-身体健康分析',
        '八字命理-总评分析',
        '八字命理-AI问答'
    ]
    
    positions = {}
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 1' and para.text in interfaces:
            interface_name = para.text
            if interface_name not in positions:
                positions[interface_name] = []
            positions[interface_name].append(i)
    
    return positions


def find_section_end(doc, start_pos):
    """找到接口章节的结束位置（下一个一级标题或文档结尾）"""
    for i in range(start_pos + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style.name == 'Heading 1':
            return i
    return len(doc.paragraphs)


def has_examples_section(doc, start_pos, end_pos):
    """检查接口章节是否包含示例数据部分"""
    for i in range(start_pos, end_pos):
        if doc.paragraphs[i].text == '示例数据':
            return True
    return False


def remove_duplicate_sections(doc):
    """删除重复的接口章节（保留带示例数据的版本）"""
    positions = find_interface_sections(doc)
    
    # 找到需要删除的位置（保留最后一个，删除之前的）
    to_delete = []
    for interface_name, pos_list in positions.items():
        if len(pos_list) > 1:
            # 检查每个版本是否有示例数据
            versions = []
            for pos in pos_list:
                end_pos = find_section_end(doc, pos)
                has_examples = has_examples_section(doc, pos, end_pos)
                versions.append({
                    'start': pos,
                    'end': end_pos,
                    'has_examples': has_examples
                })
            
            # 保留有示例数据的版本，如果没有则保留最后一个
            keep_version = None
            for v in versions:
                if v['has_examples']:
                    keep_version = v
                    break
            
            if not keep_version:
                keep_version = versions[-1]  # 保留最后一个
            
            # 标记其他版本为删除
            for v in versions:
                if v != keep_version:
                    to_delete.append((v['start'], v['end']))
    
    # 按位置倒序删除（避免索引变化问题）
    to_delete.sort(reverse=True)
    
    deleted_count = 0
    for start_pos, end_pos in to_delete:
        # 删除段落（从后往前删除）
        for i in range(end_pos - 1, start_pos - 1, -1):
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        deleted_count += 1
        print(f"  删除重复接口章节（位置 {start_pos}-{end_pos}）")
    
    return deleted_count


def main():
    """主函数"""
    print("正在清理文档中的重复接口...")
    
    # 桌面路径
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
    doc_path = os.path.join(desktop_path, 'bazi_shengong_minggong_api.docx')
    
    # 检查文件是否存在
    if not os.path.exists(doc_path):
        print(f"错误：找不到文件 {doc_path}")
        sys.exit(1)
    
    # 打开文档
    doc = Document(doc_path)
    
    # 删除重复章节
    deleted_count = remove_duplicate_sections(doc)
    
    if deleted_count > 0:
        # 保存文档
        print("正在保存文档...")
        doc.save(doc_path)
        print(f"✅ 已删除 {deleted_count} 个重复的接口章节")
    else:
        print("✅ 未发现重复的接口章节")


if __name__ == '__main__':
    main()

