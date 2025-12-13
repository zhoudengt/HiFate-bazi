#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成前端接口文档（docx格式）
只包含前端使用的接口，不包括对内的服务间接口
"""

import sys
import os
import re
import ast
from typing import Dict, Any, List
from datetime import datetime

# 添加项目根目录到路径
project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, project_root)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


# 从 grpc_gateway.py 读取已注册的端点
def get_supported_endpoints() -> List[str]:
    """从 grpc_gateway.py 读取已注册的端点，以及直接调用 REST API 的接口"""
    gateway_file = os.path.join(project_root, 'server', 'api', 'grpc_gateway.py')
    endpoints = []
    
    with open(gateway_file, 'r', encoding='utf-8') as f:
        content = f.read()
        # 查找所有 @_register 装饰器
        pattern = r'@_register\("([^"]+)"\)'
        matches = re.findall(pattern, content)
        endpoints.extend(matches)
    
    # 添加直接调用 REST API 的接口（不在 gRPC-Web 网关中）
    direct_api_endpoints = [
        '/fortune/hand/analyze/stream',
        '/fortune/face/analyze/stream',
        '/smart-fortune/smart-analyze-stream',
    ]
    endpoints.extend(direct_api_endpoints)
    
    return sorted(set(endpoints))


# 接口信息定义（基于已读取的代码文件）
ENDPOINT_INFO = {
    '/bazi/pan/display': {
        'description': '排盘展示（前端优化）',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'pan', 'type': 'dict', 'description': '排盘数据', 'nested': [
                {'name': 'basic', 'type': 'dict', 'description': '基本信息（包含阳历、农历、性别等）'},
                {'name': 'pillars', 'type': 'list', 'description': '四柱数组（年柱、月柱、日柱、时柱）', 'nested': [
                    {'name': 'type', 'type': 'str', 'description': '柱类型（year/month/day/hour）'},
                    {'name': 'label', 'type': 'str', 'description': '柱标签（年柱/月柱/日柱/时柱）'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'element', 'type': 'str', 'description': '五行属性'},
                        {'name': 'yinyang', 'type': 'str', 'description': '阴阳属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'element', 'type': 'str', 'description': '五行属性'},
                        {'name': 'yinyang', 'type': 'str', 'description': '阴阳属性'},
                    ]},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'hidden_stars', 'type': 'list', 'description': '副星列表'},
                    {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                ]},
                {'name': 'wuxing', 'type': 'dict', 'description': '五行统计（包含各五行数量和百分比）', 'nested': [
                    {'name': 'counts', 'type': 'dict', 'description': '五行数量统计'},
                    {'name': 'percentages', 'type': 'dict', 'description': '五行百分比统计'},
                ]},
                {'name': 'rizhu_analysis', 'type': 'dict', 'description': '日柱解析（性格与命运解析）', 'nested': [
                    {'name': 'rizhu', 'type': 'str', 'description': '日柱干支'},
                    {'name': 'gender', 'type': 'str', 'description': '性别'},
                    {'name': 'descriptions', 'type': 'list', 'description': '描述列表'},
                    {'name': 'summary', 'type': 'str', 'description': '总结'},
                ]},
                {'name': 'marriage_rules', 'type': 'list', 'description': '婚姻规则列表', 'nested': [
                    {'name': 'rule_id', 'type': 'str', 'description': '规则ID'},
                    {'name': 'rule_type', 'type': 'str', 'description': '规则类型'},
                    {'name': 'content', 'type': 'dict', 'description': '规则内容'},
                    {'name': 'priority', 'type': 'int', 'description': '优先级'},
                ]},
            ]},
        ],
    },
    '/bazi/fortune/display': {
        'description': '大运流年流月统一接口（性能优化）',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'current_time', 'type': 'str', 'required': False, 'description': '当前时间（可选），格式：YYYY-MM-DD HH:MM', 'example': '2024-01-01 12:00'},
            {'name': 'dayun_year_start', 'type': 'int', 'required': False, 'description': '大运起始年份（可选）', 'example': 2020},
            {'name': 'dayun_year_end', 'type': 'int', 'required': False, 'description': '大运结束年份（可选）', 'example': 2030},
            {'name': 'target_year', 'type': 'int', 'required': False, 'description': '目标年份（可选），用于计算该年份的流月', 'example': 2024},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'pillars', 'type': 'dict', 'description': '四柱详细信息'},
            {'name': 'dayun', 'type': 'dict', 'description': '大运数据', 'nested': [
                {'name': 'current', 'type': 'dict', 'description': '当前大运', 'nested': [
                    {'name': 'index', 'type': 'int', 'description': '大运序号'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'age_range', 'type': 'dict', 'description': '年龄范围', 'nested': [
                        {'name': 'start', 'type': 'int', 'description': '起始年龄'},
                        {'name': 'end', 'type': 'int', 'description': '结束年龄'},
                    ]},
                    {'name': 'year_range', 'type': 'dict', 'description': '年份范围', 'nested': [
                        {'name': 'start', 'type': 'int', 'description': '起始年份'},
                        {'name': 'end', 'type': 'int', 'description': '结束年份'},
                    ]},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前大运'},
                ]},
                {'name': 'list', 'type': 'list', 'description': '大运列表（数组格式，每个元素结构同current）'},
                {'name': 'qiyun', 'type': 'dict', 'description': '起运信息', 'nested': [
                    {'name': 'date', 'type': 'str', 'description': '起运日期'},
                    {'name': 'age_display', 'type': 'str', 'description': '起运年龄显示'},
                    {'name': 'description', 'type': 'str', 'description': '描述'},
                ]},
                {'name': 'jiaoyun', 'type': 'dict', 'description': '交运信息', 'nested': [
                    {'name': 'date', 'type': 'str', 'description': '交运日期'},
                    {'name': 'age_display', 'type': 'str', 'description': '交运年龄显示'},
                    {'name': 'description', 'type': 'str', 'description': '描述'},
                ]},
            ]},
            {'name': 'liunian', 'type': 'dict', 'description': '流年数据', 'nested': [
                {'name': 'current', 'type': 'dict', 'description': '当前流年', 'nested': [
                    {'name': 'year', 'type': 'int', 'description': '年份'},
                    {'name': 'age', 'type': 'int', 'description': '年龄'},
                    {'name': 'age_display', 'type': 'str', 'description': '年龄显示'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前流年'},
                ]},
                {'name': 'list', 'type': 'list', 'description': '流年列表（数组格式，每个元素结构同current）'},
            ]},
            {'name': 'liuyue', 'type': 'dict', 'description': '流月数据', 'nested': [
                {'name': 'current', 'type': 'dict', 'description': '当前流月', 'nested': [
                    {'name': 'month', 'type': 'int', 'description': '月份（1-12）'},
                    {'name': 'solar_term', 'type': 'str', 'description': '节气'},
                    {'name': 'term_date', 'type': 'str', 'description': '节气日期'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                ]},
                {'name': 'list', 'type': 'list', 'description': '流月列表（数组格式，每个元素结构同current）'},
                {'name': 'target_year', 'type': 'int', 'description': '目标年份'},
            ]},
        ],
    },
    '/bazi/dayun/display': {
        'description': '大运展示（前端优化）',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'current_time', 'type': 'str', 'required': False, 'description': '当前时间（可选），格式：YYYY-MM-DD HH:MM', 'example': '2024-01-01 12:00'},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'dayun', 'type': 'dict', 'description': '大运数据', 'nested': [
                {'name': 'current', 'type': 'dict', 'description': '当前大运', 'nested': [
                    {'name': 'index', 'type': 'int', 'description': '大运序号'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息'},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息'},
                    {'name': 'age_range', 'type': 'dict', 'description': '年龄范围'},
                    {'name': 'year_range', 'type': 'dict', 'description': '年份范围'},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前大运'},
                ]},
                {'name': 'list', 'type': 'list', 'description': '大运列表（数组格式，每个元素结构同current）'},
                {'name': 'qiyun', 'type': 'dict', 'description': '起运信息', 'nested': [
                    {'name': 'date', 'type': 'str', 'description': '起运日期'},
                    {'name': 'age_display', 'type': 'str', 'description': '起运年龄显示'},
                    {'name': 'description', 'type': 'str', 'description': '描述'},
                ]},
                {'name': 'jiaoyun', 'type': 'dict', 'description': '交运信息', 'nested': [
                    {'name': 'date', 'type': 'str', 'description': '交运日期'},
                    {'name': 'age_display', 'type': 'str', 'description': '交运年龄显示'},
                    {'name': 'description', 'type': 'str', 'description': '描述'},
                ]},
            ]},
        ],
    },
    '/bazi/liunian/display': {
        'description': '流年展示（前端优化）',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'year_range', 'type': 'dict', 'required': False, 'description': '年份范围（可选），如：{"start": 2020, "end": 2030}', 'example': {'start': 2020, 'end': 2030}},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'liunian', 'type': 'dict', 'description': '流年数据', 'nested': [
                {'name': 'current', 'type': 'dict', 'description': '当前流年', 'nested': [
                    {'name': 'year', 'type': 'int', 'description': '年份'},
                    {'name': 'age', 'type': 'int', 'description': '年龄'},
                    {'name': 'age_display', 'type': 'str', 'description': '年龄显示'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息'},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息'},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前流年'},
                ]},
                {'name': 'list', 'type': 'list', 'description': '流年列表（数组格式，每个元素结构同current）'},
            ]},
        ],
    },
    '/bazi/liuyue/display': {
        'description': '流月展示（前端优化）',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'target_year', 'type': 'int', 'required': False, 'description': '目标年份（可选），用于计算该年份的流月', 'example': 2024},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'liuyue', 'type': 'dict', 'description': '流月数据', 'nested': [
                {'name': 'current', 'type': 'dict', 'description': '当前流月', 'nested': [
                    {'name': 'month', 'type': 'int', 'description': '月份（1-12）'},
                    {'name': 'solar_term', 'type': 'str', 'description': '节气'},
                    {'name': 'term_date', 'type': 'str', 'description': '节气日期'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                ]},
                {'name': 'list', 'type': 'list', 'description': '流月列表（数组格式，每个元素结构同current）'},
                {'name': 'target_year', 'type': 'int', 'description': '目标年份'},
            ]},
        ],
    },
    '/bazi/wangshuai': {
        'description': '计算旺衰',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '旺衰分析结果', 'nested': [
                {'name': 'wangshuai', 'type': 'str', 'description': '旺衰类型（身旺/身弱等）'},
                {'name': 'analysis', 'type': 'str', 'description': '旺衰分析说明'},
            ]},
        ],
    },
    '/bazi/yigua/divinate': {
        'description': '易卦占卜',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'question', 'type': 'str', 'required': False, 'description': '占卜问题（可选）', 'example': '我今年的事业运势如何？'},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '易卦占卜结果', 'nested': [
                {'name': 'gua', 'type': 'str', 'description': '卦名'},
                {'name': 'gua_ci', 'type': 'str', 'description': '卦辞'},
                {'name': 'yao_ci', 'type': 'list', 'description': '爻辞列表'},
                {'name': 'interpretation', 'type': 'str', 'description': '解卦说明'},
            ]},
        ],
    },
    '/bazi/interface': {
        'description': '生成八字界面信息（包含命宫、身宫、胎元、胎息、命卦等）',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'name', 'type': 'str', 'required': False, 'description': '姓名（可选）', 'example': '张三'},
            {'name': 'location', 'type': 'str', 'required': False, 'description': '出生地点（可选）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': False, 'description': '纬度（可选）', 'example': 39.90},
            {'name': 'longitude', 'type': 'float', 'required': False, 'description': '经度（可选）', 'example': 116.40},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '完整的八字界面信息', 'nested': [
                {'name': 'palaces', 'type': 'dict', 'description': '宫位信息', 'nested': [
                    {'name': 'life_palace', 'type': 'dict', 'description': '命宫'},
                    {'name': 'body_palace', 'type': 'dict', 'description': '身宫'},
                ]},
                {'name': 'taiyuan', 'type': 'str', 'description': '胎元'},
                {'name': 'taixi', 'type': 'str', 'description': '胎息'},
                {'name': 'minggua', 'type': 'str', 'description': '命卦'},
                {'name': 'bazi_pillars', 'type': 'dict', 'description': '四柱信息'},
                {'name': 'basic_info', 'type': 'dict', 'description': '基本信息'},
            ]},
        ],
    },
    '/bazi/shengong-minggong': {
        'description': '获取身宫命宫详细信息（主星、藏干、星运、自坐、空亡、纳音、神煞等）',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '身宫命宫详细信息', 'nested': [
                {'name': 'shengong', 'type': 'dict', 'description': '身宫信息', 'nested': [
                    {'name': 'stem', 'type': 'dict', 'description': '天干', 'nested': [{'name': 'char', 'type': 'str', 'description': '天干字符'}]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支', 'nested': [{'name': 'char', 'type': 'str', 'description': '地支字符'}]},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'hidden_stems', 'type': 'list', 'description': '藏干'},
                    {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                    {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                    {'name': 'kongwang', 'type': 'bool', 'description': '空亡'},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'deities', 'type': 'list', 'description': '神煞'},
                ]},
                {'name': 'minggong', 'type': 'dict', 'description': '命宫信息', 'nested': [
                    {'name': 'stem', 'type': 'dict', 'description': '天干', 'nested': [{'name': 'char', 'type': 'str', 'description': '天干字符'}]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支', 'nested': [{'name': 'char', 'type': 'str', 'description': '地支字符'}]},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'hidden_stems', 'type': 'list', 'description': '藏干'},
                    {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                    {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                    {'name': 'kongwang', 'type': 'bool', 'description': '空亡'},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'deities', 'type': 'list', 'description': '神煞'},
                ]},
                {'name': 'pillars', 'type': 'dict', 'description': '四柱详细信息', 'nested': [
                    {'name': 'year', 'type': 'dict', 'description': '年柱详细信息'},
                    {'name': 'month', 'type': 'dict', 'description': '月柱详细信息'},
                    {'name': 'day', 'type': 'dict', 'description': '日柱详细信息'},
                    {'name': 'hour', 'type': 'dict', 'description': '时柱详细信息'},
                ]},
            ]},
        ],
    },
    '/bazi/formula-analysis': {
        'description': '算法公式规则分析',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '阳历时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male/female', 'example': 'male'},
            {'name': 'rule_types', 'type': 'list[str]', 'required': False, 'description': '规则类型列表，可选值：wealth/marriage/career/children/character/summary/health/peach_blossom/shishen/parents', 'example': ['wealth', 'career']},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '公式分析结果', 'nested': [
                {'name': 'bazi_info', 'type': 'dict', 'description': '八字基本信息', 'nested': [
                    {'name': 'solar_date', 'type': 'str', 'description': '阳历日期'},
                    {'name': 'solar_time', 'type': 'str', 'description': '出生时间'},
                    {'name': 'gender', 'type': 'str', 'description': '性别'},
                    {'name': 'bazi_pillars', 'type': 'dict', 'description': '四柱信息'},
                    {'name': 'day_stem', 'type': 'str', 'description': '日干'},
                    {'name': 'day_branch', 'type': 'str', 'description': '日支'},
                ]},
                {'name': 'bazi_data', 'type': 'dict', 'description': '完整八字数据'},
                {'name': 'matched_rules', 'type': 'dict', 'description': '匹配的规则（按类型分组）', 'nested': [
                    {'name': 'wealth', 'type': 'list', 'description': '财富规则ID列表'},
                    {'name': 'marriage', 'type': 'list', 'description': '婚姻规则ID列表'},
                    {'name': 'career', 'type': 'list', 'description': '事业规则ID列表'},
                    {'name': 'children', 'type': 'list', 'description': '子女规则ID列表'},
                    {'name': 'character', 'type': 'list', 'description': '性格规则ID列表'},
                    {'name': 'summary', 'type': 'list', 'description': '总评规则ID列表'},
                    {'name': 'health', 'type': 'list', 'description': '身体规则ID列表'},
                    {'name': 'peach_blossom', 'type': 'list', 'description': '桃花规则ID列表'},
                    {'name': 'shishen', 'type': 'list', 'description': '十神命格规则ID列表'},
                    {'name': 'parents', 'type': 'list', 'description': '父母规则ID列表'},
                ]},
                {'name': 'rule_details', 'type': 'dict', 'description': '规则详情（以规则ID为键）'},
                {'name': 'statistics', 'type': 'dict', 'description': '统计信息', 'nested': [
                    {'name': 'total_matched', 'type': 'int', 'description': '总匹配数'},
                    {'name': 'wealth_count', 'type': 'int', 'description': '财富规则数量'},
                    {'name': 'marriage_count', 'type': 'int', 'description': '婚姻规则数量'},
                    {'name': 'career_count', 'type': 'int', 'description': '事业规则数量'},
                    {'name': 'children_count', 'type': 'int', 'description': '子女规则数量'},
                    {'name': 'character_count', 'type': 'int', 'description': '性格规则数量'},
                    {'name': 'summary_count', 'type': 'int', 'description': '总评规则数量'},
                    {'name': 'health_count', 'type': 'int', 'description': '身体规则数量'},
                    {'name': 'peach_blossom_count', 'type': 'int', 'description': '桃花规则数量'},
                    {'name': 'shishen_count', 'type': 'int', 'description': '十神命格规则数量'},
                    {'name': 'parents_count', 'type': 'int', 'description': '父母规则数量'},
                ]},
            ]},
            {'name': 'error', 'type': 'str', 'description': '错误信息（失败时）'},
        ],
    },
    '/smart-analyze': {
        'description': '智能运势分析 - 自动识别用户问题意图，返回针对性的分析结果',
        'request_fields': [
            {'name': 'question', 'type': 'str', 'required': True, 'description': '用户问题', 'example': '我今年的事业运势如何？'},
            {'name': 'year', 'type': 'int', 'required': True, 'description': '出生年份', 'example': 1990},
            {'name': 'month', 'type': 'int', 'required': True, 'description': '出生月份', 'example': 5},
            {'name': 'day', 'type': 'int', 'required': True, 'description': '出生日期', 'example': 15},
            {'name': 'hour', 'type': 'int', 'required': False, 'description': '出生时辰（0-23）', 'example': 12, 'default': 12},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别（male/female）', 'example': 'male'},
            {'name': 'user_id', 'type': 'str', 'required': False, 'description': '用户ID', 'example': None},
            {'name': 'include_fortune_context', 'type': 'bool', 'required': False, 'description': '是否包含流年大运分析（实验性功能，默认关闭）', 'example': False, 'default': False},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'question', 'type': 'str', 'description': '用户问题'},
            {'name': 'intent_result', 'type': 'dict', 'description': '意图识别结果', 'nested': [
                {'name': 'intents', 'type': 'list', 'description': '识别的意图列表'},
                {'name': 'confidence', 'type': 'float', 'description': '置信度（0-1）'},
                {'name': 'rule_types', 'type': 'list', 'description': '规则类型列表'},
                {'name': 'time_intent', 'type': 'dict', 'description': '时间意图', 'nested': [
                    {'name': 'type', 'type': 'str', 'description': '时间类型（today/this_year/next_year/future_years等）'},
                    {'name': 'target_years', 'type': 'list', 'description': '目标年份列表'},
                    {'name': 'description', 'type': 'str', 'description': '时间意图描述'},
                ]},
                {'name': 'is_ambiguous', 'type': 'bool', 'description': '是否模糊'},
                {'name': 'method', 'type': 'str', 'description': '识别方法（local_model/keyword/llm）'},
            ]},
            {'name': 'bazi_info', 'type': 'dict', 'description': '八字信息', 'nested': [
                {'name': '四柱', 'type': 'dict', 'description': '四柱信息'},
                {'name': '十神', 'type': 'dict', 'description': '十神统计'},
                {'name': '五行', 'type': 'dict', 'description': '五行统计'},
            ]},
            {'name': 'matched_rules_count', 'type': 'int', 'description': '匹配的规则数量'},
            {'name': 'response', 'type': 'str', 'description': '分析结果文本'},
            {'name': 'performance', 'type': 'dict', 'description': '性能摘要'},
        ],
    },
    '/bazi/daily-fortune': {
        'description': '今日运势分析（类似 FateTell 的日运日签）',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '用户出生日期（阳历），格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '用户出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'target_date', 'type': 'str', 'required': False, 'description': '目标日期（可选，默认为今天），格式：YYYY-MM-DD', 'example': '2025-01-17'},
            {'name': 'use_llm', 'type': 'bool', 'required': False, 'description': '是否使用 LLM 生成（可选，默认使用规则匹配）', 'example': False, 'default': False},
            {'name': 'access_token', 'type': 'str', 'required': False, 'description': 'Coze Access Token（可选，use_llm=True 时需要）', 'example': None},
            {'name': 'bot_id', 'type': 'str', 'required': False, 'description': 'Coze Bot ID（可选，use_llm=True 时需要）', 'example': None},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'target_date', 'type': 'str', 'description': '目标日期'},
            {'name': 'fortune', 'type': 'dict', 'description': '运势分析结果', 'nested': [
                {'name': 'summary', 'type': 'str', 'description': '运势总结'},
                {'name': 'career', 'type': 'str', 'description': '事业运势'},
                {'name': 'wealth', 'type': 'str', 'description': '财运'},
                {'name': 'health', 'type': 'str', 'description': '健康运势'},
                {'name': 'emotion', 'type': 'str', 'description': '感情运势'},
            ]},
            {'name': 'liuri_info', 'type': 'dict', 'description': '流日信息', 'nested': [
                {'name': 'date', 'type': 'str', 'description': '日期'},
                {'name': 'liuri', 'type': 'dict', 'description': '流日干支', 'nested': [
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'element', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                    ]},
                ]},
                {'name': 'liuyue', 'type': 'dict', 'description': '流月干支', 'nested': [
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息'},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息'},
                ]},
                {'name': 'liunian', 'type': 'dict', 'description': '流年干支', 'nested': [
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息'},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息'},
                ]},
            ]},
            {'name': 'bazi_data', 'type': 'dict', 'description': '八字数据（包含基本信息、四柱、十神、五行等）'},
            {'name': 'matched_rules_count', 'type': 'int', 'description': '匹配的规则数量'},
            {'name': 'error', 'type': 'str', 'description': '错误信息（失败时）'},
        ],
    },
    '/bazi/monthly-fortune': {
        'description': '计算月运势（基于八字）',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历出生日期，格式：YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male/female', 'example': 'male'},
            {'name': 'target_month', 'type': 'str', 'required': False, 'description': '目标月份，格式：YYYY-MM，默认为本月', 'example': '2025-01'},
            {'name': 'use_llm', 'type': 'bool', 'required': False, 'description': '是否使用LLM生成', 'example': False, 'default': False},
            {'name': 'access_token', 'type': 'str', 'required': False, 'description': 'Coze Access Token（use_llm=True时需要）', 'example': None},
            {'name': 'bot_id', 'type': 'str', 'required': False, 'description': 'Coze Bot ID（use_llm=True时需要）', 'example': None},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '月运势分析结果', 'nested': [
                {'name': 'target_month', 'type': 'str', 'description': '目标月份'},
                {'name': 'fortune', 'type': 'dict', 'description': '运势分析', 'nested': [
                    {'name': 'summary', 'type': 'str', 'description': '运势总结'},
                    {'name': 'career', 'type': 'str', 'description': '事业运势'},
                    {'name': 'wealth', 'type': 'str', 'description': '财运'},
                    {'name': 'health', 'type': 'str', 'description': '健康运势'},
                    {'name': 'emotion', 'type': 'str', 'description': '感情运势'},
                ]},
                {'name': 'liuyue_info', 'type': 'dict', 'description': '流月信息'},
                {'name': 'bazi_data', 'type': 'dict', 'description': '八字数据'},
            ]},
        ],
    },
    '/payment/create-session': {
        'description': '创建Stripe支付会话',
        'request_fields': [
            {'name': 'amount', 'type': 'str', 'required': True, 'description': '金额，格式：19.90', 'example': '19.90'},
            {'name': 'currency', 'type': 'str', 'required': False, 'description': '货币代码，默认：USD', 'example': 'USD', 'default': 'USD'},
            {'name': 'product_name', 'type': 'str', 'required': True, 'description': '产品名称，如：月订阅会员', 'example': '月订阅会员'},
            {'name': 'customer_email', 'type': 'str', 'required': True, 'description': '客户邮箱', 'example': 'user@example.com'},
            {'name': 'metadata', 'type': 'dict', 'required': False, 'description': '元数据（可选）', 'example': None},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'session_id', 'type': 'str', 'description': '支付会话ID'},
            {'name': 'checkout_url', 'type': 'str', 'description': '支付页面URL'},
            {'name': 'status', 'type': 'str', 'description': '状态'},
            {'name': 'message', 'type': 'str', 'description': '消息'},
        ],
    },
    '/payment/verify': {
        'description': '验证支付状态',
        'request_fields': [
            {'name': 'session_id', 'type': 'str', 'required': True, 'description': 'Stripe Session ID', 'example': 'cs_test_...'},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'status', 'type': 'str', 'description': '支付状态'},
            {'name': 'payment_intent_id', 'type': 'str', 'description': '支付意图ID'},
            {'name': 'amount', 'type': 'str', 'description': '金额'},
            {'name': 'currency', 'type': 'str', 'description': '货币代码'},
            {'name': 'customer_email', 'type': 'str', 'description': '客户邮箱'},
            {'name': 'created_at', 'type': 'int', 'description': '创建时间戳'},
            {'name': 'metadata', 'type': 'dict', 'description': '元数据'},
            {'name': 'message', 'type': 'str', 'description': '消息'},
        ],
    },
    '/payment/unified/create': {
        'description': '统一创建支付 - 根据provider路由到不同支付渠道（Stripe/PayPal/支付宝/微信）',
        'request_fields': [
            {'name': 'provider', 'type': 'str', 'required': True, 'description': '支付渠道：stripe/paypal/alipay/wechat', 'example': 'stripe'},
            {'name': 'amount', 'type': 'str', 'required': True, 'description': '金额，格式：19.90', 'example': '19.90'},
            {'name': 'currency', 'type': 'str', 'required': False, 'description': '货币代码', 'example': 'USD', 'default': 'USD'},
            {'name': 'product_name', 'type': 'str', 'required': True, 'description': '产品名称', 'example': '月订阅会员'},
            {'name': 'customer_email', 'type': 'str', 'required': False, 'description': '客户邮箱（Stripe必需）', 'example': 'user@example.com'},
            {'name': 'openid', 'type': 'str', 'required': False, 'description': '微信用户openid（微信JSAPI支付必需）', 'example': None},
            {'name': 'payment_type', 'type': 'str', 'required': False, 'description': '微信支付类型：native/jsapi', 'example': 'native', 'default': 'native'},
            {'name': 'metadata', 'type': 'dict', 'required': False, 'description': '元数据', 'example': None},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'provider', 'type': 'str', 'description': '支付渠道'},
            {'name': 'payment_id', 'type': 'str', 'description': '支付ID'},
            {'name': 'order_id', 'type': 'str', 'description': '订单号'},
            {'name': 'payment_url', 'type': 'str', 'description': '支付URL'},
            {'name': 'checkout_url', 'type': 'str', 'description': '支付页面URL'},
            {'name': 'approval_url', 'type': 'str', 'description': 'PayPal批准URL'},
            {'name': 'code_url', 'type': 'str', 'description': '微信支付二维码URL'},
            {'name': 'jsapi_params', 'type': 'dict', 'description': '微信JSAPI支付参数'},
            {'name': 'status', 'type': 'str', 'description': '状态'},
            {'name': 'message', 'type': 'str', 'description': '消息'},
        ],
    },
    '/payment/unified/verify': {
        'description': '统一验证支付状态',
        'request_fields': [
            {'name': 'provider', 'type': 'str', 'required': True, 'description': '支付渠道', 'example': 'stripe'},
            {'name': 'payment_id', 'type': 'str', 'required': False, 'description': '支付ID（Stripe/PayPal）', 'example': None},
            {'name': 'order_id', 'type': 'str', 'required': False, 'description': '订单号（支付宝/微信）', 'example': None},
            {'name': 'session_id', 'type': 'str', 'required': False, 'description': 'Stripe Session ID', 'example': None},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'provider', 'type': 'str', 'description': '支付渠道'},
            {'name': 'status', 'type': 'str', 'description': '支付状态'},
            {'name': 'payment_id', 'type': 'str', 'description': '支付ID'},
            {'name': 'order_id', 'type': 'str', 'description': '订单号'},
            {'name': 'amount', 'type': 'str', 'description': '金额'},
            {'name': 'currency', 'type': 'str', 'description': '货币代码'},
            {'name': 'customer_email', 'type': 'str', 'description': '客户邮箱'},
            {'name': 'paid_time', 'type': 'str', 'description': '支付时间'},
            {'name': 'message', 'type': 'str', 'description': '消息'},
        ],
    },
    '/payment/providers': {
        'description': '获取所有可用的支付渠道及其状态',
        'request_fields': [],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'providers', 'type': 'list', 'description': '支付渠道列表', 'nested': [
                {'name': 'id', 'type': 'str', 'description': '渠道ID'},
                {'name': 'name', 'type': 'str', 'description': '渠道名称'},
                {'name': 'enabled', 'type': 'bool', 'description': '是否启用'},
                {'name': 'regions', 'type': 'list', 'description': '适用地区'},
                {'name': 'currencies', 'type': 'list', 'description': '支持的货币'},
                {'name': 'description', 'type': 'str', 'description': '渠道描述'},
            ]},
        ],
    },
    '/auth/login': {
        'description': '用户登录（返回JWT Token）',
        'request_fields': [
            {'name': 'username', 'type': 'str', 'required': True, 'description': '用户名', 'example': 'admin'},
            {'name': 'password', 'type': 'str', 'required': True, 'description': '密码', 'example': 'admin123'},
        ],
        'response_fields': [
            {'name': 'access_token', 'type': 'str', 'description': 'JWT访问令牌'},
            {'name': 'token_type', 'type': 'str', 'description': '令牌类型（bearer）'},
            {'name': 'expires_in_minutes', 'type': 'int', 'description': '过期时间（分钟）'},
        ],
    },
    '/calendar/query': {
        'description': '查询万年历信息',
        'request_fields': [
            {'name': 'date', 'type': 'str', 'required': False, 'description': '日期（可选，默认为今天），格式：YYYY-MM-DD', 'example': '2024-11-14'},
            {'name': 'provider', 'type': 'str', 'required': False, 'description': 'API提供商（可选），jisuapi/tianapi/6api，默认自动选择', 'example': 'jisuapi'},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'provider', 'type': 'str', 'description': '使用的API提供商'},
            {'name': 'date', 'type': 'str', 'description': '日期'},
            {'name': 'solar_date', 'type': 'str', 'description': '公历日期'},
            {'name': 'weekday', 'type': 'str', 'description': '星期几（中文）'},
            {'name': 'weekday_en', 'type': 'str', 'description': '星期几（英文）'},
            {'name': 'lunar_date', 'type': 'str', 'description': '农历日期'},
            {'name': 'ganzhi', 'type': 'dict', 'description': '干支（年、月、日）', 'nested': [
                {'name': 'year', 'type': 'dict', 'description': '年干支'},
                {'name': 'month', 'type': 'dict', 'description': '月干支'},
                {'name': 'day', 'type': 'dict', 'description': '日干支'},
            ]},
            {'name': 'yi', 'type': 'list', 'description': '宜做什么'},
            {'name': 'ji', 'type': 'list', 'description': '忌做什么'},
            {'name': 'luck_level', 'type': 'str', 'description': '吉凶（大凶/大吉等）'},
            {'name': 'deities', 'type': 'dict', 'description': '神煞方位（福神、财神、喜神）', 'nested': [
                {'name': 'fushén', 'type': 'str', 'description': '福神方位'},
                {'name': 'cáishén', 'type': 'str', 'description': '财神方位'},
                {'name': 'xǐshén', 'type': 'str', 'description': '喜神方位'},
            ]},
            {'name': 'chong_he_sha', 'type': 'dict', 'description': '冲合煞（冲、合、煞）', 'nested': [
                {'name': 'chong', 'type': 'str', 'description': '冲'},
                {'name': 'he', 'type': 'str', 'description': '合'},
                {'name': 'sha', 'type': 'str', 'description': '煞'},
            ]},
            {'name': 'error', 'type': 'str', 'description': '错误信息（失败时）'},
        ],
    },
    '/api/v2/face/analyze': {
        'description': '面相分析V2 - 支持文件上传（base64编码）',
        'request_fields': [
            {'name': 'image_base64', 'type': 'str', 'required': True, 'description': '图片base64编码', 'example': 'data:image/jpeg;base64,...'},
            {'name': 'filename', 'type': 'str', 'required': False, 'description': '文件名', 'example': 'face.jpg'},
            {'name': 'content_type', 'type': 'str', 'required': False, 'description': '内容类型', 'example': 'image/jpeg'},
            {'name': 'analysis_types', 'type': 'str', 'required': False, 'description': '分析类型', 'example': 'gongwei,liuqin,shishen', 'default': 'gongwei,liuqin,shishen'},
            {'name': 'birth_year', 'type': 'int', 'required': False, 'description': '出生年份', 'example': 1990},
            {'name': 'birth_month', 'type': 'int', 'required': False, 'description': '出生月份', 'example': 5},
            {'name': 'birth_day', 'type': 'int', 'required': False, 'description': '出生日期', 'example': 15},
            {'name': 'birth_hour', 'type': 'int', 'required': False, 'description': '出生时辰', 'example': 12},
            {'name': 'gender', 'type': 'str', 'required': False, 'description': '性别', 'example': 'male'},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '面相分析结果', 'nested': [
                {'name': 'gongwei', 'type': 'dict', 'description': '宫位分析'},
                {'name': 'liuqin', 'type': 'dict', 'description': '六亲分析'},
                {'name': 'shishen', 'type': 'dict', 'description': '十神分析'},
                {'name': 'detected_items', 'type': 'list', 'description': '检测到的物品列表', 'nested': [
                    {'name': 'name', 'type': 'str', 'description': '物品名称'},
                    {'name': 'category', 'type': 'str', 'description': '物品类别'},
                    {'name': 'position', 'type': 'dict', 'description': '位置信息', 'nested': [
                        {'name': 'x', 'type': 'float', 'description': 'X坐标'},
                        {'name': 'y', 'type': 'float', 'description': 'Y坐标'},
                        {'name': 'width', 'type': 'float', 'description': '宽度'},
                        {'name': 'height', 'type': 'float', 'description': '高度'},
                    ]},
                    {'name': 'confidence', 'type': 'float', 'description': '置信度'},
                ]},
            ]},
        ],
    },
    '/api/v2/desk-fengshui/analyze': {
        'description': '办公桌风水分析 - 支持文件上传（base64编码）',
        'request_fields': [
            {'name': 'image_base64', 'type': 'str', 'required': True, 'description': '图片base64编码', 'example': 'data:image/jpeg;base64,...'},
            {'name': 'filename', 'type': 'str', 'required': False, 'description': '文件名', 'example': 'desk.jpg'},
            {'name': 'content_type', 'type': 'str', 'required': False, 'description': '内容类型', 'example': 'image/jpeg'},
            {'name': 'solar_date', 'type': 'str', 'required': False, 'description': '阳历日期（可选，用于结合八字分析）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': False, 'description': '出生时间（可选）', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': False, 'description': '性别（可选）', 'example': 'male'},
            {'name': 'use_bazi', 'type': 'bool', 'required': False, 'description': '是否结合八字分析', 'example': True, 'default': True},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '办公桌风水分析结果', 'nested': [
                {'name': 'detected_items', 'type': 'list', 'description': '检测到的物品列表', 'nested': [
                    {'name': 'name', 'type': 'str', 'description': '物品名称'},
                    {'name': 'category', 'type': 'str', 'description': '物品类别'},
                    {'name': 'position', 'type': 'dict', 'description': '位置信息', 'nested': [
                        {'name': 'x', 'type': 'float', 'description': 'X坐标'},
                        {'name': 'y', 'type': 'float', 'description': 'Y坐标'},
                        {'name': 'width', 'type': 'float', 'description': '宽度'},
                        {'name': 'height', 'type': 'float', 'description': '高度'},
                    ]},
                    {'name': 'confidence', 'type': 'float', 'description': '置信度'},
                ]},
                {'name': 'fengshui_analysis', 'type': 'dict', 'description': '风水分析', 'nested': [
                    {'name': 'qinglong', 'type': 'dict', 'description': '青龙位分析'},
                    {'name': 'baihu', 'type': 'dict', 'description': '白虎位分析'},
                    {'name': 'zhuque', 'type': 'dict', 'description': '朱雀位分析'},
                    {'name': 'xuanwu', 'type': 'dict', 'description': '玄武位分析'},
                ]},
                {'name': 'bazi_analysis', 'type': 'dict', 'description': '八字分析（如果use_bazi=true）'},
                {'name': 'recommendations', 'type': 'list', 'description': '建议列表', 'nested': [
                    {'name': 'type', 'type': 'str', 'description': '建议类型'},
                    {'name': 'content', 'type': 'str', 'description': '建议内容'},
                    {'name': 'priority', 'type': 'str', 'description': '优先级'},
                ]},
            ]},
        ],
    },
    # 直接调用 REST API 的接口（不通过 gRPC-Web 网关）
    '/fortune/hand/analyze/stream': {
        'description': '手相分析（流式接口，Server-Sent Events）',
        'request_method': 'POST (FormData)',
        'request_fields': [
            {'name': 'image', 'type': 'File', 'required': True, 'description': '手掌照片文件', 'example': None},
            {'name': 'solar_date', 'type': 'str', 'required': False, 'description': '阳历日期 YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': False, 'description': '出生时间 HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': False, 'description': '性别 male/female', 'example': 'male'},
            {'name': 'use_bazi', 'type': 'bool', 'required': False, 'description': '是否结合八字分析', 'example': True, 'default': True},
        ],
        'response_fields': [
            {'name': 'type', 'type': 'str', 'description': '事件类型（status/result/error）'},
            {'name': 'stage', 'type': 'str', 'description': '处理阶段'},
            {'name': 'statusText', 'type': 'str', 'description': '状态文本'},
            {'name': 'progress', 'type': 'float', 'description': '进度（0-100）'},
            {'name': 'data', 'type': 'dict', 'description': '分析结果（type=result时）', 'nested': [
                {'name': 'success', 'type': 'bool', 'description': '是否成功'},
                {'name': 'report', 'type': 'dict', 'description': '手相分析报告'},
                {'name': 'features', 'type': 'dict', 'description': '手相特征', 'nested': [
                    {'name': 'hand_shape', 'type': 'str', 'description': '手型'},
                    {'name': 'palm_lines', 'type': 'list', 'description': '掌纹列表'},
                ]},
            ]},
            {'name': 'error', 'type': 'str', 'description': '错误信息（type=error时）'},
        ],
        'note': '这是一个流式接口，使用 Server-Sent Events (SSE) 实时返回分析进度和结果。需要 X-API-Key 认证。',
    },
    '/fortune/face/analyze/stream': {
        'description': '面相分析（流式接口，Server-Sent Events）',
        'request_method': 'POST (FormData)',
        'request_fields': [
            {'name': 'image', 'type': 'File', 'required': True, 'description': '头部面相照片文件', 'example': None},
            {'name': 'solar_date', 'type': 'str', 'required': False, 'description': '阳历日期 YYYY-MM-DD', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': False, 'description': '出生时间 HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': False, 'description': '性别 male/female', 'example': 'male'},
            {'name': 'use_bazi', 'type': 'bool', 'required': False, 'description': '是否结合八字分析', 'example': True, 'default': True},
        ],
        'response_fields': [
            {'name': 'type', 'type': 'str', 'description': '事件类型（status/result/error）'},
            {'name': 'stage', 'type': 'str', 'description': '处理阶段'},
            {'name': 'statusText', 'type': 'str', 'description': '状态文本'},
            {'name': 'progress', 'type': 'float', 'description': '进度（0-100）'},
            {'name': 'data', 'type': 'dict', 'description': '分析结果（type=result时）'},
            {'name': 'error', 'type': 'str', 'description': '错误信息（type=error时）'},
        ],
        'note': '这是一个流式接口，使用 Server-Sent Events (SSE) 实时返回分析进度和结果。需要 X-API-Key 认证。',
    },
    '/smart-fortune/smart-analyze-stream': {
        'description': '智能运势分析（流式接口，Server-Sent Events）',
        'request_method': 'GET (Query Parameters)',
        'request_fields': [
            {'name': 'question', 'type': 'str', 'required': True, 'description': '用户问题', 'example': '我今年的事业运势如何？'},
            {'name': 'year', 'type': 'int', 'required': True, 'description': '出生年份', 'example': 1990},
            {'name': 'month', 'type': 'int', 'required': True, 'description': '出生月份', 'example': 5},
            {'name': 'day', 'type': 'int', 'required': True, 'description': '出生日期', 'example': 15},
            {'name': 'hour', 'type': 'int', 'required': False, 'description': '出生时辰', 'example': 12, 'default': 12},
            {'name': 'gender', 'type': 'str', 'required': False, 'description': '性别', 'example': 'male', 'default': 'male'},
        ],
        'response_fields': [
            {'name': 'type', 'type': 'str', 'description': '事件类型（status/result/error）'},
            {'name': 'stage', 'type': 'str', 'description': '处理阶段（intent_recognition/bazi_calculation/rule_matching/llm_analysis等）'},
            {'name': 'message', 'type': 'str', 'description': '状态消息'},
            {'name': 'data', 'type': 'dict', 'description': '分析结果（type=result时）', 'nested': [
                {'name': 'success', 'type': 'bool', 'description': '是否成功'},
                {'name': 'response', 'type': 'str', 'description': '分析结果文本'},
                {'name': 'performance', 'type': 'dict', 'description': '性能摘要'},
            ]},
            {'name': 'error', 'type': 'str', 'description': '错误信息（type=error时）'},
        ],
        'note': '这是一个流式接口，使用 Server-Sent Events (SSE) 实时返回分析进度和结果。',
    },
}


def create_document() -> Document:
    """创建文档"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('HiFate-bazi 前端接口文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 文档信息
    info_para = doc.add_paragraph()
    info_para.add_run('生成时间：').bold = True
    info_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    info_para.add_run('\n')
    info_para.add_run('版本：').bold = True
    info_para.add_run('v1.0')
    info_para.add_run('\n')
    info_para.add_run('说明：').bold = True
    info_para.add_run('本文档仅包含前端使用的接口，不包括对内的服务间接口。所有接口通过 gRPC-Web 网关调用。')
    info_para.add_run('\n')
    info_para.add_run('调用方式：').bold = True
    info_para.add_run('POST /grpc-web/frontend.gateway.FrontendGateway/Call')
    info_para.add_run('\n')
    info_para.add_run('请求体格式：').bold = True
    info_para.add_run('{"endpoint": "/接口路径", "payload_json": "JSON字符串"}')
    
    doc.add_paragraph()  # 空行
    
    # 目录
    doc.add_heading('目录', 1)
    toc_para = doc.add_paragraph()
    categories = {
        '八字相关': ['/bazi/pan/display', '/bazi/fortune/display', '/bazi/dayun/display', 
                   '/bazi/liunian/display', '/bazi/liuyue/display', '/bazi/interface',
                   '/bazi/shengong-minggong', '/bazi/wangshuai', '/bazi/yigua/divinate',
                   '/bazi/formula-analysis'],
        '运势分析': ['/smart-analyze', '/bazi/daily-fortune', '/bazi/monthly-fortune'],
        '支付相关': ['/payment/create-session', '/payment/verify', '/payment/unified/create',
                   '/payment/unified/verify', '/payment/providers'],
        '其他功能': ['/auth/login', '/calendar/query', '/api/v2/face/analyze', '/api/v2/desk-fengshui/analyze'],
    }
    
    for category, endpoints in categories.items():
        toc_para.add_run(f'{category}：').bold = True
        endpoint_names = [ep.replace('/', '').replace('-', '_') for ep in endpoints if ep in ENDPOINT_INFO]
        toc_para.add_run('、'.join(endpoint_names))
        toc_para.add_run('\n')
    
    doc.add_page_break()
    
    # 接口详情
    doc.add_heading('接口详情', 1)
    
    # 按分类组织接口
    for category, endpoints in categories.items():
        doc.add_heading(category, 2)
        
        for endpoint in endpoints:
            if endpoint not in ENDPOINT_INFO:
                continue
            
            endpoint_info = ENDPOINT_INFO[endpoint]
            
            # 接口标题
            doc.add_heading(endpoint, 3)
            
            # 基本信息
            request_method = endpoint_info.get('request_method', 'POST')
            note = endpoint_info.get('note', '')
            
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_table.cell(0, 0).text = '接口路径'
            info_table.cell(0, 1).text = endpoint
            info_table.cell(1, 0).text = '请求方法'
            info_table.cell(1, 1).text = request_method
            info_table.cell(2, 0).text = '接口描述'
            info_table.cell(2, 1).text = endpoint_info['description'] or '无'
            
            # 如果有特殊说明，添加一行
            if note:
                row = info_table.add_row()
                row.cells[0].text = '特殊说明'
                row.cells[1].text = note
            
            doc.add_paragraph()  # 空行
            
            # 请求参数
            if endpoint_info['request_fields']:
                doc.add_heading('请求参数', 4)
                req_table = doc.add_table(rows=1, cols=5)
                req_table.style = 'Light Grid Accent 1'
                
                # 表头
                headers = ['字段名', '类型', '必填', '描述', '示例']
                for i, header in enumerate(headers):
                    req_table.cell(0, i).text = header
                    req_table.cell(0, i).paragraphs[0].runs[0].bold = True
                
                # 数据行
                for field in endpoint_info['request_fields']:
                    row = req_table.add_row()
                    row.cells[0].text = field['name']
                    row.cells[1].text = field['type']
                    row.cells[2].text = '是' if field['required'] else '否'
                    row.cells[3].text = field.get('description', '')
                    example = field.get('example', '')
                    if field.get('default') is not None:
                        example = f"{example} (默认: {field['default']})" if example else f"默认: {field['default']}"
                    row.cells[4].text = str(example) if example else ''
                
                doc.add_paragraph()  # 空行
            
            # 响应格式
            if endpoint_info['response_fields']:
                doc.add_heading('响应格式', 4)
                resp_table = doc.add_table(rows=1, cols=3)
                resp_table.style = 'Light Grid Accent 1'
                
                # 表头
                headers = ['字段名', '类型', '描述']
                for i, header in enumerate(headers):
                    resp_table.cell(0, i).text = header
                    resp_table.cell(0, i).paragraphs[0].runs[0].bold = True
                
                # 递归添加字段（包括嵌套字段）
                def add_field_to_table(field, prefix='', level=0):
                    """递归添加字段到表格"""
                    indent = '  ' * level
                    field_name = f"{indent}{prefix}{field['name']}" if prefix else field['name']
                    
                    row = resp_table.add_row()
                    row.cells[0].text = field_name
                    row.cells[1].text = field['type']
                    row.cells[2].text = field.get('description', '')
                    
                    # 如果有嵌套字段，递归添加
                    if 'nested' in field and field['nested']:
                        # 如果是列表类型，说明嵌套字段是列表元素的结构
                        if field['type'].startswith('list'):
                            prefix_str = f"{field['name']}[]."
                        else:
                            prefix_str = f"{field['name']}."
                        
                        for nested_field in field['nested']:
                            add_field_to_table(nested_field, prefix=prefix_str, level=level+1)
                
                # 数据行
                for field in endpoint_info['response_fields']:
                    add_field_to_table(field)
                
                doc.add_paragraph()  # 空行
            
            # 注意事项
            doc.add_heading('注意事项', 4)
            
            # 判断是否通过 gRPC-Web 网关
            is_grpc_web = not endpoint.startswith('/fortune/') and not endpoint.startswith('/smart-fortune/')
            
            if is_grpc_web:
                notes = [
                    '通过 gRPC-Web 网关调用，路径：/grpc-web/frontend.gateway.FrontendGateway/Call',
                    '请求体格式：{"endpoint": "' + endpoint + '", "payload_json": "{...}"}',
                    '所有接口均返回 JSON 格式数据',
                ]
            else:
                # 直接调用 REST API 的接口
                if endpoint.endswith('/stream'):
                    notes = [
                        '直接调用 REST API，路径：' + endpoint,
                        '使用 Server-Sent Events (SSE) 流式返回数据',
                        '需要设置请求头：Content-Type: multipart/form-data（文件上传）或 application/json',
                        '需要 X-API-Key 认证（手相/面相分析）',
                        '使用 EventSource 或 fetch + ReadableStream 接收流式数据',
                    ]
                else:
                    notes = [
                        '直接调用 REST API，路径：' + endpoint,
                        '请求方法：' + request_method,
                        '所有接口均返回 JSON 格式数据',
                    ]
            
            for note in notes:
                para = doc.add_paragraph(note, style='List Bullet')
            
            doc.add_paragraph()  # 空行
            doc.add_paragraph('=' * 80)  # 分隔线
            doc.add_paragraph()  # 空行
    
    # 附录：调用示例
    doc.add_page_break()
    doc.add_heading('调用示例', 1)
    
    example_code = '''// JavaScript 调用示例
const api = new ApiClient('http://your-domain.com');

// 调用八字排盘接口
const result = await api.post('/bazi/pan/display', {
    solar_date: '1990-05-15',
    solar_time: '14:30',
    gender: 'male'
});

// 调用智能分析接口
const analysis = await api.post('/smart-analyze', {
    question: '我今年的事业运势如何？',
    year: 1990,
    month: 5,
    day: 15,
    hour: 12,
    gender: 'male'
});'''
    
    code_para = doc.add_paragraph(example_code)
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    
    return doc


def main():
    """主函数"""
    print("正在生成前端接口文档...")
    
    # 获取已注册的端点
    endpoints = get_supported_endpoints()
    print(f"找到 {len(endpoints)} 个已注册的端点")
    
    # 创建文档
    doc = create_document()
    
    # 保存文档
    output_path = os.path.join(project_root, 'docs', '前端接口文档.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    print(f"文档已生成：{output_path}")
    print(f"共包含 {len([ep for ep in ENDPOINT_INFO.keys() if ep in endpoints])} 个接口")


if __name__ == '__main__':
    main()
