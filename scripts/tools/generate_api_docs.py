#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成API接口文档（Word格式）

使用方法:
    python scripts/tools/generate_api_docs.py

依赖:
    pip install python-docx
"""

import os
import sys
import re
import ast
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 请先安装 python-docx")
    print("安装命令: pip install python-docx")
    sys.exit(1)

# ⭐ 设置测试环境（自动扩展虚拟环境路径）
from test_utils import setup_test_environment
project_root = Path(setup_test_environment())

# 颜色定义
HEADER_COLOR = RGBColor(0, 51, 102)
TITLE_COLOR = RGBColor(0, 102, 204)
SUBTITLE_COLOR = RGBColor(51, 51, 51)


class APIDocGenerator:
    """API文档生成器"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        self.apis = []
        self.grpc_services = []
        
    def setup_document(self):
        """设置文档样式"""
        # 设置中文字体
        self.doc.styles['Normal'].font.name = '微软雅黑'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        self.doc.styles['Normal'].font.size = Pt(10.5)
        
    def add_title(self, text: str, level: int = 1):
        """添加标题"""
        if level == 1:
            heading = self.doc.add_heading(text, level=1)
            heading.runs[0].font.color.rgb = HEADER_COLOR
            heading.runs[0].font.size = Pt(24)
        elif level == 2:
            heading = self.doc.add_heading(text, level=2)
            heading.runs[0].font.color.rgb = TITLE_COLOR
            heading.runs[0].font.size = Pt(18)
        else:
            heading = self.doc.add_heading(text, level=3)
            heading.runs[0].font.color.rgb = SUBTITLE_COLOR
            heading.runs[0].font.size = Pt(14)
        return heading
    
    def add_paragraph(self, text: str, bold: bool = False, color: RGBColor = None):
        """添加段落"""
        p = self.doc.add_paragraph(text)
        if bold:
            p.runs[0].bold = True
        if color:
            p.runs[0].font.color.rgb = color
        return p
    
    def add_table(self, headers: List[str], rows: List[List[str]]):
        """添加表格"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.color.rgb = HEADER_COLOR
        
        # 添加数据行
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        return table
    
    def parse_api_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """解析API文件"""
        apis = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 解析路由装饰器
            route_pattern = r'@router\.(get|post|put|delete|patch)\("([^"]+)"[^)]*\)'
            routes = re.findall(route_pattern, content)
            
            # 解析函数定义
            func_pattern = r'async def (\w+)\([^)]*\):'
            functions = re.findall(func_pattern, content)
            
            # 解析请求模型
            request_pattern = r'class (\w+Request)\(BaseModel\):'
            requests = re.findall(request_pattern, content)
            
            # 解析响应模型
            response_pattern = r'class (\w+Response)\(BaseModel\):'
            responses = re.findall(response_pattern, content)
            
            # 解析文档字符串
            docstring_pattern = r'"""(.*?)"""'
            docstrings = re.findall(docstring_pattern, content, re.DOTALL)
            
            # 解析字段描述
            field_pattern = r'(\w+):\s*\w+\s*=\s*Field\([^)]*description="([^"]+)"'
            fields = re.findall(field_pattern, content)
            
            for method, path in routes:
                api_info = {
                    'method': method.upper(),
                    'path': path,
                    'file': file_path.name,
                    'description': '',
                    'request_fields': [],
                    'response_fields': []
                }
                
                # 查找对应的文档字符串
                for doc in docstrings:
                    if path in doc or method in doc.lower():
                        api_info['description'] = doc.strip()
                        break
                
                # 查找字段描述
                for field_name, field_desc in fields:
                    api_info['request_fields'].append({
                        'name': field_name,
                        'description': field_desc
                    })
                
                apis.append(api_info)
                
        except Exception as e:
            print(f"解析文件 {file_path} 失败: {e}")
        
        return apis
    
    def parse_proto_file(self, file_path: Path) -> Dict[str, Any]:
        """解析proto文件"""
        service_info = {
            'name': file_path.stem,
            'services': [],
            'messages': []
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 解析服务定义
            service_pattern = r'service (\w+) \{([^}]+)\}'
            services = re.findall(service_pattern, content, re.DOTALL)
            
            for service_name, service_body in services:
                service = {
                    'name': service_name,
                    'methods': []
                }
                
                # 解析RPC方法
                rpc_pattern = r'rpc (\w+)\(([^)]+)\) returns \(([^)]+)\);'
                rpcs = re.findall(rpc_pattern, service_body)
                
                for method_name, request_type, response_type in rpcs:
                    service['methods'].append({
                        'name': method_name,
                        'request': request_type.strip(),
                        'response': response_type.strip()
                    })
                
                service_info['services'].append(service)
            
            # 解析消息定义
            message_pattern = r'message (\w+) \{([^}]+)\}'
            messages = re.findall(message_pattern, content, re.DOTALL)
            
            for message_name, message_body in messages:
                message = {
                    'name': message_name,
                    'fields': []
                }
                
                # 解析字段
                field_pattern = r'(\w+)\s+(\w+)\s+(\w+)\s*=\s*(\d+);'
                fields = re.findall(field_pattern, message_body)
                
                for field_type, field_name, field_label, field_num in fields:
                    message['fields'].append({
                        'type': field_type,
                        'name': field_name,
                        'label': field_label,
                        'number': field_num
                    })
                
                service_info['messages'].append(message)
                
        except Exception as e:
            print(f"解析proto文件 {file_path} 失败: {e}")
        
        return service_info
    
    def scan_api_files(self):
        """扫描所有API文件"""
        api_dir = project_root / 'server' / 'api' / 'v1'
        if not api_dir.exists():
            return
        
        for file_path in api_dir.glob('*.py'):
            if file_path.name.startswith('__'):
                continue
            
            apis = self.parse_api_file(file_path)
            self.apis.extend(apis)
    
    def scan_proto_files(self):
        """扫描所有proto文件"""
        proto_dir = project_root / 'proto'
        if not proto_dir.exists():
            return
        
        for file_path in proto_dir.glob('*.proto'):
            service_info = self.parse_proto_file(file_path)
            if service_info['services']:
                self.grpc_services.append(service_info)
    
    def generate_rest_api_docs(self):
        """生成REST API文档"""
        self.add_title("REST API 接口文档", level=1)
        
        # 按路径分组
        api_groups = {}
        for api in self.apis:
            prefix = api['path'].split('/')[1] if '/' in api['path'] else 'other'
            if prefix not in api_groups:
                api_groups[prefix] = []
            api_groups[prefix].append(api)
        
        for group_name, apis in sorted(api_groups.items()):
            self.add_title(f"{group_name.upper()} 接口", level=2)
            
            for api in apis:
                # 接口标题
                title = f"{api['method']} {api['path']}"
                self.add_title(title, level=3)
                
                # 基本信息表格
                info_rows = [
                    ['请求方法', api['method']],
                    ['接口路径', api['path']],
                    ['接口说明', api['description'] or '暂无说明']
                ]
                self.add_table(['属性', '值'], info_rows)
                
                # 请求参数
                if api['request_fields']:
                    self.add_paragraph("请求参数：", bold=True)
                    param_rows = []
                    for field in api['request_fields']:
                        param_rows.append([
                            field['name'],
                            field['description']
                        ])
                    self.add_table(['参数名', '说明'], param_rows)
                
                # 请求示例
                self.add_paragraph("请求示例：", bold=True)
                example = f"""
```http
{api['method']} /api/v1{api['path']} HTTP/1.1
Host: localhost:8001
Content-Type: application/json

{{
  "solar_date": "1990-01-15",
  "solar_time": "12:00",
  "gender": "male"
}}
```
"""
                self.add_paragraph(example)
                
                # 响应示例
                self.add_paragraph("响应示例：", bold=True)
                response_example = """
```json
{
  "success": true,
  "data": {
    ...
  }
}
```
"""
                self.add_paragraph(response_example)
                
                self.doc.add_paragraph("")  # 空行
    
    def generate_grpc_docs(self):
        """生成gRPC接口文档"""
        self.add_title("gRPC 接口文档", level=1)
        
        for service_info in self.grpc_services:
            self.add_title(f"{service_info['name']} 服务", level=2)
            
            for service in service_info['services']:
                self.add_title(f"Service: {service['name']}", level=3)
                
                # 方法列表
                method_rows = []
                for method in service['methods']:
                    method_rows.append([
                        method['name'],
                        method['request'],
                        method['response']
                    ])
                self.add_table(['方法名', '请求类型', '响应类型'], method_rows)
                
                # 消息定义
                if service_info['messages']:
                    self.add_paragraph("消息定义：", bold=True)
                    for message in service_info['messages']:
                        self.add_title(f"Message: {message['name']}", level=4)
                        
                        if message['fields']:
                            field_rows = []
                            for field in message['fields']:
                                field_rows.append([
                                    field['name'],
                                    field['type'],
                                    field['label']
                                ])
                            self.add_table(['字段名', '类型', '标签'], field_rows)
                
                self.doc.add_paragraph("")  # 空行
    
    def generate_grpc_web_docs(self):
        """生成gRPC-Web接口文档"""
        self.add_title("gRPC-Web 接口文档", level=1)
        
        grpc_gateway_file = project_root / 'server' / 'api' / 'grpc_gateway.py'
        if not grpc_gateway_file.exists():
            return
        
        try:
            with open(grpc_gateway_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 解析注册的端点
            register_pattern = r'@_register\("([^"]+)"\)'
            endpoints = re.findall(register_pattern, content)
            
            self.add_paragraph("gRPC-Web 网关支持以下端点：", bold=True)
            
            endpoint_rows = []
            for endpoint in sorted(endpoints):
                endpoint_rows.append([endpoint])
            
            self.add_table(['端点路径'], endpoint_rows)
            
            self.add_paragraph("")
            self.add_paragraph("调用方式：", bold=True)
            self.add_paragraph("""
所有 gRPC-Web 接口都通过 POST 方法调用，请求体为 JSON 格式。

请求格式：
```http
POST /api/grpc-web/{endpoint} HTTP/1.1
Host: localhost:8001
Content-Type: application/json

{
  "solar_date": "1990-01-15",
  "solar_time": "12:00",
  "gender": "male"
}
```

响应格式：
```json
{
  "success": true,
  "data": {
    ...
  }
}
```
""")
            
        except Exception as e:
            print(f"解析gRPC-Web网关失败: {e}")
    
    def generate_summary(self):
        """生成接口汇总"""
        self.add_title("接口汇总", level=1)
        
        # REST API统计
        self.add_title("REST API 统计", level=2)
        rest_stats = {
            'GET': len([a for a in self.apis if a['method'] == 'GET']),
            'POST': len([a for a in self.apis if a['method'] == 'POST']),
            'PUT': len([a for a in self.apis if a['method'] == 'PUT']),
            'DELETE': len([a for a in self.apis if a['method'] == 'DELETE']),
        }
        
        stat_rows = [[method, count] for method, count in rest_stats.items()]
        self.add_table(['请求方法', '接口数量'], stat_rows)
        
        # gRPC服务统计
        self.add_title("gRPC 服务统计", level=2)
        grpc_rows = []
        for service_info in self.grpc_services:
            for service in service_info['services']:
                grpc_rows.append([
                    service['name'],
                    len(service['methods'])
                ])
        self.add_table(['服务名', '方法数量'], grpc_rows)
    
    def generate(self, output_path: str):
        """生成完整文档"""
        print("开始扫描API文件...")
        self.scan_api_files()
        print(f"找到 {len(self.apis)} 个REST API接口")
        
        print("开始扫描proto文件...")
        self.scan_proto_files()
        print(f"找到 {len(self.grpc_services)} 个gRPC服务")
        
        print("生成文档...")
        
        # 添加封面
        self.add_title("HiFate-bazi API 接口文档", level=1)
        self.add_paragraph("")
        self.add_paragraph("版本: 1.0.0")
        self.add_paragraph("生成时间: " + str(Path(__file__).stat().st_mtime))
        self.doc.add_page_break()
        
        # 生成汇总
        self.generate_summary()
        self.doc.add_page_break()
        
        # 生成REST API文档
        self.generate_rest_api_docs()
        self.doc.add_page_break()
        
        # 生成gRPC文档
        self.generate_grpc_docs()
        self.doc.add_page_break()
        
        # 生成gRPC-Web文档
        self.generate_grpc_web_docs()
        
        # 保存文档
        print(f"保存文档到: {output_path}")
        self.doc.save(output_path)
        print("文档生成完成！")


def main():
    """主函数"""
    output_path = project_root / 'docs' / 'API接口文档.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    generator = APIDocGenerator()
    generator.generate(str(output_path))
    
    print(f"\n✅ 文档已生成: {output_path}")


if __name__ == '__main__':
    main()

