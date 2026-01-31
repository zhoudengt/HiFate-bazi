#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
临时脚本：生成合并的支付接口文档（包含 create 和 verify 两个接口）
"""

import sys
from pathlib import Path
from datetime import datetime

# 添加项目根目录到路径
project_root = Path(__file__).resolve().parent.parent.parent
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 导入现有的文档生成工具
from scripts.tools.generate_api_doc_to_feishu import FastAPIParser, DocGenerator

def create_merged_document():
    """创建合并的支付接口文档"""
    
    # 解析文件
    file_path = project_root / "server" / "api" / "v1" / "unified_payment.py"
    parser = FastAPIParser(file_path)
    
    # 查找两个路由
    routes = parser.find_routes()
    create_route = None
    verify_route = None
    
    for route in routes:
        if route.get('path') == '/payment/unified/create':
            create_route = route
        elif route.get('path') == '/payment/unified/verify':
            verify_route = route
    
    if not create_route or not verify_route:
        print("❌ 无法找到支付接口路由")
        return
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # 主标题
    title = doc.add_heading('统一支付接口文档', 0)
    title.runs[0].font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # ========== 第一部分：支付业务整体流程 ==========
    doc.add_heading('一、支付业务整体流程', 1)
    
    para = doc.add_paragraph()
    para.add_run('统一支付系统基于插件化架构，支持多种支付渠道（Stripe、PayPal、PayerMax、支付宝、微信、Line Pay 等）。')
    doc.add_paragraph()
    
    doc.add_heading('1.1 业务流程', 2)
    
    flow_steps = [
        '用户在前端选择商品和服务，点击支付',
        '前端调用 `/payment/unified/create` 接口创建支付订单',
        '后端根据支付渠道创建订单，返回支付链接（payment_url 或 checkout_url）',
        '前端跳转到支付平台页面（如 Stripe Checkout），用户填写银行卡信息',
        '**重要**：用户点击"支付"按钮后，银行卡信息（卡号、CVC、过期日期）直接提交给支付平台的 API（如 Stripe API），不经过我们的后端服务器（PCI 合规要求）',
        '支付平台验证并处理支付，可能触发 3D Secure 验证（需要输入银行发送的验证码）',
        '支付完成后，支付平台重定向用户回到我们的成功页面，或通过 webhook 异步通知后端',
        '前端定期轮询或用户返回时，调用 `/payment/unified/verify` 接口验证支付状态',
        '后端查询支付平台和数据库，返回支付状态（paid/pending/failed）',
        '前端根据支付状态更新界面，完成订单流程'
    ]
    
    for i, step in enumerate(flow_steps, 1):
        para = doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_paragraph()
    
    doc.add_heading('1.2 关键时间节点', 2)
    
    time_table = doc.add_table(rows=3, cols=2)
    time_table.style = 'Light Grid Accent 1'
    
    # 表头
    time_table.rows[0].cells[0].text = '时间节点'
    time_table.rows[0].cells[1].text = '说明'
    for cell in time_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 数据行
    time_table.rows[1].cells[0].text = '订单创建时间'
    time_table.rows[1].cells[1].text = '调用 create 接口时，系统自动记录 created_at'
    
    time_table.rows[2].cells[0].text = '订单过期时间'
    time_table.rows[2].cells[1].text = 'expires_at = created_at + 30分钟，超时后订单失效，需重新创建'
    
    doc.add_paragraph()
    
    doc.add_heading('1.3 支付渠道说明', 2)
    
    provider_table = doc.add_table(rows=8, cols=3)
    provider_table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ['支付渠道', '适用地区', '主要特点']
    for i, header in enumerate(headers):
        cell = provider_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 数据行
    providers = [
        ['Stripe', '全球（美洲、欧洲、香港、菲律宾等）', '全球通用，支持信用卡，适合国际支付'],
        ['PayPal', '全球', '认知度高，备选方案'],
        ['PayerMax', '全球（除台湾 LINE Pay）', '600+ 支付方式聚合，收银台模式'],
        ['Alipay', '中国客户', '支付宝国际版'],
        ['WeChat', '中国客户', '微信支付'],
        ['Line Pay', '台湾、日本、泰国等', '直接集成'],
        ['Payssion', '台湾', 'LINE Pay 中转']
    ]
    
    for i, provider_info in enumerate(providers, 1):
        for j, value in enumerate(provider_info):
            provider_table.rows[i].cells[j].text = value
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ========== 第二部分：接口关系说明 ==========
    doc.add_heading('二、接口关系说明', 1)
    
    doc.add_heading('2.1 接口协作关系', 2)
    
    para = doc.add_paragraph()
    para.add_run('`/payment/unified/create` 和 `/payment/unified/verify` 两个接口必须配合使用，共同完成支付流程。')
    doc.add_paragraph()
    
    doc.add_heading('2.2 字段关联关系', 2)
    
    field_table = doc.add_table(rows=7, cols=3)
    field_table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ['支付渠道', 'create 接口返回字段', 'verify 接口使用参数']
    for i, header in enumerate(headers):
        cell = field_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 数据行
    field_mappings = [
        ['Stripe', 'payment_id (session_id)', 'session_id'],
        ['PayerMax', 'transaction_id 或 order_id', 'transaction_id 或 order_id'],
        ['PayPal', 'payment_id', 'payment_id'],
        ['Alipay', 'order_id', 'order_id'],
        ['WeChat', 'order_id', 'order_id'],
        ['Line Pay', 'transaction_id', 'transaction_id']
    ]
    
    for i, mapping in enumerate(field_mappings, 1):
        for j, value in enumerate(mapping):
            field_table.rows[i].cells[j].text = value
    
    doc.add_paragraph()
    
    doc.add_heading('2.3 调用时机', 2)
    
    timing_list = [
        '支付成功后：通过 webhook 或前端回调，立即调用 verify 接口确认支付状态',
        '定期轮询：前端每 30 秒调用一次 verify 接口，检查支付状态（适用于待支付订单）',
        '用户返回：用户从支付页面返回时，主动调用 verify 接口查询支付结果',
        '订单查询：需要查询历史订单状态时，使用 verify 接口'
    ]
    
    doc.add_paragraph()
    doc.add_heading('2.4 银行卡信息提交说明（重要）', 2)
    
    para = doc.add_paragraph()
    para.add_run('当用户在支付页面（如 Stripe Checkout）填写银行卡信息并点击"支付"按钮后：').bold = True
    doc.add_paragraph()
    
    security_points = [
        '银行卡信息（卡号、CVC、过期日期、持卡人姓名）直接提交给支付平台的 API（如 Stripe API），不经过我们的后端服务器',
        '这是 PCI DSS 合规要求：敏感卡信息不能存储或传输经过我们的服务器，必须由支付平台直接处理',
        '如果触发 3D Secure 验证，用户需要输入银行发送的验证码（短信或邮件），这个验证也在支付平台页面完成',
        '支付平台处理完成后，会通过以下方式通知我们：',
        '  - 重定向：将用户重定向到我们配置的 success_url（携带 session_id 等参数）',
        '  - Webhook：异步发送支付结果到我们的 webhook 接口（如果配置了）',
        '我们的后端接口 `/payment/unified/verify` 用于查询支付状态，不接收银行卡信息'
    ]
    
    doc.add_paragraph()
    doc.add_heading('2.5 点击"支付"按钮对应的接口', 2)
    
    para = doc.add_paragraph()
    para.add_run('当用户在 Stripe Checkout 页面点击"支付"按钮后：').bold = True
    doc.add_paragraph()
    
    doc.add_paragraph('1. 不是我们的后端接口')
    doc.add_paragraph('   - 用户在 Stripe Checkout 托管页面上点击"支付"按钮', style='List Bullet')
    doc.add_paragraph('   - Stripe Checkout 页面是 Stripe 完全托管的页面，由 Stripe 的前端代码控制', style='List Bullet')
    doc.add_paragraph('   - 银行卡信息直接提交给 Stripe 的服务器（Stripe API），不经过我们的后端', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('2. Stripe 内部处理的 API')
    doc.add_paragraph('   - Stripe Checkout Session 的内部确认接口（Stripe 内部处理）', style='List Bullet')
    doc.add_paragraph('   - 银行卡信息通过 Stripe.js 或 Stripe Checkout 直接提交到 Stripe 的服务器', style='List Bullet')
    doc.add_paragraph('   - 这是 Stripe 托管页面的内部流程，我们无法控制', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('3. 我们的接口调用时机')
    doc.add_paragraph('   - 在显示支付表单之前：前端调用 `/payment/unified/create` 创建支付会话，获取 checkout_url', style='List Bullet')
    doc.add_paragraph('   - 支付完成后：前端调用 `/payment/unified/verify` 查询支付状态（不接收银行卡信息）', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('总结：用户点击"支付"按钮后，银行卡信息提交给 Stripe 的 API（不是我们的接口），这是 PCI 合规的安全设计。')
    
    for point in security_points:
        doc.add_paragraph(point, style='List Bullet')
    
    for item in timing_list:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('2.6 验证结果说明', 2)
    
    result_table = doc.add_table(rows=4, cols=2)
    result_table.style = 'Light Grid Accent 1'
    
    # 表头
    result_table.rows[0].cells[0].text = '返回状态'
    result_table.rows[0].cells[1].text = '说明'
    for cell in result_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # 数据行
    results = [
        ['success: true, status: "paid"', '支付成功，订单已完成'],
        ['success: true, status: "pending"', '待支付，订单仍在有效期内'],
        ['success: false', '支付失败或订单已过期，需要重新创建订单']
    ]
    
    for i, result in enumerate(results, 1):
        for j, value in enumerate(result):
            result_table.rows[i].cells[j].text = value
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ========== 第三部分：接口1 - create ==========
    doc.add_heading('三、接口1：创建支付订单', 1)
    
    # 使用现有工具生成 create 接口文档
    create_generator = DocGenerator(parser, create_route, provider=None)
    create_doc_content = create_generator.generate_doc()
    
    # 将 Markdown 转换为 Word（简化版）
    _markdown_to_word(doc, create_doc_content)
    
    doc.add_page_break()
    
    # ========== 第四部分：接口2 - verify ==========
    doc.add_heading('四、接口2：验证支付状态', 1)
    
    # 使用现有工具生成 verify 接口文档
    verify_generator = DocGenerator(parser, verify_route, provider=None)
    verify_doc_content = verify_generator.generate_doc()
    
    # 将 Markdown 转换为 Word（简化版）
    _markdown_to_word(doc, verify_doc_content)
    
    doc.add_page_break()
    
    # ========== 第五部分：完整业务流程示例 ==========
    doc.add_heading('五、完整业务流程示例', 1)
    
    doc.add_heading('5.1 端到端 curl 示例', 2)
    
    doc.add_paragraph('以下示例展示从创建订单到验证支付的完整流程：')
    doc.add_paragraph()
    
    # 步骤1：创建订单
    doc.add_heading('步骤1：创建支付订单', 3)
    doc.add_paragraph('```bash')
    curl_create = '''curl -X POST http://localhost:8001/api/v1/payment/unified/create \\
  -H "Content-Type: application/json" \\
  -d '{
    "provider": "stripe",
    "amount": "19.90",
    "currency": "USD",
    "product_name": "月订阅会员",
    "customer_email": "test@example.com"
  }'''
    para = doc.add_paragraph(curl_create)
    para.style = 'No Spacing'
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph('```')
    doc.add_paragraph()
    
    doc.add_paragraph('响应示例：')
    doc.add_paragraph('```json')
    response_create = '''{
  "success": true,
  "provider": "stripe",
  "payment_id": "cs_test_a1B2c3D4...",
  "checkout_url": "https://checkout.stripe.com/c/pay/cs_test_...",
  "order_id": "STRIPE_1706342400000",
  "status": "created",
  "expires_at": "2024-01-27T18:30:00Z",
  "created_at": "2024-01-27T18:00:00Z"
}'''
    para = doc.add_paragraph(response_create)
    para.style = 'No Spacing'
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph('```')
    doc.add_paragraph()
    
    # 步骤2：验证支付
    doc.add_heading('步骤2：验证支付状态', 3)
    doc.add_paragraph('```bash')
    curl_verify = '''curl -X POST http://localhost:8001/api/v1/payment/unified/verify \\
  -H "Content-Type: application/json" \\
  -d '{
    "provider": "stripe",
    "session_id": "cs_test_a1B2c3D4..."
  }'''
    para = doc.add_paragraph(curl_verify)
    para.style = 'No Spacing'
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph('```')
    doc.add_paragraph()
    
    doc.add_paragraph('响应示例（支付成功）：')
    doc.add_paragraph('```json')
    response_verify = '''{
  "success": true,
  "provider": "stripe",
  "status": "paid",
  "payment_id": "cs_test_a1B2c3D4...",
  "order_id": "STRIPE_1706342400000",
  "amount": "19.90",
  "currency": "USD",
  "customer_email": "test@example.com",
  "paid_time": "2024-01-27T18:05:00Z"
}'''
    para = doc.add_paragraph(response_verify)
    para.style = 'No Spacing'
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    doc.add_paragraph('```')
    doc.add_paragraph()
    
    doc.add_heading('5.2 前端集成建议', 2)
    
    frontend_tips = [
        '创建订单后，立即保存返回的 payment_id、order_id 和 expires_at',
        '使用 checkout_url 或 payment_url 跳转到支付页面',
        '设置定时器，每 30 秒调用一次 verify 接口检查支付状态',
        '监听支付页面的返回事件，用户返回时立即调用 verify 接口',
        '如果订单过期（当前时间 > expires_at），提示用户重新创建订单',
        '根据 verify 接口返回的 status 更新 UI：paid 显示成功，pending 显示处理中，failed 显示失败'
    ]
    
    for tip in frontend_tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    # 保存文档
    desktop_path = Path.home() / "Desktop"
    if not desktop_path.exists():
        desktop_path = Path.home() / "桌面"
        if not desktop_path.exists():
            desktop_path = project_root / "docs"
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    word_file = desktop_path / f"统一支付接口文档_合并版_{timestamp}.docx"
    
    doc.save(str(word_file))
    print(f"✅ 合并文档已生成: {word_file}")
    print(f"📋 文档已保存到桌面，可以直接打开查看")
    
    return word_file


def _markdown_to_word(doc, markdown_content: str):
    """将 Markdown 内容写入 Word 文档（简化版）"""
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue
        
        # 处理标题
        if stripped.startswith('## '):
            title = doc.add_heading(stripped[3:].strip(), level=1)
            title.runs[0].font.size = Pt(16)
        elif stripped.startswith('### '):
            title = doc.add_heading(stripped[4:].strip(), level=2)
            title.runs[0].font.size = Pt(14)
        elif stripped.startswith('#### '):
            title = doc.add_heading(stripped[5:].strip(), level=3)
            title.runs[0].font.size = Pt(12)
        # 处理表格
        elif stripped.startswith('|'):
            # 收集表格行
            table_rows = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                row_line = lines[j].strip()
                if not row_line.startswith('|---'):  # 跳过分隔行
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    table_rows.append(cells)
                j += 1
            
            if table_rows:
                # 创建表格
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # 第一行设为粗体
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                i = j - 1
        # 处理代码块
        elif stripped.startswith('```'):
            code_lines = []
            j = i + 1
            while j < len(lines) and not lines[j].strip().startswith('```'):
                code_lines.append(lines[j])
                j += 1
            
            if code_lines:
                code_para = doc.add_paragraph()
                code_para.style = 'No Spacing'
                run = code_para.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                code_para.paragraph_format.left_indent = Pt(20)
                i = j
        # 处理列表项
        elif stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:].strip()
            # 移除粗体标记
            if content.startswith('**') and content.endswith('**'):
                content = content[2:-2]
                para = doc.add_paragraph(content, style='List Bullet')
                para.runs[0].bold = True
            else:
                doc.add_paragraph(content, style='List Bullet')
        # 处理普通文本
        else:
            para = doc.add_paragraph()
            # 处理粗体
            import re
            parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    para.add_run(part)
        
        i += 1


if __name__ == '__main__':
    try:
        create_merged_document()
        print("✅ 完成！")
    except Exception as e:
        print(f"❌ 生成文档失败: {e}")
        import traceback
        traceback.print_exc()
