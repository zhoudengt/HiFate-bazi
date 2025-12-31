#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成6个分析功能的接口文档，并追加到现有的 bazi_shengong_minggong_api.docx 文件中
"""

import sys
import os
import json
from datetime import datetime

# 添加项目根目录到路径
project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, project_root)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


def add_api_section(doc, api_config):
    """为文档添加一个API接口章节"""
    # 标题
    title = doc.add_heading(api_config['title'], 1)
    
    # 接口路径
    doc.add_heading('接口路径', 2)
    path_para = doc.add_paragraph()
    path_para.add_run(api_config['path']).bold = True
    
    # 请求方法
    doc.add_heading('请求方法', 2)
    method_para = doc.add_paragraph()
    method_para.add_run(api_config['method']).bold = True
    
    # 接口描述
    doc.add_heading('接口描述', 2)
    desc_para = doc.add_paragraph()
    desc_para.add_run(api_config['description'])
    
    # 请求参数
    doc.add_heading('请求参数', 2)
    req_table = doc.add_table(rows=1, cols=5)
    req_table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ['字段名', '类型', '必填', '描述', '示例']
    for i, header in enumerate(headers):
        req_table.cell(0, i).text = header
        req_table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    # 请求参数字段
    for field in api_config['request_fields']:
        row = req_table.add_row()
        row.cells[0].text = field['name']
        row.cells[1].text = field['type']
        row.cells[2].text = '是' if field['required'] else '否'
        row.cells[3].text = field['description']
        row.cells[4].text = str(field.get('example', ''))
    
    doc.add_paragraph()  # 空行
    
    # 响应格式
    doc.add_heading('响应格式', 2)
    
    # 如果是流式接口，添加SSE格式说明
    if api_config.get('is_stream'):
        resp_desc_para = doc.add_paragraph()
        resp_desc_para.add_run(api_config['response_description'])
        doc.add_paragraph()  # 空行
        
        # SSE格式说明表格
        resp_table = doc.add_table(rows=1, cols=3)
        resp_table.style = 'Light Grid Accent 1'
        
        resp_headers = ['字段名', '类型', '描述']
        for i, header in enumerate(resp_headers):
            resp_table.cell(0, i).text = header
            resp_table.cell(0, i).paragraphs[0].runs[0].bold = True
        
        for field in api_config['response_fields']:
            row = resp_table.add_row()
            row.cells[0].text = field['name']
            row.cells[1].text = field['type']
            row.cells[2].text = field.get('description', '')
    else:
        # 非流式接口使用标准表格
        resp_table = doc.add_table(rows=1, cols=3)
        resp_table.style = 'Light Grid Accent 1'
        
        resp_headers = ['字段名', '类型', '描述']
        for i, header in enumerate(resp_headers):
            resp_table.cell(0, i).text = header
            resp_table.cell(0, i).paragraphs[0].runs[0].bold = True
        
        for field in api_config['response_fields']:
            row = resp_table.add_row()
            row.cells[0].text = field['name']
            row.cells[1].text = field['type']
            row.cells[2].text = field.get('description', '')
    
    doc.add_paragraph()  # 空行
    
    # 添加JSON示例数据
    if api_config.get('examples'):
        doc.add_heading('示例数据', 2)
        examples = api_config['examples']
        
        for example_type, example_data in examples.items():
            # 示例标题
            example_title_map = {
                'progress': 'progress事件示例（增量内容）',
                'complete': 'complete事件示例（完整内容）',
                'error': 'error事件示例（错误消息）',
                'status': 'status事件示例（状态更新）',
                'basic_analysis': 'basic_analysis事件示例（基础分析结果）',
                'llm_start': 'llm_start事件示例（LLM开始）',
                'llm_chunk': 'llm_chunk事件示例（LLM增量内容）',
                'llm_end': 'llm_end事件示例（LLM结束）',
                'llm_error': 'llm_error事件示例（LLM错误）',
                'performance': 'performance事件示例（性能摘要）',
                'end': 'end事件示例（结束）'
            }
            example_title = example_title_map.get(example_type, f'{example_type}事件示例')
            
            # 添加示例标题（使用三级标题）
            example_heading = doc.add_heading(example_title, 3)
            
            # 添加SSE格式说明（如果是SSE事件）
            if example_type in ['progress', 'complete', 'error', 'llm_chunk', 'status', 'basic_analysis', 'llm_start', 'llm_end', 'llm_error', 'performance', 'end']:
                sse_format_para = doc.add_paragraph()
                sse_format_para.add_run('SSE格式（实际HTTP响应中的格式）：').bold = True
                sse_code_para = doc.add_paragraph(example_data.get('sse_format', ''), style='Intense Quote')
                # 设置等宽字体
                try:
                    for run in sse_code_para.runs:
                        run.font.name = 'Courier New'
                except:
                    pass  # 如果设置字体失败，忽略错误
                
                json_para = doc.add_paragraph()
                json_para.add_run('JSON数据（解析后的内容）：').bold = True
            else:
                json_para = doc.add_paragraph()
                json_para.add_run('JSON数据：').bold = True
            
            # 添加JSON代码块
            json_str = json.dumps(example_data.get('json', {}), ensure_ascii=False, indent=2)
            json_code_para = doc.add_paragraph(json_str, style='Intense Quote')
            # 设置等宽字体（Courier New）
            try:
                for run in json_code_para.runs:
                    run.font.name = 'Courier New'
            except:
                pass  # 如果设置字体失败，忽略错误
            
            # 如果有说明，添加说明
            if example_data.get('description'):
                desc_para = doc.add_paragraph()
                desc_para.add_run(example_data['description'])
            
            doc.add_paragraph()  # 空行
    
    # 注意事项
    doc.add_heading('注意事项', 2)
    for note in api_config['notes']:
        para = doc.add_paragraph(note, style='List Bullet')
    
    # 添加分隔线
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()


def get_bazi_base_request_fields():
    """获取BaziBaseRequest的标准字段"""
    return [
        {
            'name': 'solar_date',
            'type': 'str',
            'required': True,
            'description': '阳历日期，格式：YYYY-MM-DD。当 calendar_type=lunar 时，此字段为农历日期',
            'example': '1990-05-15'
        },
        {
            'name': 'solar_time',
            'type': 'str',
            'required': True,
            'description': '出生时间，格式：HH:MM',
            'example': '14:30'
        },
        {
            'name': 'gender',
            'type': 'str',
            'required': True,
            'description': '性别：male(男) 或 female(女)',
            'example': 'male'
        },
        {
            'name': 'calendar_type',
            'type': 'str',
            'required': False,
            'description': '历法类型：solar(阳历) 或 lunar(农历)，默认 solar',
            'example': 'solar'
        },
        {
            'name': 'location',
            'type': 'str',
            'required': False,
            'description': '出生地点（可选），用于时区转换，优先级1',
            'example': '北京'
        },
        {
            'name': 'latitude',
            'type': 'float',
            'required': False,
            'description': '纬度（可选），用于时区转换，优先级2',
            'example': 39.90
        },
        {
            'name': 'longitude',
            'type': 'float',
            'required': False,
            'description': '经度（可选），用于时区转换和真太阳时计算，优先级2',
            'example': 116.40
        }
    ]


def get_stream_api_configs():
    """获取所有流式接口的配置"""
    base_fields = get_bazi_base_request_fields()
    
    configs = [
        {
            'title': '八字命理-感情婚姻',
            'path': '/api/v1/bazi/marriage-analysis/stream',
            'method': 'POST',
            'description': '基于用户生辰数据，使用 Coze Bot 流式生成感情婚姻分析。分析内容包括：命盘总论、配偶特征、感情走势、神煞点睛和建议方向。',
            'request_fields': base_fields + [
                {
                    'name': 'bot_id',
                    'type': 'str',
                    'required': False,
                    'description': 'Coze Bot ID（可选，优先级：参数 > MARRIAGE_ANALYSIS_BOT_ID 环境变量 > COZE_BOT_ID 环境变量）',
                    'example': '7584766797639958555'
                }
            ],
            'is_stream': True,
            'response_description': '本接口使用 SSE (Server-Sent Events) 流式输出，响应格式为 text/event-stream。',
            'response_fields': [
                {
                    'name': 'data',
                    'type': 'string (SSE格式)',
                    'description': 'SSE事件流，每个事件格式为：data: {"type": "progress|complete|error", "content": "..."}\\n\\n'
                },
                {
                    'name': 'type',
                    'type': 'string',
                    'description': '事件类型：progress（增量内容）、complete（完整内容）、error（错误消息）'
                },
                {
                    'name': 'content',
                    'type': 'string',
                    'description': '事件内容（分析文本或错误消息）'
                }
            ],
            'notes': [
                '通过 gRPC-Web 网关调用，路径：/grpc-web/frontend.gateway.FrontendGateway/Call',
                '请求体格式：{"endpoint": "/api/v1/bazi/marriage-analysis/stream", "payload_json": "{...}"}',
                '响应格式为 SSE 流式输出，前端需要使用 EventSource 或 fetch API 的流式读取方式处理',
                'bot_id 参数可选，如果未提供则使用环境变量 MARRIAGE_ANALYSIS_BOT_ID 或 COZE_BOT_ID',
                '分析内容由 Coze Bot 实时生成，可能需要较长时间（通常10-30秒）'
            ],
            'examples': {
                'progress': {
                    'sse_format': 'data: {"type":"progress","content":"命盘总论\\n\\n根据您的八字信息，"}\n\n',
                    'json': {
                        'type': 'progress',
                        'content': '命盘总论\n\n根据您的八字信息，'
                    },
                    'description': 'progress 事件表示增量内容，content 字段包含部分分析文本，前端需要累积这些内容显示给用户。'
                },
                'complete': {
                    'sse_format': 'data: {"type":"complete","content":"命盘总论\\n\\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\\n\\n配偶特征\\n\\n您的配偶特征..."}\n\n',
                    'json': {
                        'type': 'complete',
                        'content': '命盘总论\n\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\n\n配偶特征\n\n您的配偶特征...'
                    },
                    'description': 'complete 事件表示完整内容，content 字段包含完整的分析文本，前端可以使用此内容作为最终结果。'
                },
                'error': {
                    'sse_format': 'data: {"type":"error","content":"Coze Bot ID 配置缺失: 请设置环境变量 MARRIAGE_ANALYSIS_BOT_ID 或 COZE_BOT_ID 或在请求参数中提供 bot_id。"}\n\n',
                    'json': {
                        'type': 'error',
                        'content': 'Coze Bot ID 配置缺失: 请设置环境变量 MARRIAGE_ANALYSIS_BOT_ID 或 COZE_BOT_ID 或在请求参数中提供 bot_id。'
                    },
                    'description': 'error 事件表示错误消息，content 字段包含错误描述，前端应该显示错误信息给用户。'
                }
            }
        },
        {
            'title': '八字命理-事业财富',
            'path': '/api/v1/career-wealth/stream',
            'method': 'POST',
            'description': '基于用户生辰数据，使用 Coze Bot 流式生成事业财富分析。分析内容包括：命盘事业财富总论、事业特征、财富星与财富宫、事业运势、财富运势和提运建议。',
            'request_fields': base_fields + [
                {
                    'name': 'bot_id',
                    'type': 'str',
                    'required': False,
                    'description': 'Coze Bot ID（可选，优先级：参数 > CAREER_WEALTH_BOT_ID 环境变量 > COZE_BOT_ID 环境变量）',
                    'example': '7584766797639958555'
                }
            ],
            'is_stream': True,
            'response_description': '本接口使用 SSE (Server-Sent Events) 流式输出，响应格式为 text/event-stream。',
            'response_fields': [
                {
                    'name': 'data',
                    'type': 'string (SSE格式)',
                    'description': 'SSE事件流，每个事件格式为：data: {"type": "progress|complete|error", "content": "..."}\\n\\n'
                },
                {
                    'name': 'type',
                    'type': 'string',
                    'description': '事件类型：progress（增量内容）、complete（完整内容）、error（错误消息）'
                },
                {
                    'name': 'content',
                    'type': 'string',
                    'description': '事件内容（分析文本或错误消息）'
                }
            ],
            'notes': [
                '通过 gRPC-Web 网关调用，路径：/grpc-web/frontend.gateway.FrontendGateway/Call',
                '请求体格式：{"endpoint": "/api/v1/career-wealth/stream", "payload_json": "{...}"}',
                '响应格式为 SSE 流式输出，前端需要使用 EventSource 或 fetch API 的流式读取方式处理',
                'bot_id 参数可选，如果未提供则使用环境变量 CAREER_WEALTH_BOT_ID 或 COZE_BOT_ID',
                '分析内容由 Coze Bot 实时生成，可能需要较长时间（通常10-30秒）'
            ],
            'examples': {
                'progress': {
                    'sse_format': 'data: {"type":"progress","content":"命盘事业财富总论\\n\\n根据您的八字信息，"}\n\n',
                    'json': {
                        'type': 'progress',
                        'content': '命盘事业财富总论\n\n根据您的八字信息，'
                    },
                    'description': 'progress 事件表示增量内容，content 字段包含部分分析文本。'
                },
                'complete': {
                    'sse_format': 'data: {"type":"complete","content":"命盘事业财富总论\\n\\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\\n\\n事业特征\\n\\n您的事业特征..."}\n\n',
                    'json': {
                        'type': 'complete',
                        'content': '命盘事业财富总论\n\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\n\n事业特征\n\n您的事业特征...'
                    },
                    'description': 'complete 事件表示完整内容，content 字段包含完整的分析文本。'
                },
                'error': {
                    'sse_format': 'data: {"type":"error","content":"Coze Bot ID 配置缺失: 请设置环境变量 CAREER_WEALTH_BOT_ID 或 COZE_BOT_ID。"}\n\n',
                    'json': {
                        'type': 'error',
                        'content': 'Coze Bot ID 配置缺失: 请设置环境变量 CAREER_WEALTH_BOT_ID 或 COZE_BOT_ID。'
                    },
                    'description': 'error 事件表示错误消息，前端应该显示错误信息给用户。'
                }
            }
        },
        {
            'title': '八字命理-子女学习',
            'path': '/api/v1/children-study/stream',
            'method': 'POST',
            'description': '基于用户生辰数据，使用 Coze Bot 流式生成子女学习分析。分析内容包括：命盘总论、子女特征、学习能力、运势分析和建议方向。',
            'request_fields': base_fields + [
                {
                    'name': 'bot_id',
                    'type': 'str',
                    'required': False,
                    'description': 'Coze Bot ID（可选，默认使用环境变量配置）',
                    'example': '7584766797639958555'
                }
            ],
            'is_stream': True,
            'response_description': '本接口使用 SSE (Server-Sent Events) 流式输出，响应格式为 text/event-stream。',
            'response_fields': [
                {
                    'name': 'data',
                    'type': 'string (SSE格式)',
                    'description': 'SSE事件流，每个事件格式为：data: {"type": "progress|complete|error", "content": "..."}\\n\\n'
                },
                {
                    'name': 'type',
                    'type': 'string',
                    'description': '事件类型：progress（增量内容）、complete（完整内容）、error（错误消息）'
                },
                {
                    'name': 'content',
                    'type': 'string',
                    'description': '事件内容（分析文本或错误消息）'
                }
            ],
            'notes': [
                '通过 gRPC-Web 网关调用，路径：/grpc-web/frontend.gateway.FrontendGateway/Call',
                '请求体格式：{"endpoint": "/api/v1/children-study/stream", "payload_json": "{...}"}',
                '响应格式为 SSE 流式输出，前端需要使用 EventSource 或 fetch API 的流式读取方式处理',
                'bot_id 参数可选，如果未提供则使用环境变量配置',
                '分析内容由 Coze Bot 实时生成，可能需要较长时间（通常10-30秒）'
            ],
            'examples': {
                'progress': {
                    'sse_format': 'data: {"type":"progress","content":"命盘总论\\n\\n根据您的八字信息，"}\n\n',
                    'json': {
                        'type': 'progress',
                        'content': '命盘总论\n\n根据您的八字信息，'
                    },
                    'description': 'progress 事件表示增量内容，content 字段包含部分分析文本。'
                },
                'complete': {
                    'sse_format': 'data: {"type":"complete","content":"命盘总论\\n\\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\\n\\n子女特征\\n\\n您的子女特征..."}\n\n',
                    'json': {
                        'type': 'complete',
                        'content': '命盘总论\n\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\n\n子女特征\n\n您的子女特征...'
                    },
                    'description': 'complete 事件表示完整内容，content 字段包含完整的分析文本。'
                },
                'error': {
                    'sse_format': 'data: {"type":"error","content":"数据计算不完整: 缺失字段：mingpan_zonglun.bazi_pillars（八字排盘）。请检查生辰数据是否正确。"}\n\n',
                    'json': {
                        'type': 'error',
                        'content': '数据计算不完整: 缺失字段：mingpan_zonglun.bazi_pillars（八字排盘）。请检查生辰数据是否正确。'
                    },
                    'description': 'error 事件表示错误消息，前端应该显示错误信息给用户。'
                }
            }
        },
        {
            'title': '八字命理-身体健康分析',
            'path': '/api/v1/health/stream',
            'method': 'POST',
            'description': '基于用户生辰数据，使用 Coze Bot 流式生成身体健康分析。分析内容包括：命盘总论、健康特征、健康运势和建议方向。',
            'request_fields': base_fields + [
                {
                    'name': 'bot_id',
                    'type': 'str',
                    'required': False,
                    'description': 'Coze Bot ID（可选，默认使用环境变量配置）',
                    'example': '7584766797639958555'
                }
            ],
            'is_stream': True,
            'response_description': '本接口使用 SSE (Server-Sent Events) 流式输出，响应格式为 text/event-stream。',
            'response_fields': [
                {
                    'name': 'data',
                    'type': 'string (SSE格式)',
                    'description': 'SSE事件流，每个事件格式为：data: {"type": "progress|complete|error", "content": "..."}\\n\\n'
                },
                {
                    'name': 'type',
                    'type': 'string',
                    'description': '事件类型：progress（增量内容）、complete（完整内容）、error（错误消息）'
                },
                {
                    'name': 'content',
                    'type': 'string',
                    'description': '事件内容（分析文本或错误消息）'
                }
            ],
            'notes': [
                '通过 gRPC-Web 网关调用，路径：/grpc-web/frontend.gateway.FrontendGateway/Call',
                '请求体格式：{"endpoint": "/api/v1/health/stream", "payload_json": "{...}"}',
                '响应格式为 SSE 流式输出，前端需要使用 EventSource 或 fetch API 的流式读取方式处理',
                'bot_id 参数可选，如果未提供则使用环境变量配置',
                '分析内容由 Coze Bot 实时生成，可能需要较长时间（通常10-30秒）'
            ],
            'examples': {
                'progress': {
                    'sse_format': 'data: {"type":"progress","content":"命盘总论\\n\\n根据您的八字信息，"}\n\n',
                    'json': {
                        'type': 'progress',
                        'content': '命盘总论\n\n根据您的八字信息，'
                    },
                    'description': 'progress 事件表示增量内容，content 字段包含部分分析文本。'
                },
                'complete': {
                    'sse_format': 'data: {"type":"complete","content":"命盘总论\\n\\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\\n\\n健康特征\\n\\n您的健康特征..."}\n\n',
                    'json': {
                        'type': 'complete',
                        'content': '命盘总论\n\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\n\n健康特征\n\n您的健康特征...'
                    },
                    'description': 'complete 事件表示完整内容，content 字段包含完整的分析文本。'
                },
                'error': {
                    'sse_format': 'data: {"type":"error","content":"流式生成失败: Coze API 调用超时"}\n\n',
                    'json': {
                        'type': 'error',
                        'content': '流式生成失败: Coze API 调用超时'
                    },
                    'description': 'error 事件表示错误消息，前端应该显示错误信息给用户。'
                }
            }
        },
        {
            'title': '八字命理-总评分析',
            'path': '/api/v1/general-review/stream',
            'method': 'POST',
            'description': '基于用户生辰数据，使用 Coze Bot 流式生成总评分析。综合分析用户的八字命盘，包括：命盘总论、性格特征、事业运势、财富运势、感情运势、健康运势和总体建议。',
            'request_fields': base_fields + [
                {
                    'name': 'bot_id',
                    'type': 'str',
                    'required': False,
                    'description': 'Coze Bot ID（可选，默认使用环境变量配置）',
                    'example': '7584766797639958555'
                }
            ],
            'is_stream': True,
            'response_description': '本接口使用 SSE (Server-Sent Events) 流式输出，响应格式为 text/event-stream。',
            'response_fields': [
                {
                    'name': 'data',
                    'type': 'string (SSE格式)',
                    'description': 'SSE事件流，每个事件格式为：data: {"type": "progress|complete|error", "content": "..."}\\n\\n'
                },
                {
                    'name': 'type',
                    'type': 'string',
                    'description': '事件类型：progress（增量内容）、complete（完整内容）、error（错误消息）'
                },
                {
                    'name': 'content',
                    'type': 'string',
                    'description': '事件内容（分析文本或错误消息）'
                }
            ],
            'notes': [
                '通过 gRPC-Web 网关调用，路径：/grpc-web/frontend.gateway.FrontendGateway/Call',
                '请求体格式：{"endpoint": "/api/v1/general-review/stream", "payload_json": "{...}"}',
                '响应格式为 SSE 流式输出，前端需要使用 EventSource 或 fetch API 的流式读取方式处理',
                'bot_id 参数可选，如果未提供则使用环境变量配置',
                '分析内容由 Coze Bot 实时生成，可能需要较长时间（通常15-40秒，因为内容更全面）'
            ],
            'examples': {
                'progress': {
                    'sse_format': 'data: {"type":"progress","content":"命盘总论\\n\\n根据您的八字信息，"}\n\n',
                    'json': {
                        'type': 'progress',
                        'content': '命盘总论\n\n根据您的八字信息，'
                    },
                    'description': 'progress 事件表示增量内容，content 字段包含部分分析文本。'
                },
                'complete': {
                    'sse_format': 'data: {"type":"complete","content":"命盘总论\\n\\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\\n\\n性格特征\\n\\n您的性格特征...\\n\\n事业运势\\n\\n您的事业运势...\\n\\n财富运势\\n\\n您的财富运势..."}\n\n',
                    'json': {
                        'type': 'complete',
                        'content': '命盘总论\n\n根据您的八字信息，您的命盘为：年柱甲子，月柱乙丑，日柱丙寅，时柱丁卯。\n\n性格特征\n\n您的性格特征...\n\n事业运势\n\n您的事业运势...\n\n财富运势\n\n您的财富运势...'
                    },
                    'description': 'complete 事件表示完整内容，content 字段包含完整的分析文本，总评分析内容通常比较全面。'
                },
                'error': {
                    'sse_format': 'data: {"type":"error","content":"流式生成失败: Coze API 调用失败"}\n\n',
                    'json': {
                        'type': 'error',
                        'content': '流式生成失败: Coze API 调用失败'
                    },
                    'description': 'error 事件表示错误消息，前端应该显示错误信息给用户。'
                }
            }
        }
    ]
    
    return configs


def get_smart_analyze_config():
    """获取AI问答接口的配置"""
    return {
        'title': '八字命理-AI问答',
        'path': '/api/v1/smart-fortune/smart-analyze-stream',
        'method': 'GET',
        'description': '智能运势分析（流式输出版）。基于用户问题，使用意图识别、八字计算和规则匹配，结合 Coze Bot 流式生成个性化分析。支持多种意图类型：财富、事业、婚姻、健康、性格等。',
        'request_fields': [
            {
                'name': 'question',
                'type': 'str',
                'required': True,
                'description': '用户问题，例如："我今年的财运如何？"、"我什么时候会结婚？"',
                'example': '我今年的财运如何？'
            },
            {
                'name': 'year',
                'type': 'int',
                'required': True,
                'description': '出生年份',
                'example': 1990
            },
            {
                'name': 'month',
                'type': 'int',
                'required': True,
                'description': '出生月份（1-12）',
                'example': 5
            },
            {
                'name': 'day',
                'type': 'int',
                'required': True,
                'description': '出生日期（1-31）',
                'example': 15
            },
            {
                'name': 'hour',
                'type': 'int',
                'required': False,
                'description': '出生时辰（0-23），默认 12',
                'example': 14
            },
            {
                'name': 'gender',
                'type': 'str',
                'required': True,
                'description': '性别：male(男) 或 female(女)',
                'example': 'male'
            },
            {
                'name': 'user_id',
                'type': 'str',
                'required': False,
                'description': '用户ID（可选），用于标识用户',
                'example': 'user123'
            }
        ],
        'is_stream': True,
        'response_description': '本接口使用 SSE (Server-Sent Events) 流式输出，响应格式为 text/event-stream。响应包含多个阶段：意图识别、基础分析、LLM深度解读等。',
        'response_fields': [
            {
                'name': 'event',
                'type': 'string',
                'description': 'SSE事件类型：status（状态更新）、basic_analysis（基础分析）、llm_start（LLM开始）、llm_chunk（LLM增量内容）、llm_end（LLM结束）、llm_error（LLM错误）、performance（性能摘要）、error（错误）、end（结束）'
            },
            {
                'name': 'data',
                'type': 'string (JSON)',
                'description': '事件数据（JSON格式），根据不同事件类型包含不同的数据'
            },
            {
                'name': 'data.stage',
                'type': 'string',
                'description': '当前阶段（status事件）：intent（意图识别）、bazi（八字计算）、rules（规则匹配）、fortune（流年大运分析）、llm（LLM深度解读）'
            },
            {
                'name': 'data.message',
                'type': 'string',
                'description': '阶段消息（status事件）'
            },
            {
                'name': 'data.content',
                'type': 'string',
                'description': 'LLM生成的内容（llm_chunk事件）'
            }
        ],
        'notes': [
            '本接口使用 GET 方法，参数通过 Query String 传递',
            '通过 gRPC-Web 网关调用，路径：/grpc-web/frontend.gateway.FrontendGateway/Call',
            '请求URL格式：/api/v1/smart-fortune/smart-analyze-stream?question=...&year=...&month=...&day=...&hour=...&gender=...',
            '响应格式为 SSE 流式输出，前端需要使用 EventSource 或 fetch API 的流式读取方式处理',
            '接口流程：意图识别 → 八字计算 → 规则匹配 → 流年大运分析（可选）→ LLM深度解读（流式）',
            '意图识别支持混合架构：关键词过滤 → 本地BERT模型 → LLM兜底，响应时间 < 1秒',
            '如果问题与命理运势无关，接口会返回拒绝消息',
            '分析内容由 Coze Bot 实时生成，可能需要较长时间（通常10-30秒）'
        ],
        'examples': {
            'status': {
                'sse_format': 'event: status\ndata: {"stage":"intent","message":"正在识别意图..."}\n\n',
                'json': {
                    'stage': 'intent',
                    'message': '正在识别意图...'
                },
                'description': 'status 事件表示状态更新，stage 字段表示当前阶段（intent/bazi/rules/fortune/llm），message 字段表示阶段描述。'
            },
            'basic_analysis': {
                'sse_format': 'event: basic_analysis\ndata: {"intent":{"intents":["wealth"],"confidence":0.85},"bazi_info":{"四柱":{"年柱":"甲子","月柱":"乙丑"},"十神":{},"五行":{}},"matched_rules_count":5}\n\n',
                'json': {
                    'intent': {
                        'intents': ['wealth'],
                        'confidence': 0.85,
                        'method': 'local_model'
                    },
                    'bazi_info': {
                        '四柱': {
                            '年柱': {'天干': '甲', '地支': '子'},
                            '月柱': {'天干': '乙', '地支': '丑'},
                            '日柱': {'天干': '丙', '地支': '寅'},
                            '时柱': {'天干': '丁', '地支': '卯'}
                        },
                        '十神': {},
                        '五行': {}
                    },
                    'matched_rules_count': 5,
                    'fortune_context': None
                },
                'description': 'basic_analysis 事件表示基础分析结果，包含意图识别结果、八字信息、匹配规则数量等，前端可以立即显示这些信息。'
            },
            'llm_start': {
                'sse_format': 'event: llm_start\ndata: {}\n\n',
                'json': {},
                'description': 'llm_start 事件表示 LLM 深度解读开始，前端可以显示"正在生成深度解读..."的提示。'
            },
            'llm_chunk': {
                'sse_format': 'event: llm_chunk\ndata: {"content":"根据您的八字信息，"}\n\n',
                'json': {
                    'content': '根据您的八字信息，'
                },
                'description': 'llm_chunk 事件表示 LLM 增量内容，content 字段包含部分分析文本，前端需要累积这些内容实时显示。'
            },
            'llm_end': {
                'sse_format': 'event: llm_end\ndata: {}\n\n',
                'json': {},
                'description': 'llm_end 事件表示 LLM 深度解读完成，前端可以隐藏"正在生成..."的提示。'
            },
            'llm_error': {
                'sse_format': 'event: llm_error\ndata: {"message":"AI深度解读失败: Coze API 调用超时"}\n\n',
                'json': {
                    'message': 'AI深度解读失败: Coze API 调用超时'
                },
                'description': 'llm_error 事件表示 LLM 深度解读失败，message 字段包含错误描述，前端应该显示错误信息。'
            },
            'performance': {
                'sse_format': 'event: performance\ndata: {"total_duration_ms":2249,"stages":[{"name":"intent_recognition","duration_ms":50,"success":true}]}\n\n',
                'json': {
                    'total_duration_ms': 2249,
                    'stages': [
                        {
                            'name': 'intent_recognition',
                            'duration_ms': 50,
                            'success': True
                        },
                        {
                            'name': 'bazi_calculation',
                            'duration_ms': 23,
                            'success': True
                        },
                        {
                            'name': 'llm_analysis',
                            'duration_ms': 1202,
                            'success': True
                        }
                    ],
                    'bottlenecks': [
                        {
                            'stage': 'llm_analysis',
                            'duration_ms': 1202,
                            'reason': 'LLM深度解读'
                        }
                    ]
                },
                'description': 'performance 事件表示性能摘要，包含总耗时、各阶段耗时和性能瓶颈信息，前端可以用于性能分析（可选显示）。'
            },
            'error': {
                'sse_format': 'event: error\ndata: {"message":"处理失败: 八字计算失败"}\n\n',
                'json': {
                    'message': '处理失败: 八字计算失败',
                    'performance': {}
                },
                'description': 'error 事件表示处理错误，message 字段包含错误描述，前端应该显示错误信息给用户。'
            },
            'end': {
                'sse_format': 'event: end\ndata: {}\n\n',
                'json': {},
                'description': 'end 事件表示流式输出结束，前端可以关闭连接或进行清理工作。'
            }
        }
    }


def main():
    """主函数"""
    print("正在读取现有文档...")
    
    # 桌面路径
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
    doc_path = os.path.join(desktop_path, 'bazi_shengong_minggong_api.docx')
    
    # 检查文件是否存在
    if not os.path.exists(doc_path):
        print(f"错误：找不到文件 {doc_path}")
        print("请确保桌面上的 bazi_shengong_minggong_api.docx 文件存在")
        sys.exit(1)
    
    # 打开现有文档
    doc = Document(doc_path)
    
    # 设置中文字体（确保新增内容也有中文字体）
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    print("正在追加6个接口的文档...")
    
    # 添加流式接口文档
    stream_configs = get_stream_api_configs()
    for config in stream_configs:
        print(f"  添加接口：{config['title']}")
        add_api_section(doc, config)
    
    # 添加AI问答接口文档
    smart_config = get_smart_analyze_config()
    print(f"  添加接口：{smart_config['title']}")
    add_api_section(doc, smart_config)
    
    # 保存文档
    print("正在保存文档...")
    doc.save(doc_path)
    
    print(f"✅ 文档已更新：{doc_path}")
    print(f"✅ 已追加 {len(stream_configs) + 1} 个接口的文档")


if __name__ == '__main__':
    main()

