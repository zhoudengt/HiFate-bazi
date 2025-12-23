#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成自定义接口文档（Word格式）
包含8个指定接口的详细文档
"""

import sys
import os
from typing import Dict, Any, List
from datetime import datetime
from pathlib import Path

# 添加项目根目录到路径
project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, project_root)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


# 接口信息定义
INTERFACE_INFO = {
    '/bazi/interface': {
        'alias': '基本信息',
        'description': '生成八字界面信息（包含命宫、身宫、胎元、胎息、命卦等）',
        'request_method': 'POST',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'name', 'type': 'str', 'required': False, 'description': '姓名（可选）', 'example': '张三'},
            {'name': 'calendar_type', 'type': 'str', 'required': False, 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar', 'default': 'solar'},
            {'name': 'location', 'type': 'str', 'required': False, 'description': '出生地点（可选，用于时区转换）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': False, 'description': '纬度（可选，用于时区转换）', 'example': 39.90},
            {'name': 'longitude', 'type': 'float', 'required': False, 'description': '经度（可选，用于时区转换和真太阳时计算）', 'example': 116.40},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '界面信息数据', 'nested': [
                {'name': 'basic_info', 'type': 'dict', 'description': '基本信息', 'nested': [
                    {'name': 'name', 'type': 'str', 'description': '姓名'},
                    {'name': 'gender', 'type': 'str', 'description': '性别'},
                    {'name': 'solar_date', 'type': 'str', 'description': '阳历日期'},
                    {'name': 'solar_time', 'type': 'str', 'description': '出生时间'},
                    {'name': 'lunar_date', 'type': 'str', 'description': '农历日期'},
                    {'name': 'location', 'type': 'str', 'description': '出生地点'},
                ]},
                {'name': 'bazi_pillars', 'type': 'dict', 'description': '四柱信息', 'nested': [
                    {'name': 'year', 'type': 'str', 'description': '年柱'},
                    {'name': 'month', 'type': 'str', 'description': '月柱'},
                    {'name': 'day', 'type': 'str', 'description': '日柱'},
                    {'name': 'hour', 'type': 'str', 'description': '时柱'},
                ]},
                {'name': 'astrology', 'type': 'dict', 'description': '占星信息', 'nested': [
                    {'name': 'zodiac', 'type': 'str', 'description': '生肖'},
                    {'name': 'constellation', 'type': 'str', 'description': '星座'},
                    {'name': 'mansion', 'type': 'str', 'description': '二十八宿'},
                    {'name': 'bagua', 'type': 'str', 'description': '命卦'},
                ]},
                {'name': 'palaces', 'type': 'dict', 'description': '宫位信息', 'nested': [
                    {'name': 'life_palace', 'type': 'dict', 'description': '命宫', 'nested': [
                        {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                        {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    ]},
                    {'name': 'body_palace', 'type': 'dict', 'description': '身宫', 'nested': [
                        {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                        {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    ]},
                    {'name': 'fetal_origin', 'type': 'dict', 'description': '胎元', 'nested': [
                        {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                        {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    ]},
                    {'name': 'fetal_breath', 'type': 'dict', 'description': '胎息', 'nested': [
                        {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                        {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    ]},
                ]},
                {'name': 'solar_terms', 'type': 'dict', 'description': '节气信息'},
                {'name': 'commander', 'type': 'dict', 'description': '人元司令'},
                {'name': 'void_emptiness', 'type': 'str', 'description': '空亡'},
            ]},
            {'name': 'message', 'type': 'str', 'description': '消息（可选）'},
        ],
        'note': '返回完整的八字界面信息，包含命宫、身宫、胎元、胎息、命卦、生肖、星座等详细信息',
    },
    '/bazi/pan/display': {
        'alias': '基本排盘',
        'description': '获取排盘数据（前端优化格式）',
        'request_method': 'POST',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'str', 'required': False, 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar', 'default': 'solar'},
            {'name': 'location', 'type': 'str', 'required': False, 'description': '出生地点（可选，用于时区转换）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': False, 'description': '纬度（可选，用于时区转换）', 'example': 39.90},
            {'name': 'longitude', 'type': 'float', 'required': False, 'description': '经度（可选，用于时区转换和真太阳时计算）', 'example': 116.40},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'pan', 'type': 'dict', 'description': '排盘数据', 'nested': [
                {'name': 'basic', 'type': 'dict', 'description': '基本信息（包含阳历、农历、性别等）'},
                {'name': 'pillars', 'type': 'list', 'description': '四柱数组（年柱、月柱、日柱、时柱）', 'nested': [
                    {'name': 'type', 'type': 'str', 'description': '柱类型（year/month/day/hour）'},
                    {'name': 'label', 'type': 'str', 'description': '柱标签（年柱/月柱/日柱/时柱）'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                        {'name': 'yinyang', 'type': 'str', 'description': '阴阳属性'},
                        {'name': 'ten_god', 'type': 'str', 'description': '十神'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                        {'name': 'yinyang', 'type': 'str', 'description': '阴阳属性'},
                        {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表', 'nested': [
                            {'name': 'char', 'type': 'str', 'description': '藏干字符（如：己土）'},
                            {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                            {'name': 'yinyang', 'type': 'str', 'description': '阴阳属性'},
                            {'name': 'ten_god', 'type': 'str', 'description': '十神'},
                        ]},
                    ]},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'hidden_stars', 'type': 'list', 'description': '副星列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '副星名称（如：正官、七杀等）'},
                    ]},
                    {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                    {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                    {'name': 'kongwang', 'type': 'str', 'description': '空亡'},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'deities', 'type': 'list', 'description': '神煞列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '神煞名称（如：天乙贵人、桃花等）'},
                    ]},
                ]},
                {'name': 'wuxing', 'type': 'dict', 'description': '五行统计（包含各五行数量和百分比）', 'nested': [
                    {'name': 'counts', 'type': 'dict', 'description': '五行数量统计'},
                    {'name': 'percentages', 'type': 'dict', 'description': '五行百分比统计'},
                ]},
                {'name': 'rizhu_analysis', 'type': 'dict', 'description': '日柱解析（性格与命运解析）', 'nested': [
                    {'name': 'rizhu', 'type': 'str', 'description': '日柱干支'},
                    {'name': 'gender', 'type': 'str', 'description': '性别'},
                    {'name': 'descriptions', 'type': 'list', 'description': '描述列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '描述文本'},
                    ]},
                    {'name': 'summary', 'type': 'str', 'description': '总结'},
                ]},
                {'name': 'marriage_rules', 'type': 'list', 'description': '婚姻规则列表', 'nested': [
                    {'name': 'rule_id', 'type': 'str', 'description': '规则ID'},
                    {'name': 'rule_code', 'type': 'str', 'description': '规则编码'},
                    {'name': 'rule_name', 'type': 'str', 'description': '规则名称'},
                    {'name': 'rule_type', 'type': 'str', 'description': '规则类型（如：marriage）'},
                    {'name': 'content', 'type': 'dict', 'description': '规则内容', 'nested': [
                        {'name': 'text', 'type': 'str', 'description': '规则文本'},
                    ]},
                    {'name': 'priority', 'type': 'int', 'description': '优先级'},
                ]},
            ]},
            {'name': 'conversion_info', 'type': 'dict', 'description': '转换信息（如果进行了农历转换或时区转换）'},
        ],
        'note': '返回前端友好的排盘数据，包括基本信息、四柱数组、五行统计、日柱解析和婚姻规则',
    },
    '/bazi/fortune/display': {
        'alias': '专业排盘-大运流年流月',
        'description': '获取大运流年流月数据（统一接口，性能优化）',
        'request_method': 'POST',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'str', 'required': False, 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar', 'default': 'solar'},
            {'name': 'location', 'type': 'str', 'required': False, 'description': '出生地点（可选，用于时区转换）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': False, 'description': '纬度（可选，用于时区转换）', 'example': 39.90},
            {'name': 'longitude', 'type': 'float', 'required': False, 'description': '经度（可选，用于时区转换和真太阳时计算）', 'example': 116.40},
            {'name': 'current_time', 'type': 'str', 'required': False, 'description': '当前时间（可选），格式：YYYY-MM-DD HH:MM', 'example': '2024-01-01 12:00'},
            {'name': 'dayun_year_start', 'type': 'int', 'required': False, 'description': '大运起始年份（可选），指定要显示的大运的起始年份', 'example': 2020},
            {'name': 'dayun_year_end', 'type': 'int', 'required': False, 'description': '大运结束年份（可选），指定要显示的大运的结束年份', 'example': 2030},
            {'name': 'target_year', 'type': 'int', 'required': False, 'description': '目标年份（可选），用于计算该年份的流月', 'example': 2024},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'pillars', 'type': 'dict', 'description': '四柱详细信息'},
            {'name': 'dayun', 'type': 'dict', 'description': '大运数据', 'nested': [
                {'name': 'current', 'type': 'dict', 'description': '当前大运', 'nested': [
                    {'name': 'index', 'type': 'int', 'description': '大运序号'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'age_range', 'type': 'dict', 'description': '年龄范围', 'nested': [
                        {'name': 'start', 'type': 'int', 'description': '起始年龄'},
                        {'name': 'end', 'type': 'int', 'description': '结束年龄'},
                    ]},
                    {'name': 'age_display', 'type': 'str', 'description': '年龄显示（如：10-19岁）'},
                    {'name': 'year_range', 'type': 'dict', 'description': '年份范围', 'nested': [
                        {'name': 'start', 'type': 'int', 'description': '起始年份'},
                        {'name': 'end', 'type': 'int', 'description': '结束年份'},
                    ]},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '藏干字符'},
                    ]},
                    {'name': 'hidden_stars', 'type': 'list', 'description': '副星列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '副星名称'},
                    ]},
                    {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                    {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                    {'name': 'kongwang', 'type': 'str', 'description': '空亡'},
                    {'name': 'deities', 'type': 'list', 'description': '神煞列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '神煞名称'},
                    ]},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前大运'},
                ]},
                {'name': 'list', 'type': 'list', 'description': '大运列表（数组格式，每个元素结构同current）', 'nested': [
                    {'name': 'index', 'type': 'int', 'description': '大运序号'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'age_range', 'type': 'dict', 'description': '年龄范围', 'nested': [
                        {'name': 'start', 'type': 'int', 'description': '起始年龄'},
                        {'name': 'end', 'type': 'int', 'description': '结束年龄'},
                    ]},
                    {'name': 'age_display', 'type': 'str', 'description': '年龄显示（如：10-19岁）'},
                    {'name': 'year_range', 'type': 'dict', 'description': '年份范围', 'nested': [
                        {'name': 'start', 'type': 'int', 'description': '起始年份'},
                        {'name': 'end', 'type': 'int', 'description': '结束年份'},
                    ]},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '藏干字符'},
                    ]},
                    {'name': 'hidden_stars', 'type': 'list', 'description': '副星列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '副星名称'},
                    ]},
                    {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                    {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                    {'name': 'kongwang', 'type': 'str', 'description': '空亡'},
                    {'name': 'deities', 'type': 'list', 'description': '神煞列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '神煞名称'},
                    ]},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前大运'},
                ]},
                {'name': 'qiyun', 'type': 'dict', 'description': '起运信息'},
                {'name': 'jiaoyun', 'type': 'dict', 'description': '交运信息'},
            ]},
            {'name': 'liunian', 'type': 'dict', 'description': '流年数据', 'nested': [
                {'name': 'current', 'type': 'dict', 'description': '当前流年', 'nested': [
                    {'name': 'year', 'type': 'int', 'description': '年份'},
                    {'name': 'age', 'type': 'int', 'description': '年龄'},
                    {'name': 'age_display', 'type': 'str', 'description': '年龄显示'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '藏干字符'},
                    ]},
                    {'name': 'hidden_stars', 'type': 'list', 'description': '副星列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '副星名称'},
                    ]},
                    {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                    {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                    {'name': 'kongwang', 'type': 'str', 'description': '空亡'},
                    {'name': 'deities', 'type': 'list', 'description': '神煞列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '神煞名称'},
                    ]},
                    {'name': 'relations', 'type': 'list', 'description': '关系列表（字符串数组，如：冲、合、刑等）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '关系名称'},
                    ]},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前流年'},
                ]},
                {'name': 'list', 'type': 'list', 'description': '流年列表（数组格式，每个元素结构同current）', 'nested': [
                    {'name': 'year', 'type': 'int', 'description': '年份'},
                    {'name': 'age', 'type': 'int', 'description': '年龄'},
                    {'name': 'age_display', 'type': 'str', 'description': '年龄显示'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'hidden_stems', 'type': 'list', 'description': '藏干列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '藏干字符'},
                    ]},
                    {'name': 'hidden_stars', 'type': 'list', 'description': '副星列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '副星名称'},
                    ]},
                    {'name': 'star_fortune', 'type': 'str', 'description': '星运'},
                    {'name': 'self_sitting', 'type': 'str', 'description': '自坐'},
                    {'name': 'kongwang', 'type': 'str', 'description': '空亡'},
                    {'name': 'deities', 'type': 'list', 'description': '神煞列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '神煞名称'},
                    ]},
                    {'name': 'relations', 'type': 'list', 'description': '关系列表（字符串数组，如：冲、合、刑等）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '关系名称'},
                    ]},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前流年'},
                ]},
            ]},
            {'name': 'liuyue', 'type': 'dict', 'description': '流月数据', 'nested': [
                {'name': 'current', 'type': 'dict', 'description': '当前流月', 'nested': [
                    {'name': 'month', 'type': 'int', 'description': '月份（1-12）'},
                    {'name': 'solar_term', 'type': 'str', 'description': '节气名称'},
                    {'name': 'term_date', 'type': 'str', 'description': '节气日期'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前流月'},
                ]},
                {'name': 'list', 'type': 'list', 'description': '流月列表（数组格式，每个元素结构同current）', 'nested': [
                    {'name': 'month', 'type': 'int', 'description': '月份（1-12）'},
                    {'name': 'solar_term', 'type': 'str', 'description': '节气名称'},
                    {'name': 'term_date', 'type': 'str', 'description': '节气日期'},
                    {'name': 'ganzhi', 'type': 'str', 'description': '干支'},
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                    {'name': 'is_current', 'type': 'bool', 'description': '是否为当前流月'},
                ]},
                {'name': 'target_year', 'type': 'int', 'description': '目标年份'},
            ]},
            {'name': 'conversion_info', 'type': 'dict', 'description': '转换信息（如果进行了农历转换或时区转换）'},
        ],
        'note': '性能优化：只计算指定大运范围内的流年（约10年），而不是所有流年。一次返回大运、流年、流月所有数据。',
    },
    '/bazi/shengong-minggong': {
        'alias': '专业排盘-神宫命宫胎元',
        'description': '获取身宫和命宫的详细信息（主星、藏干、星运、自坐、空亡、纳音、神煞等）',
        'request_method': 'POST',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'str', 'required': False, 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar', 'default': 'solar'},
            {'name': 'location', 'type': 'str', 'required': False, 'description': '出生地点（可选，用于时区转换）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': False, 'description': '纬度（可选，用于时区转换）', 'example': 39.90},
            {'name': 'longitude', 'type': 'float', 'required': False, 'description': '经度（可选，用于时区转换和真太阳时计算）', 'example': 116.40},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '身宫命宫数据', 'nested': [
                {'name': 'shengong', 'type': 'dict', 'description': '身宫详细信息', 'nested': [
                    {'name': 'stem', 'type': 'dict', 'description': '天干信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '天干字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'branch', 'type': 'dict', 'description': '地支信息', 'nested': [
                        {'name': 'char', 'type': 'str', 'description': '地支字符'},
                        {'name': 'wuxing', 'type': 'str', 'description': '五行属性'},
                    ]},
                    {'name': 'main_star', 'type': 'str', 'description': '主星'},
                    {'name': 'hidden_stars', 'type': 'list', 'description': '副星列表（字符串数组）', 'nested': [
                        {'name': 'item', 'type': 'str', 'description': '副星名称'},
                    ]},
                    {'name': 'nayin', 'type': 'str', 'description': '纳音'},
                ]},
                {'name': 'minggong', 'type': 'dict', 'description': '命宫详细信息（结构同shengong）'},
                {'name': 'taiyuan', 'type': 'dict', 'description': '胎元详细信息（结构同shengong）'},
                {'name': 'pillars', 'type': 'dict', 'description': '四柱详细信息'},
            ]},
            {'name': 'conversion_info', 'type': 'dict', 'description': '转换信息（如果进行了农历转换或时区转换）'},
        ],
        'note': '返回身宫、命宫和胎元的详细信息，包括主星、藏干、星运、自坐、空亡、纳音、神煞等',
    },
    '/daily-fortune-calendar/query': {
        'alias': '八字命理-每日运势',
        'description': '查询每日运势日历信息',
        'request_method': 'POST',
        'request_fields': [
            {'name': 'date', 'type': 'str', 'required': False, 'description': '日期（可选，默认为今天），格式：YYYY-MM-DD', 'example': '2025-01-15'},
            {'name': 'solar_date', 'type': 'str', 'required': False, 'description': '用户生辰阳历日期（可选，用于十神提示），格式：YYYY-MM-DD 或农历日期（当calendar_type=lunar时）', 'example': '1990-01-15'},
            {'name': 'solar_time', 'type': 'str', 'required': False, 'description': '用户生辰时间（可选），格式：HH:MM', 'example': '12:00'},
            {'name': 'gender', 'type': 'str', 'required': False, 'description': '用户性别（可选），male/female', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'str', 'required': False, 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar', 'default': 'solar'},
            {'name': 'location', 'type': 'str', 'required': False, 'description': '出生地点（可选，用于时区转换）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': False, 'description': '纬度（可选，用于时区转换）', 'example': 39.90},
            {'name': 'longitude', 'type': 'float', 'required': False, 'description': '经度（可选，用于时区转换和真太阳时计算）', 'example': 116.40},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'solar_date', 'type': 'str', 'description': '当前阳历日期'},
            {'name': 'lunar_date', 'type': 'str', 'description': '当前阴历日期'},
            {'name': 'weekday', 'type': 'str', 'description': '星期几（中文）'},
            {'name': 'yi', 'type': 'list', 'description': '宜（列表，字符串数组）', 'nested': [
                {'name': 'item', 'type': 'str', 'description': '宜的事项（如：祭祀、祈福、开光等）'},
            ]},
            {'name': 'ji', 'type': 'list', 'description': '忌（列表，字符串数组）', 'nested': [
                {'name': 'item', 'type': 'str', 'description': '忌的事项（如：破土、安葬、入宅等）'},
            ]},
            {'name': 'luck_level', 'type': 'str', 'description': '吉凶等级'},
            {'name': 'deities', 'type': 'dict', 'description': '神煞方位（喜神、财神、福神、胎神）'},
            {'name': 'chong_he_sha', 'type': 'dict', 'description': '冲合煞（冲、合、煞）'},
            {'name': 'jianchu', 'type': 'dict', 'description': '建除信息', 'nested': [
                {'name': 'name', 'type': 'str', 'description': '建除名称（如：危）'},
                {'name': 'energy', 'type': 'int', 'description': '能量（如：90）'},
                {'name': 'summary', 'type': 'str', 'description': '能量小结内容'},
            ]},
            {'name': 'taishen', 'type': 'str', 'description': '胎神方位（如：占门厕外正东）'},
            {'name': 'taishen_explanation', 'type': 'str', 'description': '胎神解释'},
            {'name': 'jiazi_fortune', 'type': 'str', 'description': '整体运势（六十甲子）'},
            {'name': 'shishen_hint', 'type': 'str', 'description': '十神提示（需要用户生辰）'},
            {'name': 'zodiac_relations', 'type': 'str', 'description': '生肖简运'},
            {'name': 'master_info', 'type': 'dict', 'description': '命主信息', 'nested': [
                {'name': 'rizhu', 'type': 'str', 'description': '日主（如：甲木）'},
                {'name': 'today_shishen', 'type': 'str', 'description': '今日十神（如：比肩）'},
            ]},
            {'name': 'wuxing_wear', 'type': 'str', 'description': '五行穿搭（逗号分隔）'},
            {'name': 'guiren_fangwei', 'type': 'str', 'description': '贵人方位（逗号分隔）'},
            {'name': 'wenshen_directions', 'type': 'str', 'description': '瘟神方位（逗号分隔）'},
            {'name': 'error', 'type': 'str', 'description': '错误信息（可选）'},
        ],
        'note': '基于万年历接口，提供完整的每日运势信息。如果未提供用户生辰信息，十神提示将为空。',
    },
    '/bazi/wuxing-proportion': {
        'alias': '八字命理-五行占比',
        'description': '查询五行占比分析',
        'request_method': 'POST',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'str', 'required': False, 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar', 'default': 'solar'},
            {'name': 'location', 'type': 'str', 'required': False, 'description': '出生地点（可选，用于时区转换）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': False, 'description': '纬度（可选，用于时区转换）', 'example': 39.90},
            {'name': 'longitude', 'type': 'float', 'required': False, 'description': '经度（可选，用于时区转换和真太阳时计算）', 'example': 116.40},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '五行占比数据', 'nested': [
                {'name': 'bazi_pillars', 'type': 'dict', 'description': '四柱信息'},
                {'name': 'proportions', 'type': 'dict', 'description': '五行占比统计', 'nested': [
                    {'name': '金', 'type': 'dict', 'description': '金行占比', 'nested': [
                        {'name': 'count', 'type': 'int', 'description': '数量'},
                        {'name': 'percentage', 'type': 'float', 'description': '百分比'},
                        {'name': 'details', 'type': 'list', 'description': '详细列表（字符串数组，包含该五行的天干地支字符）', 'nested': [
                            {'name': 'item', 'type': 'str', 'description': '天干或地支字符（如：甲、子等）'},
                        ]},
                    ]},
                    {'name': '木', 'type': 'dict', 'description': '木行占比（结构同金）'},
                    {'name': '水', 'type': 'dict', 'description': '水行占比（结构同金）'},
                    {'name': '火', 'type': 'dict', 'description': '火行占比（结构同金）'},
                    {'name': '土', 'type': 'dict', 'description': '土行占比（结构同金）'},
                ]},
                {'name': 'ten_gods', 'type': 'dict', 'description': '四柱十神信息（主星和副星）', 'nested': [
                    {'name': 'year', 'type': 'dict', 'description': '年柱十神', 'nested': [
                        {'name': 'main_star', 'type': 'str', 'description': '主星'},
                        {'name': 'hidden_stars', 'type': 'list', 'description': '副星列表（字符串数组）', 'nested': [
                            {'name': 'item', 'type': 'str', 'description': '副星名称'},
                        ]},
                    ]},
                    {'name': 'month', 'type': 'dict', 'description': '月柱十神（结构同年）'},
                    {'name': 'day', 'type': 'dict', 'description': '日柱十神（结构同年）'},
                    {'name': 'hour', 'type': 'dict', 'description': '时柱十神（结构同年）'},
                ]},
                {'name': 'wangshuai', 'type': 'dict', 'description': '旺衰分析结果'},
                {'name': 'element_relations', 'type': 'dict', 'description': '相生相克关系'},
            ]},
            {'name': 'error', 'type': 'str', 'description': '错误信息（可选）'},
        ],
        'note': '基于生辰八字统计五行占比（金木水火土），包括五行占比统计、四柱十神信息、旺衰分析结果和相生相克关系',
    },
    '/bazi/rizhu-liujiazi': {
        'alias': '八字命理-日元-六十甲子',
        'description': '根据用户生辰查询日柱对应的六十甲子解析',
        'request_method': 'POST',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'str', 'required': False, 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar', 'default': 'solar'},
            {'name': 'location', 'type': 'str', 'required': False, 'description': '出生地点（可选，用于时区转换）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': False, 'description': '纬度（可选，用于时区转换）', 'example': 39.90},
            {'name': 'longitude', 'type': 'float', 'required': False, 'description': '经度（可选，用于时区转换和真太阳时计算）', 'example': 116.40},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '日柱解析数据', 'nested': [
                {'name': 'id', 'type': 'int', 'description': '记录ID'},
                {'name': 'rizhu', 'type': 'str', 'description': '日柱（如：乙丑）'},
                {'name': 'analysis', 'type': 'str', 'description': '解析内容（包含【基础信息】、【深度解读】、【断语展示】等）'},
                {'name': 'enabled', 'type': 'bool', 'description': '是否启用'},
            ]},
            {'name': 'error', 'type': 'str', 'description': '错误信息（可选）'},
        ],
        'note': '流程：1. 调用八字排盘服务获取日柱；2. 根据日柱查询数据库获取解析内容；3. 返回ID、日柱、解析内容',
    },
    '/bazi/xishen-jishen': {
        'alias': '八字命理-喜神忌神',
        'description': '获取喜神五行、忌神五行和十神命格',
        'request_method': 'POST',
        'request_fields': [
            {'name': 'solar_date', 'type': 'str', 'required': True, 'description': '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', 'example': '1990-05-15'},
            {'name': 'solar_time', 'type': 'str', 'required': True, 'description': '出生时间，格式：HH:MM', 'example': '14:30'},
            {'name': 'gender', 'type': 'str', 'required': True, 'description': '性别：male(男) 或 female(女)', 'example': 'male'},
            {'name': 'calendar_type', 'type': 'str', 'required': False, 'description': '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'example': 'solar', 'default': 'solar'},
            {'name': 'location', 'type': 'str', 'required': False, 'description': '出生地点（可选，用于时区转换）', 'example': '北京'},
            {'name': 'latitude', 'type': 'float', 'required': False, 'description': '纬度（可选，用于时区转换）', 'example': 39.90},
            {'name': 'longitude', 'type': 'float', 'required': False, 'description': '经度（可选，用于时区转换和真太阳时计算）', 'example': 116.40},
        ],
        'response_fields': [
            {'name': 'success', 'type': 'bool', 'description': '是否成功'},
            {'name': 'data', 'type': 'dict', 'description': '喜神忌神数据', 'nested': [
                {'name': 'solar_date', 'type': 'str', 'description': '阳历日期'},
                {'name': 'solar_time', 'type': 'str', 'description': '出生时间'},
                {'name': 'gender', 'type': 'str', 'description': '性别'},
                {'name': 'xi_shen_elements', 'type': 'list', 'description': '喜神五行列表（包含名称和ID）', 'nested': [
                    {'name': 'name', 'type': 'str', 'description': '五行名称（如：金）'},
                    {'name': 'id', 'type': 'int', 'description': '五行ID（如：4）'},
                ]},
                {'name': 'ji_shen_elements', 'type': 'list', 'description': '忌神五行列表（包含名称和ID）', 'nested': [
                    {'name': 'name', 'type': 'str', 'description': '五行名称（如：水）'},
                    {'name': 'id', 'type': 'int', 'description': '五行ID（如：5）'},
                ]},
                {'name': 'shishen_mingge', 'type': 'list', 'description': '十神命格列表（包含名称和ID）', 'nested': [
                    {'name': 'name', 'type': 'str', 'description': '命格名称（如：正官格）'},
                    {'name': 'id', 'type': 'int', 'description': '命格ID（如：2001）'},
                ]},
                {'name': 'wangshuai', 'type': 'str', 'description': '旺衰状态'},
                {'name': 'total_score', 'type': 'float', 'description': '总分'},
            ]},
            {'name': 'error', 'type': 'str', 'description': '错误信息（可选）'},
        ],
        'note': '根据用户的生辰：1. 从旺衰分析中获取喜神五行和忌神五行；2. 从公式分析中获取十神命格；3. 查询配置表获取对应的ID',
    },
}


def create_document() -> Document:
    """创建文档"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('HiFate-bazi 接口文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 文档信息
    info_para = doc.add_paragraph()
    info_para.add_run('生成时间：').bold = True
    info_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    info_para.add_run('\n')
    info_para.add_run('版本：').bold = True
    info_para.add_run('v1.0')
    info_para.add_run('\n')
    info_para.add_run('说明：').bold = True
    info_para.add_run('本文档包含8个指定接口的详细说明，包括请求参数、响应格式和注意事项。所有接口通过 gRPC-Web 网关调用。')
    info_para.add_run('\n')
    info_para.add_run('调用方式：').bold = True
    info_para.add_run('POST /grpc-web/frontend.gateway.FrontendGateway/Call')
    info_para.add_run('\n')
    info_para.add_run('请求体格式：').bold = True
    info_para.add_run('{"endpoint": "/接口路径", "payload_json": "JSON字符串"}')
    
    doc.add_paragraph()  # 空行
    doc.add_page_break()
    
    # 接口详情
    doc.add_heading('接口详情', 1)
    
    for endpoint, info in INTERFACE_INFO.items():
        # 接口标题（带别名）
        title_text = f"{endpoint} ({info['alias']})"
        doc.add_heading(title_text, 2)
        
        # 基本信息表格
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_table.cell(0, 0).text = '接口路径'
        info_table.cell(0, 1).text = endpoint
        info_table.cell(1, 0).text = '接口别名'
        info_table.cell(1, 1).text = info['alias']
        info_table.cell(2, 0).text = '请求方法'
        info_table.cell(2, 1).text = info['request_method']
        info_table.cell(3, 0).text = '接口描述'
        info_table.cell(3, 1).text = info['description']
        
        # 如果有备注，添加一行
        if info.get('note'):
            row = info_table.add_row()
            row.cells[0].text = '备注'
            row.cells[1].text = info['note']
        
        doc.add_paragraph()  # 空行
        
        # 请求参数
        if info['request_fields']:
            doc.add_heading('请求参数', 3)
            req_table = doc.add_table(rows=1, cols=5)
            req_table.style = 'Light Grid Accent 1'
            
            # 表头
            headers = ['字段名', '类型', '必填', '描述', '示例']
            for i, header in enumerate(headers):
                req_table.cell(0, i).text = header
                req_table.cell(0, i).paragraphs[0].runs[0].bold = True
            
            # 数据行
            for field in info['request_fields']:
                row = req_table.add_row()
                row.cells[0].text = field['name']
                row.cells[1].text = field['type']
                row.cells[2].text = '是' if field['required'] else '否'
                row.cells[3].text = field.get('description', '')
                example = field.get('example', '')
                if field.get('default') is not None:
                    example = f"{example} (默认: {field['default']})" if example else f"默认: {field['default']}"
                row.cells[4].text = str(example) if example else ''
            
            doc.add_paragraph()  # 空行
        
        # 响应格式
        if info['response_fields']:
            doc.add_heading('响应格式', 3)
            resp_table = doc.add_table(rows=1, cols=3)
            resp_table.style = 'Light Grid Accent 1'
            
            # 表头
            headers = ['字段名', '类型', '描述']
            for i, header in enumerate(headers):
                resp_table.cell(0, i).text = header
                resp_table.cell(0, i).paragraphs[0].runs[0].bold = True
            
            # 递归添加字段（包括嵌套字段）
            def add_field_to_table(field, prefix='', level=0):
                """递归添加字段到表格"""
                indent = '  ' * level
                field_name = f"{indent}{prefix}{field['name']}" if prefix else f"{indent}{field['name']}"
                
                row = resp_table.add_row()
                row.cells[0].text = field_name
                row.cells[1].text = field['type']
                row.cells[2].text = field.get('description', '')
                
                # 如果有嵌套字段，递归添加
                if 'nested' in field and field['nested']:
                    # 如果是列表类型，说明嵌套字段是列表元素的结构
                    if field['type'].startswith('list'):
                        prefix_str = f"{field['name']}[]."
                    else:
                        prefix_str = f"{field['name']}."
                    
                    for nested_field in field['nested']:
                        add_field_to_table(nested_field, prefix=prefix_str, level=level+1)
            
            # 数据行
            for field in info['response_fields']:
                add_field_to_table(field)
            
            doc.add_paragraph()  # 空行
        
        # 注意事项
        doc.add_heading('注意事项', 3)
        notes = [
            '通过 gRPC-Web 网关调用，路径：/grpc-web/frontend.gateway.FrontendGateway/Call',
            '请求体格式：{"endpoint": "' + endpoint + '", "payload_json": "{...}"}',
            '所有接口均返回 JSON 格式数据',
            '所有日期时间参数支持农历输入（当calendar_type=lunar时）',
            '所有接口支持时区转换（通过location或latitude/longitude参数）',
        ]
        
        for note in notes:
            para = doc.add_paragraph(note, style='List Bullet')
        
        doc.add_paragraph()  # 空行
        doc.add_paragraph('=' * 80)  # 分隔线
        doc.add_paragraph()  # 空行
    
    return doc


def main():
    """主函数"""
    print("开始生成接口文档...")
    
    # 创建文档
    doc = create_document()
    
    # 保存到桌面
    desktop_path = Path.home() / 'Desktop'
    output_path = desktop_path / '接口文档.docx'
    
    print(f"保存文档到: {output_path}")
    doc.save(str(output_path))
    print("文档生成完成！")


if __name__ == '__main__':
    main()

