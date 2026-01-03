#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成完整的 input_data 文档（包含所有大运流年数据）
"""

import json
import sys
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 添加项目根目录到路径
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, project_root)


def format_value(value, indent=0, max_length=None):
    """格式化值，递归处理嵌套结构"""
    prefix = "  " * indent
    
    if value is None:
        return f"{prefix}None"
    elif isinstance(value, bool):
        return f"{prefix}{value}"
    elif isinstance(value, (int, float)):
        return f"{prefix}{value}"
    elif isinstance(value, str):
        # 如果字符串太长，截断
        limit = max_length if max_length else 200
        if len(value) > limit:
            return f"{prefix}{value[:limit]}..."
        return f"{prefix}{value}"
    elif isinstance(value, dict):
        lines = [f"{prefix}{{"]
        for k, v in value.items():
            lines.append(f"{prefix}  {k}:")
            lines.append(format_value(v, indent + 2, max_length))
        lines.append(f"{prefix}}}")
        return "\n".join(lines)
    elif isinstance(value, list):
        lines = [f"{prefix}[长度: {len(value)}]"]
        for i, item in enumerate(value):
            lines.append(f"{prefix}  [{i}]:")
            lines.append(format_value(item, indent + 2, max_length))
        lines.append(f"{prefix}]")
        return "\n".join(lines)
    else:
        return f"{prefix}{str(value)}"


def add_section(doc, title, data, level=1):
    """添加章节"""
    heading = doc.add_heading(title, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加数据内容
    if isinstance(data, dict):
        for key, value in data.items():
            # 添加字段名
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            
            # 添加字段值
            if isinstance(value, (dict, list)):
                # 复杂结构，递归处理
                value_str = format_value(value, indent=0, max_length=500)
                doc.add_paragraph(value_str, style='List Bullet')
            else:
                p.add_run(str(value))
    elif isinstance(data, list):
        for i, item in enumerate(data):
            p = doc.add_paragraph()
            p.add_run(f"[{i}]: ").bold = True
            if isinstance(item, (dict, list)):
                value_str = format_value(item, indent=0, max_length=500)
                doc.add_paragraph(value_str, style='List Bullet')
            else:
                p.add_run(str(item))
    else:
        doc.add_paragraph(str(data))


def generate_full_input_data_doc(input_data_file, output_file):
    """生成完整的 input_data 文档"""
    
    # 读取 input_data
    with open(input_data_file, 'r', encoding='utf-8') as f:
        debug_data = json.load(f)
    
    input_data = debug_data.get('input_data', {})
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加标题
    title = doc.add_heading('子女学习分析 - 完整 input_data 数据', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加用户信息
    doc.add_heading('用户信息', 1)
    user_info = {
        'solar_date': '1987-09-16',
        'solar_time': '05:00',
        'gender': 'male',
        'calendar_type': 'solar'
    }
    for key, value in user_info.items():
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(str(value))
    
    # 添加各个部分
    sections = [
        ('mingpan_zinv_zonglun', '1. 命盘子女总论'),
        ('zinvxing_zinvgong', '2. 子女星与子女宫'),
        ('shengyu_shiji', '3. 生育时机'),
        ('yangyu_jianyi', '4. 养育建议'),
        ('children_rules', '5. 子女规则')
    ]
    
    for section_key, section_title in sections:
        if section_key in input_data:
            doc.add_heading(section_title, 1)
            section_data = input_data[section_key]
            
            # 特殊处理 shengyu_shiji（大运流年）
            if section_key == 'shengyu_shiji':
                # 当前大运
                current_dayun = section_data.get('current_dayun')
                if current_dayun:
                    doc.add_heading('当前大运', 2)
                    add_section(doc, '当前大运信息', current_dayun, level=3)
                    
                    # 当前大运的流年
                    liunians = current_dayun.get('liunians', [])
                    if liunians:
                        doc.add_heading(f'当前大运流年（共{len(liunians)}个）', 2)
                        for i, liunian in enumerate(liunians, 1):
                            year = liunian.get("year", "")
                            stem = liunian.get("stem", "")
                            branch = liunian.get("branch", "")
                            type_display = liunian.get("type_display", liunian.get("type", ""))
                            doc.add_heading(f'流年 {i}: {year}年 {stem}{branch} ({type_display})', 3)
                            # 完整输出所有字段
                            for key, value in liunian.items():
                                p = doc.add_paragraph()
                                p.add_run(f"{key}: ").bold = True
                                if isinstance(value, (dict, list)):
                                    value_str = format_value(value, indent=0, max_length=500)
                                    doc.add_paragraph(value_str, style='List Bullet')
                                else:
                                    p.add_run(str(value))
                    else:
                        doc.add_paragraph("⚠️ 当前大运没有流年数据")
                
                # 关键大运
                key_dayuns = section_data.get('key_dayuns', [])
                if key_dayuns:
                    doc.add_heading(f'关键大运（共{len(key_dayuns)}个）', 2)
                    for i, key_dayun in enumerate(key_dayuns, 1):
                        step = key_dayun.get('step', '')
                        stem = key_dayun.get('stem', '')
                        branch = key_dayun.get('branch', '')
                        doc.add_heading(f'关键大运 {i}: {step}运 {stem}{branch}', 3)
                        add_section(doc, '大运信息', key_dayun, level=4)
                        
                        # 关键大运的流年
                        liunians = key_dayun.get('liunians', [])
                        if liunians:
                            doc.add_heading(f'流年（共{len(liunians)}个）', 4)
                            for j, liunian in enumerate(liunians, 1):
                                year = liunian.get("year", "")
                                stem = liunian.get("stem", "")
                                branch = liunian.get("branch", "")
                                type_display = liunian.get("type_display", liunian.get("type", ""))
                                doc.add_heading(f'流年 {j}: {year}年 {stem}{branch} ({type_display})', 5)
                                # 完整输出所有字段
                                for key, value in liunian.items():
                                    p = doc.add_paragraph()
                                    p.add_run(f"{key}: ").bold = True
                                    if isinstance(value, (dict, list)):
                                        value_str = format_value(value, indent=0, max_length=500)
                                        doc.add_paragraph(value_str, style='List Bullet')
                                    else:
                                        p.add_run(str(value))
                        else:
                            doc.add_paragraph("⚠️ 此大运没有流年数据", style='List Bullet')
                
                # 所有大运
                all_dayuns = section_data.get('all_dayuns', [])
                if all_dayuns:
                    doc.add_heading(f'所有大运（共{len(all_dayuns)}个）', 2)
                    for i, dayun in enumerate(all_dayuns, 1):
                        step = dayun.get('step', '')
                        stem = dayun.get('stem', '')
                        branch = dayun.get('branch', '')
                        doc.add_heading(f'大运 {i}: {step}运 {stem}{branch}', 3)
                        add_section(doc, '大运详情', dayun, level=4)
            else:
                # 其他部分正常处理
                add_section(doc, section_title, section_data, level=2)
    
    # 保存文档
    doc.save(output_file)
    print(f"✅ 文档已生成: {output_file}")


if __name__ == '__main__':
    input_data_file = '/tmp/input_data_1987.json'
    output_file = os.path.expanduser('~/Desktop/子女学习分析-完整input_data数据-1987-09-16.docx')
    
    generate_full_input_data_doc(input_data_file, output_file)

