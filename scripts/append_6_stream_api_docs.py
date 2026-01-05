#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
追加 6 个流式分析接口文档到 Word 文档
1. /bazi/marriage-analysis/stream - 感情婚姻流式分析
2. /career-wealth/stream - 事业财富流式分析
3. /children-study/stream - 子女学习流式分析
4. /health/stream - 身体健康流式分析
5. /general-review/stream - 总评流式分析
6. /smart-fortune/smart-analyze-stream - 智能运势流式分析
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), kwargs[edge].get('sz', '4'))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """设置表格边框"""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'}
            )


def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading


def add_table_with_data(doc, headers, data, col_widths=None):
    """添加带数据的表格"""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10)
    
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data) if cell_data else ''
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)
    
    set_table_borders(table)
    return table


def add_bold_text(doc, text):
    """添加加粗文本"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_stream_api_doc(doc, title, path, method, description, request_params, has_bot_id=True):
    """添加流式接口文档（通用模板）"""
    
    add_heading(doc, f'{title} - {path}', level=1)
    
    # 接口详情
    doc.add_paragraph()
    add_heading(doc, '接口详情', level=2)
    
    interface_details = [
        ('接口路径', path),
        ('接口别名', title),
        ('请求方法', method),
        ('接口描述', description),
        ('备注', 'SSE流式响应：1. 流式返回大模型分析（type: "progress"）2. 返回完成标记（type: "complete"）'),
    ]
    add_table_with_data(doc, ['属性', '值'], interface_details, col_widths=[3, 12])
    
    # 请求参数
    doc.add_paragraph()
    add_heading(doc, '请求参数', level=2)
    add_table_with_data(doc, ['字段名', '类型', '必填', '描述', '示例'], request_params, col_widths=[2.5, 1.5, 1, 7, 2])
    
    # 响应格式
    doc.add_paragraph()
    add_heading(doc, '响应格式', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('SSE流式响应，每行格式：data: {"type": "progress|complete|error", "content": ...}')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    
    # SSE事件类型
    doc.add_paragraph()
    add_bold_text(doc, 'SSE事件类型：')
    
    sse_types = [
        ('progress', '大模型分析进度（字符串片段）'),
        ('complete', '大模型分析完成（完整的分析内容）'),
        ('error', '错误信息'),
    ]
    add_table_with_data(doc, ['类型', '描述'], sse_types, col_widths=[2, 13])
    
    # 响应字段
    doc.add_paragraph()
    add_bold_text(doc, '响应字段说明：')
    
    response_fields = [
        ('type', 'str', '事件类型（progress/complete/error）'),
        ('content', 'str', '内容（分析文本或错误信息）'),
    ]
    add_table_with_data(doc, ['字段名', '类型', '描述'], response_fields, col_widths=[3, 2, 10])


def add_marriage_analysis_stream_doc(doc):
    """添加感情婚姻流式分析接口文档"""
    request_params = [
        ('solar_date', 'str', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', '1990-01-15'),
        ('solar_time', 'str', '是', '出生时间，格式：HH:MM', '12:00'),
        ('gender', 'str', '是', '性别：male(男) 或 female(女)', 'male'),
        ('calendar_type', 'str', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'solar'),
        ('location', 'str', '否', '出生地点（用于时区转换，优先级1）', '北京'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）', '39.90'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）', '116.40'),
        ('bot_id', 'str', '否', 'Coze Bot ID（可选，优先级：参数 > 数据库配置）', '758xxx'),
    ]
    add_stream_api_doc(
        doc,
        title='感情婚姻流式分析',
        path='/bazi/marriage-analysis/stream',
        method='POST',
        description='流式生成感情婚姻分析。包括：命盘总论、配偶特征、感情走势、神煞点睛、建议方向等5个部分',
        request_params=request_params
    )


def add_career_wealth_stream_doc(doc):
    """添加事业财富流式分析接口文档"""
    request_params = [
        ('solar_date', 'str', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', '1990-01-15'),
        ('solar_time', 'str', '是', '出生时间，格式：HH:MM', '12:00'),
        ('gender', 'str', '是', '性别：male(男) 或 female(女)', 'male'),
        ('calendar_type', 'str', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'solar'),
        ('location', 'str', '否', '出生地点（用于时区转换，优先级1）', '北京'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）', '39.90'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）', '116.40'),
        ('bot_id', 'str', '否', 'Coze Bot ID（可选，优先级：参数 > 数据库配置）', '758xxx'),
    ]
    add_stream_api_doc(
        doc,
        title='事业财富流式分析',
        path='/career-wealth/stream',
        method='POST',
        description='流式生成事业财富分析。基于八字数据分析事业发展、财富运势、适合行业等内容',
        request_params=request_params
    )


def add_children_study_stream_doc(doc):
    """添加子女学习流式分析接口文档"""
    request_params = [
        ('solar_date', 'str', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', '1990-01-15'),
        ('solar_time', 'str', '是', '出生时间，格式：HH:MM', '12:00'),
        ('gender', 'str', '是', '性别：male(男) 或 female(女)', 'male'),
        ('calendar_type', 'str', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'solar'),
        ('location', 'str', '否', '出生地点（用于时区转换，优先级1）', '北京'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）', '39.90'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）', '116.40'),
        ('bot_id', 'str', '否', 'Coze Bot ID（可选，优先级：参数 > 数据库配置）', '758xxx'),
    ]
    add_stream_api_doc(
        doc,
        title='子女学习流式分析',
        path='/children-study/stream',
        method='POST',
        description='流式生成子女学习分析。基于八字数据分析子女运势、学业发展、教育建议等内容',
        request_params=request_params
    )


def add_health_stream_doc(doc):
    """添加身体健康流式分析接口文档"""
    request_params = [
        ('solar_date', 'str', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', '1990-01-15'),
        ('solar_time', 'str', '是', '出生时间，格式：HH:MM', '12:00'),
        ('gender', 'str', '是', '性别：male(男) 或 female(女)', 'male'),
        ('calendar_type', 'str', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'solar'),
        ('location', 'str', '否', '出生地点（用于时区转换，优先级1）', '北京'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）', '39.90'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）', '116.40'),
        ('bot_id', 'str', '否', 'Coze Bot ID（可选，优先级：参数 > 数据库配置）', '758xxx'),
    ]
    add_stream_api_doc(
        doc,
        title='身体健康流式分析',
        path='/health/stream',
        method='POST',
        description='流式生成身体健康分析。基于八字数据分析健康状况、易发疾病、养生建议等内容',
        request_params=request_params
    )


def add_general_review_stream_doc(doc):
    """添加总评流式分析接口文档"""
    request_params = [
        ('solar_date', 'str', '是', '阳历日期，格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', '1990-01-15'),
        ('solar_time', 'str', '是', '出生时间，格式：HH:MM', '12:00'),
        ('gender', 'str', '是', '性别：male(男) 或 female(女)', 'male'),
        ('calendar_type', 'str', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'solar'),
        ('location', 'str', '否', '出生地点（用于时区转换，优先级1）', '北京'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）', '39.90'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）', '116.40'),
        ('bot_id', 'str', '否', 'Coze Bot ID（可选，优先级：参数 > 数据库配置）', '758xxx'),
    ]
    add_stream_api_doc(
        doc,
        title='总评流式分析',
        path='/general-review/stream',
        method='POST',
        description='流式生成总评分析。综合八字数据进行全面分析，包括命格特点、运势总览、人生建议等内容',
        request_params=request_params
    )


def add_smart_fortune_stream_doc(doc):
    """添加智能运势流式分析接口文档"""
    
    add_heading(doc, '智能运势流式分析 - /smart-fortune/smart-analyze-stream', level=1)
    
    # 接口详情
    doc.add_paragraph()
    add_heading(doc, '接口详情', level=2)
    
    interface_details = [
        ('接口路径', '/smart-fortune/smart-analyze-stream'),
        ('接口别名', '智能运势流式分析'),
        ('请求方法', 'GET'),
        ('接口描述', '智能运势分析（流式输出版），支持问答式交互'),
        ('备注', '支持两种场景：场景1（点击选择项）返回简短答复+预设问题列表；场景2（点击预设问题/输入问题）返回详细流式回答+相关问题列表'),
    ]
    add_table_with_data(doc, ['属性', '值'], interface_details, col_widths=[3, 12])
    
    # 请求参数
    doc.add_paragraph()
    add_heading(doc, '请求参数（Query Parameters）', level=2)
    
    request_params = [
        ('question', 'str', '否*', '用户问题（场景2必填）', '我的财运如何'),
        ('year', 'int', '是', '出生年份', '1990'),
        ('month', 'int', '是', '出生月份', '1'),
        ('day', 'int', '是', '出生日期', '15'),
        ('hour', 'int', '否', '出生时辰（0-23），默认12', '12'),
        ('gender', 'str', '是', '性别：male(男) 或 female(女)', 'male'),
        ('user_id', 'str', '否*', '用户ID（场景1和场景2必填）', 'user123'),
        ('category', 'str', '否', '分类（如：事业财富、婚姻、健康等）', '事业财富'),
    ]
    add_table_with_data(doc, ['字段名', '类型', '必填', '描述', '示例'], request_params, col_widths=[2.5, 1.5, 1, 7, 2])
    
    # 场景说明
    doc.add_paragraph()
    add_bold_text(doc, '场景说明：')
    
    p = doc.add_paragraph()
    run = p.add_run('场景1：category有值，question为空或等于category → 返回简短答复（100字内）+ 预设问题列表（10-15个）')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    run = p.add_run('场景2：category有值，question有值且不等于category → 返回详细流式回答 + 3个相关问题列表')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    
    # 响应格式
    doc.add_paragraph()
    add_heading(doc, '响应格式', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('SSE流式响应，每行格式：data: {"type": "...", "data": {...}}')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    
    # SSE事件类型
    doc.add_paragraph()
    add_bold_text(doc, 'SSE事件类型：')
    
    sse_types = [
        ('bazi_summary', '八字摘要（基础信息）'),
        ('matched_rules', '匹配的规则列表'),
        ('llm_chunk', 'LLM分析文本片段（流式）'),
        ('llm_complete', 'LLM分析完成'),
        ('preset_questions', '预设问题列表（场景1）'),
        ('related_questions', '相关问题列表（场景2）'),
        ('error', '错误信息'),
        ('end', '流结束标记'),
    ]
    add_table_with_data(doc, ['类型', '描述'], sse_types, col_widths=[3, 12])


def main():
    doc_path = '/Users/zhoudt/Desktop/66666.docx'
    doc = Document(doc_path)
    
    doc.add_page_break()
    
    # 1. 感情婚姻流式分析
    add_marriage_analysis_stream_doc(doc)
    doc.add_page_break()
    
    # 2. 事业财富流式分析
    add_career_wealth_stream_doc(doc)
    doc.add_page_break()
    
    # 3. 子女学习流式分析
    add_children_study_stream_doc(doc)
    doc.add_page_break()
    
    # 4. 身体健康流式分析
    add_health_stream_doc(doc)
    doc.add_page_break()
    
    # 5. 总评流式分析
    add_general_review_stream_doc(doc)
    doc.add_page_break()
    
    # 6. 智能运势流式分析
    add_smart_fortune_stream_doc(doc)
    
    doc.save(doc_path)
    print(f'✅ 6个流式接口文档已成功追加到 {doc_path}')
    print('   1. /bazi/marriage-analysis/stream (感情婚姻流式分析)')
    print('   2. /career-wealth/stream (事业财富流式分析)')
    print('   3. /children-study/stream (子女学习流式分析)')
    print('   4. /health/stream (身体健康流式分析)')
    print('   5. /general-review/stream (总评流式分析)')
    print('   6. /smart-fortune/smart-analyze-stream (智能运势流式分析)')


if __name__ == '__main__':
    main()

