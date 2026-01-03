#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成大模型 input_data 参数文档
从源代码解析所有给大模型传递的 input_data 结构
完整展开所有嵌套字段，并标注每个字段的来源
"""

import sys
import os
import re
import ast
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Tuple

# 导入接口定义
try:
    from llm_input_data_definitions import ALL_INTERFACES
except ImportError:
    # 如果导入失败，使用空列表
    ALL_INTERFACES = []

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    sys.exit(1)


# ==================== 接口定义 ====================

INTERFACES = [
    {
        'name': '八字命理-子女学习',
        'endpoint': '/api/v1/bazi/children-study/stream',
        'function': 'build_children_study_input_data',
        'file': 'server/api/v1/children_study_analysis.py',
        'type': 'structured'  # structured 或 prompt
    },
    {
        'name': '八字命理-身体健康',
        'endpoint': '/api/v1/bazi/health/stream',
        'function': 'build_health_input_data',
        'file': 'server/api/v1/health_analysis.py',
        'type': 'structured'
    },
    {
        'name': '八字命理-事业财富',
        'endpoint': '/api/v1/bazi/career-wealth/stream',
        'function': 'stream_career_wealth_analysis_generator',
        'file': 'server/api/v1/career_wealth_analysis.py',
        'type': 'structured',
        'input_data_location': 'inline'  # inline 表示在函数内部直接构建
    },
    {
        'name': '八字命理-总评',
        'endpoint': '/api/v1/bazi/general-review/stream',
        'function': 'build_general_review_input_data',
        'file': 'server/api/v1/general_review_analysis.py',
        'type': 'structured'
    },
    {
        'name': '八字命理-婚姻',
        'endpoint': '/api/v1/bazi/marriage-analysis/stream',
        'function': 'marriage_analysis_stream_generator',
        'file': 'server/api/v1/marriage_analysis.py',
        'type': 'structured',
        'input_data_location': 'inline'
    },
    {
        'name': '智能运势分析',
        'endpoint': '/api/v1/bazi/fortune/analyze',
        'function': '_build_input_data',
        'file': 'server/services/fortune_llm_client.py',
        'type': 'structured',
        'class': 'FortuneLLMClient'
    },
    {
        'name': '五行占比分析',
        'endpoint': '/api/v1/bazi/wuxing-proportion/stream',
        'function': 'build_llm_prompt',
        'file': 'server/services/wuxing_proportion_service.py',
        'type': 'prompt'
    },
    {
        'name': '喜神忌神分析',
        'endpoint': '/api/v1/bazi/xishen-jishen/stream',
        'function': 'xishen_jishen_stream_generator',
        'file': 'server/api/v1/xishen_jishen.py',
        'type': 'prompt'
    },
    {
        'name': 'LLM生成完整报告',
        'endpoint': '/api/v1/bazi/llm-generate',
        'function': 'build_comprehensive_prompt',
        'file': 'server/services/llm_generate_service.py',
        'type': 'prompt'
    },
    {
        'name': 'Coze AI分析八字',
        'endpoint': '/api/v1/bazi/ai-analyze',
        'function': 'analyze_bazi_with_ai',
        'file': 'server/services/bazi_ai_service.py',
        'type': 'prompt'
    }
]


def extract_input_data_from_source(file_path: str, function_name: str, class_name: str = None) -> Dict[str, Any]:
    """
    从源代码中提取 input_data 结构
    
    Args:
        file_path: 源代码文件路径
        function_name: 函数名
        class_name: 类名（如果是类方法）
        
    Returns:
        dict: 提取的 input_data 结构信息
    """
    project_root = Path(__file__).parent.parent
    full_path = project_root / file_path
    
    if not full_path.exists():
        return {'error': f'文件不存在: {file_path}'}
    
    try:
        with open(full_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        return {'error': f'读取文件失败: {e}'}
    
    # 查找函数定义
    if class_name:
        pattern = rf'class {class_name}.*?def {function_name}\(.*?\):.*?(?=def |class |\Z)'
    else:
        pattern = rf'def {function_name}\(.*?\):.*?(?=def |class |\Z)'
    
    match = re.search(pattern, content, re.DOTALL)
    if not match:
        return {'error': f'未找到函数: {function_name}'}
    
    func_content = match.group(0)
    
    # 查找 input_data = { 的位置
    input_data_pattern = r'input_data\s*=\s*\{'
    input_data_match = re.search(input_data_pattern, func_content)
    
    if not input_data_match:
        # 尝试查找 prompt 构建
        prompt_pattern = r'prompt\s*=.*?["\']'
        prompt_match = re.search(prompt_pattern, func_content, re.DOTALL)
        if prompt_match:
            return {
                'type': 'prompt',
                'content': func_content[prompt_match.start():prompt_match.end() + 500]
            }
        return {'error': '未找到 input_data 或 prompt'}
    
    # 提取 input_data 字典内容（简单提取，不完整解析）
    start_pos = input_data_match.end() - 1
    brace_count = 0
    in_string = False
    string_char = None
    end_pos = start_pos
    
    for i, char in enumerate(func_content[start_pos:], start_pos):
        if char in ('"', "'") and (i == start_pos or func_content[i-1] != '\\'):
            if not in_string:
                in_string = True
                string_char = char
            elif char == string_char:
                in_string = False
                string_char = None
        elif not in_string:
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0:
                    end_pos = i + 1
                    break
    
    input_data_code = func_content[start_pos:end_pos]
    
    # 提取字段和注释
    fields = extract_fields_from_code(input_data_code)
    
    return {
        'type': 'structured',
        'fields': fields,
        'code_snippet': input_data_code[:1000]  # 只保留前1000字符
    }


def extract_fields_from_code(code: str) -> List[Dict[str, Any]]:
    """
    从代码中提取字段信息
    
    Args:
        code: input_data 的代码片段
        
    Returns:
        list: 字段列表
    """
    fields = []
    
    # 匹配字段定义和注释
    # 格式: 'field_name': { 或 'field_name': value,  # 注释
    pattern = r"['\"](\w+)['\"]\s*:\s*([^,}]+)(?:,\s*#\s*(.+?)(?=\n|$))?"
    
    matches = re.finditer(pattern, code, re.MULTILINE)
    
    for match in matches:
        field_name = match.group(1)
        field_value = match.group(2).strip()
        comment = match.group(3) if match.group(3) else ''
        
        # 推断类型
        field_type = infer_type(field_value)
        
        fields.append({
            'name': field_name,
            'type': field_type,
            'description': comment,
            'value': field_value[:100]  # 只保留前100字符
        })
    
    return fields


def infer_type(value: str) -> str:
    """
    推断字段类型
    
    Args:
        value: 字段值（代码片段）
        
    Returns:
        str: 类型
    """
    value = value.strip()
    
    if value.startswith('{'):
        return 'dict'
    elif value.startswith('['):
        return 'list'
    elif value.startswith("'") or value.startswith('"'):
        return 'string'
    elif value in ('True', 'False'):
        return 'boolean'
    elif value.replace('.', '').replace('-', '').isdigit():
        return 'number'
    elif '.' in value and '(' in value:
        return 'function_call'
    else:
        return 'unknown'


def expand_nested_fields(structure: Dict[str, Any], parent_path: str = '') -> List[Dict[str, Any]]:
    """
    递归展开嵌套字段
    
    Args:
        structure: 结构定义
        parent_path: 父路径
        
    Returns:
        list: 扁平化的字段列表
    """
    fields = []
    
    if isinstance(structure, dict):
        if 'fields' in structure:
            # 这是有结构的定义
            for field_name, field_data in structure['fields'].items():
                field_path = f"{parent_path}.{field_name}" if parent_path else field_name
                fields.append({
                    'path': field_path,
                    'type': field_data.get('type', 'unknown'),
                    'description': field_data.get('description', ''),
                    'source': field_data.get('source', '')
                })
                
                # 递归处理嵌套
                if 'nested' in field_data:
                    nested_fields = expand_nested_fields(
                        {'fields': field_data['nested']},
                        field_path
                    )
                    fields.extend(nested_fields)
        else:
            # 这是普通的字典，直接遍历
            for key, value in structure.items():
                field_path = f"{parent_path}.{key}" if parent_path else key
                fields.append({
                    'path': field_path,
                    'type': infer_type(str(value)),
                    'description': '',
                    'source': ''
                })
                
                if isinstance(value, dict):
                    nested_fields = expand_nested_fields(value, field_path)
                    fields.extend(nested_fields)
    
    return fields


def create_word_document(output_path: str):
    """
    创建 Word 文档
    
    Args:
        output_path: 输出文件路径
    """
    doc = Document()
    
    # 设置中文字体
    def set_chinese_font(run, font_name='宋体', size=12):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size)
    
    # 标题
    title = doc.add_heading('大模型 input_data 参数文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    doc.add_paragraph('本文档详细列出了所有给大模型传递的 input_data 结构，包括所有嵌套字段及其数据来源。')
    doc.add_paragraph('生成时间：' + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    doc.add_paragraph('')
    
    # 一、结构化 input_data 接口
    doc.add_heading('一、结构化 input_data 接口', 1)
    
    # 优先使用配置文件中的定义
    if ALL_INTERFACES:
        structured_interfaces = [i for i in ALL_INTERFACES if i.get('type') == 'structured']
    else:
        structured_interfaces = [i for i in INTERFACES if i['type'] == 'structured']
    
    for idx, interface in enumerate(structured_interfaces, 1):
        doc.add_heading(f'{idx}. {interface["name"]}', 2)
        doc.add_paragraph(f'接口路径：{interface["endpoint"]}')
        doc.add_paragraph(f'构建函数：{interface["function"]}')
        doc.add_paragraph(f'文件位置：{interface["file"]}')
        doc.add_paragraph('')
        
        # 如果配置文件中有字段定义，直接使用
        if 'fields' in interface and interface['fields']:
            # 创建表格
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # 表头
            header_cells = table.rows[0].cells
            header_cells[0].text = '字段路径'
            header_cells[1].text = '字段类型'
            header_cells[2].text = '字段说明'
            header_cells[3].text = '数据来源'
            
            # 设置表头样式
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run, '宋体', 12)
                        run.bold = True
            
            # 添加数据行
            for field in interface['fields']:
                row = table.add_row()
                row.cells[0].text = field.get('path', '')
                row.cells[1].text = field.get('type', 'unknown')
                row.cells[2].text = field.get('description', '')
                row.cells[3].text = field.get('source', '')
                
                # 设置单元格字体
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_chinese_font(run, '宋体', 10)
        else:
            # 如果没有配置文件定义，尝试从源代码提取
            result = extract_input_data_from_source(
                interface['file'],
                interface['function'],
                interface.get('class')
            )
            
            if 'error' in result:
                doc.add_paragraph(f'⚠️ 提取失败：{result["error"]}')
            elif result.get('type') == 'structured' and 'fields' in result:
                # 创建表格（使用提取的字段）
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                
                # 表头
                header_cells = table.rows[0].cells
                header_cells[0].text = '字段路径'
                header_cells[1].text = '字段类型'
                header_cells[2].text = '字段说明'
                header_cells[3].text = '数据来源/说明'
                
                # 设置表头样式
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_chinese_font(run, '宋体', 12)
                            run.bold = True
                
                # 添加数据行
                for field in result['fields']:
                    row = table.add_row()
                    row.cells[0].text = field.get('name', '')
                    row.cells[1].text = field.get('type', 'unknown')
                    row.cells[2].text = field.get('description', '')
                    row.cells[3].text = field.get('source', field.get('value', ''))
                    
                    # 设置单元格字体
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                set_chinese_font(run, '宋体', 10)
        
        doc.add_paragraph('')
    
    # 二、文本 prompt 接口
    doc.add_heading('二、文本 prompt 接口', 1)
    
    prompt_interfaces = [i for i in INTERFACES if i['type'] == 'prompt']
    
    for idx, interface in enumerate(prompt_interfaces, 1):
        doc.add_heading(f'{idx}. {interface["name"]}', 2)
        doc.add_paragraph(f'接口路径：{interface["endpoint"]}')
        doc.add_paragraph(f'构建函数：{interface["function"]}')
        doc.add_paragraph(f'文件位置：{interface["file"]}')
        doc.add_paragraph('')
        
        # 提取 prompt 模板
        result = extract_input_data_from_source(
            interface['file'],
            interface['function'],
            interface.get('class')
        )
        
        if 'error' in result:
            doc.add_paragraph(f'⚠️ 提取失败：{result["error"]}')
        elif result.get('type') == 'prompt':
            doc.add_paragraph('Prompt 模板：')
            doc.add_paragraph(result.get('content', '')[:2000], style='Intense Quote')
        
        doc.add_paragraph('')
    
    # 保存文档
    doc.save(output_path)
    print(f"✅ 文档已生成：{output_path}")


if __name__ == '__main__':
    # 输出路径
    desktop_path = Path.home() / 'Desktop'
    output_file = desktop_path / 'HiFate-bazi大模型input_data参数文档.docx'
    
    create_word_document(str(output_file))
    print(f"✅ 文档生成完成：{output_file}")
