#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
追加 /daily-fortune-calendar/stream 接口文档到 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), kwargs[edge].get('sz', '4'))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """设置表格边框"""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 
                top={'val': 'single', 'sz': '4', 'color': '000000'},
                bottom={'val': 'single', 'sz': '4', 'color': '000000'},
                left={'val': 'single', 'sz': '4', 'color': '000000'},
                right={'val': 'single', 'sz': '4', 'color': '000000'}
            )


def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    # 设置中文字体
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading


def add_table_with_data(doc, headers, data, col_widths=None):
    """添加带数据的表格"""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # 设置表头样式
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10)
    
    # 填充数据
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data) if cell_data else ''
            # 设置单元格样式
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)
    
    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)
    
    # 设置边框
    set_table_borders(table)
    
    return table


def main():
    # 打开现有文档
    doc_path = '/Users/zhoudt/Desktop/66666.docx'
    doc = Document(doc_path)
    
    # 添加分页符（可选，确保新接口从新页开始）
    doc.add_page_break()
    
    # ==================== 接口标题 ====================
    add_heading(doc, '流式每日运势日历 - /daily-fortune-calendar/stream', level=1)
    
    # ==================== 接口详情 ====================
    doc.add_paragraph()
    add_heading(doc, '接口详情', level=2)
    
    interface_details = [
        ('接口路径', '/daily-fortune-calendar/stream'),
        ('接口别名', '流式每日运势日历'),
        ('请求方法', 'POST'),
        ('接口描述', '流式查询每日运势日历信息（SSE流式响应）'),
        ('备注', '与 /query 接口相同的输入，但以SSE流式方式返回数据：1. 首先返回完整的每日运势数据（type: "data"）2. 然后流式返回行动建议（type: "progress"）3. 最后返回完成标记（type: "complete"）。返回完整的万年历信息、胎神信息、运势内容、命主信息、五行穿搭、贵人方位、瘟神方位等'),
    ]
    
    add_table_with_data(doc, ['属性', '值'], interface_details, col_widths=[3, 12])
    
    # ==================== 请求参数 ====================
    doc.add_paragraph()
    add_heading(doc, '请求参数', level=2)
    
    request_params = [
        ('date', 'str', '否', '查询日期（可选，默认为今天），格式：YYYY-MM-DD', '2025-01-15'),
        ('solar_date', 'str', '否', '用户生辰阳历日期（可选），格式：YYYY-MM-DD（当calendar_type=lunar时，可为农历日期）', '1990-05-15'),
        ('solar_time', 'str', '否', '用户生辰时间（可选），格式：HH:MM', '14:30'),
        ('gender', 'str', '否', '用户性别（可选）：male(男) 或 female(女)', 'male'),
        ('calendar_type', 'str', '否', '历法类型：solar(阳历) 或 lunar(农历)，默认solar', 'solar'),
        ('location', 'str', '否', '出生地点（用于时区转换，优先级1）', '北京'),
        ('latitude', 'float', '否', '纬度（用于时区转换，优先级2）', '39.90'),
        ('longitude', 'float', '否', '经度（用于时区转换和真太阳时计算，优先级2）', '116.40'),
    ]
    
    add_table_with_data(doc, ['字段名', '类型', '必填', '描述', '示例'], request_params, col_widths=[2.5, 1.5, 1, 7, 2])
    
    # ==================== 响应格式 ====================
    doc.add_paragraph()
    add_heading(doc, '响应格式', level=2)
    
    # 添加说明文本
    p = doc.add_paragraph()
    run = p.add_run('SSE流式响应，每行格式：data: {"type": "data|progress|complete|error", "content": ...}')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    
    # SSE事件类型表格
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('SSE事件类型：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    sse_types = [
        ('data', '完整的每日运势数据（JSON对象）'),
        ('progress', '行动建议生成进度（字符串片段）'),
        ('complete', '行动建议完成（完整的行动建议内容）'),
        ('error', '错误信息'),
    ]
    add_table_with_data(doc, ['类型', '描述'], sse_types, col_widths=[2, 13])
    
    # data类型的content字段结构
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('data类型的content字段结构：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    response_fields = [
        ('success', 'bool', '是否成功'),
        ('solar_date', 'str', '当前阳历日期（如：2025年1月15日）'),
        ('lunar_date', 'str', '当前阴历日期（如：腊月十六）'),
        ('weekday', 'str', '星期几（中文）（如：星期三）'),
        ('weekday_en', 'str', '星期几（英文）（如：Wednesday）'),
        ('year_pillar', 'str', '年柱（如：甲辰年）'),
        ('month_pillar', 'str', '月柱（如：戊子月）'),
        ('day_pillar', 'str', '日柱（如：乙卯日）'),
        ('yi', 'list', '宜（如：["解除", "扫舍"]）'),
        ('ji', 'list', '忌（如：["嫁娶", "入殓"]）'),
        ('luck_level', 'str', '吉凶等级（如：中等）'),
        ('deities', 'dict', '神煞方位信息'),
        ('  deities.xishen', 'str', '喜神方位'),
        ('  deities.caishen', 'str', '财神方位'),
        ('  deities.fushen', 'str', '福神方位'),
        ('  deities.taishen', 'str', '胎神方位'),
        ('chong_he_sha', 'dict', '冲合煞信息'),
        ('  chong_he_sha.chong', 'str', '冲'),
        ('  chong_he_sha.he', 'str', '合'),
        ('  chong_he_sha.sha', 'str', '煞'),
        ('jianchu', 'dict', '建除信息'),
        ('  jianchu.name', 'str', '建除名称（如：危）'),
        ('  jianchu.energy', 'int', '能量分数（如：90）'),
        ('  jianchu.summary', 'str', '能量小结内容'),
        ('taishen', 'str', '胎神方位（如：占门厕外正东）'),
        ('taishen_explanation', 'str', '胎神解释'),
        ('jiazi_fortune', 'str', '整体运势（六十甲子运势内容）'),
        ('shishen_hint', 'str', '十神提示（需要用户生辰信息才有值）'),
        ('zodiac_relations', 'str', '生肖简运（生肖刑冲破害关系）'),
        ('master_info', 'dict', '命主信息'),
        ('  master_info.rizhu', 'str', '日主（如：甲木）'),
        ('  master_info.today_shishen', 'str', '今日十神（如：比肩）'),
        ('wuxing_wear', 'str', '五行穿搭（逗号分隔，如：绿色,青色）'),
        ('guiren_fangwei', 'str', '贵人方位（逗号分隔，如：东北,正南）'),
        ('wenshen_directions', 'str', '瘟神方位（逗号分隔，如：正西,西南）'),
        ('error', 'str', '错误信息（可选，失败时返回）'),
    ]
    
    add_table_with_data(doc, ['字段名', '类型', '描述'], response_fields, col_widths=[4, 1.5, 9.5])
    
    # progress/complete/error 类型说明
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('progress/complete/error类型的content字段：')
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    other_types = [
        ('progress.content', 'str', '行动建议生成的文本片段（如："宜："、"今日适合..."）'),
        ('complete.content', 'str', '完整的行动建议内容'),
        ('error.content', 'str', '错误信息（如："获取每日运势失败"）'),
    ]
    add_table_with_data(doc, ['字段名', '类型', '描述'], other_types, col_widths=[4, 1.5, 9.5])
    
    # 保存文档
    doc.save(doc_path)
    print(f'✅ 接口文档已成功追加到 {doc_path}')


if __name__ == '__main__':
    main()

