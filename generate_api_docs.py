#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_interface_doc(doc, interface_num, interface_path, interface_alias, method, description, remarks):
    title = doc.add_paragraph()
    title.add_run(f"{interface_num}. {interface_path} ({interface_alias})").bold = True
    doc.add_paragraph("接口详情")
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Light Grid Accent 1'
    details = [
        ("接口路径", interface_path),
        ("接口别名", interface_alias),
        ("请求方法", method),
        ("接口描述", description),
        ("备注", remarks)
    ]
    for i, (label, value) in enumerate(details):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = value
    doc.add_paragraph()

def add_request_params_table(doc, params):
    doc.add_paragraph("请求参数")
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = "字段名"
    header_cells[1].text = "类型"
    header_cells[2].text = "必填"
    header_cells[3].text = "描述"
    header_cells[4].text = "示例"
    for param in params:
        row_cells = table.add_row().cells
        row_cells[0].text = param['name']
        row_cells[1].text = param['type']
        row_cells[2].text = param['required']
        row_cells[3].text = param['description']
        row_cells[4].text = param.get('example', '')
    doc.add_paragraph()

def add_response_table(doc, response_fields):
    doc.add_paragraph("响应格式")
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = "字段名"
    header_cells[1].text = "类型"
    header_cells[2].text = "描述"
    for field in response_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field['name']
        row_cells[1].text = field['type']
        row_cells[2].text = field['description']
    doc.add_paragraph()

doc = Document('/Users/zhoudt/Desktop/66666.docx')
doc.add_page_break()

# 接口1: 面相分析 V2
add_interface_doc(doc, 2, "/api/v2/face/analyze", "AI智能面相分析V2", "POST",
    "上传人脸照片进行AI智能面相分析，支持宫位、六亲、十神等多种分析类型，可结合八字信息进行综合分析",
    "返回完整的面相分析结果，包括三停分析、五眼分析、十三宫位分析、六亲分析、十神分析等详细信息")

face_request_params = [
    {'name': 'image_base64', 'type': 'str', 'required': '是', 'description': 'base64编码的图片数据（必填）', 'example': 'iVBORw0KGgoAAAANSUhEUgAA...'},
    {'name': 'filename', 'type': 'str', 'required': '否', 'description': '文件名（可选），默认 "face.jpg"', 'example': 'face.jpg'},
    {'name': 'content_type', 'type': 'str', 'required': '否', 'description': '内容类型（可选），默认 "image/jpeg"', 'example': 'image/jpeg'},
    {'name': 'analysis_types', 'type': 'str', 'required': '否', 'description': '分析类型（可选），逗号分隔，默认 "gongwei,liuqin,shishen"。可选值：gongwei(宫位)、liuqin(六亲)、shishen(十神)、features(特征)', 'example': 'gongwei,liuqin,shishen'},
    {'name': 'birth_year', 'type': 'int', 'required': '否', 'description': '出生年份（可选），用于结合八字分析', 'example': '1990'},
    {'name': 'birth_month', 'type': 'int', 'required': '否', 'description': '出生月份（可选），用于结合八字分析', 'example': '1'},
    {'name': 'birth_day', 'type': 'int', 'required': '否', 'description': '出生日期（可选），用于结合八字分析', 'example': '15'},
    {'name': 'birth_hour', 'type': 'int', 'required': '否', 'description': '出生时辰（可选），用于结合八字分析', 'example': '12'},
    {'name': 'gender', 'type': 'str', 'required': '否', 'description': '性别（可选），用于结合八字分析。可选值：male(男)、female(女)', 'example': 'male'}
]
add_request_params_table(doc, face_request_params)

face_response_fields = [
    {'name': 'success', 'type': 'bool', 'description': '是否成功'},
    {'name': 'data', 'type': 'dict', 'description': '分析结果数据'},
    {'name': 'data.face_detected', 'type': 'bool', 'description': '是否检测到人脸'},
    {'name': 'data.landmarks', 'type': 'dict', 'description': '关键点信息，包含MediaPipe关键点数量和映射特征点数量'},
    {'name': 'data.santing_analysis', 'type': 'dict', 'description': '三停分析结果，包含上停、中停、下停的比例和评价'},
    {'name': 'data.wuyan_analysis', 'type': 'dict', 'description': '五眼分析结果，包含五个眼睛宽度和脸宽的测量值'},
    {'name': 'data.gongwei_analysis', 'type': 'list', 'description': '十三宫位分析列表，每个宫位包含名称、位置、特征、断语等信息'},
    {'name': 'data.liuqin_analysis', 'type': 'list', 'description': '六亲分析列表，包含父、母、兄弟、子女等关系的分析'},
    {'name': 'data.shishen_analysis', 'type': 'list', 'description': '十神分析列表，包含比肩、劫财、食神、伤官等十神的分析'},
    {'name': 'data.overall_summary', 'type': 'str', 'description': '综合面相分析总结'},
    {'name': 'data.birth_info', 'type': 'dict', 'description': '生辰信息（如果提供了生辰参数），包含year、month、day、hour、gender等字段'}
]
add_response_table(doc, face_response_fields)

doc.add_page_break()

# 接口2: 办公桌风水分析
add_interface_doc(doc, 3, "/api/v2/desk-fengshui/analyze", "办公桌风水分析", "POST",
    "上传办公桌照片进行AI智能风水分析，自动识别物品和位置，匹配风水规则，为每个物品生成详细分析，提供三级建议体系",
    "返回完整的办公桌风水分析结果，包括检测到的物品列表、每个物品的详细分析、三级建议体系（必须调整/建议添加/可选优化）、综合评分等")

desk_request_params = [
    {'name': 'image_base64', 'type': 'str', 'required': '是', 'description': 'base64编码的办公桌照片数据（必填）', 'example': 'iVBORw0KGgoAAAANSUhEUgAA...'},
    {'name': 'filename', 'type': 'str', 'required': '否', 'description': '文件名（可选），默认 "desk.jpg"', 'example': 'desk.jpg'},
    {'name': 'content_type', 'type': 'str', 'required': '否', 'description': '内容类型（可选），默认 "image/jpeg"', 'example': 'image/jpeg'}
]
add_request_params_table(doc, desk_request_params)

desk_response_fields = [
    {'name': 'success', 'type': 'bool', 'description': '是否成功'},
    {'name': 'data', 'type': 'dict', 'description': '分析结果数据'},
    {'name': 'data.items', 'type': 'list', 'description': '检测到的物品列表，每个物品包含name(英文名)、label(中文名)、confidence(置信度)、bbox(边界框)、position(位置信息)等'},
    {'name': 'data.item_analyses', 'type': 'list', 'description': '每个物品的详细风水分析列表，包含name、label、current_position、is_position_ideal、analysis(详细分析)、suggestion(调整建议)等'},
    {'name': 'data.recommendations', 'type': 'dict', 'description': '三级建议体系，包含must_adjust(必须调整)、should_add(建议添加)、optional_optimize(可选优化)三个分类'},
    {'name': 'data.recommendations.must_adjust', 'type': 'list', 'description': '必须调整的建议列表（违反禁忌的物品）'},
    {'name': 'data.recommendations.should_add', 'type': 'list', 'description': '建议添加的物品列表'},
    {'name': 'data.recommendations.optional_optimize', 'type': 'list', 'description': '可选优化的建议列表'},
    {'name': 'data.adjustments', 'type': 'list', 'description': '调整建议列表（需要移动的物品），包含item、item_label、current_position、ideal_position、reason、priority、action等'},
    {'name': 'data.additions', 'type': 'list', 'description': '增加建议列表（建议添加的物品），包含item、item_label、position、reason、suggestion、priority、action、element等'},
    {'name': 'data.removals', 'type': 'list', 'description': '删除建议列表（不宜摆放的物品），包含item、item_label、current_position、reason、priority、action、suggestion等'},
    {'name': 'data.score', 'type': 'int', 'description': '综合评分（0-100），表示办公桌风水布局的整体评分'},
    {'name': 'data.summary', 'type': 'str', 'description': '分析总结，包含整体评价和主要建议'},
    {'name': 'error', 'type': 'str', 'description': '错误信息（仅在success为false时返回）'}
]
add_response_table(doc, desk_response_fields)

doc.save('/Users/zhoudt/Desktop/66666.docx')
print("完成！接口文档已追加到文档")
