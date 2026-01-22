#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成办公桌风水流式接口文档并追加到Word文档
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_interface_doc(doc, interface_num, interface_path, interface_alias, method, description, remarks):
    title = doc.add_paragraph()
    title.add_run(f"{interface_num}. {interface_path} ({interface_alias})").bold = True
    doc.add_paragraph("接口详情")
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Light Grid Accent 1'
    details = [
        ("接口路径", interface_path),
        ("接口别名", interface_alias),
        ("请求方法", method),
        ("接口描述", description),
        ("备注", remarks)
    ]
    for i, (label, value) in enumerate(details):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = value
    doc.add_paragraph()

def add_request_params_table(doc, params):
    doc.add_paragraph("请求参数")
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = "字段名"
    header_cells[1].text = "类型"
    header_cells[2].text = "必填"
    header_cells[3].text = "描述"
    header_cells[4].text = "示例"
    for param in params:
        row_cells = table.add_row().cells
        row_cells[0].text = param['name']
        row_cells[1].text = param['type']
        row_cells[2].text = param['required']
        row_cells[3].text = param['description']
        row_cells[4].text = param.get('example', '')
    doc.add_paragraph()

def add_response_table(doc, response_fields):
    doc.add_paragraph("响应格式")
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = "字段名"
    header_cells[1].text = "类型"
    header_cells[2].text = "描述"
    for field in response_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field['name']
        row_cells[1].text = field['type']
        row_cells[2].text = field['description']
    doc.add_paragraph()

# 打开现有文档
doc = Document('/Users/zhoudt/Desktop/66666.docx')
doc.add_page_break()

# 办公桌风水流式接口
add_interface_doc(doc, 5, "/api/v2/desk-fengshui/analyze/stream", "办公桌风水分析-流式", "POST",
    "上传办公桌照片进行AI智能风水分析，先返回基础分析数据，然后流式返回大模型整合分析结果。使用Server-Sent Events (SSE)格式实时推送分析内容",
    "流式接口，响应格式为text/event-stream。先返回完整的基础分析数据（type: 'data'），然后流式返回大模型整合分析（type: 'progress'），最后返回完成信号（type: 'complete'）")

# 请求参数
desk_stream_request_params = [
    {'name': 'image', 'type': 'File (multipart/form-data)', 'required': '是', 'description': '办公桌照片文件（必填），支持JPG、PNG、WEBP格式', 'example': 'desk.jpg'},
    {'name': 'bot_id', 'type': 'str', 'required': '否', 'description': 'Bot ID（可选），用于指定使用的大模型Bot，默认从数据库配置读取', 'example': '7597409425955127336'}
]
add_request_params_table(doc, desk_stream_request_params)

# 响应格式说明
doc.add_paragraph("响应说明")
doc.add_paragraph("本接口使用Server-Sent Events (SSE)格式，响应类型为text/event-stream。响应数据以'data: '开头的行，每行包含一个JSON对象。")
doc.add_paragraph()

# 响应字段
desk_stream_response_fields = [
    {'name': 'type: data', 'type': 'object', 'description': '基础分析数据（第一个返回的数据块），包含完整的办公桌风水分析结果'},
    {'name': 'type: data.content.success', 'type': 'bool', 'description': '是否成功'},
    {'name': 'type: data.content.items', 'type': 'list', 'description': '检测到的物品列表，每个物品包含name(英文名)、label(中文名)、confidence(置信度)、bbox(边界框)、position(位置信息)等'},
    {'name': 'type: data.content.item_analyses', 'type': 'list', 'description': '每个物品的详细风水分析列表，包含name、label、current_position、is_position_ideal、analysis(详细分析)、suggestion(调整建议)等'},
    {'name': 'type: data.content.recommendations', 'type': 'dict', 'description': '三级建议体系，包含must_adjust(必须调整)、should_add(建议添加)、optional_optimize(可选优化)三个分类'},
    {'name': 'type: data.content.recommendations.must_adjust', 'type': 'list', 'description': '必须调整的建议列表（违反禁忌的物品）'},
    {'name': 'type: data.content.recommendations.should_add', 'type': 'list', 'description': '建议添加的物品列表'},
    {'name': 'type: data.content.recommendations.optional_optimize', 'type': 'list', 'description': '可选优化的建议列表'},
    {'name': 'type: data.content.adjustments', 'type': 'list', 'description': '调整建议列表（需要移动的物品），包含item、item_label、current_position、ideal_position、reason、priority、action等'},
    {'name': 'type: data.content.additions', 'type': 'list', 'description': '增加建议列表（建议添加的物品），包含item、item_label、position、reason、suggestion、priority、action、element等'},
    {'name': 'type: data.content.removals', 'type': 'list', 'description': '删除建议列表（不宜摆放的物品），包含item、item_label、current_position、reason、priority、action、suggestion等'},
    {'name': 'type: data.content.score', 'type': 'int', 'description': '综合评分（0-100），表示办公桌风水布局的整体评分'},
    {'name': 'type: data.content.summary', 'type': 'str', 'description': '分析总结，包含整体评价和主要建议'},
    {'name': 'type: progress', 'type': 'object', 'description': '流式LLM输出数据块（多个），实时推送大模型生成的分析内容'},
    {'name': 'type: progress.content', 'type': 'str', 'description': 'LLM生成的分析文本片段（增量内容）'},
    {'name': 'type: progress.statusText', 'type': 'str', 'description': '状态文本（可选），用于显示当前处理状态'},
    {'name': 'type: complete', 'type': 'object', 'description': '完成信号（最后一个数据块），表示流式输出已完成'},
    {'name': 'type: complete.content', 'type': 'str', 'description': '完整的LLM分析内容（可选），包含所有流式输出的完整文本'},
    {'name': 'type: error', 'type': 'object', 'description': '错误信息（如果发生错误）'},
    {'name': 'type: error.content', 'type': 'str', 'description': '错误描述信息'}
]
add_response_table(doc, desk_stream_response_fields)

# 添加使用示例
doc.add_paragraph("使用示例")
doc.add_paragraph("1. 使用curl调用（multipart/form-data）：")
example_code = doc.add_paragraph()
example_code.add_run('curl -X POST "http://localhost:8001/api/v2/desk-fengshui/analyze/stream" \\\n  -F "image=@desk.jpg"')
example_code.style = 'No Spacing'
example_code_font = example_code.runs[0].font
example_code_font.name = 'Courier New'
example_code_font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph("2. 前端JavaScript调用示例：")
js_example = doc.add_paragraph()
js_example.add_run('const formData = new FormData();\nformData.append("image", file);\n\nconst response = await fetch("/api/v2/desk-fengshui/analyze/stream", {\n  method: "POST",\n  body: formData\n});\n\nconst reader = response.body.getReader();\nconst decoder = new TextDecoder();\nlet buffer = "";\n\nwhile (true) {\n  const { done, value } = await reader.read();\n  if (done) break;\n  \n  buffer += decoder.decode(value, { stream: true });\n  const lines = buffer.split("\\n");\n  buffer = lines.pop() || "";\n  \n  for (const line of lines) {\n    if (line.startsWith("data: ")) {\n      const data = JSON.parse(line.substring(6));\n      if (data.type === "data") {\n        // 处理基础分析数据\n        console.log("基础数据:", data.content);\n        console.log("评分:", data.content.score);\n        console.log("物品数量:", data.content.items.length);\n      } else if (data.type === "progress") {\n        // 处理流式LLM输出\n        console.log("LLM内容:", data.content);\n      } else if (data.type === "complete") {\n        // 流式输出完成\n        console.log("完成");\n      }\n    }\n  }\n}')
js_example.style = 'No Spacing'
js_example_font = js_example.runs[0].font
js_example_font.name = 'Courier New'
js_example_font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph("注意事项：")
notes = [
    "1. 本接口使用SSE格式，需要客户端支持流式读取",
    "2. 响应数据以'data: '开头，每行一个JSON对象",
    "3. 第一个数据块type为'data'，包含完整的基础分析结果（与普通接口相同）",
    "4. 后续数据块type为'progress'，包含流式LLM输出片段",
    "5. 最后一个数据块type为'complete'，表示流式输出完成",
    "6. 如果发生错误，会返回type为'error'的数据块",
    "7. 基础分析数据与普通接口(/api/v2/desk-fengshui/analyze)返回的数据结构相同",
    "8. LLM整合分析基于基础分析数据生成，提供更深入的综合解读和调整建议",
    "9. 分析超时时间为90秒，建议上传图片大小不超过10MB",
    "10. 本接口不需要提供八字信息，专注于办公桌风水布局分析"
]
for note in notes:
    doc.add_paragraph(note, style='List Bullet')

doc.save('/Users/zhoudt/Desktop/66666.docx')
print("✅ 完成！办公桌风水流式接口文档已追加到文档")
